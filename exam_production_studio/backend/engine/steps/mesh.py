"""细目表步骤（阶段五 steps/mesh，仅考纲百套卷，D5）。

依据规划表分组产出：
- 每个专题 → 一份 专题训练卷_第x卷_{省份}_{专题名}.docx，覆盖该专题下全部考点；
- 每门课程 → 课程综合卷_第x卷_{省份}_{课程名}.docx（默认 1 份骨架，覆盖该课程全部考点）。

细目表字段（§4.2）：卷别/卷号/地区/课程名称/专题名称/题号/题型/难度/考查内容/
对应考点/对应知识点/考纲要求/出题意图。逐题蓝图由后续 CD/LLM 富化，这里产出可追溯骨架。
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from engine import repo

_FIELDS = ["题号", "题型", "难度", "考查内容", "对应考点", "对应知识点", "考纲要求", "出题意图"]


def _clean(name: str) -> str:
    for ch in '\\/:*?"<>|':
        name = name.replace(ch, "")
    return name.strip()


def _write_mesh_doc(out_dir: Path, kind: str, vol_no: Any, province: str, course: str,
                    theme: str, points: list[dict[str, Any]]) -> Path:
    doc = Document()
    head = doc.add_paragraph()
    head.add_run(f"{kind}细目表").bold = True
    doc.add_paragraph(f"卷别：{kind}　卷号：第{vol_no}卷　地区：{province}")
    doc.add_paragraph(f"课程名称：{course}　专题名称：{theme}")

    table = doc.add_table(rows=1, cols=len(_FIELDS))
    table.style = "Table Grid"
    for i, f in enumerate(_FIELDS):
        table.rows[0].cells[i].text = f

    qno = 0
    for p in points:
        point_name = p.get("topic") or p.get("point_name") or ""
        knowledge = p.get("point_name") or ""
        qno += 1
        cells = table.add_row().cells
        cells[0].text = str(qno)
        cells[1].text = "综合"        # 题型：待 CD 逐题蓝图细化
        cells[2].text = "适中"
        cells[3].text = f"考查「{point_name}」相关内容"
        cells[4].text = point_name    # 对应考点 = 规划表 C
        cells[5].text = knowledge     # 对应知识点 = 规划表 D
        cells[6].text = "理解"
        cells[7].text = "覆盖该考点核心出题角度"

    fname = _clean(f"{kind}_第{vol_no}卷_{province}_{theme or course}")
    p = out_dir / f"{fname}.docx"
    doc.save(str(p))
    return p


def gen_mesh(ctx, rows: list[dict[str, Any]] | None = None) -> list[Path]:
    rows = rows or repo.get_papers(ctx.project_id)
    # 仅按考点训练卷分组（专题卷/综合卷为其聚合，不作为细目表分组来源，避免重复计入）
    rows = [r for r in rows
            if (r.get("meta") or {}).get("paper_subtype", r.get("paper_type")) in ("考点训练卷", "", None)]
    out_dir = ctx.dir("生产规划")
    out_dir.mkdir(parents=True, exist_ok=True)
    province = ctx.province or ""
    paths: list[Path] = []

    # 按 (课程, 专题) 分组 → 专题训练卷细目表
    themes: dict[tuple, dict[str, Any]] = {}
    courses: dict[str, dict[str, Any]] = {}
    for r in rows:
        meta = r.get("meta") or {}
        course = meta.get("course") or r.get("module") or ctx.course or "课程"
        theme = meta.get("theme") or "综合"
        tvol = meta.get("theme_vol_no") or ""
        crange = meta.get("course_vol_range")
        tk = (course, theme)
        themes.setdefault(tk, {"vol": tvol, "points": []})["points"].append(r)
        c = courses.setdefault(course, {"range": crange, "points": []})
        c["points"].append(r)
        if crange and not c["range"]:
            c["range"] = crange

    for (course, theme), info in themes.items():
        paths.append(_write_mesh_doc(out_dir, "专题训练卷", info["vol"], province, course, theme, info["points"]))

    for course, info in courses.items():
        crange = info.get("range")
        vol = f"{crange[0]}-{crange[1]}" if crange and crange[1] > crange[0] else (crange[0] if crange else "")
        paths.append(_write_mesh_doc(out_dir, "课程综合卷", vol, province, course, "课程综合", info["points"]))

    # 导出细目表到 桌面/输出结果/生产规划/{产品名}/{省份}_{考类}/
    from engine import archive
    for p in paths:
        archive.export_planning_artifact(ctx, p)
    return paths
