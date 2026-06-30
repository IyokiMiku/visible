# -*- coding: utf-8 -*-
"""
文档模板套用工具 v3.2 — 支持多课程/子目录/主题匹配

用法: python _apply_template.py
"""

import copy, re, os, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ═════════ XML 常量 ═════════
W_R = qn('w:r'); W_T = qn('w:t'); W_PPR = qn('w:pPr'); W_RPR = qn('w:rPr')

# ═════════ 正则 ═════════
FILENAME_RE = re.compile(
    r'第(\d+)练\s+(.+?)\s+(.+?)[（(](高职分类考试|高职分类考试)[）)]'
    r'《(.+?)》[（(](.+?)[）)]\s*一课一练'
)
CN_NUM = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
SECTION_RE = re.compile(r'^[（(]?\s*([一二三四五六七八九十]+)\s*[）)、．.]?\s*(.+)$')


# ═════════ 规划表 ═════════
def parse_section(text):
    text = text.strip()
    m = SECTION_RE.match(text)
    if m: return CN_NUM.get(m.group(1), 0), m.group(2).strip()
    return None, text if text else ''

def normalize_topic(s):
    s = re.sub(r'[（(][一二三四五六七八九十]+[)）]', '', s)
    s = re.sub(r'[、/·\s]', '', s).strip()
    s = s.replace('等速万向节', '等角速万向节')
    return s

def _fuzzy_key(s):
    """超模糊查找键：去所有连接词和符号"""
    return re.sub(r'[与和及、/·\s]', '', s)

def load_planning_table(xlsx_path):
    from openpyxl import load_workbook
    wb = load_workbook(xlsx_path)
    ws = wb[wb.sheetnames[0]]
    rows, cur_course, cur_sec_num, cur_sec_name = [], '', None, ''
    for row in ws.iter_rows(min_row=6, max_row=ws.max_row, values_only=True):
        vals = [str(v).strip() if v else '' for v in row[:8]]
        v0 = vals[0]
        if v0.startswith('课程') and '：' in v0: cur_course = v0; continue
        try: int(v0)
        except ValueError:
            if v0 and v0 != '序号' and '考纲' not in v0: cur_sec_num, cur_sec_name = parse_section(v0)
            continue
        rows.append({'序号':int(v0),'考纲知识点':vals[1],'试卷主题':vals[2],'级别':vals[3],
                     '题型':vals[4],'难度':vals[5],'套数':int(vals[6] or 1),'考纲标号':vals[7],
                     '节号':cur_sec_num,'节主题':cur_sec_name,'课程':cur_course})
    # 累计练号映射
    lian_map, rl = {}, 1
    for r in rows:
        for _ in range(r['套数']): lian_map[rl] = r; rl += 1
    # 主题映射
    topic_map = {}
    for r in rows:
        k = normalize_topic(r['试卷主题'])
        if k not in topic_map: topic_map[k] = r
    return lian_map, topic_map, rows


# ═════════ 查表 ═════════
def lookup_plan(seq, topic_raw, lian_map, topic_map, allow_lian=True):
    """查规划表：主题匹配优先，可选练号兜底"""
    k = normalize_topic(topic_raw)
    fk = _fuzzy_key(k)
    # 1. 精确主题匹配
    row = topic_map.get(k)
    # 2. 模糊子串
    if row is None:
        for tk, tv in topic_map.items():
            if k in tk or tk in k: row = tv; break
    # 2b. 子串匹配（去掉连接词后）
    if row is None:
        for tk, tv in topic_map.items():
            nk = _fuzzy_key(k); ntk = _fuzzy_key(tk)
            if nk and ntk and (nk in ntk or ntk in nk): row = tv; break
    # 2c. 查别名表
    if row is None:
        for tk, tv in topic_map.items():
            tfk = _fuzzy_key(tk)
            if fk and tfk and (fk in tfk or tfk in fk): row = tv; break
    # 4. 练号兜底
    if row is None and allow_lian:
        row = lian_map.get(seq)
    if row:
        sn, nm = row['节号'], row['节主题']
        st = f'第{sn}节 {nm}' if sn is not None else nm
        return row, st, row['考纲知识点']
    return None, '', '相关知识点'


# ═════════ 文件名解析 ═════════
def parse_filename(fp):
    fn = os.path.basename(fp)
    m = FILENAME_RE.search(fn)
    if not m: raise ValueError(f'无法解析: {fn}')
    ef = m.group(6).strip()
    return (int(m.group(1)), m.group(2).strip(), m.group(3).strip(),
            m.group(4), m.group(5).strip(), ef,
            ef.split('·')[0].strip() if '·' in ef else ef)


# ═════════ XML工具 ═════════
def _set_run_text(re, txt):
    ts = re.findall(W_T)
    if ts:
        ts[0].text = txt; ts[0].set(qn('xml:space'),'preserve')
        for e in ts[1:]: re.remove(e)
    else:
        t = OxmlElement('w:t'); t.text = txt; t.set(qn('xml:space'),'preserve'); re.append(t)

def _clear_runs(p):
    rs = p.findall(W_R)
    if not rs: r = OxmlElement('w:r'); p.append(r); return r
    for e in rs[1:]: p.remove(e)
    return rs[0]


# ═════════ 编写说明文本 ═════════
def _ed_p1(province, exam_type, textbook, edition_full):
    return (f'编写说明：考虑到中职学生普遍基础知识相对薄弱的情况，我们依据支架式教学理念，'
            f'精心编制了{province}（{exam_type}）《{textbook}》（{edition_full}）一课一练。'
            f'专辑里的每一份练习，都与课堂所授知识点紧密相关，题目围绕课堂所学知识点呈现。'
            f'目的在于激发学生的学习兴趣，培养他们的学习自觉性，帮助学生扎实掌握课程的基本概念与基本方法，'
            f'为他们后续的逐步提升奠定坚实基础。')

def _ed_p2(province, exam_type, textbook, edition_full, seq, exam_points):
    return (f'本卷是{province}（{exam_type}）《{textbook}》（{edition_full}）'
            f'一课一练的第{seq}练，内容涵盖{exam_points}。')


# ═════════ 模板更新 ═════════
def _update_table(doc, province, exam_type, textbook, edition_full, seq, exam_points):
    if not doc.tables: return
    cell = doc.tables[0].rows[0].cells[0]; paras = cell.paragraphs
    p1 = _ed_p1(province, exam_type, textbook, edition_full)
    if paras:
        pe = paras[0]._element; _clear_runs(pe); _set_run_text(pe.find(W_R), p1)
    p2 = _ed_p2(province, exam_type, textbook, edition_full, seq, exam_points)
    if len(paras) >= 2:
        pe = paras[1]._element; _clear_runs(pe); _set_run_text(pe.find(W_R), p2)
    elif paras:
        p1e = paras[0]._element; rpr = p1e.find('.//'+W_RPR); ppr = p1e.find(W_PPR)
        p2e = OxmlElement('w:p')
        if ppr is not None: p2e.append(copy.deepcopy(ppr))
        nr = OxmlElement('w:r')
        if rpr is not None: nr.append(copy.deepcopy(rpr))
        _set_run_text(nr, p2); p2e.append(nr); cell._element.append(p2e)


def _update_titles(body, province, exam_type, textbook, edition_full, seq, topic, section_title):
    ch = list(body); p2i = None
    for i,c in enumerate(ch):
        if etree.QName(c.tag).localname == 'p' and '一课一练' in ''.join(list(c.itertext())):
            p2i = i; break
    if p2i is None: return
    p = ch[p2i]; r = _clear_runs(p); _set_run_text(r, f'{province}（{exam_type}）一课一练')
    if p2i+1 < len(ch):
        p3 = ch[p2i+1]
        if etree.QName(p3.tag).localname == 'p':
            r = _clear_runs(p3); _set_run_text(r, f'《{textbook}》（{edition_full}）')
    if p2i+2 < len(ch):
        p4 = ch[p2i+2]
        if etree.QName(p4.tag).localname == 'p':
            r = _clear_runs(p4); _set_run_text(r, section_title if section_title else '')
            rpr_e = r.find(W_RPR)
            for txt in ['  ', f'第{seq}练', f'  {topic}']:
                nr = OxmlElement('w:r')
                if rpr_e is not None: nr.append(copy.deepcopy(rpr_e))
                _set_run_text(nr, txt); p4.append(nr)


def _insert_separator(doc, img_path):
    if not os.path.exists(img_path): return
    body = doc.element.body; ch = list(body)
    ti = next((i for i,c in enumerate(ch) if etree.QName(c.tag).localname == 'tbl'), None)
    if ti is None: return
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ip.paragraph_format.space_before = Pt(0); ip.paragraph_format.space_after = Pt(0)
    ip.add_run().add_picture(img_path, width=Cm(4.60))
    ie = ip._element; body.remove(ie); list(body)[ti].addnext(ie)
    # 删空行
    ch = list(body)
    for i,c in enumerate(ch):
        if etree.QName(c.tag).localname == 'p':
            if list(c.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')):
                while i+1 < len(body):
                    nx = list(body)[i+1]
                    if etree.QName(nx.tag).localname == 'p' and not ''.join(list(nx.itertext())).strip():
                        body.remove(nx)
                    else: break
                break


# ═════════ 内容追加 ═════════
HEADER_RE = re.compile(r'^[一二三四五六七八九十]+\s*[、.．]')
QNUM_RE = re.compile(r'^(\d+)\s*[.、．)）]')

def _classify_paragraphs(doc):
    ps = doc.paragraphs; first = None
    for i,p in enumerate(ps):
        if HEADER_RE.match(p.text.strip()): first = i; break
    if first is None:
        for i,p in enumerate(ps):
            if QNUM_RE.match(p.text.strip()): first = i; break
    if first is None: first = 0
    keep = list(range(first, len(ps)))
    while keep and not ps[keep[-1]].text.strip(): keep.pop()
    return keep

def _remap_images(elem, sd, dd, rc):
    from docx.opc.part import Part; from docx.opc.packuri import PackURI
    AB = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    RE = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    sp,dp = sd.part, dd.part; dkg = dp.package
    for blip,attr,rid in [(b,RE,b.get(RE)) for b in elem.findall('.//'+AB) if b.get(RE)]:
        if rid in rc: blip.set(attr, rc[rid]); continue
        if rid not in sp.rels: continue
        r = sp.rels[rid]; st = r.target_part
        nm = str(st.partname); ex = {str(p.partname) for p in dkg.iter_parts()}
        c,cd = 1,nm
        while cd in ex: c+=1; cd = f'{os.path.splitext(nm)[0]}_c{c}{os.path.splitext(nm)[1]}'
        np = Part(PackURI(cd), st.content_type, st.blob, dkg)
        nrid = dp.relate_to(np, r.reltype); rc[rid] = nrid; blip.set(attr, nrid)

def _append_content(td, sd):
    ki = _classify_paragraphs(sd)
    if not ki: return 0
    body = td.element.body; sp = body.find(qn('w:sectPr')); rc = {}; k = 0
    for i in ki:
        np = copy.deepcopy(sd.paragraphs[i]._element)
        _remap_images(np, sd, td, rc)
        (sp.addprevious(np) if sp is not None else body.append(np)); k += 1
    return k


# ═════════ 单文件 ═════════
def process_one(inp, out, tpl, sep_img, lm, tm, allow_lian=True):
    try:
        seq, topic, province, exam_type, textbook, ef, es = parse_filename(inp)
    except ValueError as e:
        print(f'  跳过: {e}'); return False
    row, st, ep = lookup_plan(seq, topic, lm, tm, allow_lian)
    print(f'  第{seq}练 | {province} | {exam_type} | 《{textbook}》({ef}) | {topic}')
    if row:
        print(f'    → 规划: {row["级别"]} | 节:{st} | {ep[:60]}...')
    else:
        print(f'    → 规划: 未匹配，使用默认')

    td = Document(tpl); sd = Document(inp)
    _update_table(td, province, exam_type, textbook, ef, seq, ep)
    _update_titles(td.element.body, province, exam_type, textbook, ef, seq, topic, st)
    _insert_separator(td, sep_img)
    n = _append_content(td, sd)
    print(f'  追加{n}段 → {os.path.basename(out)}')
    td.save(out); return True


# ═════════ 主程序 ═════════
def run(cfg):
    """cfg = {template, separator_image, planning_table, input_dir, output_dir}"""
    tpl = cfg['template']; sep = cfg.get('separator_image', '')
    base = cfg['input_dir']; out_dir = cfg.get('output_dir', base)

    if not os.path.exists(tpl): print(f'错误: 模板 {tpl}'); return
    if not os.path.exists(base): print(f'错误: 目录 {base}'); return

    pt = cfg.get('planning_table', '')
    if pt and os.path.exists(pt):
        lm, tm, _ = load_planning_table(pt)
    else:
        if pt: print(f'警告: 规划表 {pt} 不存在')
        lm, tm = {}, {}

    # 收集文件（支持子目录）
    files = []
    for root, dirs, fns in os.walk(base):
        for fn in fns:
            if fn.endswith('.docx') and fn.startswith('第') and '原卷' not in fn:
                files.append(os.path.join(root, fn))

    if not files:
        print('未找到待处理文件'); return

    # 有子目录 → 练号是每课程独立的，不允许练号兜底
    has_subdirs = any(os.path.isdir(os.path.join(base, d)) for d in os.listdir(base) if not d.startswith('.'))
    allow_lian = not has_subdirs

    print(f'模板: {os.path.basename(tpl)}')
    print(f'目录: {base}')
    print(f'共 {len(files)} 个文件  (练号兜底: {"是" if allow_lian else "否"})\n')

    ok = err = skip = 0
    for fp in sorted(files):
        rel = os.path.relpath(fp, base)
        outp = os.path.join(out_dir, rel) if out_dir != base else fp
        if out_dir != base: os.makedirs(os.path.dirname(outp), exist_ok=True)
        print(f'处理: {rel[:80]}')
        try:
            if process_one(fp, outp, tpl, sep, lm, tm, allow_lian): ok += 1
            else: skip += 1
        except Exception as e:
            print(f'  失败: {e}'); import traceback; traceback.print_exc(); err += 1
    print(f'\n完成: {ok}成功 {skip}跳过 {err}失败')


if __name__ == '__main__':
    # 示例：按当前 wyy 项目结构套用模板。
    # 如需处理其他省份/考类，请改 planning_table、input_dir、output_dir。
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    run({
        'template':         os.path.join(project_root, '02_配置资源', '模板和资源', 'template.docx'),
        'separator_image':  os.path.join(project_root, '02_配置资源', '模板和资源', 'separator.png'),
        'planning_table':   os.path.join(project_root, '04_生成输出', '考点规划表', '重庆市', '机械加工类', '重庆市机械加工类_机械基础_一课一练考点规划表.xlsx'),
        'input_dir':        os.path.join(project_root, '04_生成输出', '生成结果', '重庆市 机械加工类', '机械基础'),
        'output_dir':       os.path.join(project_root, '04_生成输出', '生成结果', '重庆市 机械加工类', '机械基础'),
    })
