"""DOCX 生成。"""
import os
import re
import shutil
import sys

from .paths import BASE_DIR, TEMPLATE_PATH, SEPARATOR_PATH
from .planning import _CN_TO_DIGIT, _get_topic_output_base, _parse_edition, _normalize_province_name
from .text_processing import _clean_paper_text

COMMON_SCRIPT_NAMES = ("class.py", "answer2none.py", "zip.py", "fix.py")
RESIDUAL_FORMULA_RE = re.compile(r"\{\{?math:|\\(?:frac|times|Omega|rho|Phi|Delta|mu|sqrt)\b")


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def _find_residual_formula_text(doc):
    """扫描 DOCX 可见文本中的公式标记/LaTeX 残留，防止乱码进入成品。"""
    hits = []
    paragraphs = list(doc.paragraphs) + list(_iter_table_paragraphs(doc.tables))
    for paragraph in paragraphs:
        text = paragraph.text or ""
        if RESIDUAL_FORMULA_RE.search(text):
            hits.append(text.strip()[:120])
    return hits


def _copy_common_scripts_to_output(sub_dir):
    """将常用后处理脚本复制到当前考类/教材输出目录，便于直接双击使用。"""
    source_dir = BASE_DIR / "01_工具脚本" / "通用脚本"
    for script_name in COMMON_SCRIPT_NAMES:
        src = source_dir / script_name
        if not src.exists():
            print(f"  → 警告：未找到通用脚本 {src}")
            continue
        dst = sub_dir / script_name
        try:
            shutil.copy2(src, dst)
        except Exception as exc:
            print(f"  → 警告：复制通用脚本失败 {script_name}: {exc}")

def _number_to_cn(num):
    """将小整数转为中文序号，用于“第X单元”。"""
    try:
        num = int(num)
    except (TypeError, ValueError):
        return str(num or "")
    if num <= 0:
        return str(num)
    cn_digits = "零一二三四五六七八九"
    if num < 10:
        return cn_digits[num]
    if num == 10:
        return "十"
    if num < 20:
        return "十" + cn_digits[num % 10]
    if num < 100:
        tens, ones = divmod(num, 10)
        return cn_digits[tens] + "十" + (cn_digits[ones] if ones else "")
    return str(num)


def _strip_unit_title(unit_str):
    """提取单元名称，去掉“第1章/第1单元”等前缀。"""
    text = str(unit_str or "").strip()
    m = re.match(r"第\s*[\d一二三四五六七八九十百]+\s*[章节单元]\s*(.+)", text)
    return m.group(1).strip() if m else text


def _parse_section_title(section_str):
    """返回章节的阿拉伯数字序号和名称。"""
    text = str(section_str or "").strip()
    if not text:
        return 0, ""

    # 格式: “（一）土方工程” / “(一) 土方工程”
    m = re.match(r'\s*[（(]([一二三四五六七八九十百]+)[）)]\s*(.+)', text)
    if m:
        return _CN_TO_DIGIT.get(m.group(1), 0), m.group(2).strip()

    # 格式: “一、制图基本知识” / “一. 制图基本知识”
    m = re.match(r'\s*([一二三四五六七八九十百]+)[、.．]\s*(.+)', text)
    if m:
        return _CN_TO_DIGIT.get(m.group(1), 0), m.group(2).strip()

    # 格式: “1．力系与平衡” / “3、材料与选用”
    m = re.match(r'\s*(\d+)[．.、]\s*(.+)', text)
    if m:
        return int(m.group(1)), m.group(2).strip()

    return 0, text


def _format_section_title(topic, set_suffix):
    """按三级标题格式化：单元=一级标题，章=二级标题，节=试卷主题。"""
    theme = topic["theme"]
    section_num, section_name = _parse_section_title(topic.get("section", ""))
    theme_num = topic.get("theme_num") or 1

    if topic.get("unit"):
        unit_num = _number_to_cn(topic.get("unit_num") or 1)
        unit_name = _strip_unit_title(topic.get("unit", ""))
        chapter_num = _number_to_cn(section_num or 1)
        chapter_part = f"第{chapter_num}章 {section_name}".strip()
        section_part = f"第{_number_to_cn(theme_num)}节 {theme}"
        return f"第{unit_num}单元 {unit_name} {chapter_part} {section_part}{set_suffix}"

    # 两级模板：节 + 第x练 + 试卷主题。
    section_part = f"第{section_num}节 {section_name}" if section_num else section_name
    seq = topic["seq"]
    if section_part:
        return f"{section_part} 第{seq}练  {theme}{set_suffix}"
    return f"第{seq}练  {theme}{set_suffix}"

def generate_docx(meta, topic, set_idx, paper_text, output_dir, needs_manual_review=False):
    """将试卷文本生成为格式化的 DOCX 文件"""
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本" / "核心脚本"))
    from docx_utils1 import (
        copy_template, set_margins, add_editorial_note,
        add_paragraph_with_style, add_question_options_table,
        add_labeled_text, save_docx
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 清理 AI 返回文本
    paper_text = _clean_paper_text(paper_text)

    # 确定文件名
    seq = topic["seq"]
    theme = topic["theme"]
    province = _normalize_province_name(meta["province"])
    exam_type = meta.get("exam_type_name") or "高职分类考试"
    set_suffix = f"({set_idx})" if topic["sets"] > 1 else ""

    # 从教材详情中提取规范化的教材名、出版社和版次
    # 优先使用课程→教材映射（多课程表中每个课程对应不同教材）
    textbook_name = ""
    publisher = "高教版"
    edition = "第一版"
    edition_display = "高教版·第一版"

    course_map = meta.get("course_textbook_map", {})
    if course_map and topic.get("course") in course_map:
        detail = course_map[topic["course"]]
        textbook_name = detail["name"]
        publisher = detail["publisher"]
        edition = detail["edition"]
        edition_display = detail["display"]
    elif meta.get("textbook_details"):
        detail = meta["textbook_details"][0]
        textbook_name = detail["name"]
        publisher = detail["publisher"]
        edition = detail["edition"]
        edition_display = detail["display"]
    elif meta.get("textbook_list"):
        for tb_name, tb_edition_raw in meta["textbook_list"]:
            textbook_name = tb_name
            publisher, edition = _parse_edition(tb_edition_raw)
            edition_display = f"{publisher}·{edition}"
            break

    if not textbook_name:
        tb_match = re.search(r"《(.+?)》", meta.get("textbooks", ""))
        if tb_match:
            textbook_name = tb_match.group(1)

    filename = f"第{seq}练 {theme}{set_suffix} {province}（{exam_type}）《{textbook_name}》（{edition_display}） 一课一练 （解析版）.docx"
    if needs_manual_review:
        filename = f"（待人工审核）{filename}"

    # 按"省份 类别/课程或教材名/"组织子目录，与规划表位置保持一致
    sub_dir = _get_topic_output_base(meta, topic, output_dir)
    os.makedirs(sub_dir, exist_ok=True)
    _copy_common_scripts_to_output(sub_dir)
    output_path = sub_dir / filename

    # 复制模板
    tpl = str(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else None
    if tpl:
        doc = copy_template(tpl, str(output_path))
    else:
        from docx import Document
        doc = Document()
        doc.save(str(output_path))
        doc = Document(str(output_path))

    set_margins(doc)

    # 编写说明
    sep_img = str(SEPARATOR_PATH) if SEPARATOR_PATH.exists() else None
    add_editorial_note(
        doc,
        textbook_name=textbook_name,
        edition=edition,
        chapter_seq=seq,
        knowledge_scope=topic["knowledge"],
        province=province,
        publisher=publisher,
        separator_image=sep_img,
        exam_type=exam_type,
    )

    # 标题
    title1 = f"{province}（{exam_type}）一课一练"
    title2 = f"《{textbook_name}》（{edition_display}）"
    if topic.get("unit"):
        title2 = f"{title2} 第{seq}练{set_suffix}"
    title3 = _format_section_title(topic, set_suffix)

    add_paragraph_with_style(doc, title1, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title2, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title3, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 答案/解析使用红色字体
    RED_COLOR = (255, 0, 0)

    # 逐行解析并输出
    lines = paper_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 大题标题（加粗，黑体）
        if re.match(r"^[一二三四五六七八九十][、.．]", line):
            add_paragraph_with_style(doc, line, font_name="黑体", font_size=12,
                                     bold=True, space_after=6)
            i += 1
            continue

        # 题目开头（数字编号）
        if re.match(r"^\d+[\.．、]", line):
            # 收集题目完整文本（直到下一个【答案】或下一题）
            q_lines = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if next_line.startswith("【答案】") or next_line.startswith("【解析】"):
                    break
                if not next_line:
                    # 防止题目文本和【答案】之间的偶发空行进入 DOCX。
                    lookahead = j + 1
                    while lookahead < len(lines) and not lines[lookahead].strip():
                        lookahead += 1
                    if lookahead < len(lines) and lines[lookahead].strip().startswith(("【答案】", "【解析】")):
                        j = lookahead
                        continue
                    j += 1
                    continue
                if re.match(r"^\d+[\.．、]", next_line):
                    break
                if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                    break
                q_lines.append(lines[j])
                j += 1

            q_text = "\n".join(q_lines)

            # 判断是否有选项（含 A. B. C. D.）
            if re.search(r"[A-D][\.．]\s*\S", q_text):
                add_question_options_table(doc, q_text)
            else:
                add_paragraph_with_style(doc, q_text, font_size=10.5, space_after=2)

            i = j
            continue

        # 【答案】— 红色字体；多行答案继续归入同一红色段落
        if line.startswith("【答案】"):
            answer_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith("【解析】") or next_line.startswith("【答案】")
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    answer_lines.append(next_line)
                j += 1
            answer_content = "\n".join(answer_lines)
            add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            i = j
            continue

        # 【解析】— 红色字体；多行解析继续归入同一红色段落
        if line.startswith("【解析】"):
            explanation_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith("【答案】") or next_line.startswith("【解析】")
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    explanation_lines.append(next_line)
                j += 1
            explanation_content = "\n".join(explanation_lines)
            add_labeled_text(doc, "【解析】", explanation_content, color=RED_COLOR)
            i = j
            continue

        # 普通文本行
        if line:
            add_paragraph_with_style(doc, line, font_size=10.5, space_after=2)

        i += 1

    save_docx(doc, str(output_path))
    residuals = _find_residual_formula_text(doc)
    if residuals:
        print(f"  [!] 公式转换后仍发现 {len(residuals)} 处疑似 LaTeX/math 残留，请检查 DOCX：{output_path}")
        for sample in residuals[:3]:
            print(f"      - {sample}")
    return str(output_path)
