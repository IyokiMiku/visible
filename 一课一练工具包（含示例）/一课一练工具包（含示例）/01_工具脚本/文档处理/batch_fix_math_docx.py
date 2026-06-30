"""批量修复 DOCX 中残留的 {math:...}/{{math:...}} 为 Word 原生公式。

用途：
  早期生成的 Word 文档里如果直接显示了 ``{math:...}`` 或 ``{{math:...}}``，
  说明公式标记没有在生成阶段转换成 Word 公式。运行本脚本可批量扫描并重写
  含有公式标记的段落，把公式标记转换为 Word 原生 OMML 公式。

默认行为：
  - 扫描 ``04_生成输出/生成结果`` 下所有 .docx 文件；
  - 处理前为每个将被修改的文档生成 ``.mathbak.docx`` 备份；
  - 递归处理正文段落和表格单元格内段落；
  - 顺带修复【答案】/【解析】段落正文被整体加粗的问题，只保留标签加粗；
  - 顺带删除题干/选项与【答案】之间的空白段落；
  - 跳过 Word 临时文件 ``~$*.docx`` 和备份文件 ``*.mathbak.docx``。

依赖：
  - python-docx
  - lxml
  - latex2mathml
  - 本机 Microsoft Office 自带的 MML2OMML.XSL

如果缺少公式转换依赖，底层工具会提示并降级为普通文本；为保证专业排版，
建议先执行：

    pip install lxml latex2mathml

示例：

    python 01_工具脚本/文档处理/batch_fix_math_docx.py
    python 01_工具脚本/文档处理/batch_fix_math_docx.py --path "04_生成输出/生成结果/重庆市 机械加工类"
    python 01_工具脚本/文档处理/batch_fix_math_docx.py --path "某个文件.docx"
    python 01_工具脚本/文档处理/batch_fix_math_docx.py --dry-run
    python 01_工具脚本/文档处理/batch_fix_math_docx.py --no-backup
"""

from __future__ import annotations

import argparse
import copy
import re
import shutil
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]
CORE_DIR = PROJECT_ROOT / "01_工具脚本" / "核心脚本"
if str(CORE_DIR) not in sys.path:
    sys.path.insert(0, str(CORE_DIR))

from docx_utils1 import _add_rich_text_run  # noqa: E402  内部工具：负责把 math 标记转为 OMML
from docx.shared import Pt  # noqa: E402

MATH_MARKER_RE = re.compile(r"\{\{math:.*?\}\}|\{math:.*?\}", re.DOTALL)
LABEL_RE = re.compile(r"【(?:答案|解析|详解)】")
ANSWER_LABEL_RE = re.compile(r"^\s*【答案】")


def _iter_table_paragraphs(tables) -> Iterable:
    """递归遍历表格单元格中的段落。"""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def _iter_all_paragraphs(doc: Document) -> Iterable:
    """遍历正文和表格中的全部段落。"""
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)


def _paragraph_has_math_marker(paragraph) -> bool:
    return bool(MATH_MARKER_RE.search(paragraph.text or ""))


def _rgb_to_tuple(rgb) -> tuple[int, int, int] | None:
    if rgb is None:
        return None
    text = str(rgb)
    if len(text) != 6:
        return None
    return int(text[0:2], 16), int(text[2:4], 16), int(text[4:6], 16)


def _paragraph_default_color(paragraph) -> tuple[int, int, int] | None:
    """读取段落默认 run 颜色。"""
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return None
    p_r_pr = p_pr.find(qn("w:rPr"))
    if p_r_pr is None:
        return None
    color = p_r_pr.find(qn("w:color"))
    if color is None:
        return None
    value = color.get(qn("w:val"))
    if not value or value.lower() == "auto" or len(value) != 6:
        return None
    return int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)


def _first_text_run(paragraph):
    for run in paragraph.runs:
        if run.text:
            return run
    return paragraph.runs[0] if paragraph.runs else None


def _detect_paragraph_style(paragraph) -> dict:
    """从旧段落中提取尽量保守的文本样式，用于重建段落内容。"""
    run = _first_text_run(paragraph)

    font_name = "宋体"
    font_size = 10.5
    # 修复公式时不要把旧段落首个 run 的加粗状态扩散到整段。
    # 解析/答案段落通常只有【解析】/【答案】标签应加粗，正文不应加粗。
    bold = False
    font_color = _paragraph_default_color(paragraph)

    if run is not None:
        if run.font.name:
            font_name = run.font.name
        if run.font.size:
            font_size = run.font.size.pt
        if run.font.color is not None and run.font.color.rgb is not None:
            font_color = _rgb_to_tuple(run.font.color.rgb)

    # 如果段落中已有红色 run，通常是【答案】/【解析】段落；重建时保持整段红色。
    for candidate in paragraph.runs:
        if candidate.font.color is not None and candidate.font.color.rgb is not None:
            rgb = _rgb_to_tuple(candidate.font.color.rgb)
            if rgb == (255, 0, 0):
                font_color = rgb
                break

    text = paragraph.text or ""
    if text.lstrip().startswith(("【答案】", "【解析】", "【详解】")):
        font_color = (255, 0, 0)

    return {
        "font_name": font_name,
        "font_size": font_size,
        "bold": bold,
        "font_color": font_color,
    }


def _clear_paragraph_content(paragraph) -> None:
    """清空段落内容，保留段落属性 pPr。"""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _remove_paragraph(paragraph) -> None:
    """从文档 XML 中删除段落。"""
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _is_blank_paragraph(paragraph) -> bool:
    """判断是否为空白段落。"""
    return not (paragraph.text or "").strip()


def _remove_blank_paragraphs_before_answer(paragraphs) -> int:
    """删除紧挨在【答案】段落前的空白段落。"""
    removed = 0
    before_answer = False
    for paragraph in reversed(list(paragraphs)):
        text = (paragraph.text or "").strip()
        if ANSWER_LABEL_RE.match(text):
            before_answer = True
            continue
        if before_answer and _is_blank_paragraph(paragraph):
            _remove_paragraph(paragraph)
            removed += 1
            continue
        before_answer = False
    return removed


def _fix_label_bold(paragraph) -> bool:
    """修复【答案】/【解析】段落：只让标签加粗，正文不加粗。"""
    text = paragraph.text or ""
    match = LABEL_RE.search(text)
    if not match:
        return False

    # 若正文没有被加粗，且没有公式标记需要重建，则不动，避免无谓改写。
    body_start = match.end()
    pos = 0
    has_bold_body = False
    for run in paragraph.runs:
        run_text = run.text or ""
        run_start = pos
        run_end = pos + len(run_text)
        if run_end > body_start and run.bold:
            has_bold_body = True
            break
        pos = run_end
    if not has_bold_body and not MATH_MARKER_RE.search(text):
        return False

    style = _detect_paragraph_style(paragraph)
    prefix = text[:match.start()]
    label = match.group(0)
    rest = text[match.end():]

    _clear_paragraph_content(paragraph)

    if prefix:
        _add_rich_text_run(
            paragraph,
            prefix,
            style["font_name"],
            style["font_size"],
            False,
            style["font_color"],
        )

    label_run = paragraph.add_run(label)
    label_run.font.name = "黑体"
    label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    label_run.font.size = Pt(style["font_size"])
    label_run.bold = True
    if style["font_color"]:
        label_run.font.color.rgb = RGBColor(*style["font_color"])

    _add_rich_text_run(
        paragraph,
        rest,
        "宋体",
        style["font_size"],
        False,
        style["font_color"],
    )
    return True


def _fix_paragraph_math(paragraph) -> bool:
    """将单个段落中的 math 标记替换为 Word 公式。返回是否修改。"""
    text = paragraph.text or ""
    if not MATH_MARKER_RE.search(text):
        return False

    style = _detect_paragraph_style(paragraph)
    _clear_paragraph_content(paragraph)

    # 保持生成器原本的标签样式：只有【答案】/【解析】/【详解】标签加粗，
    # 后面的正文和公式不加粗。之前直接用首个 run 的 bold 重建整段，会导致
    # 解析部分全部加粗。
    label_match = re.match(r"^(【(?:答案|解析|详解)】)(.*)$", text, flags=re.DOTALL)
    if label_match:
        label, rest = label_match.groups()
        label_run = paragraph.add_run(label)
        label_run.font.name = "黑体"
        label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        label_run.font.size = Pt(style["font_size"])
        label_run.bold = True
        if style["font_color"]:
            label_run.font.color.rgb = RGBColor(*style["font_color"])

        _add_rich_text_run(
            paragraph,
            rest,
            "宋体",
            style["font_size"],
            False,
            style["font_color"],
        )
    else:
        _add_rich_text_run(
            paragraph,
            text,
            style["font_name"],
            style["font_size"],
            False,
            style["font_color"],
        )
    return True


def _backup_path(path: Path) -> Path:
    """生成不覆盖旧备份的备份路径。"""
    base = path.with_suffix(".mathbak.docx")
    if not base.exists():
        return base
    idx = 2
    while True:
        candidate = path.with_name(f"{path.stem}.mathbak{idx}.docx")
        if not candidate.exists():
            return candidate
        idx += 1


def _iter_docx_files(path: Path) -> list[Path]:
    if path.is_file():
        return [path] if path.suffix.lower() == ".docx" else []
    if not path.exists():
        return []
    return sorted(
        p
        for p in path.rglob("*.docx")
        if p.is_file()
        and not p.name.startswith("~$")
        and not p.name.endswith(".mathbak.docx")
        and ".mathbak" not in p.stem
    )


def fix_docx(path: Path, *, dry_run: bool = False, backup: bool = True) -> int:
    """修复一个 DOCX 文件，返回修改项数量。"""
    doc = Document(path)
    paragraphs = list(_iter_all_paragraphs(doc))

    math_count = sum(1 for paragraph in paragraphs if _paragraph_has_math_marker(paragraph))
    bold_count = 0
    for paragraph in paragraphs:
        text = paragraph.text or ""
        match = LABEL_RE.search(text)
        if not match:
            continue
        body_start = match.end()
        pos = 0
        for run in paragraph.runs:
            run_text = run.text or ""
            run_start = pos
            run_end = pos + len(run_text)
            if run_end > body_start and run.bold:
                bold_count += 1
                break
            pos = run_end

    blank_count = 0
    before_answer = False
    for paragraph in reversed(paragraphs):
        text = (paragraph.text or "").strip()
        if ANSWER_LABEL_RE.match(text):
            before_answer = True
            continue
        if before_answer and _is_blank_paragraph(paragraph):
            blank_count += 1
            continue
        before_answer = False

    changed = math_count + bold_count + blank_count
    if changed == 0 or dry_run:
        return changed

    if backup:
        shutil.copy2(path, _backup_path(path))

    # 先修复公式标记；公式修复会重建含公式的段落，并顺带按正确规则处理标签加粗。
    for paragraph in paragraphs:
        _fix_paragraph_math(paragraph)

    # 再修复不含公式、但【答案】/【解析】正文被整体加粗的段落。
    paragraphs = list(_iter_all_paragraphs(doc))
    for paragraph in paragraphs:
        _fix_label_bold(paragraph)

    # 最后删除【答案】前的空白段落，解决题干/选项与答案之间多空行。
    paragraphs = list(_iter_all_paragraphs(doc))
    _remove_blank_paragraphs_before_answer(paragraphs)

    doc.save(path)
    return changed


def main() -> None:
    parser = argparse.ArgumentParser(description="批量修复 DOCX 中残留的 {math:...}/{{math:...}} 为 Word 原生公式")
    parser.add_argument(
        "--path",
        "-p",
        default=str(PROJECT_ROOT / "04_生成输出" / "生成结果"),
        help="要处理的 DOCX 文件或目录，默认处理 04_生成输出/生成结果",
    )
    parser.add_argument("--dry-run", action="store_true", help="只统计将被修复的文件和段落，不写入文件")
    parser.add_argument("--no-backup", action="store_true", help="不生成 .mathbak.docx 备份，直接覆盖原文件")
    args = parser.parse_args()

    target = Path(args.path)
    if not target.is_absolute():
        target = PROJECT_ROOT / target

    files = _iter_docx_files(target)
    if not files:
        print(f"未找到 DOCX 文件：{target}")
        return

    total_files_changed = 0
    total_paragraphs_changed = 0
    print(f"扫描路径：{target}")
    print(f"发现 DOCX 文件：{len(files)} 个")
    if args.dry_run:
        print("当前为 dry-run 模式，不会修改文件。")
    elif not args.no_backup:
        print("将为被修改的文件生成 .mathbak.docx 备份。")

    for path in files:
        try:
            changed = fix_docx(path, dry_run=args.dry_run, backup=not args.no_backup)
        except Exception as exc:
            print(f"处理失败：{path} -> {exc}")
            continue

        if changed:
            total_files_changed += 1
            total_paragraphs_changed += changed
            rel = path.relative_to(PROJECT_ROOT) if path.is_relative_to(PROJECT_ROOT) else path
            action = "将修复" if args.dry_run else "已修复"
            print(f"{action}：{rel}（段落 {changed} 个）")

    print(
        f"完成：共扫描 {len(files)} 个文档，"
        f"{'将修改' if args.dry_run else '已修改'} {total_files_changed} 个文档，"
        f"涉及段落 {total_paragraphs_changed} 个。"
    )


if __name__ == "__main__":
    main()
