"""将当前目录内所有解析版 DOCX 的答案和解析内容标红。

默认只处理本脚本所在目录，不递归子目录；直接覆写原文档。
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

RED = RGBColor(255, 0, 0)
RED_HEX = "FF0000"
LABEL_RE = re.compile(r"【(?:答案|解析|详解)】")
SECTION_RE = re.compile(r"^[一二三四五六七八九十][、.．]")
QUESTION_RE = re.compile(r"^\d+[\.．、]")


def _set_run_red(run) -> None:
    run.font.color.rgb = RED
    r_pr = run._r.get_or_add_rPr()
    color = r_pr.color
    if color is None:
        color = r_pr._add_color()
    color.set(qn("w:val"), RED_HEX)
    color.attrib.pop(qn("w:themeColor"), None)
    color.attrib.pop(qn("w:themeTint"), None)
    color.attrib.pop(qn("w:themeShade"), None)

    # Some WPS/Word files use w14:textFill, which can override w:color.
    for child in list(r_pr):
        if child.tag.endswith("}textFill"):
            r_pr.remove(child)


def _set_paragraph_default_red(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_r_pr = p_pr.find(qn("w:rPr"))
    if p_r_pr is None:
        p_r_pr = OxmlElement("w:rPr")
        p_pr.append(p_r_pr)
    color = p_r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        p_r_pr.append(color)
    color.set(qn("w:val"), RED_HEX)
    color.attrib.pop(qn("w:themeColor"), None)
    color.attrib.pop(qn("w:themeTint"), None)
    color.attrib.pop(qn("w:themeShade"), None)
    for child in list(p_r_pr):
        if child.tag.endswith("}textFill"):
            p_r_pr.remove(child)


def _mark_from_label(paragraph) -> bool:
    """标红段落中从【答案】/【解析】/【详解】开始到段落末尾的内容。"""
    full_text = paragraph.text or ""
    match = LABEL_RE.search(full_text)
    if not match:
        return False

    label_start = match.start()
    pos = 0
    for run in paragraph.runs:
        text = run.text or ""
        run_start = pos
        run_end = pos + len(text)
        if run_end > label_start:
            _set_run_red(run)
        pos = run_end
    return True


def _mark_paragraphs(paragraphs) -> int:
    changed = 0
    in_answer_area = False

    for paragraph in paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue

        has_label = bool(LABEL_RE.search(text))

        # 新大题或新题号表示上一段答案/解析区域结束；若本段自己含标签，则仍从标签处标红。
        if SECTION_RE.match(text) or (QUESTION_RE.match(text) and not has_label):
            in_answer_area = False

        if has_label:
            if _mark_from_label(paragraph):
                _set_paragraph_default_red(paragraph)
                changed += 1
            in_answer_area = True
            continue

        if in_answer_area:
            _set_paragraph_default_red(paragraph)
            for run in paragraph.runs:
                _set_run_red(run)
            changed += 1

    return changed


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def mark_docx(path: Path) -> int:
    doc = Document(path)
    changed = _mark_paragraphs(doc.paragraphs)
    changed += _mark_paragraphs(list(_iter_table_paragraphs(doc.tables)))

    if changed:
        doc.save(path)

    return changed


def main() -> None:
    base_dir = Path(__file__).resolve().parent
    files = sorted(
        p for p in base_dir.iterdir()
        if p.is_file()
        and p.suffix.lower() == ".docx"
        and not p.name.startswith("~$")
        and ("（解析版）" in p.name or "(解析版)" in p.name or "解析版" in p.name)
    )

    if not files:
        print("当前目录未找到解析版 DOCX 文件。")
        return

    total_changed = 0
    for path in files:
        try:
            changed = mark_docx(path)
            total_changed += changed
            if changed:
                print(f"已处理: {path.name}（标红段落 {changed} 个）")
            else:
                print(f"无需修改: {path.name}")
        except Exception as exc:
            print(f"处理失败: {path.name} -> {exc}")

    print(f"完成：共扫描 {len(files)} 个解析版文档，标红段落 {total_changed} 个。")


if __name__ == "__main__":
    main()
