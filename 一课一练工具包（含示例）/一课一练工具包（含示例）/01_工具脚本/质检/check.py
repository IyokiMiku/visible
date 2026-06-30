"""一课一练试卷质检工具 — 对已生成试卷进行严格的质量审查

使用方法：
  python 01_工具脚本/质检/check.py                              # 交互式选择文件
  python 01_工具脚本/质检/check.py --file 第1练_xxx.txt         # 检查单个原始文本
  python 01_工具脚本/质检/check.py --dir 04_生成输出/生成结果/_原始文本      # 批量检查目录
  python 01_工具脚本/质检/check.py --file xxx.docx              # 检查 DOCX 文件
"""

import json
import os
import re
import sys
import time
from pathlib import Path

import openpyxl
from openai import OpenAI

BASE_DIR = Path(__file__).resolve().parents[2]
CONFIG_PATH = BASE_DIR / "02_配置资源" / "config.json"
SPEC_PATH = BASE_DIR / "02_配置资源" / "编写规范" / "编写规范.md"


def load_config():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def load_spec():
    with open(SPEC_PATH, "r", encoding="utf-8") as f:
        return f.read()


def call_api(client, model, system_prompt, user_prompt, max_tokens=8000):
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=max_tokens,
                temperature=0.1,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"  [!] API 调用失败 (第{attempt+1}次): {e}")
            if attempt < 2:
                time.sleep((attempt + 1) * 10)
            else:
                raise


def extract_text_from_docx(docx_path):
    """从 DOCX 提取纯文本内容"""
    from docx import Document
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)
    return "\n".join(lines)


def _normalize_math_markers_for_check(text):
    """质检前把 {{math:...}} 还原为内部公式文本，避免标记噪声干扰查重和长度判断。"""
    if not text:
        return text
    return re.sub(r"\{\{math:(.*?)\}\}", lambda m: m.group(1), text, flags=re.DOTALL)


def _normalize_choice_answer(raw_answer):
    """规范化选择题答案，兼容 A、C / A C / AC / A,B,D 等多选写法。

    返回值：
    - 单选/多选：按 A-B-C-D 顺序去重后的字母串，如 A、AC、ABD、ABCD；
    - 判断题：√ 或 ×；
    - 无法识别：去除首尾空白后的原文本。
    """
    raw = (raw_answer or "").strip()
    if not raw:
        return ""

    if "√" in raw or re.search(r"\b(对|正确)\b", raw):
        return "√"
    if "×" in raw or "x" == raw.lower() or re.search(r"\b(错|错误)\b", raw):
        return "×"

    upper = raw.upper().strip()
    # 只把纯选项答案识别为 A-D，避免把“2 A”“10 A”等电流单位误判为选择题答案。
    if re.fullmatch(r"[A-D](?:\s*[,，、/ ]\s*[A-D])*", upper) or re.fullmatch(r"[A-D]{1,4}", upper):
        letters = re.findall(r"[A-D]", upper)
        present = set(letters)
        return "".join(letter for letter in "ABCD" if letter in present)
    return raw


def _is_choice_answer(answer):
    return bool(answer) and all(ch in "ABCD" for ch in answer)


def _is_single_choice_answer(answer):
    return _is_choice_answer(answer) and len(answer) == 1


def _is_multi_choice_answer(answer):
    return _is_choice_answer(answer) and len(answer) >= 2


def _format_answer_dist(dist):
    return f"A={dist['A']}, B={dist['B']}, C={dist['C']}, D={dist['D']}"


def local_check(paper_text):
    """本地快速检查（不依赖 API），返回发现的问题列表"""
    paper_text = _normalize_math_markers_for_check(paper_text)
    issues = []

    # 提取选择题
    choice_pattern = re.compile(
        r"(\d+)\.\s*(.+?)（\s*）\s*\n"
        r"A[\.．]\s*(.+?)[\t\n]"
        r"B[\.．]\s*(.+?)[\t\n]"
        r"C[\.．]\s*(.+?)[\t\n]"
        r"D[\.．]\s*(.+?)\n"
        r"【答案】\s*([A-D])",
        re.MULTILINE
    )

    # 更宽松的选项提取
    questions = []
    lines = paper_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        q_match = re.match(r"^(\d+)\.\s*(.+)", line)
        if q_match:
            q_num = int(q_match.group(1))
            q_stem = q_match.group(2)

            # 收集选项
            options = {}
            j = i + 1
            while j < len(lines) and j < i + 10:
                opt_line = lines[j]
                for opt_match in re.finditer(r"([A-D])[\.．]\s*(.+?)(?=\s*[A-D][\.．]|\s*$|\t)", opt_line):
                    opt_letter = opt_match.group(1)
                    opt_text = opt_match.group(2).strip()
                    if opt_text:
                        options[opt_letter] = opt_text
                if lines[j].strip().startswith("【答案】"):
                    answer_match = re.match(r"【答案】\s*(.+)", lines[j].strip())
                    if answer_match:
                        answer = _normalize_choice_answer(answer_match.group(1))
                        # 只有实际收集到 A-D 选项的题才纳入选择题质检，
                        # 避免计算题答案中的单位 A 被误算为单选答案。
                        if options:
                            questions.append({
                                "num": q_num,
                                "stem": q_stem,
                                "options": options,
                                "answer": answer,
                            })
                    break
                j += 1
            i = j + 1
        else:
            i += 1

    # 检查1：选项长度比
    for q in questions:
        if len(q["options"]) == 4:
            lengths = [len(v) for v in q["options"].values()]
            max_len = max(lengths)
            min_len = max(min(lengths), 1)
            ratio = max_len / min_len
            if ratio > 2.0:
                opts_display = " / ".join(f"{k}({len(v)}字)" for k, v in q["options"].items())
                issues.append({
                    "type": "选项长度失衡",
                    "severity": "严重",
                    "question": q["num"],
                    "detail": f"长度比={ratio:.1f}（超过2.0限制）。各选项: {opts_display}",
                })

    # 检查2：废选项
    waste_words = {"正常", "无影响", "无变化", "不确定", "不动", "任意", "无要求", "更省油", "更快", "更好", "装饰", "没什么"}
    for q in questions:
        for letter, text in q["options"].items():
            if text in waste_words or (len(text) <= 3 and text in waste_words):
                issues.append({
                    "type": "废选项",
                    "severity": "严重",
                    "question": q["num"],
                    "detail": f"选项{letter}=\"{text}\"是无意义填充词",
                })

    # 检查3：答案分布
    choice_questions = [q for q in questions if _is_single_choice_answer(q["answer"])]
    if choice_questions:
        answer_dist = {"A": 0, "B": 0, "C": 0, "D": 0}
        for q in choice_questions:
            answer_dist[q["answer"]] += 1
        max_count = max(answer_dist.values())
        total = len(choice_questions)
        dist_str = _format_answer_dist(answer_dist)

        if total <= 5:
            # 少量选择题（≤5道）：只有出现4道及以上相同答案才报严重
            if max_count >= 4:
                dominant = [k for k, v in answer_dist.items() if v == max_count][0]
                issues.append({
                    "type": "答案分布失衡",
                    "severity": "严重",
                    "question": "全卷",
                    "detail": f"答案分布: {dist_str}。{dominant}出现{max_count}次/{total}题",
                })
            else:
                # 不报错，但记录分布供展示
                issues.append({
                    "type": "答案分布",
                    "severity": "信息",
                    "question": "全卷",
                    "detail": f"答案分布: {dist_str}",
                })
        else:
            # 多选择题（>5道）：使用比例判断
            if max_count > total * 0.5:
                dominant = [k for k, v in answer_dist.items() if v == max_count][0]
                issues.append({
                    "type": "答案分布失衡",
                    "severity": "严重",
                    "question": "全卷",
                    "detail": f"答案分布: {dist_str}。{dominant}占{max_count}/{total}={max_count/total:.0%}",
                })
            elif max_count >= total * 0.4:
                issues.append({
                    "type": "答案分布不均",
                    "severity": "警告",
                    "question": "全卷",
                    "detail": f"答案分布: {dist_str}",
                })
            else:
                issues.append({
                    "type": "答案分布",
                    "severity": "信息",
                    "question": "全卷",
                    "detail": f"答案分布: {dist_str}",
                })

    # 检查3.1：多选题全选与答案分布
    multi_choice_questions = [q for q in questions if _is_multi_choice_answer(q["answer"])]
    for q in multi_choice_questions:
        if q["answer"] == "ABCD":
            issues.append({
                "type": "多选题全选",
                "severity": "严重",
                "question": q["num"],
                "detail": "多选题答案为ABCD全选，必须重出该题或将至少一个选项改为合理错误项；修改后仍需满足选项长度比≤2.0、干扰项有效、解析同步更新。",
            })

    if multi_choice_questions:
        multi_dist = {"A": 0, "B": 0, "C": 0, "D": 0}
        for q in multi_choice_questions:
            for letter in q["answer"]:
                multi_dist[letter] += 1
        total_multi = len(multi_choice_questions)
        max_multi = max(multi_dist.values())
        min_multi = min(multi_dist.values())
        dist_str = _format_answer_dist(multi_dist)

        if total_multi <= 2:
            issues.append({
                "type": "多选答案分布",
                "severity": "信息",
                "question": "全卷",
                "detail": f"多选答案分布: {dist_str}（多选题{total_multi}题，样本较少仅提示）",
            })
        elif total_multi <= 5:
            if max_multi == total_multi or min_multi == 0:
                issues.append({
                    "type": "多选答案分布不均",
                    "severity": "警告",
                    "question": "全卷",
                    "detail": f"多选答案分布: {dist_str}。存在某个选项在全部多选题中均为正确项，或某个选项从未作为正确项",
                })
            else:
                issues.append({
                    "type": "多选答案分布",
                    "severity": "信息",
                    "question": "全卷",
                    "detail": f"多选答案分布: {dist_str}",
                })
        else:
            max_ratio = max_multi / total_multi
            min_ratio = min_multi / total_multi
            if max_ratio > 0.85 or min_multi == 0:
                issues.append({
                    "type": "多选答案分布失衡",
                    "severity": "严重",
                    "question": "全卷",
                    "detail": f"多选答案分布: {dist_str}。最高出现率{max_ratio:.0%}，最低出现率{min_ratio:.0%}",
                })
            elif max_ratio > 0.75 or min_ratio < 0.25:
                issues.append({
                    "type": "多选答案分布不均",
                    "severity": "警告",
                    "question": "全卷",
                    "detail": f"多选答案分布: {dist_str}",
                })
            else:
                issues.append({
                    "type": "多选答案分布",
                    "severity": "信息",
                    "question": "全卷",
                    "detail": f"多选答案分布: {dist_str}",
                })

    # 检查4：禁用符号与解析完整性
    prohibited = ["→", "↑", "↓", "=>", "≫"]
    # 解析可能是多行格式：
    # 【解析】
    # （1）……
    # （2）……
    # 旧逻辑只用 re.findall(r"【解析】(.+)") 抓同一行内容，会把这种合法解析误判为空。
    analysis_blocks = []
    lines_for_analysis = paper_text.split("\n")
    for line_idx, line in enumerate(lines_for_analysis):
        stripped = line.strip()
        if not stripped.startswith("【解析】"):
            continue

        content_lines = [stripped[4:].strip()]
        j = line_idx + 1
        while j < len(lines_for_analysis):
            next_line = lines_for_analysis[j].strip()
            if re.match(r"^\d+[\.．、]\s*", next_line):
                break
            if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                break
            if next_line.startswith("【答案】") or next_line.startswith("【解析】"):
                break
            if next_line:
                content_lines.append(next_line)
            j += 1

        analysis_blocks.append("\n".join(content_lines).strip())

    for idx, analysis in enumerate(analysis_blocks, 1):
        for sym in prohibited:
            if sym in analysis:
                issues.append({
                    "type": "禁用符号",
                    "severity": "警告",
                    "question": idx,
                    "detail": f"解析中出现禁用符号 \"{sym}\"",
                })
        # 检查解析是否为单个短语（过短且无主谓结构）
        clean = analysis.strip().rstrip("。")
        if len(clean) <= 8 and "=" not in clean:
            issues.append({
                "type": "解析过短",
                "severity": "严重",
                "question": idx,
                "detail": f"解析仅\"{analysis.strip()}\"（{len(clean)}字），必须是完整因果句",
            })

    # 检查5：答案自暴露（弱化版：仅当题干直接包含正确选项中的较长连续片段时提示）
    for q in questions:
        if q["answer"] in q["options"]:
            correct_text = q["options"][q["answer"]]
            # 旧规则会检查3~5字片段，容易把正常知识词重复误判为自暴露。
            # 现在只检查≥6字的连续片段，且最多检查到10字，降低误报。
            max_segment_len = min(len(correct_text), 10)
            for length in range(max_segment_len, 5, -1):
                for start in range(len(correct_text) - length + 1):
                    segment = correct_text[start:start + length]
                    if segment in q["stem"]:
                        issues.append({
                            "type": "答案自暴露",
                            "severity": "警告",
                            "question": q["num"],
                            "detail": f"题干包含正确答案较长片段\"{segment}\"",
                        })
                        break
                else:
                    continue
                break

    # 检查6：全部为容易题（选项都很短可能意味着全是简单概念题）
    all_short = all(
        max(len(v) for v in q["options"].values()) <= 8
        for q in questions if len(q["options"]) == 4
    )
    if all_short and len(questions) >= 5:
        issues.append({
            "type": "难度偏低",
            "severity": "信息",
            "question": "全卷",
            "detail": "所有选择题选项均≤8字，可能全部为简单概念题，缺少适中/困难题",
        })

    # 检查7：同卷自重复（全题型题干相似度检测）
    all_question_stems = _extract_all_question_stems(paper_text)
    for i in range(len(all_question_stems)):
        for j in range(i + 1, len(all_question_stems)):
            duplicate, score, reason = _is_duplicate_stem_pair(
                all_question_stems[i]["stem"], all_question_stems[j]["stem"]
            )
            if duplicate:
                issues.append({
                    "type": "题干重复",
                    "severity": "严重",
                    "question": f"{all_question_stems[i]['num']}&{all_question_stems[j]['num']}",
                    "detail": (
                        f"第{all_question_stems[i]['num']}题与第{all_question_stems[j]['num']}题疑似重复"
                        f"（{reason}，重复度={score:.0%}）"
                    ),
                })

    return issues, questions


def _extract_all_question_stems(paper_text):
    """提取全题型题干，用于同卷查重。

    不依赖 A/B/C/D 选项，因此选择题、判断题、填空题、综合题都会参与查重。
    仅提取题干部分，遇到选项、答案、解析、下一题或下一大题即停止。
    """
    stems = []
    lines = paper_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        q_match = re.match(r"^(\d+)[\.．、]\s*(.+)", line)
        if not q_match:
            i += 1
            continue

        q_num = int(q_match.group(1))
        stem_lines = [q_match.group(2).strip()]
        j = i + 1
        while j < len(lines):
            next_line = lines[j].strip()
            if not next_line:
                j += 1
                continue
            if re.match(r"^\d+[\.．、]\s*", next_line):
                break
            if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                break
            if re.match(r"^[A-D][\.．]\s*", next_line):
                break
            if next_line.startswith("【答案】") or next_line.startswith("【解析】"):
                break
            stem_lines.append(next_line)
            j += 1

        stem = " ".join(stem_lines).strip()
        stem = re.sub(r"（\s*）\s*$", "", stem).strip()
        if stem:
            stems.append({"num": q_num, "stem": stem})
        i = max(j, i + 1)

    return stems


_TEMPLATE_PATTERNS = [
    r"下列", r"关于", r"有关", r"对于", r"在.*?中", r"根据.*?可知",
    r"说法", r"表述", r"选项", r"哪一项", r"哪项", r"的是", r"是\s*\(\s*\)",
    r"正确的是", r"错误的是", r"不正确的是", r"不属于", r"属于",
    r"主要", r"基本", r"一般", r"通常", r"应当", r"可以", r"能够", r"需要",
    r"作用", r"原因", r"特点", r"目的", r"要求", r"方法", r"措施",
]

_STOP_WORDS = {
    "下列", "关于", "有关", "对于", "说法", "表述", "选项", "哪一项", "哪项",
    "正确", "错误", "不正确", "属于", "不属于", "主要", "基本", "一般", "通常",
    "应当", "可以", "能够", "需要", "作用", "原因", "特点", "目的", "要求", "方法", "措施",
}


def _normalize_stem_for_dup(text):
    """去掉出题模板词和标点，保留核心知识词用于查重。"""
    if not text:
        return ""
    text = re.sub(r"^\d+[\.．、]\s*", "", text)
    text = re.sub(r"（\s*）", "", text)
    text = re.sub(r"[，,。；;：:？！?、（）()《》<>\[\]【】\s]", "", text)
    for pattern in _TEMPLATE_PATTERNS:
        text = re.sub(pattern, "", text)
    return text.strip()


def _longest_common_substring_len(text_a, text_b):
    """最长连续公共片段长度。"""
    if not text_a or not text_b:
        return 0
    prev = [0] * (len(text_b) + 1)
    best = 0
    for ca in text_a:
        curr = [0]
        for idx, cb in enumerate(text_b, 1):
            val = prev[idx - 1] + 1 if ca == cb else 0
            curr.append(val)
            if val > best:
                best = val
        prev = curr
    return best


def _keyword_units(text):
    """用去模板后的连续二字片段近似表示核心关键词。"""
    text = _normalize_stem_for_dup(text)
    if len(text) < 2:
        return {text} if text else set()
    units = {text[i:i + 2] for i in range(len(text) - 1)}
    return {u for u in units if u and u not in _STOP_WORDS}


def _char_set_similarity(text_a, text_b):
    """字符集合相似度：共同字符数 ÷ 较短文本字符数。"""
    if not text_a or not text_b:
        return 0.0
    set_a = set(text_a)
    set_b = set(text_b)
    shorter = min(len(set_a), len(set_b))
    if shorter == 0:
        return 0.0
    return len(set_a & set_b) / shorter


def _keyword_similarity(text_a, text_b):
    """核心关键词相似度：共同关键词片段数 ÷ 较少关键词片段数。"""
    units_a = _keyword_units(text_a)
    units_b = _keyword_units(text_b)
    shorter = min(len(units_a), len(units_b))
    if shorter == 0:
        return 0.0
    return len(units_a & units_b) / shorter


def _is_duplicate_stem_pair(text_a, text_b):
    """判断两道题题干是否重复：连续短语 + 关键词相似度 + 去模板字符相似度。"""
    norm_a = _normalize_stem_for_dup(text_a)
    norm_b = _normalize_stem_for_dup(text_b)
    if not norm_a or not norm_b:
        return False, 0.0, "核心题干为空"

    shorter_len = min(len(norm_a), len(norm_b))
    phrase_len = _longest_common_substring_len(norm_a, norm_b)
    phrase_score = phrase_len / max(shorter_len, 1)
    keyword_score = _keyword_similarity(text_a, text_b)
    char_score = _char_set_similarity(norm_a, norm_b)
    score = max(phrase_score, keyword_score, char_score)

    if phrase_len >= 8 and phrase_score >= 0.55:
        return True, score, f"连续相同片段{phrase_len}字"
    if keyword_score >= 0.75 and char_score >= 0.60:
        return True, score, "核心关键词高度重合"
    if char_score >= 0.85 and shorter_len >= 6:
        return True, score, "去模板后题干高度相似"
    return False, score, "未达到重复阈值"


def _text_similarity(text_a, text_b):
    """兼容旧调用：返回改进后的题干重复度分数。"""
    return _is_duplicate_stem_pair(text_a, text_b)[1]


def cross_paper_check(paper_text_1, paper_text_2):
    """极重要双卷查重检测：检查（一）（二）两练之间的重复率

    Returns:
        issues: 问题列表
        stats: 统计信息 dict
    """
    issues = []

    # 提取两份试卷的题干
    stems_1 = re.findall(r"^\d+\.\s*(.+?)(?:（\s*）|\s*$)", paper_text_1, re.MULTILINE)
    stems_2 = re.findall(r"^\d+\.\s*(.+?)(?:（\s*）|\s*$)", paper_text_2, re.MULTILINE)

    if not stems_1 or not stems_2:
        return issues, {"error": "无法提取题目"}

    # 计算跨卷题干相似度
    high_sim_pairs = 0
    total_comparisons = min(len(stems_1), len(stems_2))

    for s1 in stems_1:
        for s2 in stems_2:
            sim = _text_similarity(s1, s2)
            if sim > 0.6:
                high_sim_pairs += 1
                break  # 每题只算一次

    repeat_rate = high_sim_pairs / max(total_comparisons, 1)

    stats = {
        "paper1_questions": len(stems_1),
        "paper2_questions": len(stems_2),
        "high_similarity_pairs": high_sim_pairs,
        "repeat_rate": repeat_rate,
    }

    if repeat_rate > 0.2:
        issues.append({
            "type": "双卷重复率过高",
            "severity": "严重",
            "question": "跨卷",
            "detail": f"两练之间相似题目{high_sim_pairs}对，重复率={repeat_rate:.0%}（超过20%限制）",
        })
    elif repeat_rate > 0.1:
        issues.append({
            "type": "双卷重复率偏高",
            "severity": "警告",
            "question": "跨卷",
            "detail": f"两练之间相似题目{high_sim_pairs}对，重复率={repeat_rate:.0%}",
        })

    return issues, stats


def ai_check(client, model, paper_text, spec_text):
    """调用 AI 进行深度质检"""
    system_prompt = f"""你是一位严格的试题质检专家。你的任务是根据以下质量标准，对试卷进行逐题审查，找出所有不合格的题目。

质量标准：
{spec_text}

你必须检查以下维度：
1. 【选项长度比】最长选项字数÷最短选项字数是否≤2.0
2. 【干扰项实质性】是否有"正常""无影响"等废选项
3. 【选项句式一致性】同题四选项是否结构一致
4. 【答案位置分布】单选ABCD是否均匀；多选答案A/B/C/D出现是否均衡，且多选题不得ABCD全选
5. 【难度分布】是否有适中题和困难题（计算/推理/综合）
6. 【计算题覆盖】是否至少有1题涉及数值
7. 【解析质量】是否为完整句，是否有禁用符号
8. 【知识正确性】答案是否正确，解析是否准确
9. 【题干质量】是否有自暴露/模糊表述/对话式用语
10. 【情景化比例】是否≥30%题目有具体场景"""

    user_prompt = f"""请对以下试卷进行严格质检，逐题审查并给出详细报告。

【试卷内容】
{paper_text}

【输出格式要求】
请按以下格式输出质检报告：

## 总体评价
- 总分（满分100）：X分
- 等级：优秀(≥90)/良好(80-89)/合格(70-79)/不合格(<70)

## 答案分布统计
A: X题  B: X题  C: X题  D: X题

## 难度分布
- 容易：X题
- 适中：X题
- 困难：X题

## 逐题问题清单
| 题号 | 问题类型 | 严重程度 | 具体问题 | 修改建议 |
|------|---------|---------|---------|---------|
（每个有问题的题目一行，无问题的跳过）

## 突出问题汇总
（列出最严重的3-5个问题）

## 修改建议
（给出具体的整改方向）"""

    return call_api(client, model, system_prompt, user_prompt, max_tokens=6000)


def print_local_report(issues, questions, filename):
    """打印本地检查报告"""
    print(f"\n{'='*60}")
    print(f"本地快速检查报告: {filename}")
    print(f"{'='*60}")
    print(f"提取到选择题: {len(questions)}道")

    if questions:
        choice_qs = [q for q in questions if _is_single_choice_answer(q["answer"])]
        if choice_qs:
            dist = {"A": 0, "B": 0, "C": 0, "D": 0}
            for q in choice_qs:
                dist[q["answer"]] += 1
            print(f"单选答案分布: A={dist['A']} B={dist['B']} C={dist['C']} D={dist['D']}")
        multi_qs = [q for q in questions if _is_multi_choice_answer(q["answer"])]
        if multi_qs:
            dist = {"A": 0, "B": 0, "C": 0, "D": 0}
            for q in multi_qs:
                for letter in q["answer"]:
                    dist[letter] += 1
            print(f"多选答案分布: A={dist['A']} B={dist['B']} C={dist['C']} D={dist['D']}")

    severe = [i for i in issues if i["severity"] == "严重"]
    warning = [i for i in issues if i["severity"] == "警告"]

    print(f"\n发现问题: 严重 {len(severe)} 个, 警告 {len(warning)} 个")

    if severe:
        print(f"\n{'─'*40}")
        print("【严重问题】")
        for issue in severe:
            print(f"  ✗ [第{issue['question']}题] {issue['type']}: {issue['detail']}")

    if warning:
        print(f"\n{'─'*40}")
        print("【警告】")
        for issue in warning:
            print(f"  ⚠ [第{issue['question']}题] {issue['type']}: {issue['detail']}")

    if not issues:
        print("\n  ✓ 本地快速检查未发现明显问题")

    # 评分
    score = 100
    score -= len(severe) * 15
    score -= len(warning) * 5
    score = max(0, score)
    print(f"\n本地检查评分: {score}/100", end="")
    if score >= 90:
        print(" (优秀)")
    elif score >= 80:
        print(" (良好)")
    elif score >= 70:
        print(" (合格)")
    else:
        print(" (不合格)")

    return score


def main():
    import argparse
    parser = argparse.ArgumentParser(description="一课一练试卷质检工具")
    parser.add_argument("--file", "-f", help="待检查的文件路径（.txt 或 .docx）")
    parser.add_argument("--dir", "-d", help="批量检查目录")
    parser.add_argument("--ai", action="store_true", default=True, help="启用 AI 深度检查（默认启用）")
    parser.add_argument("--local-only", action="store_true", help="仅本地检查，不调用 API")
    parser.add_argument("--output", "-o", help="报告输出目录")
    args = parser.parse_args()

    config = load_config()
    spec_text = load_spec()

    client = None
    if not args.local_only:
        client = OpenAI(api_key=config["api_key"], base_url=config["base_url"])

    # 确定要检查的文件
    files_to_check = []

    if args.file:
        files_to_check.append(Path(args.file))
    elif args.dir:
        target_dir = Path(args.dir)
        files_to_check = sorted(target_dir.glob("*.txt")) + sorted(target_dir.glob("*.docx"))
    else:
        # 交互式选择
        default_dir = BASE_DIR / "04_生成输出" / config.get("output_dir", "生成结果") / "_原始文本"
        if default_dir.exists():
            files = sorted(default_dir.glob("*.txt"))
        else:
            default_dir = BASE_DIR / "04_生成输出" / config.get("output_dir", "生成结果")
            files = sorted(default_dir.glob("*.docx"))

        if not files:
            print("没有找到可检查的文件。请使用 --file 或 --dir 指定路径。")
            return

        print("\n可检查的文件：")
        for i, f in enumerate(files, 1):
            print(f"  {i}. {f.name}")
        print(f"  {len(files)+1}. 全部检查")

        choice = input("\n请选择（编号或范围如 1-3）: ").strip()
        if choice == str(len(files) + 1) or choice.lower() == "all":
            files_to_check = files
        elif "-" in choice:
            s, e = choice.split("-")
            files_to_check = files[int(s)-1:int(e)]
        elif "," in choice:
            indices = [int(x.strip()) - 1 for x in choice.split(",")]
            files_to_check = [files[i] for i in indices]
        else:
            files_to_check = [files[int(choice) - 1]]

    if not files_to_check:
        print("没有文件可检查")
        return

    # 报告输出目录
    report_dir = Path(args.output) if args.output else BASE_DIR / "04_生成输出" / "质检报告"
    os.makedirs(report_dir, exist_ok=True)

    print(f"\n即将检查 {len(files_to_check)} 个文件")
    print(f"检查模式: {'本地+AI深度' if not args.local_only else '仅本地'}")
    print(f"报告目录: {report_dir}")
    print("=" * 60)

    results_summary = []

    for file_path in files_to_check:
        file_path = Path(file_path)
        print(f"\n▶ 检查: {file_path.name}")

        # 读取文件内容
        if file_path.suffix == ".docx":
            paper_text = extract_text_from_docx(str(file_path))
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                paper_text = f.read()

        if not paper_text.strip():
            print("  文件内容为空，跳过")
            continue

        # 本地快速检查
        issues, questions = local_check(paper_text)
        local_score = print_local_report(issues, questions, file_path.name)

        # AI 深度检查
        ai_report = ""
        if not args.local_only and client:
            print(f"\n  正在进行 AI 深度质检...")
            try:
                start_time = time.time()
                ai_report = ai_check(client, config["model"], paper_text, spec_text)
                elapsed = time.time() - start_time
                print(f"  AI 质检完成 ({elapsed:.1f}s)")
                print(f"\n{'─'*40}")
                print("【AI 深度质检报告】")
                print(ai_report)
            except Exception as e:
                print(f"  AI 质检失败: {e}")

        # 保存报告
        report_name = f"质检_{file_path.stem}.md"
        report_path = report_dir / report_name
        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"# 质检报告: {file_path.name}\n\n")
            f.write(f"检查时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"## 本地快速检查\n\n")
            f.write(f"- 评分: {local_score}/100\n")
            f.write(f"- 提取选择题: {len(questions)}道\n")
            if questions:
                choice_qs = [q for q in questions if _is_single_choice_answer(q["answer"])]
                if choice_qs:
                    dist = {"A": 0, "B": 0, "C": 0, "D": 0}
                    for q in choice_qs:
                        dist[q["answer"]] += 1
                    f.write(f"- 单选答案分布: A={dist['A']} B={dist['B']} C={dist['C']} D={dist['D']}\n")
                multi_qs = [q for q in questions if _is_multi_choice_answer(q["answer"])]
                if multi_qs:
                    dist = {"A": 0, "B": 0, "C": 0, "D": 0}
                    for q in multi_qs:
                        for letter in q["answer"]:
                            dist[letter] += 1
                    f.write(f"- 多选答案分布: A={dist['A']} B={dist['B']} C={dist['C']} D={dist['D']}\n")
            f.write(f"- 严重问题: {len([i for i in issues if i['severity']=='严重'])}个\n")
            f.write(f"- 警告: {len([i for i in issues if i['severity']=='警告'])}个\n\n")
            if issues:
                f.write("### 问题明细\n\n")
                f.write("| 题号 | 严重程度 | 类型 | 详情 |\n")
                f.write("|------|---------|------|------|\n")
                for issue in issues:
                    f.write(f"| {issue['question']} | {issue['severity']} | {issue['type']} | {issue['detail']} |\n")
                f.write("\n")
            if ai_report:
                f.write(f"## AI 深度质检报告\n\n")
                f.write(ai_report)
                f.write("\n")

        results_summary.append({
            "file": file_path.name,
            "local_score": local_score,
            "issues_severe": len([i for i in issues if i["severity"] == "严重"]),
            "issues_warning": len([i for i in issues if i["severity"] == "警告"]),
        })

        # 避免 API 限流
        if not args.local_only and len(files_to_check) > 1:
            time.sleep(3)

    # 批量汇总
    if len(results_summary) > 1:
        print(f"\n\n{'='*60}")
        print("批量质检汇总")
        print(f"{'='*60}")
        print(f"{'文件':<40} {'评分':>6} {'严重':>6} {'警告':>6}")
        print("─" * 60)
        for r in results_summary:
            name = r["file"][:38]
            print(f"{name:<40} {r['local_score']:>4}/100 {r['issues_severe']:>4} {r['issues_warning']:>4}")
        avg_score = sum(r["local_score"] for r in results_summary) / len(results_summary)
        print("─" * 60)
        print(f"{'平均':<40} {avg_score:>7.1f}")

        # 保存汇总
        summary_path = report_dir / "质检汇总.md"
        with open(summary_path, "w", encoding="utf-8") as f:
            f.write("# 批量质检汇总\n\n")
            f.write(f"检查时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"| 文件 | 评分 | 严重问题 | 警告 |\n")
            f.write(f"|------|------|---------|------|\n")
            for r in results_summary:
                f.write(f"| {r['file']} | {r['local_score']}/100 | {r['issues_severe']} | {r['issues_warning']} |\n")
            f.write(f"\n平均分: {avg_score:.1f}/100\n")

    print(f"\n质检报告已保存至: {report_dir}")


if __name__ == "__main__":
    main()
