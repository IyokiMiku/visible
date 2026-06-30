"""Word 文档操作公共工具（基于 python-docx）

统一管理一课一练 DOCX 的排版样式，新对话中生成试卷只需调用这里的功能即可保持一致格式。

核心功能：
  copy_template(tpl, out)        — 复制模板 docx，返回 Document
  set_margins(doc, t,b,l,r)      — 设页边距（默认 2.54/2.54/3.18/3.18 cm）
  add_editorial_note(doc, ...)    — 编写说明蓝框+分隔图（核心排版入口）
  add_paragraph_with_style(doc, ...) — 段落（中英自动分字体）
  add_heading_with_style(doc, ...)   — 标题（兼容中文模板缺样式）
  add_labeled_text(doc, ...)         — 【答案】/【解析】标签段落
  add_table_with_style(doc, ...)     — 带样式表格
  save_docx(doc, out)            — 保存文档

中英混排：所有文字输出函数通过 _add_mixed_run() 自动将拉丁字符设为
Times New Roman，中文字符保持指定字体（宋体/楷体/黑体）。

编写说明格式：
  - 蓝色虚线单格表格 (#4472C4, 1.5pt)
  - 楷体 10.5pt 加粗，字体颜色 #4472C4
  - 单倍行距，段前段后 4pt，首行缩进 2 字符
  - 可选分隔图：4.60cm 宽，左对齐，紧贴后续标题
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import copy
import os
import re
from pathlib import Path

try:
    from lxml import etree
except ImportError:  # 允许没有依赖时降级为普通文本公式
    etree = None

try:
    import latex2mathml.converter
except ImportError:  # 允许没有依赖时降级为普通文本公式
    latex2mathml = None

# -- 字体规则：中文保持指定字体，数字/字母统一 Times New Roman --
LATIN_FONT = 'Times New Roman'

def _has_latin(text):
    return bool(re.search(r'[a-zA-Z0-9]', text))

def _split_cjk_latin(text):
    """将文本拆分为中文/拉丁文交替的片段列表"""
    tokens = []
    current = ''
    current_type = None
    latin_pattern = re.compile(r'[a-zA-Z0-9\.\,\;\:\!\?\-\+\=\/\*\(\)\[\]\{\}\<\>\~\@\#\$\%\^\&\|\\\'\u03c0\u03b1-\u03c9\u0391-\u03a9\u2126\u00b0\u00d7\u00b7\u221a\u2264\u2265\u2248\u2260\u221e\u03bc\u2103]+')
    
    for ch in text:
        ch_type = 'latin' if latin_pattern.match(ch) else 'cjk'
        if current_type is None:
            current_type = ch_type
            current = ch
        elif ch_type == current_type:
            current += ch
        else:
            tokens.append((current, current_type))
            current = ch
            current_type = ch_type
    if current:
        tokens.append((current, current_type))
    return tokens

def _add_formula_run(para, formula_text, font_size, bold=False, font_color=None):
    """添加简易 LaTeX 公式片段：字母/数字 Times New Roman，并处理常见上下标。"""
    text = formula_text.strip()
    replacements = {
        r'\times': '×', r'\cdot': '·', r'\div': '÷', r'\pi': 'π',
        r'\le': '≤', r'\ge': '≥', r'\neq': '≠', r'\approx': '≈',
        r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\Delta': 'Δ',
        r'\Omega': 'Ω', r'\mu': 'μ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r'\\mathrm\{([^{}]+)\}', r'\1', text)
    text = re.sub(r'\\text\{([^{}]+)\}', r'\1', text)
    text = re.sub(r'\s+', ' ', text)

    i = 0
    while i < len(text):
        ch = text[i]
        if ch in ('_', '^') and i + 1 < len(text):
            is_sub = ch == '_'
            j = i + 1
            if text[j] == '{':
                k = text.find('}', j + 1)
                if k != -1:
                    payload = text[j + 1:k]
                    i = k + 1
                else:
                    payload = text[j + 1:]
                    i = len(text)
            else:
                payload = text[j]
                i = j + 1
            run = para.add_run(payload)
            run.font.name = LATIN_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), LATIN_FONT)
            run.font.size = Pt(font_size)
            run.bold = bold
            run.font.subscript = is_sub
            run.font.superscript = not is_sub
            if font_color:
                run.font.color.rgb = RGBColor(*font_color)
            continue
        run = para.add_run(ch)
        run.font.name = LATIN_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), LATIN_FONT)
        run.font.size = Pt(font_size)
        run.bold = bold
        if font_color:
            run.font.color.rgb = RGBColor(*font_color)
        i += 1


_MATH_WARNED = False
_MML2OMML_XSL = None


def _warn_math_fallback(reason):
    """公式转换不可用时只提示一次，并降级为普通文本。"""
    global _MATH_WARNED
    if not _MATH_WARNED:
        print(f"  [!] Word 原生公式转换不可用，公式将降级为普通文本：{reason}")
        _MATH_WARNED = True


def _split_math_markers(text):
    """拆分 {math:...}/{{math:...}} 公式标记和普通文本片段。

    公式内容本身经常含有 LaTeX 花括号，如 X_{L甲}、\frac{U}{R}。
    不能用简单的非贪婪正则匹配到第一个 } 就结束，否则会把公式截断。
    """
    source = text or ''
    parts = []
    pos = 0
    length = len(source)

    while pos < length:
        starts = []
        double_start = source.find('{{math:', pos)
        single_start = source.find('{math:', pos)
        if double_start != -1:
            starts.append((double_start, '{{math:', '}}'))
        if single_start != -1:
            starts.append((single_start, '{math:', '}'))
        if not starts:
            parts.append(('text', source[pos:]))
            break

        start, opener, closer = min(starts, key=lambda item: item[0])
        if start > pos:
            parts.append(('text', source[pos:start]))

        content_start = start + len(opener)
        i = content_start
        depth = 0
        end = -1
        if closer == '}}':
            # {{math:...}}：公式内部允许单个 { } 成对出现；只有 depth=0 的 }} 才结束。
            while i < length:
                if source.startswith('}}', i) and depth == 0:
                    end = i
                    break
                if source[i] == '{':
                    depth += 1
                elif source[i] == '}' and depth > 0:
                    depth -= 1
                i += 1
            marker_end = end + 2 if end != -1 else -1
        else:
            # {math:...}：公式内部允许 LaTeX 单花括号成对出现；depth=0 的 } 才结束。
            while i < length:
                ch = source[i]
                if ch == '{':
                    depth += 1
                elif ch == '}':
                    if depth == 0:
                        end = i
                        break
                    depth -= 1
                i += 1
            marker_end = end + 1 if end != -1 else -1

        if end == -1:
            # 标记不完整时保守当普通文本处理，避免误删内容。
            parts.append(('text', source[start:]))
            break

        parts.append(('math', source[content_start:end].strip()))
        pos = marker_end

    return [(kind, chunk) for kind, chunk in parts if chunk]


def _find_mml2omml_xsl():
    """查找 Microsoft Office 自带的 MathML→OMML 转换样式表。"""
    global _MML2OMML_XSL
    if _MML2OMML_XSL is not None:
        return _MML2OMML_XSL

    candidates = []
    program_roots = [
        os.environ.get('ProgramFiles'),
        os.environ.get('ProgramFiles(x86)'),
    ]
    for root in program_roots:
        if not root:
            continue
        root_path = Path(root) / 'Microsoft Office'
        candidates.extend(root_path.glob('**/MML2OMML.XSL'))

    for path in candidates:
        if path.exists():
            _MML2OMML_XSL = path
            return path

    _MML2OMML_XSL = False
    return None


def _normalize_latex_for_mathml(latex):
    """将少量中文命题常用写法规范为 latex2mathml 更易识别的形式。"""
    text = (latex or '').strip()
    text = text.strip('$')
    replacements = {
        '×': r'\times ',
        '·': r'\cdot ',
        '÷': r'\div ',
        '≈': r'\approx ',
        '≤': r'\le ',
        '≥': r'\ge ',
        '≠': r'\ne ',
        'ρ': r'\rho ',
        'Ω': r'\Omega ',
        'Φ': r'\Phi ',
        'φ': r'\phi ',
        'Δ': r'\Delta ',
        'μ': r'\mu ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _latex_to_plain_text(latex):
    """公式转换失败时，把常见 LaTeX 降级成可读纯文本，避免 DOCX 出现反斜杠乱码。"""
    text = (latex or '').strip().strip('$')
    text = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'\1/\2', text)
    text = re.sub(r'\\(?:mathrm|text)\{([^{}]+)\}', r'\1', text)
    replacements = {
        r'\times': '×',
        r'\cdot': '·',
        r'\div': '÷',
        r'\le': '≤',
        r'\ge': '≥',
        r'\neq': '≠',
        r'\ne': '≠',
        r'\approx': '≈',
        r'\pi': 'π',
        r'\rho': 'ρ',
        r'\Omega': 'Ω',
        r'\Phi': 'Φ',
        r'\phi': 'φ',
        r'\Delta': 'Δ',
        r'\mu': 'μ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r'\\(?:left|right)\s*', '', text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])_\{([^{}]+)\}', r'\1_\2', text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])\^\{([^{}]+)\}', r'\1^\2', text)
    text = text.replace('{', '').replace('}', '')
    text = re.sub(r'\\+', '', text)
    return re.sub(r'\s+', ' ', text).strip()


def _latex_to_omml(latex):
    """把 LaTeX/线性公式转换为 Word 原生 OMML XML 元素。失败返回 None。"""
    if etree is None:
        _warn_math_fallback('缺少 lxml，请先安装 pip install lxml latex2mathml')
        return None
    if latex2mathml is None:
        _warn_math_fallback('缺少 latex2mathml，请先安装 pip install latex2mathml')
        return None

    xsl_path = _find_mml2omml_xsl()
    if not xsl_path:
        _warn_math_fallback('未找到 Microsoft Office 的 MML2OMML.XSL')
        return None

    try:
        mathml = latex2mathml.converter.convert(_normalize_latex_for_mathml(latex))
        mathml_root = etree.fromstring(mathml.encode('utf-8'))
        transform = etree.XSLT(etree.parse(str(xsl_path)))
        omml_tree = transform(mathml_root)
        omml_root = omml_tree.getroot()
        # MML2OMML.XSL 常返回 m:oMathPara；段落内混排时只插入其中的 m:oMath。
        if omml_root.tag == qn('m:oMathPara'):
            inline_math = omml_root.find(qn('m:oMath'))
            if inline_math is not None:
                omml_root = inline_math
        return parse_xml(etree.tostring(omml_root, encoding='unicode'))
    except Exception as exc:
        _warn_math_fallback(f'公式“{latex}”转换失败：{exc}')
        return None


def _set_omml_color(omml_element, font_color):
    """给 OMML 内部 run 设置字体颜色。"""
    if not font_color:
        return
    color = ''.join(f'{value:02X}' for value in font_color)
    # python-docx 的 BaseOxmlElement.xpath() 不接收 namespaces 参数；直接按
    # Clark name 遍历，避免 lxml 抛出 Undefined namespace prefix。
    for run in omml_element.iter(qn('m:r')):
        rpr = run.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run.insert(0, rpr)
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = OxmlElement('w:color')
            rpr.append(color_el)
        color_el.set(qn('w:val'), color)


def _append_omml_to_paragraph(para, omml_element, font_color=None):
    """向段落追加 Word 原生公式节点。"""
    _set_omml_color(omml_element, font_color)
    para._p.append(omml_element)


def _add_math_marker_run(para, formula_text, cjk_font, font_size, bold, font_color):
    """添加 {{math:...}} 公式；不可转换时降级为清洗后的普通文本。"""
    omml = _latex_to_omml(formula_text)
    if omml is not None:
        _append_omml_to_paragraph(para, omml, font_color)
    else:
        _add_formula_run(para, _latex_to_plain_text(formula_text), font_size, bold, font_color)


def _add_rich_text_run(para, text, cjk_font, font_size, bold, font_color, latin_font=LATIN_FONT):
    """添加文本；{{math:...}} 片段生成 Word 原生公式，$...$ 片段按简易公式处理。"""
    text = re.sub(r'([A-Za-z0-9\)）])\s*乘以?\s*([A-Za-z0-9\(（])', r'\1×\2', text or '')
    for kind, chunk in _split_math_markers(text):
        if not chunk:
            continue
        if kind == 'math':
            _add_math_marker_run(para, chunk, cjk_font, font_size, bold, font_color)
            continue
        parts = re.split(r'(\$[^$]+\$)', chunk)
        for part in parts:
            if not part:
                continue
            if len(part) >= 2 and part.startswith('$') and part.endswith('$'):
                _add_formula_run(para, part[1:-1], font_size, bold, font_color)
            else:
                _add_mixed_run(para, part, cjk_font, font_size, bold, font_color, latin_font)


def _add_mixed_run(para, text, cjk_font, font_size, bold, font_color, latin_font=LATIN_FONT):
    """添加段落文本，自动分中英文设置字体"""
    tokens = _split_cjk_latin(text)
    for seg_text, seg_type in tokens:
        run = para.add_run(seg_text)
        font = latin_font if seg_type == 'latin' else cjk_font
        run.font.name = font
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        run.font.size = Pt(font_size)
        run.bold = bold
        if font_color:
            run.font.color.rgb = RGBColor(*font_color)


def copy_template(template_path: str, output_path: str) -> Document:
    """复制 Word 模板到输出路径，返回 Document 对象（保留页眉页脚）"""
    doc = Document(template_path)
    doc.save(output_path)
    return Document(output_path)


def set_margins(doc: Document, top: float = 2.54, bottom: float = 2.54,
                left: float = 3.18, right: float = 3.18):
    """设置页边距（单位：厘米）"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_heading_with_style(doc: Document, text: str, level: int = 1,
                           font_name: str = "黑体", font_size: int = 16):
    """添加带样式的标题（兼容中文Word模板缺少Heading样式的场景）"""
    try:
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
        return heading
    except KeyError:
        # 中文Word模板没有"Heading N"样式，退化为加粗段落
        return add_paragraph_with_style(
            doc, text, font_name=font_name, font_size=font_size,
            bold=True, space_after=12
        )


def add_paragraph_with_style(doc: Document, text: str, font_name: str = "宋体",
                             font_size: float = 12, bold: bool = False,
                             font_color: tuple = None,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 6):
    """添加带样式的段落，中英文自动分字体（中文=font_name，数字字母=Times New Roman）

    Args:
        font_color: RGB 颜色元组，如 (255,0,0) 为红色，None 为默认黑色
    """
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.5
    _add_rich_text_run(para, text, font_name, font_size, bold, font_color)
    return para


def add_table_with_style(doc: Document, headers: list, rows: list,
                         style: str = "Table Grid"):
    """添加带样式的表格，headers 为表头列表，rows 为数据行列表（每行一个 list）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
                run.font.size = Pt(10)
                run.bold = True

    # 数据行
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                    run.font.size = Pt(10)

    return table


def append_docx_after_template(template_path: str, content_path: str) -> str:
    """将内容文档追加到模板后面，返回合并后的文档路径"""
    template_doc = Document(template_path)
    content_doc = Document(content_path)

    for element in content_doc.element.body:
        template_doc.element.body.append(copy.deepcopy(element))

    merged_path = content_path.replace(".docx", "_merged.docx")
    template_doc.save(merged_path)
    return merged_path


def save_docx(doc: Document, output_path: str) -> str:
    """保存 Document 到文件"""
    doc.save(output_path)
    return output_path


def add_labeled_text(doc: Document, label: str, content: str,
                     label_font: str = "黑体", label_size: float = 10.5,
                     content_font: str = "宋体", content_size: float = 10.5,
                     color: tuple = None, bold_label: bool = True,
                     bold_content: bool = False):
    """添加标签+正文混合样式段落（标签和正文可分别设字体/颜色）

    用于【答案】xxx 和【解析】xxx 等场景

    Args:
        label: 标签文字（如"【答案】"）
        content: 正文内容
        label_font: 标签字体
        label_size: 标签字号（pt）
        content_font: 正文字体
        content_size: 正文字号（pt）
        color: RGB 颜色元组，如 (255,0,0)，None 为黑色
        bold_label: 标签是否加粗
        bold_content: 正文是否加粗
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.5

    # 标签 run
    label_run = para.add_run(label)
    label_run.font.name = label_font
    label_run._element.rPr.rFonts.set(qn('w:eastAsia'), label_font)
    label_run.font.size = Pt(label_size)
    label_run.bold = bold_label
    if color:
        label_run.font.color.rgb = RGBColor(*color)

    # 正文 run —— 自动分中英文设置字体；$...$ 片段按简易 LaTeX 公式处理
    _add_rich_text_run(para, content, content_font, content_size, bold_content, color)

    return para


def add_editorial_note(doc: Document, textbook_name: str, edition: str,
                       chapter_seq: int, knowledge_scope: str,
                       province: str = "重庆市",
                       publisher: str = "高教版",
                       separator_image: str = None,
                       exam_type: str = "高职分类考试"):
    """添加编写说明表格 — 蓝色虚线边框表格，单倍行距，段前段后4pt，首行缩进2字符

    Args:
        province: 省份名，如"重庆市"；自治区应使用全称，如"内蒙古自治区"、"新疆维吾尔自治区"、"西藏自治区"
        publisher: 出版社简称，如"高教版"、"机工版"，默认"高教版"
        separator_image: 可选的分隔图片路径，插入在表格与标题之间
        exam_type: 考试类型，如"高职分类考试"、"高职分类考试"
    """
    prefix = (
        f'编写说明：考虑到中职学生普遍基础知识相对薄弱的情况，我们依据支架式教学理念，'
        f'精心编制了{province}（{exam_type}）《{textbook_name}》（{publisher}·{edition}）一课一练。'
        f'专辑里的每一份练习，都与课堂所授知识点紧密相关，题目围绕课堂所学知识点呈现。'
        f'目的在于激发学生的学习兴趣，培养他们的学习自觉性，帮助学生扎实掌握课程的基本概念与基本方法，'
        f'为他们后续的逐步提升奠定坚实基础。'
    )
    info = (
        f'本卷是{province}（{exam_type}）《{textbook_name}》（{publisher}·{edition}）一课一练的第{chapter_seq}练，'
        f'内容涵盖{knowledge_scope}。'
    )

    # 创建 1行1列表格
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True

    # 去除模板自带的顶部空行（表格前第一个空白段落）
    body = doc.element.body
    first_p = body.find(qn('w:p'))
    if first_p is not None:
        first_t = first_p.find(qn('w:r'))
        if first_t is None or not (first_t.text or '').strip():
            body.remove(first_p)

    # 设置表格蓝色虚线边框 1.5pt
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'dashed')
        border.set(qn('w:sz'), '12')  # 1.5pt = 12/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        borders.append(border)
    tblPr.append(borders)

    # 设置表格宽度为页面宽度
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    cell = table.rows[0].cells[0]

    # 第一段
    p1 = cell.paragraphs[0]
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.first_line_indent = Cm(0.74)
    _add_mixed_run(p1, prefix, "楷体", 10.5, True, (0x44, 0x72, 0xC4))

    # 第二段
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.first_line_indent = Cm(0.74)
    _add_mixed_run(p2, info, "楷体", 10.5, True, (0x44, 0x72, 0xC4))

    # -- 可选分隔图片（表格与标题之间，左对齐，紧贴标题）--
    if separator_image and os.path.exists(separator_image):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        img_para.paragraph_format.space_before = Pt(0)
        img_para.paragraph_format.space_after = Pt(0)
        run = img_para.add_run()
        run.add_picture(separator_image, width=Cm(4.60))
    else:
        doc.add_paragraph()  # 无图片时表格与标题之间空行


# -- 选项布局自动适配阈值 --
_OPTION_SINGLE_COL_THRESHOLD = 18    # 任一选项超此字数 → 单列（双列表格单格约能放18个中文字符）
_OPTION_ROW_BALANCE_RATIO = 1.6       # 同排两选项长度比超此值 → 单列（避免短选项下方大量空白）


def _should_use_single_col(options):
    """判断是否应该使用单列布局（非双列）。

    触发条件（任一满足即单列）：
      1. 任一选项长度 > _OPTION_SINGLE_COL_THRESHOLD
      2. 同排两个选项长度悬殊（比例 > _OPTION_ROW_BALANCE_RATIO），
         双列会导致短选项下方留大片空白
    """
    n = len(options)
    if n <= 1:
        return True

    # 条件1：长选项
    if max(len(o) for o in options) > _OPTION_SINGLE_COL_THRESHOLD:
        return True

    # 条件2：同排选项长度不平衡
    # 对于A/B/C/D四选项，配对为 (A,B) 和 (C,D)
    for i in range(0, n - 1, 2):
        a_len = len(options[i])
        b_len = len(options[i + 1])
        shorter = max(min(a_len, b_len), 1)
        if max(a_len, b_len) / shorter > _OPTION_ROW_BALANCE_RATIO:
            return True

    return False


def add_question_options_table(doc: Document, q_text: str,
                               font_name: str = "宋体", font_size: float = 10.5):
    """解析题目文本，用无边框表格输出选项，确保列对齐。

    输入格式示例：
      '1. 题干文本\\nA. 选项A\\t\\tB. 选项B\\nC. 选项C\\t\\tD. 选项D'

    自动适配布局：
      - 短且均衡 → 双列表格（紧凑省空间）
      - 长选项 或 同排长度悬殊 → 单列表格（避免挤窄/空白）
      - 非选项文本 → 普通段落
    """
    # -- 不包含制表符分隔的选项 → 普通段落 --
    if '\t' not in q_text:
        add_paragraph_with_style(doc, q_text, font_name, font_size, space_after=0)
        return

    # -- 按行分割 --
    lines = q_text.split('\n')
    stem = lines[0].strip()

    # -- 从剩余行提取选项（按制表符分割） --
    options = []
    for line in lines[1:]:
        parts = re.split(r'\t+', line)
        for p in parts:
            p = p.strip()
            if p:
                options.append(p)

    if not options:
        add_paragraph_with_style(doc, q_text, font_name, font_size, space_after=0)
        return

    # -- 题干段落 --
    add_paragraph_with_style(doc, stem, font_name, font_size, space_after=0)

    # -- 自动选择布局 --
    n = len(options)
    if _should_use_single_col(options):
        # 单列布局：每行一个选项
        rows, cols = n, 1
    elif n <= 2:
        rows, cols = 1, n
    else:
        # 双列布局：每行两个选项
        rows = (n + 1) // 2
        cols = 2

    table = doc.add_table(rows=rows, cols=cols)

    # -- 表格宽度100% --
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # -- 无边框 --
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)

    # -- 填充选项 --
    for i, opt in enumerate(options):
        if cols == 1:
            row_idx, col_idx = i, 0
        else:
            row_idx = i // cols
            col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        # 清空默认段落
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        _add_rich_text_run(p, opt, font_name, font_size, False, None)

    return table
