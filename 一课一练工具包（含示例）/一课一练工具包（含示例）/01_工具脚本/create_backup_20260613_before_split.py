"""一课一练试卷生成器 — 读取考点规划表，调用 Claude API 生成高质量试卷 DOCX

使用方法：
  python create.py                        # 交互式选择规划表和范围
  python create.py --file 规划表.xlsx      # 指定规划表
  python create.py --file 规划表.xlsx --range 1-5  # 生成第1~5个主题
"""

import json
import os
import re
import sys
import time
import traceback
from pathlib import Path

# 强制 stdout/stderr 使用 UTF-8，避免 Windows GBK 编码问题导致特殊字符打印崩溃
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding.lower() != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

import openpyxl
from openai import OpenAI

# === 路径常量 ===
BASE_DIR = Path(__file__).parent
CONFIG_PATH = BASE_DIR / "02_配置资源" / "config.json"
SPEC_PATH = BASE_DIR / "02_配置资源" / "编写规范" / "编写规范.md"
TEMPLATE_PATH = BASE_DIR / "02_配置资源" / "模板和资源" / "template.docx"
SEPARATOR_PATH = BASE_DIR / "02_配置资源" / "模板和资源" / "separator.png"
QUESTION_TYPES_DIR = BASE_DIR / "02_配置资源" / "题型定义"
REF_DIR = BASE_DIR / "03_项目数据" / "参考资料"

# === 加载配置 ===
def load_config():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

# === 加载编写规范 ===
def load_spec():
    with open(SPEC_PATH, "r", encoding="utf-8") as f:
        return f.read()


# === 加载参考资料 ===
MAX_EXAM_REF_CHARS = 5000      # 历年真题原文最大字符数
MAX_TEXTBOOK_REF_CHARS = 6000  # 教材参考最大字符数
MAX_STYLE_REF_CHARS = 5000     # 真题风格库最大字符数
MAX_REF_CHARS = 12000          # 全部参考资料合并后的最大字符数


def _read_limited_text(path, limit):
    """读取文本并按字符数截断。"""
    text = path.read_text(encoding="utf-8")
    if len(text) > limit:
        text = text[:limit] + "\n\n...（内容已截断，以上内容足够参考）"
    return text


def _style_dir_candidates(style_dir, meta):
    """按优先级返回真题风格库目录候选。"""
    province = meta.get("province", "")
    category = meta.get("category", "")
    exam_type = "高职分类考试" if "高职" in meta.get("config_line", "") else "高职分类考试"

    candidates = []
    if province and category:
        candidates.append(style_dir / province / category)
    if province:
        candidates.append(style_dir / province / "通用")
    candidates.append(style_dir / "通用" / exam_type)
    candidates.append(style_dir / "通用")

    # 去重并保序
    seen = set()
    result = []
    for path in candidates:
        key = str(path)
        if key not in seen:
            seen.add(key)
            result.append(path)
    return result


def _collect_style_files_from_dir(directory):
    """按固定顺序收集某个风格库目录下的 txt 文件。"""
    preferred = [
        "风格总则.txt",
        "单选题风格.txt",
        "多选题风格.txt",
        "判断题风格.txt",
        "填空题风格.txt",
        "综合题风格.txt",
        "代表样题.txt",
    ]
    files = []
    for name in preferred:
        path = directory / name
        if path.exists():
            files.append(path)
    for path in sorted(directory.glob("*.txt")):
        if path not in files:
            files.append(path)
    return files


def _load_structured_style_reference(style_dir, meta):
    """加载 省份/考类 结构化真题风格库。"""
    for directory in _style_dir_candidates(style_dir, meta):
        if not directory.exists():
            continue
        files = _collect_style_files_from_dir(directory)
        if not files:
            continue

        parts = []
        remaining = MAX_STYLE_REF_CHARS
        for path in files:
            if remaining <= 0:
                break
            text = _read_limited_text(path, remaining)
            parts.append(f"【{path.stem}】\n{text}")
            remaining -= len(text)
        if parts:
            return f"【风格库目录：{directory.relative_to(style_dir)}】\n" + "\n\n".join(parts)
    return ""


def _load_reference_materials(topic, meta):
    """加载与当前主题相关的参考资料（真题+真题风格+教材）

    匹配策略：
      - 真题：按类别关键词匹配文件名（如"汽车类"）
      - 真题风格：优先按省份+类别匹配，其次按类别/课程/主题关键词匹配
      - 教材：按主题名/节名/课程名模糊匹配文件名
    """
    if not REF_DIR.exists():
        return ""

    materials = []
    total_len = 0

    # 1. 匹配真题文件（按类别匹配）
    exam_dir = REF_DIR / "真题"
    if exam_dir.exists():
        category = meta.get("category", "")
        province = meta.get("province", "")
        matched_files = []
        for f in exam_dir.glob("*.txt"):
            score = 0
            if category and category in f.name:
                score += 2
            if province and province in f.name:
                score += 1
            if score > 0:
                matched_files.append((score, f))

        if matched_files:
            matched_files.sort(key=lambda x: x[0], reverse=True)
            best_file = matched_files[0][1]
            text = _read_limited_text(best_file, MAX_EXAM_REF_CHARS)
            materials.append(f"【历年真题参考——请模仿以下真题的出题风格、难度和措辞习惯，但不得照搬题干、选项或情境】\n{text}")
            total_len += len(text)

    # 2. 匹配真题风格库（优先 省份/考类 目录结构，其次兼容旧版扁平 txt）
    style_dir = REF_DIR / "真题风格"
    if style_dir.exists() and total_len < MAX_REF_CHARS:
        style_text = _load_structured_style_reference(style_dir, meta)

        if not style_text:
            category = meta.get("category", "")
            province = meta.get("province", "")
            theme = topic.get("theme", "")
            section = topic.get("section", "")
            course = topic.get("course", "")

            style_keywords = []
            if province and category:
                style_keywords.append((f"{province}{category}", 5))
                style_keywords.append((f"{province} {category}", 5))
            if category:
                style_keywords.append((category, 3))
            if province:
                style_keywords.append((province, 1))
            for kw in (course, section, theme):
                if kw:
                    clean_kw = re.sub(r'^[\s（(一二三四五六七八九十）)\d．.、]+', '', kw).strip()
                    if clean_kw:
                        style_keywords.append((clean_kw[:8], 1))

            matched_files = []
            for f in style_dir.glob("*.txt"):
                fname = f.stem
                score = sum(weight for kw, weight in style_keywords if kw and kw in fname)
                if score > 0:
                    matched_files.append((score, f))

            if matched_files:
                matched_files.sort(key=lambda x: x[0], reverse=True)
                selected_files = [f for _, f in matched_files[:2]]
                style_parts = []
                remaining = MAX_STYLE_REF_CHARS
                for style_file in selected_files:
                    if remaining <= 0:
                        break
                    text = _read_limited_text(style_file, remaining)
                    style_parts.append(f"【{style_file.stem}】\n{text}")
                    remaining -= len(text)
                style_text = "\n\n".join(style_parts)

        if style_text:
            materials.append(
                "【真题风格参考——只模仿口吻、设问方式、选项长度、干扰项风格和解析简洁程度；"
                "不得照搬样题内容，知识准确性仍以当前考纲和教材为准】\n"
                f"{style_text}"
            )
            total_len += len(style_text)

    # 3. 匹配教材文件（按课程/章节/主题名匹配）
    textbook_dir = REF_DIR / "教材"
    if textbook_dir.exists() and total_len < MAX_REF_CHARS:
        theme = topic.get("theme", "")
        section = topic.get("section", "")
        course = topic.get("course", "")

        # 提取关键词用于匹配
        keywords = []
        if theme:
            keywords.append(theme)
        if section:
            sec_clean = re.sub(r'^[\s（(一二三四五六七八九十）)\d．.、]+', '', section)
            if sec_clean:
                keywords.append(sec_clean[:6])
        if course:
            course_clean = re.sub(r'课程[一二三四五六七八九十\d]+[：:]', '', course)
            course_clean = re.sub(r'[（(].+[）)]', '', course_clean).strip()
            if course_clean:
                keywords.append(course_clean)

        matched_files = []
        for f in textbook_dir.glob("*.txt"):
            fname = f.stem
            score = sum(1 for kw in keywords if kw and kw in fname)
            if score > 0:
                matched_files.append((score, f))

        if matched_files:
            matched_files.sort(key=lambda x: x[0], reverse=True)
            best_file = matched_files[0][1]
            text = _read_limited_text(best_file, MAX_TEXTBOOK_REF_CHARS)
            materials.append(f"【教材参考内容——出题知识点必须与以下教材内容一致】\n{text}")
            total_len += len(text)

    if not materials:
        return ""

    combined = "\n\n".join(materials)

    # 超长截断
    if len(combined) > MAX_REF_CHARS:
        combined = combined[:MAX_REF_CHARS] + "\n\n...（参考资料已截断，以上内容足够参考）"

    return combined


# === Token 用量跟踪 ===
_USAGE_FILE = BASE_DIR / ".token_usage.json"


def _load_daily_usage():
    """加载当天的累计 token 用量"""
    today = time.strftime("%Y-%m-%d")
    if _USAGE_FILE.exists():
        try:
            with open(_USAGE_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("date") == today:
                return data
        except (json.JSONDecodeError, KeyError):
            pass
    return {"date": today, "prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}


def _save_daily_usage(usage):
    """保存当天的累计 token 用量"""
    with open(_USAGE_FILE, "w", encoding="utf-8") as f:
        json.dump(usage, f, ensure_ascii=False)


def _print_token_summary(session_usage, daily_usage):
    """打印 token 用量汇总"""
    print(f"\n{'─' * 40}")
    print(f"  本次会话 Token 消耗:")
    print(f"    输入: {session_usage['prompt_tokens']:,} tokens")
    print(f"    输出: {session_usage['completion_tokens']:,} tokens")
    print(f"    合计: {session_usage['total_tokens']:,} tokens ({session_usage['api_calls']} 次调用)")
    print(f"  今日累计 Token 消耗:")
    print(f"    输入: {daily_usage['prompt_tokens']:,} tokens")
    print(f"    输出: {daily_usage['completion_tokens']:,} tokens")
    print(f"    合计: {daily_usage['total_tokens']:,} tokens ({daily_usage['api_calls']} 次调用)")
    print(f"{'─' * 40}")


# === API 调用 ===
def call_api(client, model, system_prompt, user_prompt, max_tokens=8000, temperature=0.7):
    """调用 OpenAI 兼容 API（支持 Claude 代理），返回 (文本, usage_dict)"""
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            # 提取 token 用量
            usage = None
            if hasattr(response, 'usage') and response.usage:
                usage = {
                    "prompt_tokens": getattr(response.usage, 'prompt_tokens', 0) or 0,
                    "completion_tokens": getattr(response.usage, 'completion_tokens', 0) or 0,
                    "total_tokens": getattr(response.usage, 'total_tokens', 0) or 0,
                }
            return response.choices[0].message.content, usage
        except Exception as e:
            print(f"  [!] API 调用失败 (第{attempt+1}次): {e}")
            if attempt < 2:
                wait = (attempt + 1) * 10
                print(f"      等待 {wait}s 后重试...")
                time.sleep(wait)
            else:
                raise

# === 版次格式化辅助函数 ===
_DIGIT_TO_CN = {"1": "一", "2": "二", "3": "三", "4": "四", "5": "五",
                "6": "六", "7": "七", "8": "八", "9": "九", "10": "十"}

_CN_TO_DIGIT = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                "六": 6, "七": 7, "八": 8, "九": 9, "十": 10,
                "十一": 11, "十二": 12, "十三": 13, "十四": 14, "十五": 15}


def _parse_edition(raw_edition: str) -> tuple:
    """将原始版次字符串解析为 (出版社简称, 中文版次)。

    示例：
        '高教第3版' → ('高教版', '第三版')
        '机工第2版' → ('机工版', '第二版')
        '第1版'     → ('高教版', '第一版')  # 默认出版社
    """
    pub_match = re.match(r"([\u4e00-\u9fa5]+?)第(\d+)版", raw_edition)
    if pub_match:
        pub_short = pub_match.group(1)
        ver_num = pub_match.group(2)
    else:
        ver_match = re.search(r"第(\d+)版", raw_edition)
        pub_short = "高教"
        ver_num = ver_match.group(1) if ver_match else "1"

    if not pub_short.endswith("版"):
        pub_short += "版"
    edition_cn = f"第{_DIGIT_TO_CN.get(ver_num, ver_num)}版"
    return pub_short, edition_cn


# === 解析规划表 ===
def parse_planning_table(xlsx_path):
    """解析考点规划表，返回元数据和主题列表"""
    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    ws = wb[wb.sheetnames[0]]

    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    # 提取头部元数据
    meta = {
        "title": str(rows[0][0] or "").strip(),
        "config_line": str(rows[1][0] or "").strip(),
        "textbooks": str(rows[2][0] or "").strip(),
    }

    # 解析题型配置
    config_line = meta["config_line"]
    qt_match = re.search(r"题型[：:]\s*(.+?)\s*[|｜]", config_line)
    diff_match = re.search(r"难度[：:]\s*([\d:]+)", config_line)
    meta["question_types_str"] = qt_match.group(1) if qt_match else ""
    meta["difficulty_str"] = diff_match.group(1) if diff_match else "80:10:10"

    # 解析教材（优先从 row2 中提取，若无则留空后面从课程行补充）
    tb_text = meta["textbooks"]
    textbooks = re.findall(r"《(.+?)》([^\s、《》]+第\d+版)", tb_text)
    meta["textbook_list"] = textbooks

    # 解析出版社和版次（用于文件名和标题中的规范格式）
    meta["textbook_details"] = []
    for tb_name, tb_edition_raw in textbooks:
        publisher, edition_cn = _parse_edition(tb_edition_raw)
        meta["textbook_details"].append({
            "name": tb_name,
            "publisher": publisher,
            "edition": edition_cn,
            "display": f"{publisher}·{edition_cn}",
        })

    # 也尝试从 row1 中查找教材信息（有些表把教材放在配置行）
    if not textbooks:
        for check_row in rows[:5]:
            check_text = str(check_row[0] or "")
            found = re.findall(r"《(.+?)》([^\s、《》]+第\d+版)", check_text)
            if found:
                textbooks = found
                meta["textbook_list"] = textbooks
                for tb_name, tb_edition_raw in textbooks:
                    publisher, edition_cn = _parse_edition(tb_edition_raw)
                    meta["textbook_details"].append({
                        "name": tb_name,
                        "publisher": publisher,
                        "edition": edition_cn,
                        "display": f"{publisher}·{edition_cn}",
                    })
                break

    # 解析省份
    province_match = re.search(r"([\u4e00-\u9fa5]+(?:省|市))", meta["title"])
    meta["province"] = province_match.group(1) if province_match else "重庆市"

    # 解析类别
    cat_match = re.search(r"(?:省|市)([\u4e00-\u9fa5]+类)", meta["title"])
    meta["category"] = cat_match.group(1) if cat_match else ""

    # 找到表头行
    header_idx = None
    for i, row in enumerate(rows):
        if row[0] == "序号":
            header_idx = i
            break

    if header_idx is None:
        raise ValueError("找不到表头行（'序号'列）")

    # 解析主题
    topics = []
    current_course = ""
    current_section = ""

    for row in rows[header_idx + 1:]:
        seq = row[0]
        knowledge = row[1]
        theme = row[2]
        level = row[3]
        q_types = row[4]
        difficulty = row[5]
        sets = row[6]
        exam_ref = row[7]

        # 课程行
        if isinstance(seq, str) and "课程" in seq:
            current_course = seq.strip()
            continue
        # 章节行 — 格式1: 序号列为空，知识点列含"一、xxx"
        if seq is None and knowledge and knowledge.strip().startswith(("一", "二", "三", "四", "五", "六", "七", "八", "九", "十")):
            current_section = knowledge.strip()
            continue
        # 章节行 — 格式2: 序号列含"(一) xxx"或"（一）xxx"
        if isinstance(seq, str) and re.match(r'\s*[（(][一二三四五六七八九十百]+[）)]\s*\S', seq):
            current_section = seq.strip()
            continue
        # 章节行 — 格式3: 序号列为缩进文本（如"  绪论""  1．力系与平衡"），知识点列为空
        if isinstance(seq, str) and seq.startswith(" ") and not knowledge:
            current_section = seq.strip()
            continue
        # 空行或汇总行
        if seq is None or not isinstance(seq, (int, float)):
            continue

        sets_count = int(sets) if sets else 1
        topics.append({
            "seq": int(seq),
            "knowledge": str(knowledge or "").strip(),
            "theme": str(theme or "").strip(),
            "level": str(level or "标准").strip(),
            "question_types": str(q_types or "").strip(),
            "difficulty": str(difficulty or "80:10:10").strip(),
            "sets": sets_count,
            "exam_ref": str(exam_ref or "").strip(),
            "course": current_course,
            "section": current_section,
        })

    # 构建课程→教材映射（按出现顺序，课程一对应教材一，依次类推）
    course_names = []
    for row in rows[header_idx + 1:]:
        val = row[0]
        if isinstance(val, str) and "课程" in val:
            course_names.append(val.strip())
    meta["course_textbook_map"] = {}
    if meta["textbook_details"] and course_names:
        for idx, course in enumerate(course_names):
            if idx < len(meta["textbook_details"]):
                meta["course_textbook_map"][course] = meta["textbook_details"][idx]

    return meta, topics


# === 构建出题 prompt ===
def build_generation_prompt(meta, topic, set_idx, spec_text, existing_summaries=None):
    """构建发送给 Claude 的出题提示词

    Args:
        existing_summaries: 已生成的题目摘要列表（用于分批生成时避免重复）
    """

    # 解析题型数量
    qt_str = topic["question_types"]
    single_match = re.search(r"单选(\d+)", qt_str)
    judge_match = re.search(r"判断(\d+)", qt_str)
    comp_match = re.search(r"综合(\d+)", qt_str)
    fill_match = re.search(r"填空(\d+)", qt_str)
    multi_match = re.search(r"多选(\d+)", qt_str)
    # "选择X"视为单选（与"多选X"区分）
    if not single_match:
        single_match = re.search(r"选择(\d+)", qt_str)

    single_count = int(single_match.group(1)) if single_match else 0
    judge_count = int(judge_match.group(1)) if judge_match else 0
    comp_count = int(comp_match.group(1)) if comp_match else 0
    fill_count = int(fill_match.group(1)) if fill_match else 0
    multi_count = int(multi_match.group(1)) if multi_match else 0

    # 难度比例
    diff_str = topic["difficulty"]
    diff_parts = diff_str.split(":")
    if len(diff_parts) == 3:
        easy_pct, mid_pct, hard_pct = int(diff_parts[0]), int(diff_parts[1]), int(diff_parts[2])
    else:
        easy_pct, mid_pct, hard_pct = 80, 10, 10

    # 计算各难度的题目数
    total_choice = single_count + multi_count
    hard_count = max(1, round(total_choice * hard_pct / 100))
    mid_count = max(1, round(total_choice * mid_pct / 100))
    easy_count = total_choice - hard_count - mid_count

    # 从教材列表获取当前课程对应的教材
    textbook_info = meta["textbooks"]

    # 主题后缀
    set_suffix = f"(第{set_idx}套)" if topic["sets"] > 1 else ""

    # 构建输出格式示例（大题编号动态计算）
    _CN_NUMS = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    sec_idx = 0
    format_parts = []

    # 单项选择题
    format_parts.append(f"""{_CN_NUMS[sec_idx]}、单项选择题
1. 题干文本（   ）
A. 选项A\t\t\tB. 选项B
C. 选项C\t\t\tD. 选项D
【答案】X
【解析】解析文本（完整陈述句，禁止使用→=↑↓等符号）

2. ...（以此类推）""")
    sec_idx += 1
    q_start = single_count + 1

    # 多项选择题（如有）
    if multi_count > 0:
        format_parts.append(f"""{_CN_NUMS[sec_idx]}、多项选择题
{q_start}. 题干文本（   ）
A. 选项A\t\t\tB. 选项B
C. 选项C\t\t\tD. 选项D
【答案】XY（两个或多个字母）
【解析】解析文本""")
        sec_idx += 1
        q_start += multi_count

    # 判断题
    format_parts.append(f"""{_CN_NUMS[sec_idx]}、判断题
{q_start}. 判断题题干。（   ）
【答案】√ 或 ×
【解析】解析文本""")
    sec_idx += 1
    q_start += judge_count

    # 填空题（如有）
    if fill_count > 0:
        format_parts.append(f"""{_CN_NUMS[sec_idx]}、填空题
{q_start}. 填空题题干______。
【答案】答案文本
【解析】解析文本""")
        sec_idx += 1
        q_start += fill_count

    # 综合题
    format_parts.append(f"""{_CN_NUMS[sec_idx]}、综合题
{q_start}. 综合题题干
【答案】完整答案
【解析】解析文本""")

    format_example = "\n\n".join(format_parts)

    # 综合题难度要求：规划表中本章节/主题综合题达到3道及以上时，至少1道必须为困难题
    comp_difficulty_requirement = ""
    if comp_count >= 3:
        comp_difficulty_requirement = "\n- 综合题：至少1道必须为困难难度（多步计算/故障推理链/跨知识点综合），并在题干或解析中体现多步骤分析过程"

    # 窄考点/极重要主题防重复：生成阶段先分配不同考查角度，减少生成后靠质检返修。
    angle_templates = [
        "概念识别：考查基本定义、适用条件或规范名称，题干避免直接复述正确选项",
        "结构功能：考查部件/电路/机构的组成、作用或相互关系",
        "工作过程：考查动作顺序、信号/电流/力的传递路径或状态变化",
        "故障判断：给出异常现象，判断可能原因、检测部位或处理方向",
        "参数变化：考查条件改变后性能、读数、状态或结果如何变化",
        "应用场景：给出生产、维修、测量或操作情境，判断正确做法",
        "对比辨析：比较相近概念、结构、工况或方法的区别",
        "计算/判断：结合简单数据、图示描述或逻辑条件进行计算、判断或推理",
    ]
    angle_count = single_count + multi_count + judge_count + fill_count + comp_count
    angle_lines = []
    if angle_count >= 6 or "极重要" in topic.get("level", ""):
        for idx in range(1, angle_count + 1):
            angle_lines.append(f"- 第{idx}题：{angle_templates[(idx - 1) % len(angle_templates)]}")
    angle_section = ""
    if angle_lines:
        angle_section = f"""
【考查角度分配——窄考点防重复，必须执行】
本主题如果考点较窄，不要反复用“下列关于……正确的是”这类同一问法。请按题号尽量采用以下不同角度命题；同一核心词可以出现，但题干情境、设问任务和考查动作必须不同：
{chr(10).join(angle_lines)}
"""

    # 加载参考资料（真题+教材）
    ref_text = _load_reference_materials(topic, meta)
    ref_section = ""
    if ref_text:
        ref_section = f"\n【参考资料——请模仿真题的出题风格和难度，基于教材内容确保知识准确】\n\n{ref_text}\n"

    system_prompt = f"""你是一位经验丰富的中职教育命题专家，严格按照高职高考/高职分类考试真题标准出题。

以下是你必须严格遵守的出题规范（任何违反都是不合格的）：

{spec_text}
{ref_section}
【特别强调——最容易犯的错误】：
1. 选项长度比：最长选项字数÷最短选项字数必须≤2.0，绝对禁止一个选项20字其他选项2-3字
2. 干扰项禁止使用"正常""无影响""不动""更省油""无要求""任意"等废词
3. 选择题答案位置由后续程序自动调整；你不要为了ABCD均匀而牺牲题目质量，只需确保每题答案唯一、选项合理、解析正确
4. 每卷必须有≥1道适中难度题+≥1道困难/计算题
5. 四个选项必须句式结构一致（都是短语、或都是完整句）
6. 解析禁止写单个短语（如"耐磨处理。""精密配合。"），必须写1-3句完整因果句，直击题目要点
7. 括号（ ）位置不必强制放在句末，应放在语义最自然的位置，使填入选项后读起来通顺
8. 生成题目中不要有注解，例如“中央处理器（CPU）”；应直接使用教材中的规范名称，不额外加英文缩写或括号解释
9. 不要出现仅有一个选项含有“或”“和”“且”“以及”“并且”等连接词；四个选项的语言结构必须保持同类、对称
10. 综合题或计算题包含多个小问时，答案必须按小问分行书写，如（1）……换行（2）……，不要把不同小题答案挤在同一段
11. 算式和物理公式优先使用 Word 原生公式标记：需要分式、根号、上下标、近似号、希腊字母、单位组合等公式排版时，写成 {{math:...}}，标记内部使用简洁 LaTeX/线性公式语法，如 {{math:I=\\frac{{U}}{{R}}}}、{{math:\\Phi=BS}}、{{math:R=\\rho\\frac{{L}}{{S}}}}；普通中文解释写在标记外。禁止使用 \\(...\\)、$...$ 包裹公式。简单符号仍可直接写“×、ρ、Ω、≈”。"""

    user_prompt = f"""请为以下主题生成一份完整的一课一练试卷{set_suffix}：

【基本信息】
- 省份：{meta["province"]}
- 类别：{meta["category"]}
- 课程：{topic["course"]}
- 章节：{topic["section"]}
- 主题：{topic["theme"]}
- 考纲知识点：{topic["knowledge"]}
- 考纲编号：{topic["exam_ref"]}
- 参考教材：{textbook_info}
- 重要程度：{topic["level"]}

【题型和数量要求】
- 单选题：{single_count}道（无需刻意控制ABCD分布，优先保证题目质量、答案唯一、选项合理）
- 判断题：{judge_count}道（对错比约4:6至5:5）
{f"- 多选题：{multi_count}道" if multi_count > 0 else ""}\
{f"- 填空题：{fill_count}道（空格处用6个下划线______表示，禁止用括号）" if fill_count > 0 else ""}\
- 综合题：{comp_count}道{comp_difficulty_requirement}

【难度分布（选择题）】
- 容易（识记）：{easy_count}道
- 适中（理解+简单计算/应用）：{mid_count}道
- 困难（多步计算/故障推理链/跨知识点综合）：{hard_count}道
{angle_section}
【输出格式要求——极其重要，必须严格遵守】

1. 直接输出试卷正文，禁止在开头添加任何标题、标记、说明文字（如 # ## 标题、文件名、"以下是试卷"等）
2. 大类题型标题直接写"一、单项选择题""二、判断题"等（有多选则写"X、多项选择题"），禁止添加任何 markdown 标记（如 ### #）
3. 选择题选项用制表符\\t分隔同行的两个选项，格式如下：
   A. 选项A\\t\\t\\tB. 选项B
   C. 选项C\\t\\t\\tD. 选项D
   禁止使用表格、列表或其他格式排列选项
4. 综合题或计算题包含多个小问时，答案与解析必须按小问分行书写：每个小题单独一行，以“（1）”“（2）”开头；不要把不同小题答案挤在同一段
5. 算式和物理公式优先使用 Word 原生公式标记：需要分式、根号、上下标、近似号、希腊字母、单位组合等公式排版时，写成 {{math:...}}，标记内部使用简洁 LaTeX/线性公式语法，如 {{math:I=\\frac{{U}}{{R}}}}、{{math:\\Phi=BS}}、{{math:R=\\rho\\frac{{L}}{{S}}}}；普通中文解释写在标记外。禁止使用 \\(...\\)、$...$ 包裹公式。简单符号仍可直接写“×、ρ、Ω、≈”。
6. 禁止在末尾输出自检清单、备注或任何额外内容——在输出前请在内部完成自检，不要写入输出中

严格按以下格式输出（从"一、单项选择题"直接开始，前面不要有任何内容）：

{format_example}

【内部自检（不要输出到结果中）】：
输出前请逐题确认以下各项，但不要将此清单写入输出文本：
- 选项长度比≤2.0（最长÷最短）
- 无废选项（无"正常""无影响"等）
- 句式结构一致
- 有≥1道适中题+≥1道困难题
- 解析无禁用符号
- 与已有题目不重复（若提供了已有摘要）
- 【最重要】答案自暴露检测：逐题检查正确答案选项中≥4字的关键词是否出现在题干中。如果出现，必须立即重新构思该题（修改题干措辞或更换考查角度），直到题干不再包含正确答案的特征词汇。绝不允许带着自暴露问题输出。"""

    # 注入已有题目摘要（分批防重复）
    if existing_summaries:
        avoid_section = "\n".join(f"  - {s}" for s in existing_summaries)
        user_prompt += f"""

【本卷已有题目（严禁重复）】
以下是本卷已生成的题目摘要，你生成的题目不得与下列任何一道考查相同知识点或使用相同情境：
{avoid_section}

请确保新生成的每道题都有独立的考查角度，不得只换数字或措辞。"""

    return system_prompt, user_prompt


# === 解析 API 返回的试卷文本 ===
def parse_paper_text(text):
    """将 API 返回的纯文本试卷解析为结构化数据"""
    sections = []
    current_section = None
    current_lines = []

    for line in text.split("\n"):
        stripped = line.strip()
        # 检测大题标题
        if re.match(r"^[一二三四五六七八九十][、.．]", stripped):
            if current_section:
                sections.append({"title": current_section, "content": "\n".join(current_lines)})
            current_section = stripped
            current_lines = []
        else:
            current_lines.append(line)

    if current_section:
        sections.append({"title": current_section, "content": "\n".join(current_lines)})

    return sections


# === 清理 AI 返回文本 ===
def _normalize_generated_text(text):
    """修正常见模型转义残留，避免数学符号进入 DOCX 前变成乱码。"""
    if not text:
        return text

    # 保护 {{math:...}} 内部的 LaTeX/线性公式，避免被正文清洗逻辑压平成普通文本。
    math_chunks = []

    def _stash_math(match):
        math_chunks.append(match.group(0))
        return f"@@MATH_MARKER_{len(math_chunks) - 1}@@"

    text = re.sub(r"\{\{math:.*?\}\}", _stash_math, text, flags=re.DOTALL)

    # 去掉常见 LaTeX 行内公式包裹，避免 DOCX 中残留反斜杠括号。
    text = re.sub(r"\\\s*\((.*?)\\\s*\)", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\\\s*\[(.*?)\\\s*\]", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\$(.*?)\$", r"\1", text, flags=re.DOTALL)

    # 将常见 LaTeX 分式转为普通文本，优先处理一层花括号分式。
    frac_pattern = re.compile(r"\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}")
    while True:
        new_text = frac_pattern.sub(r"(\1)/(\2)", text)
        if new_text == text:
            break
        text = new_text

    # 先处理成对定界符，避免后续把 \left 里的 \le 误替换成“≤”。
    text = re.sub(r"\\\s*l\s*e\s*f\s*t\s*([（(\[{])", r"\1", text)
    text = re.sub(r"\\\s*r\s*i\s*g\s*h\s*t\s*([）)\]}])", r"\1", text)
    text = re.sub(r"\\\s*r\s*i\s*g\s*h\s*t", "", text)

    replacements = {
        r"\pm": "±",
        r"\times": "×",
        r"\cdot": "·",
        r"\div": "÷",
        r"\le": "≤",
        r"\ge": "≥",
        r"\neq": "≠",
        r"\approx": "≈",
        r"\pi": "π",
        r"\rho": "ρ",
        r"\Omega": "Ω",
        r"\mu": "μ",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # 若已经出现“≤ft/≤f t”，通常是 \left 被误替换后的残留，直接还原为左定界符。
    text = re.sub(r"≤\s*f\s*t\s*([（(\[{])", r"\1", text)
    text = re.sub(r"≤\s*f\s*t", "", text)

    # 兼容模型把 LaTeX 符号拆成“\ p m”或“\p m”等异常形式的情况。
    text = re.sub(r"\\\s*p\s*m", "±", text)
    text = re.sub(r"\\\s*t\s*i\s*m\s*e\s*s", "×", text)
    text = re.sub(r"\\\s*r\s*h\s*o", "ρ", text)
    text = re.sub(r"\\\s*O\s*m\s*e\s*g\s*a", "Ω", text)

    # 清理 LaTeX 空格命令和简单分式外层括号，让“(L)/(S)”变成“L/S”。
    text = re.sub(r"\\[,;:!]?", "", text)
    text = re.sub(r"\(([A-Za-z0-9.]+)\)/\(([A-Za-z0-9.]+)\)", r"\1/\2", text)

    # 兼容“1 2 0 ± 0.03 mm”这类数字被逐位空格拆开的尺寸公差。
    text = re.sub(
        r"(?<!\d)((?:\d\s+){2,}\d)\s*±\s*(\d+(?:\.\d+)?)\s*(mm|cm|m|μm|um)\b",
        lambda m: re.sub(r"\s+", "", m.group(1)) + f"±{m.group(2)} {m.group(3)}",
        text,
    )

    def _restore_math(match):
        idx = int(match.group(1))
        return math_chunks[idx]

    text = re.sub(r"@@MATH_MARKER_(\d+)@@", _restore_math, text)
    return text


def _clean_paper_text(text):
    """清理 AI 返回的试卷文本：去除思考块、markdown 标记、自检清单等多余内容"""
    text = _normalize_generated_text(text)

    # 部分推理模型/代理偶尔会把隐藏思考以 <think>...</think> 泄露到正文。
    text = re.sub(r"(?is)<think\b[^>]*>.*?</think>\s*", "", text)
    text = re.sub(r"(?is)^.*?</think>\s*", "", text)

    lines = text.split("\n")
    cleaned = []

    # 标记是否进入自检清单区域
    in_checklist = False
    current_section = ""

    for line in lines:
        stripped = line.strip()

        # 跳过自检清单（从"【自检清单"或"□"开头的连续行）
        if "自检清单" in stripped or (in_checklist and stripped.startswith("□")):
            in_checklist = True
            continue
        if in_checklist and not stripped:
            continue
        in_checklist = False

        # 跳过 markdown 标题行（如 # 重庆市一课一练、## 《xxx》）
        if re.match(r"^#{1,6}\s", stripped):
            # 但保留大类题型标题（去掉 ### 前缀）
            content_after_hash = re.sub(r"^#{1,6}\s*", "", stripped)
            if re.match(r"^[一二三四五六七八九十][、.．]", content_after_hash):
                cleaned.append(content_after_hash)
            # 其他 markdown 标题（如 # 重庆市一课一练）直接丢弃
            continue

        if re.match(r"^[一二三四五六七八九十][、.．]", stripped):
            current_section = stripped

        # 判断题答案统一规范为 √ / ×，避免模型输出“对/错/正确/错误”。
        if stripped.startswith("【答案】") and "判断" in current_section:
            ans_text = stripped[4:].strip()
            if ans_text in ("对", "正确", "√", "是"):
                line = "【答案】√"
            elif ans_text in ("错", "错误", "×", "否"):
                line = "【答案】×"

        cleaned.append(line)

    # 去除首尾空行
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()

    return "\n".join(cleaned)


def _format_section_title(section_str, seq, theme, set_suffix):
    """将章节字符串格式化为标题行，如 '第1节 汽车传动系  第2练  离合器基础'"""
    if not section_str:
        return f"第{seq}练  {theme}{set_suffix}"

    # 格式2: "(一) 汽车传动系" 或 "（一）汽车传动系"
    m = re.match(r'\s*[（(]([一二三四五六七八九十百]+)[）)]\s*(.+)', section_str)
    if m:
        cn_num = m.group(1)
        section_name = m.group(2).strip()
        digit = _CN_TO_DIGIT.get(cn_num, 0)
        return f"第{digit}节 {section_name}  第{seq}练  {theme}{set_suffix}"

    # 格式1: "一、制图基本知识" 或 "  一、制图基本知识"
    m = re.match(r'\s*([一二三四五六七八九十百]+)[、.．]\s*(.+)', section_str)
    if m:
        cn_num = m.group(1)
        section_name = m.group(2).strip()
        digit = _CN_TO_DIGIT.get(cn_num, 0)
        return f"第{digit}节 {section_name}  第{seq}练  {theme}{set_suffix}"

    # 格式3: "1．力系与平衡" 或 "3．材料与选用"（阿拉伯数字+顿号/点）
    m = re.match(r'\s*(\d+)[．.、]\s*(.+)', section_str)
    if m:
        digit = int(m.group(1))
        section_name = m.group(2).strip()
        return f"第{digit}节 {section_name}  第{seq}练  {theme}{set_suffix}"

    # 其他格式（如"绪论"）：直接作为节名使用
    section_clean = section_str.strip()
    return f"{section_clean}  第{seq}练  {theme}{set_suffix}"


# === 分批生成（防重复策略） ===
BATCH_SIZE = 5  # 每批生成的题目数


def _extract_question_summaries(paper_text):
    """从试卷文本中提取每道题的摘要（题干前80字），用于传入下一批避免重复"""
    summaries = []
    for line in paper_text.split("\n"):
        line = line.strip()
        if re.match(r"^\d+[\.．、]", line):
            summaries.append(line[:80])
    return summaries


def _split_numbered_theme(theme):
    """拆分“xxx（一）/xxx(二)”这类连续主题，返回 (基础主题, 序号)。普通主题序号为 None。"""
    match = re.match(r"^(.+?)[（(]([一二三四五六七八九十]+)[）)]$", str(theme or "").strip())
    if not match:
        return str(theme or "").strip(), None
    base = match.group(1).strip()
    num = _CN_TO_DIGIT.get(match.group(2))
    return base, num


def generate_paper_in_batches(client, model, meta, topic, set_idx, spec_text, config,
                              session_usage=None, daily_usage=None):
    """分批生成试卷：每 BATCH_SIZE 道题为一批，累积摘要传入下一批防止重复

    流程：
      1. 第1批：正常生成前 N 道题
      2. 提取已生成题目的摘要
      3. 第2批：将摘要注入 prompt，AI 生成剩余题目时会避开已有内容
      4. 合并所有批次的结果

    Returns:
        (paper_text, total_usage) 或 (None, None) 失败时
    """
    # 计算总题数
    qt_str = topic["question_types"]
    total_questions = 0
    for m in re.finditer(r'(\d+)', qt_str):
        total_questions += int(m.group(1))

    # 如果总题数 <= BATCH_SIZE，不需要分批，一次生成
    if total_questions <= BATCH_SIZE:
        sys_prompt, user_prompt = build_generation_prompt(meta, topic, set_idx, spec_text)
        paper_text, usage = call_api(
            client, model, sys_prompt, user_prompt,
            max_tokens=config.get("max_tokens", 8000),
            temperature=config.get("temperature", 0.7),
        )
        if session_usage and usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                session_usage[key] += usage.get(key, 0)
            session_usage["api_calls"] += 1
        if daily_usage and usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                daily_usage[key] += usage.get(key, 0)
            daily_usage["api_calls"] += 1
            _save_daily_usage(daily_usage)
        token_info = f", {usage['total_tokens']} tokens" if usage else ""
        print(f"  API 返回成功 ({len(paper_text)}字{token_info})")
        return paper_text, usage

    # 分批生成
    all_text_parts = []
    existing_summaries = []
    total_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
    batch_num = 0

    # 第1批：生成完整试卷（AI 一次性生成所有题型）
    # 但通过摘要机制，后续重新生成时可以避开已有内容
    sys_prompt, user_prompt = build_generation_prompt(
        meta, topic, set_idx, spec_text, existing_summaries=existing_summaries or None
    )
    paper_text, usage = call_api(
        client, model, sys_prompt, user_prompt,
        max_tokens=config.get("max_tokens", 8000),
        temperature=config.get("temperature", 0.7),
    )
    batch_num += 1

    if usage:
        for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
            total_usage[key] += usage.get(key, 0)
        if session_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                session_usage[key] += usage.get(key, 0)
            session_usage["api_calls"] += 1
        if daily_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                daily_usage[key] += usage.get(key, 0)
            daily_usage["api_calls"] += 1
            _save_daily_usage(daily_usage)

    token_info = f", {usage['total_tokens']} tokens" if usage else ""
    print(f"  第{batch_num}批返回成功 ({len(paper_text)}字{token_info})")

    # 提取摘要供后续（极重要双卷的第二卷）使用
    existing_summaries = _extract_question_summaries(paper_text)

    return paper_text, total_usage


def _clean_course_folder_name(course):
    """从规划表课程行提取用于输出目录的课程/教材名称。"""
    if not course:
        return ""
    name = str(course).strip()
    name = re.sub(r"^课程[一二三四五六七八九十\d]+\s*[：:]\s*", "", name)
    name = re.sub(r"[（(]\s*约?\d+%\s*[）)]", "", name)
    return name.strip()


def _get_topic_textbook_name(meta, topic):
    """获取当前 topic 对应的教材名；没有教材信息时回退到课程名。"""
    course_map = meta.get("course_textbook_map", {})
    if course_map and topic.get("course") in course_map:
        return course_map[topic["course"]].get("name", "")
    if meta.get("textbook_details"):
        return meta["textbook_details"][0].get("name", "")
    return _clean_course_folder_name(topic.get("course", ""))


def _get_topic_output_base(meta, topic, output_dir):
    """按 生成结果/省份 类别/课程或教材名 定位当前规划表对应输出目录。"""
    province = meta.get("province", "")
    category = meta.get("category", "")
    base = Path(output_dir)
    if province and category:
        base = base / f"{province} {category}"
    elif province:
        base = base / province

    folder_name = _get_topic_textbook_name(meta, topic)
    if folder_name:
        base = base / folder_name
    return base


# === 生成 DOCX ===
def generate_docx(meta, topic, set_idx, paper_text, output_dir, needs_manual_review=False):
    """将试卷文本生成为格式化的 DOCX 文件"""
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本" / "核心脚本"))
    from docx_utils1 import (
        copy_template, set_margins, add_editorial_note,
        add_paragraph_with_style, add_question_options_table,
        add_labeled_text, save_docx
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 清理 AI 返回文本
    paper_text = _clean_paper_text(paper_text)

    # 确定文件名
    seq = topic["seq"]
    theme = topic["theme"]
    province = meta["province"]
    set_suffix = f"({set_idx})" if topic["sets"] > 1 else ""

    # 从教材详情中提取规范化的教材名、出版社和版次
    # 优先使用课程→教材映射（多课程表中每个课程对应不同教材）
    textbook_name = ""
    publisher = "高教版"
    edition = "第一版"
    edition_display = "高教版·第一版"

    course_map = meta.get("course_textbook_map", {})
    if course_map and topic.get("course") in course_map:
        detail = course_map[topic["course"]]
        textbook_name = detail["name"]
        publisher = detail["publisher"]
        edition = detail["edition"]
        edition_display = detail["display"]
    elif meta.get("textbook_details"):
        detail = meta["textbook_details"][0]
        textbook_name = detail["name"]
        publisher = detail["publisher"]
        edition = detail["edition"]
        edition_display = detail["display"]
    elif meta.get("textbook_list"):
        for tb_name, tb_edition_raw in meta["textbook_list"]:
            textbook_name = tb_name
            publisher, edition = _parse_edition(tb_edition_raw)
            edition_display = f"{publisher}·{edition}"
            break

    if not textbook_name:
        tb_match = re.search(r"《(.+?)》", meta.get("textbooks", ""))
        if tb_match:
            textbook_name = tb_match.group(1)

    filename = f"第{seq}练 {theme}{set_suffix} {province}（高职分类考试）《{textbook_name}》（{edition_display}） 一课一练 （解析版）.docx"
    if needs_manual_review:
        filename = f"（待人工审核）{filename}"

    # 按"省份 类别/课程或教材名/"组织子目录，与规划表位置保持一致
    sub_dir = _get_topic_output_base(meta, topic, output_dir)
    os.makedirs(sub_dir, exist_ok=True)
    output_path = sub_dir / filename

    # 复制模板
    tpl = str(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else None
    if tpl:
        doc = copy_template(tpl, str(output_path))
    else:
        from docx import Document
        doc = Document()
        doc.save(str(output_path))
        doc = Document(str(output_path))

    set_margins(doc)

    # 编写说明
    sep_img = str(SEPARATOR_PATH) if SEPARATOR_PATH.exists() else None
    add_editorial_note(
        doc,
        textbook_name=textbook_name,
        edition=edition,
        chapter_seq=seq,
        knowledge_scope=topic["knowledge"],
        province=province,
        publisher=publisher,
        separator_image=sep_img,
    )

    # 标题
    exam_type = "高职分类考试" if "高职" in meta.get("config_line", "") else "高职分类考试"
    title1 = f"{province}（{exam_type}）一课一练"
    title2 = f"《{textbook_name}》（{edition_display}）"
    title3 = _format_section_title(topic["section"], seq, theme, set_suffix)

    add_paragraph_with_style(doc, title1, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title2, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title3, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 答案/解析使用红色字体
    RED_COLOR = (255, 0, 0)

    # 逐行解析并输出
    lines = paper_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 大题标题（加粗，黑体）
        if re.match(r"^[一二三四五六七八九十][、.．]", line):
            add_paragraph_with_style(doc, line, font_name="黑体", font_size=12,
                                     bold=True, space_after=6)
            i += 1
            continue

        # 题目开头（数字编号）
        if re.match(r"^\d+[\.．、]", line):
            # 收集题目完整文本（直到下一个【答案】或下一题）
            q_lines = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if next_line.startswith("【答案】") or next_line.startswith("【解析】"):
                    break
                if re.match(r"^\d+[\.．、]", next_line):
                    break
                if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                    break
                q_lines.append(lines[j])
                j += 1

            q_text = "\n".join(q_lines)

            # 判断是否有选项（含 A. B. C. D.）
            if re.search(r"[A-D][\.．]\s*\S", q_text):
                add_question_options_table(doc, q_text)
            else:
                add_paragraph_with_style(doc, q_text, font_size=10.5, space_after=2)

            i = j
            continue

        # 【答案】— 红色字体；多行答案继续归入同一红色段落
        if line.startswith("【答案】"):
            answer_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith("【解析】") or next_line.startswith("【答案】")
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    answer_lines.append(next_line)
                j += 1
            answer_content = "\n".join(answer_lines)
            add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            i = j
            continue

        # 【解析】— 红色字体；多行解析继续归入同一红色段落
        if line.startswith("【解析】"):
            explanation_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith("【答案】") or next_line.startswith("【解析】")
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    explanation_lines.append(next_line)
                j += 1
            explanation_content = "\n".join(explanation_lines)
            add_labeled_text(doc, "【解析】", explanation_content, color=RED_COLOR)
            i = j
            continue

        # 普通文本行
        if line:
            add_paragraph_with_style(doc, line, font_size=10.5, space_after=2)

        i += 1

    save_docx(doc, str(output_path))
    return str(output_path)


# === 质检集成 ===
MAX_REGEN_ATTEMPTS = 3  # 严重问题最多重新生成次数


def _quick_check(paper_text):
    """对试卷文本进行快速本地质检，返回 (严重问题列表, 警告列表, 评分, 信息列表)"""
    check_dir = (BASE_DIR / "01_工具脚本" / "质检").resolve()
    check_file = check_dir / "check.py"
    if str(check_dir) not in sys.path:
        sys.path.insert(0, str(check_dir))

    try:
        from check import local_check
    except ModuleNotFoundError as exc:
        if exc.name != "check" or not check_file.exists():
            raise
        import importlib.util

        spec = importlib.util.spec_from_file_location("local_quality_check", check_file)
        if spec is None or spec.loader is None:
            raise ImportError(f"无法加载质检模块: {check_file}") from exc
        check_module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(check_module)
        local_check = check_module.local_check

    issues, questions = local_check(paper_text)
    severe = [i for i in issues if i["severity"] == "严重"]
    warnings = [i for i in issues if i["severity"] == "警告"]
    infos = [i for i in issues if i["severity"] == "信息"]

    score = 100 - len(severe) * 15 - len(warnings) * 5
    score = max(0, score)

    return severe, warnings, score, questions, infos


def _print_qc_summary(severe, warnings, score, infos=None):
    """打印质检摘要"""
    if severe:
        print(f"  质检: {score}/100 | 严重问题 {len(severe)} 个:")
        for issue in severe:
            print(f"    ✗ [第{issue['question']}题] {issue['type']}: {issue['detail']}")
    if warnings:
        print(f"  质检: 警告 {len(warnings)} 个:")
        for issue in warnings[:3]:
            print(f"    ⚠ [第{issue['question']}题] {issue['type']}: {issue['detail']}")
        if len(warnings) > 3:
            print(f"    ...及其他 {len(warnings)-3} 个警告")
    if not severe and not warnings:
        print(f"  质检: {score}/100 ✓ 通过")
    # 显示信息类提示（如答案分布）
    if infos:
        for info in infos:
            print(f"    ℹ {info['type']}: {info['detail']}")


def _ask_user_keep(score, warnings):
    """当仅有轻微问题时，询问用户是否保留"""
    print(f"\n  试卷评分 {score}/100，仅有轻微问题（{len(warnings)}个警告）。")
    while True:
        choice = input("  是否保留此试卷？(y=保留 / n=重新生成): ").strip().lower()
        if choice in ("y", "yes", "是", ""):
            return True
        elif choice in ("n", "no", "否"):
            return False
        print("  请输入 y 或 n")


def _fix_answer_distribution(paper_text):
    """通过交换选项顺序来调整答案分布，使ABCD尽可能分散

    策略：找出出现次数最多的答案字母，将部分题的选项顺序互换，
    使正确答案从高频字母变为低频字母。
    """
    lines = paper_text.split("\n")

    # 提取所有单选题的位置和答案
    choice_items = []  # [(line_idx_of_answer, current_answer, option_lines_range)]
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # 找到题目开头
        q_match = re.match(r"^(\d+)[\.．、]\s*", line)
        if q_match:
            q_start = i
            # 找该题的选项行和答案行
            j = i + 1
            option_start = None
            answer_idx = None
            while j < len(lines):
                jline = lines[j].strip()
                if re.search(r"[A-D][\.．]\s*\S", jline) and option_start is None:
                    option_start = j
                if jline.startswith("【答案】"):
                    ans_match = re.match(r"【答案】\s*([A-D])$", jline)
                    if ans_match:
                        answer_idx = j
                        choice_items.append({
                            "q_start": q_start,
                            "option_start": option_start,
                            "answer_idx": answer_idx,
                            "answer": ans_match.group(1),
                        })
                    break
                if re.match(r"^\d+[\.．、]\s*", jline) or re.match(r"^[一二三四五六七八九十][、.．]", jline):
                    break
                j += 1
            i = j if j > i else i + 1
        else:
            i += 1

    if len(choice_items) < 4:
        return paper_text

    # 统计分布
    dist = {"A": 0, "B": 0, "C": 0, "D": 0}
    for item in choice_items:
        dist[item["answer"]] += 1

    max_count = max(dist.values())
    total = len(choice_items)

    # 5题以下：4个相同才需要调整；5题以上：超50%才调整
    need_fix = False
    if total <= 5 and max_count >= 4:
        need_fix = True
    elif total > 5 and max_count > total * 0.5:
        need_fix = True

    if not need_fix:
        return paper_text

    # 找出最多和最少的字母
    sorted_letters = sorted(dist.items(), key=lambda x: x[1], reverse=True)
    most_letter = sorted_letters[0][0]
    least_letter = sorted_letters[-1][0]

    # 选取需要调整的题目（答案为最多字母的题中选一部分）
    swap_map = {most_letter: least_letter, least_letter: most_letter}
    targets = [item for item in choice_items if item["answer"] == most_letter]
    # 只调整一半（使分布更均匀而不是反转）
    num_to_fix = (max_count - dist[least_letter]) // 2
    targets_to_fix = targets[:num_to_fix]

    for item in targets_to_fix:
        if item["option_start"] is None:
            continue

        # 找到该题的选项文本（可能在1-2行内）
        opt_lines_idx = []
        j = item["option_start"]
        while j < item["answer_idx"]:
            if re.search(r"[A-D][\.．]\s*\S", lines[j]):
                opt_lines_idx.append(j)
            j += 1

        # 在选项行中交换 most_letter 和 least_letter 的文本
        old_answer = most_letter
        new_answer = least_letter

        for idx in opt_lines_idx:
            line = lines[idx]
            # 提取各选项
            opts = re.findall(r"([A-D])[\.．]\s*([^\t\n]+?)(?=\s*[A-D][\.．]|\s*$|\t)", line)
            if not opts:
                opts = re.findall(r"([A-D])[\.．]\s*(.+)", line)
            if len(opts) >= 2:
                opt_dict = {letter: text.strip() for letter, text in opts}
                if old_answer in opt_dict and new_answer in opt_dict:
                    # 交换两个选项的文本内容
                    opt_dict[old_answer], opt_dict[new_answer] = opt_dict[new_answer], opt_dict[old_answer]
                    # 重建该行
                    if "\t" in line:
                        # 制表符分隔的行
                        parts = []
                        for letter in sorted(opt_dict.keys()):
                            parts.append(f"{letter}. {opt_dict[letter]}")
                        if len(parts) == 2:
                            lines[idx] = f"{parts[0]}\t\t\t{parts[1]}"
                        else:
                            lines[idx] = "\t\t\t".join(parts)
                    else:
                        # 每行一个选项
                        lines[idx] = f"{opts[0][0]}. {opt_dict[opts[0][0]]}"
            elif len(opts) == 1:
                # 单个选项一行的格式，需要找到对应的另一行
                letter = opts[0][0]
                if letter == old_answer or letter == new_answer:
                    # 找到配对行并交换
                    for idx2 in opt_lines_idx:
                        if idx2 == idx:
                            continue
                        opts2 = re.findall(r"([A-D])[\.．]\s*(.+)", lines[idx2])
                        if opts2 and opts2[0][0] in (old_answer, new_answer):
                            # 交换两行的选项文本（保留字母前缀）
                            text1 = opts[0][1].strip()
                            text2 = opts2[0][1].strip()
                            lines[idx] = f"{letter}. {text2}"
                            lines[idx2] = f"{opts2[0][0]}. {text1}"
                            break

        # 更新答案行
        lines[item["answer_idx"]] = f"【答案】{new_answer}"

    result = "\n".join(lines)
    new_dist = {"A": 0, "B": 0, "C": 0, "D": 0}
    for item in choice_items:
        if item in targets_to_fix:
            new_dist[least_letter] += 1
        else:
            new_dist[item["answer"]] += 1
    print(f"  → 已调整答案分布: {' '.join(f'{k}={v}' for k,v in new_dist.items())}")

    return result


def _extract_question_blocks(paper_text):
    """提取每道题的完整题块、所属大题题型和摘要，用于定向修复"""
    lines = paper_text.split("\n")
    blocks = {}
    current_type = ""
    current_num = None
    current_lines = []

    def save_current():
        if current_num is None or not current_lines:
            return
        block = "\n".join(current_lines).strip()
        first_line = current_lines[0].strip()
        stem = re.sub(r"^\d+[\.．、]\s*", "", first_line)
        stem = re.sub(r"（\s*）\s*$", "", stem).strip()
        blocks[current_num] = {
            "block": block,
            "type": current_type,
            "stem": stem,
            "summary": stem[:80] if stem else block.replace("\n", " ")[:80],
        }

    for line in lines:
        stripped = line.strip()
        if re.match(r"^[一二三四五六七八九十][、.．]", stripped):
            save_current()
            current_type = re.sub(r"^[一二三四五六七八九十][、.．]\s*", "", stripped).strip()
            current_num = None
            current_lines = []
            continue

        m = re.match(r"^(\d+)[\.．、]\s*", stripped)
        if m:
            save_current()
            current_num = int(m.group(1))
            current_lines = [line]
        elif current_num is not None:
            current_lines.append(line)

    save_current()
    return blocks


def _question_format_requirement(question_type, original_block):
    """根据原题所属大题生成修复时的格式要求，防止主观题被重出成选择题。"""
    # 优先按“大题标题”判断题型，不能因为综合题题干里出现“判断/分析”等动词就误判成判断题。
    section_type = question_type or ""
    if any(keyword in section_type for keyword in ("综合", "简答", "计算", "分析", "应用", "作图", "绘图", "画图")):
        return "必须保持为原来的主观题/综合题题型：不要生成判断题句式，不要只输出“正确/错误”或“√/×”；不要生成A-D选项；必须保留综合题设问、完整【答案】和【解析】。"
    if "多项" in section_type or "多选" in section_type:
        return "必须保持为多项选择题：保留A-D四个选项，【答案】可为多个字母，并给出【解析】。"
    if "单项" in section_type or "单选" in section_type:
        return "必须保持为单项选择题：保留A-D四个选项，【答案】只能是一个字母，并给出【解析】。"
    if "判断" in section_type:
        return "必须保持为判断题：不要生成A-D选项，【答案】只能为“√”或“×”，并给出【解析】。"
    if "填空" in section_type:
        return "必须保持为填空题：不要生成A-D选项，保留填空设问形式，并给出【答案】和【解析】。"

    # 兜底：只有大题标题无法识别时，才参考原题内容。
    text = original_block or ""
    if "多项" in text or "多选" in text:
        return "必须保持为多项选择题：保留A-D四个选项，【答案】可为多个字母，并给出【解析】。"
    if "单项" in text or "单选" in text or re.search(r"[A-D][\.．]\s*\S", text):
        return "必须保持为单项选择题：保留A-D四个选项，【答案】只能是一个字母，并给出【解析】。"
    if "填空" in text:
        return "必须保持为填空题：不要生成A-D选项，保留填空设问形式，并给出【答案】和【解析】。"
    if any(keyword in text for keyword in ("综合", "简答", "计算", "分析", "应用", "作图", "绘图", "画图")):
        return "必须保持为原来的主观题题型：不要生成A-D选项，按原题型输出题干、必要小问、【答案】和【解析】。"
    if "判断" in text:
        return "必须保持为判断题：不要生成A-D选项，【答案】只能为“√”或“×”，并给出【解析】。"
    return "必须保持原题所属题型和原有格式；如果原题没有A-D选项，重出题也不得添加A-D选项。"


def _issue_question_nums(issue):
    """从质检问题中提取涉及的题号。支持单题和“1&2”这类重复题格式。"""
    q_value = issue.get("question")
    if not q_value or q_value == "全卷":
        return []
    nums = []
    for part in str(q_value).split("&"):
        try:
            nums.append(int(part.strip()))
        except (ValueError, TypeError):
            pass
    return nums


def _choose_duplicate_target_nums(duplicate_issues, question_blocks, preferred_nums=None):
    """为题干重复问题选择尽可能少的重出题号。

    将每条“题干重复”视为一条边，选择覆盖所有重复边的最小改动集合的近似解：
      - 若某题已经因答案自暴露/其他问题需要重出，则优先用它覆盖重复边；
      - 否则每轮选择覆盖剩余重复边最多的题；
      - 覆盖数相同则优先改较靠后的题，尽量保留前面的基础题。
    """
    preferred_nums = set(preferred_nums or [])
    pairs = []
    issue_by_pair = []

    for issue in duplicate_issues:
        nums = [n for n in _issue_question_nums(issue) if n in question_blocks]
        if len(nums) >= 2:
            pair = tuple(dict.fromkeys(nums[:2]))
            if len(pair) == 2:
                pairs.append(pair)
                issue_by_pair.append((pair, issue))

    if not pairs:
        return set(), []

    remaining = set(range(len(pairs)))
    target_nums = set()

    while remaining:
        candidates = set()
        for idx in remaining:
            candidates.update(pairs[idx])

        def coverage(num):
            return sum(1 for idx in remaining if num in pairs[idx])

        preferred_candidates = [n for n in candidates if n in preferred_nums]
        if preferred_candidates:
            chosen = max(preferred_candidates, key=lambda n: (coverage(n), n))
        else:
            chosen = max(candidates, key=lambda n: (coverage(n), n))

        target_nums.add(chosen)
        remaining = {idx for idx in remaining if chosen not in pairs[idx]}

    duplicate_details = []
    for pair, issue in issue_by_pair:
        chosen_in_pair = [n for n in pair if n in target_nums]
        chosen_text = "、".join(str(n) for n in chosen_in_pair) if chosen_in_pair else "未选择"
        duplicate_details.append(
            f"第{issue.get('question', '')}题重复：{issue.get('detail', '')}；本次重出第{chosen_text}题"
        )

    return target_nums, duplicate_details


def _fix_duplicate_questions(client, config, meta, topic, set_idx, spec_text,
                             paper_text, duplicate_issues, session_usage=None, daily_usage=None,
                             preferred_nums=None):
    """重复题定向修复：选择能覆盖全部重复关系的最少题目重出。"""
    question_blocks = _extract_question_blocks(paper_text)
    if not question_blocks:
        return paper_text

    target_nums, duplicate_details = _choose_duplicate_target_nums(
        duplicate_issues, question_blocks, preferred_nums=preferred_nums
    )

    if not target_nums:
        pairs = ",".join(str(i.get("question", "")) for i in duplicate_issues)
        print(f"  → 警告：重复项 {pairs} 无法定位可修复题号，请检查查重/修复逻辑。")
        return paper_text

    repair_sections = []
    for num in sorted(target_nums):
        original_block = question_blocks[num]["block"]
        question_type = question_blocks[num].get("type", "") or "未识别题型"
        format_requirement = _question_format_requirement(question_type, original_block)
        avoid_summaries = []
        for other_num, info in sorted(question_blocks.items()):
            if other_num == num:
                continue
            avoid_summaries.append(f"第{other_num}题：{info['summary']}")
        repair_sections.append(f"""【需要重出的题目：第{num}题】
原题所属题型：{question_type}
原题：
{original_block}

本题格式要求：{format_requirement}

必须避开的其他题目摘要：
{chr(10).join(avoid_summaries)}""")

    fix_prompt = f"""以下试卷中存在题干重复问题。请只重出指定题号的题目，其他题保持不变。

【重复问题】
{chr(10).join(duplicate_details)}

【当前主题信息】
- 课程：{topic.get('course', '')}
- 章节：{topic.get('section', '')}
- 主题：{topic.get('theme', '')}
- 考纲知识点：{topic.get('knowledge', '')}
- 考纲编号：{topic.get('exam_ref', '')}

【重出任务】
{chr(10).join(repair_sections)}

【硬性要求】
1. 只输出需要替换的题目，不要输出完整试卷，不要输出说明文字。
2. 必须保持原题号不变，只重出上述指定题号。
3. 必须严格保持“原题所属题型”和“本题格式要求”：单选仍为单选，多选仍为多选，判断仍为判断，综合/简答/计算等主观题不得改成选择题。
4. 新题仍需围绕当前主题、考纲知识点和原题所属题型。
5. 新题不得与“必须避开的其他题目摘要”中的任何题在题干、情境、设问角度、核心关键词组合上相似。
6. 禁止只替换数字、设备名称、选项顺序；必须换一个考查角度或应用场景。
7. 如原题是选择题，题干中不得包含正确答案的≥4字关键词，选项长度比≤2.0，四个选项句式一致。
8. 输出格式必须与原题型一致：选择题包含选项，非选择题不得添加A-D选项，均需包含【答案】、【解析】。
9. 涉及分式、根号、上下标、近似号、希腊字母等公式排版时，使用 {{math:...}} 标记，标记内部用简洁 LaTeX/线性公式语法；普通中文解释写在标记外。"""

    sys_prompt = """你是一位试卷去重修复专家。你需要在不改动其他题的前提下，只重出指定题号。
重出的题目必须避开已给出的全卷其他题摘要，避免题干、情境和设问角度重复。"""

    result, usage = call_api(
        client, config["model"], sys_prompt, fix_prompt,
        max_tokens=config.get("max_tokens", 8000),
        temperature=config.get("temperature", 0.7),
    )

    if usage:
        if session_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                session_usage[key] += usage.get(key, 0)
            session_usage["api_calls"] += 1
        if daily_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                daily_usage[key] += usage.get(key, 0)
            daily_usage["api_calls"] += 1
            _save_daily_usage(daily_usage)

    if not result:
        return paper_text

    fixed_text = _replace_questions_in_paper(paper_text, result, target_nums)
    print(f"  → 重复题仅修复第{','.join(str(n) for n in sorted(target_nums))}题（最少覆盖重复关系）")
    return fixed_text



def _fix_problem_questions(client, config, meta, topic, set_idx, spec_text,
                           paper_text, issues, session_usage=None, daily_usage=None):
    """只重新生成有问题的题目，保留其他题目不变

    策略：将有问题的题号和具体问题描述发给AI，让其只输出替换后的题目。
    然后在原文中替换对应题目块。
    """
    # 收集有问题的题号
    problem_nums = set()
    problem_details = []
    for issue in issues:
        nums = _issue_question_nums(issue)
        if nums:
            problem_nums.update(nums)
            nums_text = "&".join(str(n) for n in nums)
            problem_details.append(f"第{nums_text}题: {issue['type']} - {issue['detail']}")

    if not problem_nums:
        return paper_text

    question_blocks = _extract_question_blocks(paper_text)
    repair_sections = []
    for num in sorted(problem_nums):
        info = question_blocks.get(num, {})
        original_block = info.get("block", "")
        question_type = info.get("type", "") or "未识别题型"
        format_requirement = _question_format_requirement(question_type, original_block)
        avoid_summaries = []
        for other_num, other_info in sorted(question_blocks.items()):
            if other_num == num:
                continue
            avoid_summaries.append(f"第{other_num}题：{other_info.get('summary', '')}")
        repair_sections.append(f"""【需要修复的题目：第{num}题】
原题所属题型：{question_type}
原题：
{original_block}

本题格式要求：{format_requirement}

必须避开的其他题目摘要：
{chr(10).join(avoid_summaries)}""")

    # 构建修复 prompt
    fix_prompt = f"""以下试卷中有几道题存在质量问题，请只重新出这几道题，其他题保持不变。

【有问题的题目及原因】
{chr(10).join(problem_details)}

【需要修复的原题】
{chr(10).join(repair_sections)}

【要求】
请只输出需要替换的题目，不要输出完整试卷，不要输出其他说明文字。
必须保持原题号不变，并严格保持“原题所属题型”和“本题格式要求”：单选仍为单选，多选仍为多选，判断仍为判断，综合/简答/计算等主观题不得改成选择题。

【修复强度要求——不要小修小补】
1. 必须先针对“有问题的题目及原因”逐项彻底消除问题，不允许只换一两个字、只调整选项顺序、只改答案字母。
2. 如果原题的问题涉及题干、选项、答案、解析中的任一部分，必须重构整道题：重新设计题干表达、四个选项/答案和解析，而不是局部打补丁。
3. 新题必须明显区别于原题和其他题目摘要：更换考查角度或应用场景，不得只替换数字、设备名称或同义词。
4. 如原题是选择题：题干中不得包含正确答案的较长连续片段；选项长度比≤2.0；四个选项句式一致；正确答案唯一；干扰项必须有实质性。
5. 解析必须是1-3句完整因果句，直接说明为什么答案正确、其他关键干扰为什么不成立；禁止只写短语。
6. 不得新增任何质量问题：不得出现废选项、禁用符号、题干重复、答案自暴露、解析过短或题型改变。
7. 如原题不是选择题，不得添加A-D选项，按原题型输出题干、必要小问、【答案】和【解析】。
8. 涉及分式、根号、上下标、近似号、希腊字母等公式排版时，使用 {{math:...}} 标记，标记内部用简洁 LaTeX/线性公式语法；普通中文解释写在标记外。"""

    sys_prompt = f"""你是一位试卷修复专家。你需要修复指定题目的质量问题，只输出修复后的题目。
保持题号不变，保持原有的考查方向但换一个角度或情境重新命题。"""

    result, usage = call_api(
        client, config["model"], sys_prompt, fix_prompt,
        max_tokens=config.get("max_tokens", 8000),
        temperature=config.get("temperature", 0.7),
    )

    if usage:
        if session_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                session_usage[key] += usage.get(key, 0)
            session_usage["api_calls"] += 1
        if daily_usage:
            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                daily_usage[key] += usage.get(key, 0)
            daily_usage["api_calls"] += 1
            _save_daily_usage(daily_usage)

    if not result:
        return paper_text

    # 将修复后的题目替换到原文中
    fixed_text = _replace_questions_in_paper(paper_text, result, problem_nums)
    print(f"  → 已修复第{','.join(str(n) for n in sorted(problem_nums))}题")
    return fixed_text


def _repair_qc_issues_targeted(client, config, meta, topic, set_idx, spec_text,
                               paper_text, severe, warnings, score, infos,
                               session_usage=None, daily_usage=None,
                               max_rounds=3):
    """按优先级定向修复质检问题，尽可能少改题。

    修复顺序：
      1. 答案自暴露（即使只是警告也优先修）；
      2. 其他单题质量问题；
      3. 题干重复（用最小题号集合覆盖全部重复关系）。
    每轮只处理当前最高优先级的一组问题，修复后立即重新质检，避免继续改动已合格题目。
    """
    repaired_rounds = 0
    # 跟踪已修复题号，避免同一题反复修复无效（如 AI 持续返回空解析）
    prior_fixed_nums = set()  # 本轮已修复过的题号（防止同一修复轮次内重复修同一题）

    for repair_round in range(1, max_rounds + 1):
        blocking_issues = severe + warnings
        if score > 90 and not blocking_issues:
            break
        if not blocking_issues:
            break

        before_text = paper_text
        before_severe = severe
        before_warnings = warnings
        before_score = score
        before_infos = infos
        before_blocking_count = len(before_severe) + len(before_warnings)

        # 1. 答案自暴露优先：它在本地质检中通常是"警告"，但会直接影响命题质量。
        exposure_issues = [i for i in blocking_issues if i.get("type") == "答案自暴露"]
        # 过滤掉已修复过的
        exposure_issues = [i for i in exposure_issues
                          if not set(_issue_question_nums(i)).issubset(prior_fixed_nums)]
        if exposure_issues:
            nums = sorted({n for issue in exposure_issues for n in _issue_question_nums(issue)})
            prior_fixed_nums.update(nums)
            print(f"  → 优先修复答案自暴露：第{','.join(str(n) for n in nums)}题...")
            paper_text = _fix_problem_questions(
                client, config, meta, topic, set_idx, spec_text,
                paper_text, exposure_issues,
                session_usage, daily_usage,
            )
        else:
            # 2. 其他明确落到单题的问题：只重出这些题。
            duplicate_issues = [i for i in severe if i.get("type") == "题干重复"]
            single_question_issues = [
                i for i in blocking_issues
                if i.get("type") != "题干重复" and _issue_question_nums(i)
            ]

            if single_question_issues:
                # 过滤掉已经在本轮修复过的问题（避免AI持续返回无效修复的死循环）
                fresh_issues = []
                for issue in single_question_issues:
                    issue_nums = set(_issue_question_nums(issue))
                    if not issue_nums.issubset(prior_fixed_nums):
                        fresh_issues.append(issue)
                if fresh_issues:
                    nums = sorted({n for issue in fresh_issues for n in _issue_question_nums(issue)})
                    prior_fixed_nums.update(nums)
                    print(f"  → 定向修复单题质量问题：第{','.join(str(n) for n in nums)}题...")
                    paper_text = _fix_problem_questions(
                        client, config, meta, topic, set_idx, spec_text,
                        paper_text, fresh_issues,
                        session_usage, daily_usage,
                    )
            elif duplicate_issues:
                # 过滤掉已修复过的题号
                fresh_dup = []
                for issue in duplicate_issues:
                    issue_nums = set(_issue_question_nums(issue))
                    if not issue_nums.issubset(prior_fixed_nums):
                        fresh_dup.append(issue)
                if not fresh_dup:
                    # 所有重复题号都已修复过，停止本轮修复
                    break
                duplicate_issues = fresh_dup
                # 3. 查重最后处理；选择最少题目覆盖全部重复关系。
                # 窄考点下若重复项过多，逐题重出容易陷入“反复修同一题但不收敛”的循环，直接交给人工审核兜底。
                if len(duplicate_issues) >= 8:
                    print(f"  → 题干重复项过多（{len(duplicate_issues)}个），停止自动逐题修复，避免窄考点反复重出死循环。")
                    break
                question_blocks = _extract_question_blocks(paper_text)
                target_nums, _ = _choose_duplicate_target_nums(duplicate_issues, question_blocks)
                if not target_nums:
                    pairs = ",".join(str(i.get("question", "")) for i in duplicate_issues)
                    print(f"  → 警告：重复项 {pairs} 无法定位可修复题号，请检查查重/修复逻辑。")
                    break
                print(f"  → 修复题干重复：第{','.join(str(i['question']) for i in duplicate_issues)}题存在重复，仅重出第{','.join(str(n) for n in sorted(target_nums))}题...")
                paper_text = _fix_duplicate_questions(
                    client, config, meta, topic, set_idx, spec_text,
                    paper_text, duplicate_issues,
                    session_usage, daily_usage,
                )
            else:
                # 例如全卷级问题若无法通过选项交换解决，就交给整卷重生兜底。
                break

        if paper_text == before_text:
            print("  → 定向修复未替换到题目，停止本轮定向修复。")
            break

        cleaned = _clean_paper_text(paper_text)
        new_severe, new_warnings, new_score, _, new_infos = _quick_check(cleaned)
        new_blocking_count = len(new_severe) + len(new_warnings)

        # 安全阀：修复后如果分数降低、严重问题变多，或问题总数变多，就回滚到修复前版本。
        if (new_score < before_score
                or len(new_severe) > len(before_severe)
                or new_blocking_count > before_blocking_count):
            print(
                f"  → 本轮修复未采用：评分 {before_score}/100 → {new_score}/100，"
                f"严重问题 {len(before_severe)} → {len(new_severe)}，"
                f"严重+警告 {before_blocking_count} → {new_blocking_count}。已回滚到修复前版本。"
            )
            paper_text = before_text
            severe = before_severe
            warnings = before_warnings
            score = before_score
            infos = before_infos
            break

        repaired_rounds += 1
        severe, warnings, score, infos = new_severe, new_warnings, new_score, new_infos
        _print_qc_summary(severe, warnings, score, infos)

        if score > 90 and not severe and not warnings:
            break

    return paper_text, severe, warnings, score, infos, repaired_rounds


def _replace_questions_in_paper(original_text, fixed_text, target_nums):
    """将修复后的题目块替换到原始试卷文本中"""
    # 解析修复文本中的题目块
    fixed_blocks = {}
    lines = fixed_text.split("\n")
    current_num = None
    current_lines = []

    for line in lines:
        m = re.match(r"^(\d+)[\.．、]\s*", line.strip())
        if m:
            if current_num is not None and current_num in target_nums:
                fixed_blocks[current_num] = "\n".join(current_lines)
            current_num = int(m.group(1))
            current_lines = [line]
        elif current_num is not None:
            current_lines.append(line)

    if current_num is not None and current_num in target_nums:
        fixed_blocks[current_num] = "\n".join(current_lines)

    if not fixed_blocks:
        return original_text

    # 在原文中定位并替换对应题目块
    orig_lines = original_text.split("\n")
    result_lines = []
    i = 0
    while i < len(orig_lines):
        line = orig_lines[i]
        m = re.match(r"^(\d+)[\.．、]\s*", line.strip())
        if m:
            q_num = int(m.group(1))
            if q_num in fixed_blocks:
                # 跳过原题块（直到下一题或下一大类标题）
                j = i + 1
                while j < len(orig_lines):
                    next_line = orig_lines[j].strip()
                    if re.match(r"^\d+[\.．、]\s*", next_line):
                        break
                    if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                        break
                    j += 1
                # 插入修复后的题目
                result_lines.append(fixed_blocks[q_num])
                result_lines.append("")
                i = j
                continue
        result_lines.append(line)
        i += 1

    return "\n".join(result_lines)


# === 后处理：原卷版生成 + 打包 + 分类 ===

import shutil
import zipfile


SUBJECTIVE_TYPE_KEYWORDS = ("简答", "综合", "计算", "作图", "绘图", "画图", "分析", "应用")
OBJECTIVE_TYPE_KEYWORDS = ("选择", "判断", "填空")
QUESTION_RE = re.compile(r"^\s*\d+\s*[．.、]")
TYPE_HEADING_RE = re.compile(r"^[一二三四五六七八九十百千万]+[、.．]\s*([^（(\s]+)")


def _is_type_heading(text):
    """判断段落是否为大题标题。"""
    return bool(TYPE_HEADING_RE.match(text.strip()))


def _is_subjective_heading(text):
    """判断大题标题是否为需要留答题空间的主观题。"""
    match = TYPE_HEADING_RE.match(text.strip())
    if not match:
        return False
    name = match.group(1)
    if any(keyword in name for keyword in OBJECTIVE_TYPE_KEYWORDS):
        return False
    return any(keyword in name for keyword in SUBJECTIVE_TYPE_KEYWORDS)


def _insert_blank_paragraph_after(paragraph):
    """在指定段落后插入一个空白段落，并尽量沿用原段落格式。"""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement('w:p')
    pPr = paragraph._element.pPr
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _add_subjective_answer_spaces(doc, blank_count=3):
    """在原卷版主观题每道题后插入空白答题区域。"""
    insert_after = []
    in_subjective = False
    current_question_last_para = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if _is_type_heading(text):
            if in_subjective and current_question_last_para is not None:
                insert_after.append(current_question_last_para)
            in_subjective = _is_subjective_heading(text)
            current_question_last_para = None
            continue

        if in_subjective and QUESTION_RE.match(text):
            if current_question_last_para is not None:
                insert_after.append(current_question_last_para)
            current_question_last_para = p
            continue

        if in_subjective and current_question_last_para is not None:
            current_question_last_para = p

    if in_subjective and current_question_last_para is not None:
        insert_after.append(current_question_last_para)

    inserted = 0
    for p in reversed(insert_after):
        anchor = p
        for _ in range(blank_count):
            anchor = _insert_blank_paragraph_after(anchor)
            inserted += 1
    return inserted


def _convert_to_blank(src_path, dst_path):
    """将解析版 docx 转换为原卷版（删除答案和解析段落）"""
    from docx import Document
    from docx.oxml.ns import qn as _qn

    shutil.copy2(src_path, dst_path)
    doc = Document(dst_path)
    paras = list(doc.paragraphs)

    # 标记需要删除的段落（含【答案】【解析】【详解】）
    rm_indices = []
    in_answer = False

    for i, p in enumerate(paras):
        text = p.text.strip()

        # 大类标题或新题号 → 重置答案区域标记
        if re.match(r'^[一二三四五六七八九十][、.．]', text):
            in_answer = False
            continue
        if re.match(r'^\d+[\.．、]', text) and '【答案】' not in text:
            in_answer = False
            continue

        # 混合行（题干+【答案】同行）→ 截掉答案部分
        if re.match(r'^\d+[\.．、]', text) and '【答案】' in text:
            idx = p.text.find('【答案】')
            if idx >= 0:
                pos = 0
                for run in p.runs:
                    rt = run.text or ''
                    rs = pos
                    re_ = pos + len(rt)
                    if rs >= idx:
                        run.text = ''
                    elif re_ > idx:
                        run.text = rt[:idx - rs]
                    pos = re_
            in_answer = True
            continue

        # 答案/解析行
        if '【答案】' in text or '【解析】' in text or '【详解】' in text:
            in_answer = True
            rm_indices.append(i)
            continue

        if in_answer:
            rm_indices.append(i)
            continue

    # 删除标记的段落（从后往前）
    for i in reversed(rm_indices):
        paras[i]._element.getparent().remove(paras[i]._element)

    # 去除底纹和高亮
    W_PPR = _qn('w:pPr')
    W_RPR = _qn('w:rPr')
    W_SHD = _qn('w:shd')
    W_HL = _qn('w:highlight')
    W_R = _qn('w:r')

    for p in doc.paragraphs:
        pPr = p._element.find(W_PPR)
        if pPr is not None:
            for e in pPr.findall(W_SHD):
                pPr.remove(e)
        for r in p._element.findall(W_R):
            rPr = r.find(W_RPR)
            if rPr is None:
                continue
            for tag in (W_SHD, W_HL):
                for e in rPr.findall(tag):
                    rPr.remove(e)

    _add_subjective_answer_spaces(doc, blank_count=3)
    doc.save(dst_path)


def _find_docx_pairs(directory):
    """递归查找（解析版）/（原卷版）配对文件，兼容已分类和未分类目录。"""
    pairs = {}
    pattern = re.compile(r'^(.+?)(?:[（(](解析版|原卷版)[）)])')

    for root, dirs, files in os.walk(directory):
        if '_原始文本' in root:
            continue
        root_path = Path(root)
        folder_variant = root_path.name if root_path.name in ('解析版', '原卷版') else None
        logical_root = root_path.parent if folder_variant else root_path

        for f in files:
            path = root_path / f
            if not path.is_file():
                continue
            if not f.endswith('.docx') or f.startswith('~'):
                continue
            match = pattern.match(f)
            if not match:
                continue
            base_name = match.group(1).strip()
            variant = match.group(2)
            key = (str(logical_root), base_name)
            if key not in pairs:
                pairs[key] = {}
            pairs[key][variant] = str(path)

    return {k: v for k, v in pairs.items() if '解析版' in v and '原卷版' in v}


def _post_process(output_dir):
    """后处理三步：1.生成原卷版 2.打包zip 3.分类到子文件夹"""

    # === 第1步：解析版 → 原卷版 ===
    print("\n[1/3] 生成原卷版...")
    converted = 0
    for root, dirs, files in os.walk(output_dir):
        if '_原始文本' in root or root.endswith('原卷版'):
            continue
        root_path = Path(root)
        dst_dir = root_path.parent / '原卷版' if root_path.name == '解析版' else root_path
        for f in files:
            src = root_path / f
            if not src.is_file():
                continue
            if f.endswith('.docx') and '解析版' in f and not f.startswith('~'):
                out_name = f.replace('解析版', '原卷版')
                dst = dst_dir / out_name
                if not dst.exists():
                    try:
                        dst_dir.mkdir(parents=True, exist_ok=True)
                        _convert_to_blank(str(src), str(dst))
                        converted += 1
                        print(f"  {out_name}")
                    except Exception as e:
                        print(f"  失败: {f} → {e}")
    print(f"  共生成 {converted} 个原卷版文件")

    # === 第2步：配对打包为 zip ===
    print("\n[2/3] 打包 zip（解析版+原卷版）...")
    pairs = _find_docx_pairs(output_dir)
    zipped = 0
    for (root, base_name), variants in pairs.items():
        zip_name = f"{base_name}.zip"
        zip_path = os.path.join(root, zip_name)
        if not os.path.exists(zip_path):
            try:
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for variant_type, filepath in variants.items():
                        zf.write(filepath, os.path.basename(filepath))
                zipped += 1
                print(f"  {zip_name}")
            except Exception as e:
                print(f"  失败: {zip_name} → {e}")
    print(f"  共打包 {zipped} 个 zip 文件")

    # === 第3步：分类到解析版/原卷版子文件夹 ===
    print("\n[3/3] 分类文件...")
    moved_jiexi = 0
    moved_yuanjuan = 0

    for root, dirs, files in os.walk(output_dir):
        if '_原始文本' in root or root.endswith('解析版') or root.endswith('原卷版'):
            continue

        has_docx = any(f.endswith('.docx') and ('解析版' in f or '原卷版' in f)
                       for f in files if not f.startswith('~'))
        if not has_docx:
            continue

        dir_jiexi = os.path.join(root, '解析版')
        dir_yuanjuan = os.path.join(root, '原卷版')
        os.makedirs(dir_jiexi, exist_ok=True)
        os.makedirs(dir_yuanjuan, exist_ok=True)

        for f in files:
            if not f.endswith('.docx') or f.startswith('~'):
                continue
            src = os.path.join(root, f)
            if '解析版' in f:
                shutil.move(src, os.path.join(dir_jiexi, f))
                moved_jiexi += 1
            elif '原卷版' in f:
                shutil.move(src, os.path.join(dir_yuanjuan, f))
                moved_yuanjuan += 1

    print(f"  解析版: {moved_jiexi} 个文件")
    print(f"  原卷版: {moved_yuanjuan} 个文件")
    print("\n后处理完成！")


# === 主流程 ===
def main():
    import argparse
    parser = argparse.ArgumentParser(description="一课一练试卷生成器")
    parser.add_argument("--file", "-f", help="考点规划表 xlsx 路径")
    parser.add_argument("--range", "-r", help="生成范围，如 1-5 或 3,7,12")
    parser.add_argument("--output", "-o", help="输出目录")
    parser.add_argument("--no-check", action="store_true", help="跳过质检，直接保存")
    args = parser.parse_args()

    # 加载配置
    config = load_config()
    spec_text = load_spec()

    # 初始化 API 客户端
    client = OpenAI(
        api_key=config["api_key"],
        base_url=config["base_url"],
    )

    # 选择规划表文件
    xlsx_path = args.file
    if not xlsx_path:
        planning_dir = BASE_DIR / "04_生成输出" / "考点规划表"
        if planning_dir.exists():
            files = sorted(
                [f for f in planning_dir.rglob("*.xlsx") if not f.name.startswith(("~", "_"))],
                key=lambda p: str(p.relative_to(planning_dir))
            )
            if not files:
                print("错误：考点规划表目录为空")
                return
            print("\n可用的考点规划表：")
            for i, f in enumerate(files, 1):
                rel = f.relative_to(planning_dir)
                print(f"  {i}. {rel}")
            choice = input("\n请输入编号选择规划表: ").strip()
            try:
                xlsx_path = str(files[int(choice) - 1])
            except (ValueError, IndexError):
                print("无效选择")
                return
        else:
            print("错误：找不到考点规划表目录")
            return

    # 解析规划表
    print(f"\n正在解析规划表: {Path(xlsx_path).name}")
    meta, topics = parse_planning_table(xlsx_path)
    print(f"  省份: {meta['province']}")
    print(f"  类别: {meta['category']}")
    print(f"  教材: {meta['textbooks']}")
    print(f"  共 {len(topics)} 个主题")

    # 确定生成范围
    if args.range:
        range_str = args.range
    else:
        print(f"\n主题列表：")
        for t in topics:
            level_mark = {"极重要": "★★", "重要": "★", "标准": "○"}.get(t["level"], "")
            sets_mark = f" ×{t['sets']}套" if t["sets"] > 1 else ""
            print(f"  {t['seq']:>3}. [{level_mark}] {t['theme']}{sets_mark}")
        range_str = input("\n请输入生成范围（如 1-5 或 3,7,12 或 all）: ").strip()

    # 解析范围
    if range_str.lower() == "all":
        selected_seqs = [t["seq"] for t in topics]
    elif "-" in range_str:
        start, end = range_str.split("-")
        selected_seqs = list(range(int(start), int(end) + 1))
    elif "," in range_str:
        selected_seqs = [int(x.strip()) for x in range_str.split(",")]
    else:
        selected_seqs = [int(range_str)]

    selected_topics = [t for t in topics if t["seq"] in selected_seqs]
    if not selected_topics:
        print("没有匹配的主题")
        return

    # 输出目录
    output_dir = args.output or str(BASE_DIR / "04_生成输出" / config.get("output_dir", "生成结果"))
    os.makedirs(output_dir, exist_ok=True)

    # 逐主题生成
    total_sets = sum(t["sets"] for t in selected_topics)
    print(f"\n即将生成 {len(selected_topics)} 个主题共 {total_sets} 份试卷")
    print(f"输出目录: {output_dir}")
    print(f"使用模型: {config['model']}")
    print(f"质检模式: {'跳过' if args.no_check else '生成后自动质检'}")
    print("=" * 60)

    success_count = 0
    fail_count = 0
    regen_count = 0

    # Token 用量跟踪
    session_usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0, "api_calls": 0}
    daily_usage = _load_daily_usage()

    # 编号连续主题防重复：仅在“xxx（二）”等明确续篇中参考前一篇“xxx（一）”摘要
    prev_paper_summaries = {}  # theme_base → {part_num: summaries list}

    for topic in selected_topics:
        for set_idx in range(1, topic["sets"] + 1):
            set_label = f"(第{set_idx}套)" if topic["sets"] > 1 else ""
            print(f"\n▶ 第{topic['seq']}练 {topic['theme']}{set_label}")
            print(f"  知识点: {topic['knowledge'][:60]}...")
            print(f"  级别: {topic['level']} | 题型: {topic['question_types']}")

            try:
                # 只有明确的连续编号主题才做跨卷防重复：如“电阻定律（二）”参考“电阻定律（一）”。
                # 普通主题（如“欧姆定律”）不会参考前一个不同主题的摘要。
                theme_base, theme_part_num = _split_numbered_theme(topic['theme'])
                existing_summaries_for_dedup = None
                if theme_part_num and theme_part_num > 1:
                    existing_summaries_for_dedup = prev_paper_summaries.get(theme_base, {}).get(theme_part_num - 1)

                # 构建 prompt（传入已有摘要防重复）
                sys_prompt, user_prompt = build_generation_prompt(
                    meta, topic, set_idx, spec_text,
                    existing_summaries=existing_summaries_for_dedup
                )

                # 生成+质检循环
                paper_text = None
                accepted = False
                needs_manual_review = False
                best_paper_text = None
                best_score = -1
                for attempt in range(1, MAX_REGEN_ATTEMPTS + 1):
                    attempt_label = f"(第{attempt}次)" if attempt > 1 else ""

                    if paper_text is None:
                        # 第一次生成整卷；后续若质检不过，只在上一版基础上定向修复，不再整卷重生。
                        print(f"  正在调用 API 生成试题{attempt_label}...")
                        start_time = time.time()

                        paper_text, usage = call_api(
                            client, config["model"],
                            sys_prompt, user_prompt,
                            max_tokens=config.get("max_tokens", 8000),
                            temperature=config.get("temperature", 0.7),
                        )
                        elapsed = time.time() - start_time

                        # 累计 token 用量
                        if usage:
                            for key in ("prompt_tokens", "completion_tokens", "total_tokens"):
                                session_usage[key] += usage.get(key, 0)
                                daily_usage[key] += usage.get(key, 0)
                            session_usage["api_calls"] += 1
                            daily_usage["api_calls"] += 1
                            _save_daily_usage(daily_usage)
                            token_info = f", {usage['total_tokens']} tokens"
                        else:
                            token_info = ""

                        print(f"  API 返回成功 ({elapsed:.1f}s, {len(paper_text)}字{token_info})")
                        paper_text = _normalize_generated_text(paper_text)
                    else:
                        print(f"  继续在上一版基础上定向修复{attempt_label}...")

                    # 跳过质检模式
                    if args.no_check:
                        accepted = True
                        break

                    # 自动调整答案分布（通过交换选项，不重新生成）
                    paper_text = _fix_answer_distribution(paper_text)

                    # 本地质检
                    cleaned = _clean_paper_text(paper_text)
                    severe, warnings, score, _, infos = _quick_check(cleaned)
                    _print_qc_summary(severe, warnings, score, infos)

                    # 记录三次生成中质检分数最高的一份，供最终兜底保留
                    if score > best_score:
                        best_score = score
                        best_paper_text = paper_text

                    # 定向修复质检问题：先修答案自暴露，再修其他单题问题，最后用最少改动修复查重问题。
                    if severe or warnings:
                        paper_text, severe, warnings, score, infos, fixed_rounds = _repair_qc_issues_targeted(
                            client, config, meta, topic, set_idx, spec_text,
                            paper_text, severe, warnings, score, infos,
                            session_usage, daily_usage,
                        )
                        regen_count += fixed_rounds
                        if score > best_score:
                            best_score = score
                            best_paper_text = paper_text
                        if fixed_rounds == 0 and (severe or warnings):
                            paper_text = best_paper_text or paper_text
                            accepted = True
                            needs_manual_review = True
                            print("  → 自动修复未产生有效改进，停止反复重试，保留当前最高分版本并标记待人工审核。")
                            break

                    # 评分>90且无严重/警告问题 → 自动保存，无需人工确认
                    if score > 90 and not severe and not warnings:
                        accepted = True
                        needs_manual_review = False
                        break

                    # 未通过则继续在上一版基础上定向修复，最多尝试三轮
                    if attempt < MAX_REGEN_ATTEMPTS:
                        print(f"  → 质检仍未通过，将继续只针对剩余问题修正上一版试卷（{attempt}/{MAX_REGEN_ATTEMPTS}）...")
                        time.sleep(2)
                        continue

                    # 三次均未通过：不再询问用户，保留分数最高的一份并标记待人工审核
                    paper_text = best_paper_text or paper_text
                    accepted = True
                    needs_manual_review = True
                    print(f"\n  已生成 {MAX_REGEN_ATTEMPTS} 次，最高评分 {best_score}/100。")
                    print("  → 自动保留最高分试卷，并在文件名前添加“（待人工审核）”。")
                    break

                if not accepted:
                    print(f"  ✗ 跳过: 第{topic['seq']}练 {topic['theme']}")
                    fail_count += 1
                    continue

                # 保存原始文本（与docx同目录结构下的_原始文本子目录）
                txt_base = _get_topic_output_base(meta, topic, output_dir)

                txt_dir = txt_base / "_原始文本"
                os.makedirs(txt_dir, exist_ok=True)
                txt_name = f"第{topic['seq']}练_{topic['theme']}{set_label}.txt"
                with open(txt_dir / txt_name, "w", encoding="utf-8") as f:
                    f.write(_normalize_generated_text(paper_text))

                # 生成 DOCX
                print("  正在生成 DOCX...")
                docx_path = generate_docx(meta, topic, set_idx, paper_text, output_dir, needs_manual_review=needs_manual_review)
                print(f"  ✓ 完成: {Path(docx_path).name}")
                success_count += 1

                # 记录题目摘要：只供后续同基础主题的编号续篇使用
                paper_summaries = _extract_question_summaries(paper_text)
                if paper_summaries and theme_part_num:
                    prev_paper_summaries.setdefault(theme_base, {})[theme_part_num] = paper_summaries

            except Exception as e:
                print(f"  ✗ 失败: {e}")
                traceback.print_exc()
                fail_count += 1

            # 避免 API 限流
            if success_count + fail_count < total_sets:
                time.sleep(2)

    # 汇总
    print("\n" + "=" * 60)
    print(f"生成完成！成功 {success_count} 份，失败/跳过 {fail_count} 份")
    if regen_count > 0:
        print(f"质检修复: {regen_count} 次")
    print(f"输出目录: {output_dir}")

    # Token 用量汇总
    if session_usage["api_calls"] > 0:
        _print_token_summary(session_usage, daily_usage)

    # 询问是否继续生成
    print()
    while True:
        cont = input("是否继续生成其他练习？(y=继续 / n=结束): ").strip().lower()
        if cont in ("y", "yes", "是"):
            print()
            main()
            return
        elif cont in ("n", "no", "否", ""):
            break
        print("请输入 y 或 n")

    # 后处理：生成原卷版 → 打包zip → 分类
    print("\n" + "=" * 60)
    print("正在执行后处理（原卷版生成 → 打包 → 分类）...")
    print("=" * 60)
    _post_process(output_dir)

    # 最终总结
    print("\n" + "=" * 60)
    print("工作总结")
    print("=" * 60)
    print(f"  本次生成: {success_count} 份试卷")
    if fail_count > 0:
        print(f"  跳过/失败: {fail_count} 份")
    if regen_count > 0:
        print(f"  质检修复: {regen_count} 次")
    print(f"  输出目录: {output_dir}")
    if session_usage["api_calls"] > 0:
        print(f"  API 调用: {session_usage['api_calls']} 次")
        print(f"  Token 消耗: {session_usage['total_tokens']:,} tokens")
    print(f"  今日累计: {daily_usage['total_tokens']:,} tokens ({daily_usage['api_calls']} 次调用)")
    print("=" * 60)
    input("\n按 Enter 关闭...")


if __name__ == "__main__":
    main()
