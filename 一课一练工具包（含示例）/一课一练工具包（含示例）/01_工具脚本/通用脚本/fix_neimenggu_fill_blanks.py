"""Regenerate only fill-in-the-blank sections for Inner Mongolia papers.

This is a one-off repair for papers where the planning table mistakenly used
multiple-choice questions. It replaces only the second section in raw TXT/DOCX
outputs and keeps all other question sections unchanged.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt
from openai import OpenAI

BASE_DIR = Path(__file__).resolve().parents[2]
if str(BASE_DIR / "01_工具脚本") not in sys.path:
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本"))

from 生成器.config_io import call_api, load_config  # noqa: E402
from 生成器.postprocess import _convert_to_blank  # noqa: E402

TARGET_DIRS = [
    BASE_DIR / "04_生成输出" / "生成结果" / "内蒙古自治区 机电类" / "电工电子技术与技能",
    BASE_DIR / "04_生成输出" / "生成结果" / "内蒙古自治区 机电类" / "电气测量技术",
]

SECTION_RE = re.compile(r"^[一二三四五六七八九十百]+[、.．]\s*\S+")
QUESTION_RE = re.compile(r"^\s*(\d+)\s*[.．、]")
OPTION_LINE_RE = re.compile(r"^\s*[A-D][.．、]\s*\S", re.M)


@dataclass
class RepairResult:
    raw_path: Path
    status: str
    message: str
    docx_path: str = ""
    original_path: str = ""
    zip_path: str = ""


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip() + "\n"


def find_second_section(text: str) -> tuple[int, int, str]:
    lines = text.splitlines()
    start = None
    for i, line in enumerate(lines):
        stripped = line.strip()
        if re.match(r"^二[、.．]\s*(多项选择题|填空题)", stripped):
            start = i
            break
    if start is None:
        raise ValueError("找不到第二大题（多项选择题/填空题）")

    end = len(lines)
    for j in range(start + 1, len(lines)):
        stripped = lines[j].strip()
        if SECTION_RE.match(stripped) and not re.match(r"^二[、.．]", stripped):
            end = j
            break
    old_block = "\n".join(lines[start:end]).strip()
    return start, end, old_block


def replace_second_section(text: str, new_block: str) -> str:
    lines = text.splitlines()
    start, end, _ = find_second_section(text)
    new_lines = normalize_newlines(new_block).splitlines()
    merged = lines[:start] + new_lines + lines[end:]
    return "\n".join(merged).rstrip() + "\n"


def infer_seq_theme(raw_path: Path) -> tuple[int, str]:
    m = re.match(r"第(\d+)练_(.+)\.txt$", raw_path.name)
    if not m:
        raise ValueError(f"无法从文件名解析序号和主题: {raw_path.name}")
    return int(m.group(1)), m.group(2)


def extract_brief_context(text: str, old_block: str) -> str:
    before = text.split(old_block, 1)[0].strip()
    after = text.split(old_block, 1)[1].strip() if old_block in text else ""
    # Keep enough context for topic style while limiting prompt size.
    return "\n\n".join([
        "【前文单选题】\n" + before[-3500:],
        "【后文题目】\n" + after[:3500],
    ])


def clean_api_section(text: str) -> str:
    text = normalize_newlines(text)
    fence = re.search(r"```(?:text|markdown)?\s*(.*?)\s*```", text, re.S)
    if fence:
        text = normalize_newlines(fence.group(1))

    lines = text.splitlines()
    heading_idx = None
    for i, line in enumerate(lines):
        if re.fullmatch(r"\s*二[、.．]\s*填空题\s*", line):
            heading_idx = i
            break
    if heading_idx is not None:
        text = "\n".join(lines[heading_idx:])
    else:
        text = "二、填空题\n" + text

    text = re.sub(r"</?think>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return normalize_newlines(text)


def split_question_blocks(section: str) -> dict[int, str]:
    lines = section.splitlines()
    positions: list[tuple[int, int]] = []
    for i, line in enumerate(lines):
        m = QUESTION_RE.match(line)
        if m:
            positions.append((int(m.group(1)), i))
    blocks: dict[int, str] = {}
    for idx, (num, pos) in enumerate(positions):
        next_pos = positions[idx + 1][1] if idx + 1 < len(positions) else len(lines)
        blocks[num] = "\n".join(lines[pos:next_pos]).strip()
    return blocks


def validate_fill_section(section: str) -> None:
    if not re.match(r"^二[、.．]\s*填空题", section.strip()):
        raise ValueError("新内容必须以“二、填空题”开头")
    if OPTION_LINE_RE.search(section):
        raise ValueError("填空题区块中仍包含 A-D 选项行")
    blocks = split_question_blocks(section)
    expected = {5, 6, 7, 8}
    if set(blocks) != expected:
        raise ValueError(f"题号必须且只能为 5-8，实际为 {sorted(blocks)}")
    for num in sorted(expected):
        block = blocks[num]
        if "______" not in block:
            raise ValueError(f"第{num}题缺少 ______")
        if "【答案】" not in block:
            raise ValueError(f"第{num}题缺少【答案】")
        if "【解析】" not in block:
            raise ValueError(f"第{num}题缺少【解析】")
    if section.count("【答案】") != 4 or section.count("【解析】") != 4:
        raise ValueError("填空题区块必须包含 4 个答案和 4 个解析")


def build_prompt(raw_path: Path, paper_text: str, old_block: str) -> tuple[str, str]:
    seq, theme = infer_seq_theme(raw_path)
    textbook = raw_path.parents[1].name
    system_prompt = """你是中职对口招生机电类命题专家。请只生成真正的填空题，不能生成选择题或把选项机械改成空。所有题目必须知识准确、答案明确、解析为完整因果句。"""
    user_prompt = f"""请为内蒙古自治区机电类《{textbook}》第{seq}练“{theme}”重新生成第二大题填空题，用于替换原来误写的多项选择题。

硬性要求：
1. 只输出“二、填空题”这一大题，不要输出其他题型、说明、标题或自检清单。
2. 固定生成 4 道题，题号必须是 5、6、7、8。
3. 每题必须是真正的填空题，题干中使用 6 个下划线 ______ 表示空白；一题可以有一个或多个空。
4. 不得出现 A. B. C. D. 选项，也不得把原多选题的选项简单删字改成填空。
5. 每题后必须紧跟【答案】和【解析】；多个空的答案按题干空白顺序书写，可用分号分隔。
6. 题目应围绕本练主题，覆盖核心术语、参数、公式、操作步骤或关键结论；难度从简单到适中排列，可包含1道带数值或公式的填空。
7. 解析必须说明为什么这样填，不能只复述答案。

原错误第二大题如下，只能参考其知识范围，不能机械改写：
{old_block}

同卷其他题目上下文如下。请避免与其他题完全重复：
{extract_brief_context(paper_text, old_block)}

输出格式示例：
二、填空题
5. ……______……。
【答案】……
【解析】……

6. ……______……______……。
【答案】答案1；答案2
【解析】……
"""
    return system_prompt, user_prompt


def generate_fill_section(client: OpenAI, config: dict, raw_path: Path, paper_text: str, old_block: str) -> str:
    system_prompt, user_prompt = build_prompt(raw_path, paper_text, old_block)
    text, _usage = call_api(
        client,
        config["model"],
        system_prompt,
        user_prompt,
        max_tokens=min(int(config.get("max_tokens", 8000)), 1800),
        temperature=0.35,
    )
    section = clean_api_section(text)
    validate_fill_section(section)
    return section


def copy_to_backup(path: Path, backup_root: Path, base_root: Path) -> Path:
    rel = path.relative_to(base_root)
    dst = backup_root / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, dst)
    return dst


def find_analysis_docx(root: Path, seq: int) -> Path | None:
    candidates = [
        p for p in root.rglob(f"第{seq}练*解析版*.docx")
        if p.is_file() and not p.name.startswith("~$") and "原卷版" not in p.parts
    ]
    if not candidates:
        return None
    # Prefer files inside 解析版 if present.
    candidates.sort(key=lambda p: (0 if p.parent.name == "解析版" else 1, len(str(p))))
    return candidates[0]


def block_text(block) -> str:
    tag = block.tag.rsplit("}", 1)[-1]
    if tag == "p":
        texts = [node.text or "" for node in block.iter() if node.tag.rsplit("}", 1)[-1] == "t"]
        return "".join(texts).strip()
    if tag == "tbl":
        texts = [node.text or "" for node in block.iter() if node.tag.rsplit("}", 1)[-1] == "t"]
        return "\n".join(t for t in texts if t).strip()
    return ""


def add_plain_paragraph_before(doc: Document, anchor, text: str, *, bold=False, red=False) -> None:
    p = doc.add_paragraph()
    if text.startswith("【答案】"):
        label, content = "【答案】", text[len("【答案】"):]
        r1 = p.add_run(label)
        r2 = p.add_run(content)
        for run in (r1, r2):
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(10.5)
    elif text.startswith("【解析】"):
        label, content = "【解析】", text[len("【解析】"):]
        r1 = p.add_run(label)
        r2 = p.add_run(content)
        for run in (r1, r2):
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(12 if bold else 10.5)
        if red:
            run.font.color.rgb = RGBColor(255, 0, 0)
    body = doc._body._element
    body.remove(p._element)
    anchor.addprevious(p._element)


def update_docx_section(docx_path: Path, fill_section: str) -> None:
    doc = Document(str(docx_path))
    body = doc._body._element
    children = list(body)
    start_idx = None
    end_idx = None
    for i, child in enumerate(children):
        text = block_text(child)
        if re.match(r"^二[、.．]\s*(多项选择题|填空题)", text):
            start_idx = i
            break
    if start_idx is None:
        raise ValueError(f"DOCX 找不到第二大题: {docx_path.name}")
    for j in range(start_idx + 1, len(children)):
        text = block_text(children[j])
        if re.match(r"^三[、.．]\s*\S+", text):
            end_idx = j
            break
    if end_idx is None:
        raise ValueError(f"DOCX 找不到第三大题锚点: {docx_path.name}")
    anchor = children[end_idx]
    for child in children[start_idx:end_idx]:
        body.remove(child)
    for line in fill_section.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        add_plain_paragraph_before(
            doc,
            anchor,
            stripped,
            bold=bool(re.match(r"^二[、.．]", stripped)),
        )
    doc.save(str(docx_path))


def original_docx_path(analysis_path: Path) -> Path:
    name = analysis_path.name.replace("解析版", "原卷版")
    if analysis_path.parent.name == "解析版":
        return analysis_path.parent.parent / "原卷版" / name
    return analysis_path.with_name(name)


def update_original_and_zip(analysis_path: Path) -> tuple[str, str]:
    original_path = original_docx_path(analysis_path)
    original_path.parent.mkdir(parents=True, exist_ok=True)
    if original_path.exists():
        original_path.unlink()
    _convert_to_blank(str(analysis_path), str(original_path))

    base_name = analysis_path.name.replace("（解析版）", "").replace("(解析版)", "").strip()
    zip_root = analysis_path.parent.parent if analysis_path.parent.name == "解析版" else analysis_path.parent
    zip_path = zip_root / f"{base_name}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(analysis_path, analysis_path.name)
        zf.write(original_path, original_path.name)
    return str(original_path), str(zip_path)


def repair_one(raw_path: Path, client: OpenAI, config: dict, backup_root: Path, target_root: Path, dry_run=False) -> RepairResult:
    try:
        paper_text = read_text(raw_path)
        _start, _end, old_block = find_second_section(paper_text)
        if "多项选择题" not in old_block and "A." not in old_block and "B." not in old_block:
            # Still allow already-converted fill blanks to be regenerated if the section is malformed.
            pass
        fill_section = generate_fill_section(client, config, raw_path, paper_text, old_block)
        new_text = replace_second_section(paper_text, fill_section)
        if new_text == paper_text:
            return RepairResult(raw_path, "skipped", "替换后文本未变化")
        if dry_run:
            return RepairResult(raw_path, "dry-run", "API生成与校验通过，未写入")

        copy_to_backup(raw_path, backup_root / "raw", target_root)
        write_text(raw_path, new_text)

        seq, _theme = infer_seq_theme(raw_path)
        docx_path = find_analysis_docx(target_root, seq)
        original_path = ""
        zip_path = ""
        if docx_path:
            copy_to_backup(docx_path, backup_root / "docx", target_root)
            orig = original_docx_path(docx_path)
            if orig.exists():
                copy_to_backup(orig, backup_root / "docx", target_root)
            update_docx_section(docx_path, fill_section)
            original_path, zip_path = update_original_and_zip(docx_path)
        else:
            return RepairResult(raw_path, "partial", "原始文本已修复，但找不到解析版DOCX")
        return RepairResult(raw_path, "ok", "修复成功", str(docx_path), original_path, zip_path)
    except Exception as exc:  # keep batch running
        return RepairResult(raw_path, "failed", str(exc))


def collect_raw_files(limit: int | None = None, offset: int = 0) -> list[Path]:
    files: list[Path] = []
    for target in TARGET_DIRS:
        files.extend(sorted((target / "_原始文本").glob("*.txt"), key=lambda p: int(re.match(r"第(\d+)练", p.name).group(1))))
    return files[offset: offset + limit if limit else None]


def write_report(results: list[RepairResult], report_path: Path, backup_root: Path) -> None:
    counts = {}
    for r in results:
        counts[r.status] = counts.get(r.status, 0) + 1
    lines = [
        "# 内蒙古填空题修复报告",
        "",
        f"- 时间：{time.strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 备份目录：`{backup_root}`",
        f"- 统计：{json.dumps(counts, ensure_ascii=False)}",
        "",
        "| 状态 | 原始文本 | 说明 | 解析版DOCX | 原卷版DOCX | ZIP |",
        "|---|---|---|---|---|---|",
    ]
    for r in results:
        rel = r.raw_path.relative_to(BASE_DIR)
        lines.append(
            f"| {r.status} | `{rel}` | {str(r.message).replace('|', '｜')} | "
            f"`{Path(r.docx_path).relative_to(BASE_DIR) if r.docx_path else ''}` | "
            f"`{Path(r.original_path).relative_to(BASE_DIR) if r.original_path else ''}` | "
            f"`{Path(r.zip_path).relative_to(BASE_DIR) if r.zip_path else ''}` |"
        )
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    parser = argparse.ArgumentParser(description="Regenerate Inner Mongolia fill-in sections only")
    parser.add_argument("--limit", type=int, default=None, help="最多处理多少份")
    parser.add_argument("--offset", type=int, default=0, help="跳过前多少份")
    parser.add_argument("--dry-run", action="store_true", help="只调用API并校验，不写入")
    args = parser.parse_args()

    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_root = BASE_DIR / "04_生成输出" / "生成结果" / f"_fill_blank_fix_backup_{timestamp}"
    report_path = BASE_DIR / "04_生成输出" / "质检报告" / f"内蒙古填空题修复报告_{timestamp}.md"
    report_path.parent.mkdir(parents=True, exist_ok=True)

    config = load_config()
    client = OpenAI(api_key=config["api_key"], base_url=config["base_url"])
    files = collect_raw_files(args.limit, args.offset)
    print(f"准备处理 {len(files)} 份原始文本。")
    results: list[RepairResult] = []
    for idx, raw_path in enumerate(files, 1):
        target_root = raw_path.parents[1]
        print(f"[{idx}/{len(files)}] {raw_path.relative_to(BASE_DIR)}")
        result = repair_one(raw_path, client, config, backup_root, target_root, dry_run=args.dry_run)
        print(f"  -> {result.status}: {result.message}")
        results.append(result)
        time.sleep(0.5)

    write_report(results, report_path, backup_root)
    print(f"报告已写入: {report_path}")
    failed = [r for r in results if r.status == "failed"]
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
