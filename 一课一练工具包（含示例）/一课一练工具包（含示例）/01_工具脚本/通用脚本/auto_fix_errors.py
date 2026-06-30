"""
AI 自动修正试卷报错工具

功能：
  1. 读取爬虫生成的 AI文档质检报告汇总.xlsx
  2. 对"答案错误"类型：
     - 解析具体描述，提取题型/原答案/建议正确答案
     - 读取 docx 题干，检查是否含有无法处理的图片
     - AI 独立做题（不看期望答案），得出自己的答案
     - 对比 AI 答案与具体描述中的建议答案
     - 一致 → 修改文档中的答案 → 打 ✔
     - 不一致或无法判断 → 留给人工处理
  3. 对其他错误类型：
     - 打开对应 docx 文件，定位报错位置
     - AI 阅读"具体描述"+ 原文上下文，判断是否为真实错误
     - 真实错误 → 按描述修改文档 → 打 ✔
     - 公式图片误报（检测到图片存在） → 打 ✔
     - 无法判断 → 不打勾，留给人工
  4. 保存修正后的 docx 和更新后的 Excel

用法：
  1. 确保 AI文档质检报告汇总.xlsx 和 docx 文件在同一目录
  2. 运行: python auto_fix_errors.py
  3. 按提示操作

依赖：
  pip install openpyxl python-docx openai
"""

import os
import sys
import re
import json
import time
import shutil

try:
    from openai import OpenAI
except ImportError:
    print("缺少依赖，请执行: pip install openai")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment
except ImportError:
    print("缺少依赖，请执行: pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖，请执行: pip install python-docx")
    sys.exit(1)


# ════════════════════ 配置区域 ════════════════════

CONFIG = {
    "api_key": "sk-ufvlhvdnkafenmwoznvqstpoboxnguaxwvhetmmcmkegzwzs",
    "base_url": "https://api.siliconflow.cn/v1",
    "model": "deepseek-ai/DeepSeek-V3",

    "report_xlsx": "AI文档质检报告汇总.xlsx",
    "docx_dir": ".",  # docx 试卷文件所在目录

    "backup": True,
    "ai_delay": 1,  # API 调用间隔（秒）
}

# 资料ID 与本地 docx 文件名的映射
# 如果文件名中包含资料ID或能自动匹配则不需要手动配置
RESOURCE_FILE_MAP = {
    # "58106924": "第52卷 时序逻辑与RS触发器.docx",
}

# ════════════════════ 配置结束 ════════════════════


def parse_answer_info(position, description):
    """
    从具体位置和具体描述中全面提取答案相关信息。
    返回 dict:
      question_type: "choice" / "true_false" / "fill_blank" / "essay" / "unknown"
      old_answer: 原答案（字母/正确错误/文本）
      new_answer: 描述中建议的正确答案
      has_explicit_answers: 是否有明确的新旧答案
    """
    info = {
        "question_type": "unknown",
        "old_answer": "",
        "new_answer": "",
        "has_explicit_answers": False,
    }

    combined = f"{position} {description}"

    # ── 1. 判断题型 ──
    if re.search(r"单选|多选|选择题|选项", combined):
        info["question_type"] = "choice"
    elif re.search(r"判断题|判断为|标注为错误|标注为正确", combined):
        info["question_type"] = "true_false"
    elif re.search(r"填空|填写", combined):
        info["question_type"] = "fill_blank"
    elif re.search(r"简答|论述|问答", combined):
        info["question_type"] = "essay"

    # 如果位置中提到大题类型，也用来判断
    if info["question_type"] == "unknown":
        if re.search(r"第[一二三四五六七八九十\d]+大题", position):
            pass  # 需要结合描述内容判断
        if re.search(r"[A-D]", description) and re.search(r"选项?|选", description):
            info["question_type"] = "choice"

    # ── 2. 提取选择题答案 ──
    if info["question_type"] in ("choice", "unknown"):
        # 原答案提取：多种模式
        old_patterns = [
            r'原答案(?:为|是|给出|误写为)[：:\s]*([A-Da-d])',
            r'原(?:答案|答)[：:\s]*([A-Da-d])',
            r'原答案([A-Da-d])',
            r'原(?:答案)?(?:.*?)([A-Da-d])(?:错误|有误|不正确|不对)',
        ]
        for pat in old_patterns:
            m = re.search(pat, description)
            if m:
                info["old_answer"] = m.group(1).upper()
                if info["question_type"] == "unknown":
                    info["question_type"] = "choice"
                break

        # 正确答案提取：多种模式
        new_patterns = [
            r'正确(?:答案)?(?:应(?:为|是|改为))[：:\s]*([A-Da-d])',
            r'(?:应(?:为|改为|是))[：:\s]*([A-Da-d])(?:选项|[，,。]|$)',
            r'正确(?:答案)?(?:应)?(?:为|是)[：:\s]*([A-Da-d])',
            r'本题正确答案应为([A-Da-d])',
            r'正确应为([A-Da-d])',
            r'(?:正确|应为)[：:\s]*([A-Da-d])选项',
        ]
        for pat in new_patterns:
            m = re.search(pat, description)
            if m:
                info["new_answer"] = m.group(1).upper()
                info["has_explicit_answers"] = bool(info["old_answer"])
                if info["question_type"] == "unknown":
                    info["question_type"] = "choice"
                break

    # ── 3. 提取判断题答案 ──
    if info["question_type"] == "true_false" or re.search(
            r'(?:标注|判断)为(?:了)?["\u201c]?(?:正确|错误|对|错)', description):
        info["question_type"] = "true_false"
        if re.search(r'(?:标注|判断|原答案)(?:为|是)["\u201c]?错误', description) and \
           re.search(r'应(?:改|判断)?为["\u201c]?正确', description):
            info["old_answer"] = "错误"
            info["new_answer"] = "正确"
            info["has_explicit_answers"] = True
        elif re.search(r'(?:标注|判断|原答案)(?:为|是)["\u201c]?正确', description) and \
             re.search(r'应(?:改|判断)?为["\u201c]?错误', description):
            info["old_answer"] = "正确"
            info["new_answer"] = "错误"
            info["has_explicit_answers"] = True
        elif re.search(r'(?:标注|判断|原答案)(?:为|是)["\u201c]?正确', description) and \
             re.search(r'(?:实际|实则|应)[^\n]{0,10}错误', description):
            info["old_answer"] = "正确"
            info["new_answer"] = "错误"
            info["has_explicit_answers"] = True
        elif re.search(r'(?:标注|判断|原答案)(?:为|是)["\u201c]?错误', description) and \
             re.search(r'(?:实际|实则|应)[^\n]{0,10}正确', description):
            info["old_answer"] = "错误"
            info["new_answer"] = "正确"
            info["has_explicit_answers"] = True

    # ── 4. 提取填空题答案 ──
    if info["question_type"] == "fill_blank" or (
            not info["old_answer"] and re.search(r'原答案["\u201c]', description)):
        quoted_after_orig = re.findall(
            r'原答案["\u201c]([^"\u201d]+)["\u201d]', description)
        if quoted_after_orig:
            info["old_answer"] = quoted_after_orig[0]
            if info["question_type"] == "unknown":
                info["question_type"] = "fill_blank"
        # 也提取引号内的原文
        if not info["old_answer"]:
            all_quoted = re.findall(r'["\u201c]([^"\u201d]+)["\u201d]', description)
            if all_quoted and "原答案" in description:
                info["old_answer"] = all_quoted[0]
                if info["question_type"] == "unknown":
                    info["question_type"] = "fill_blank"

    return info


def question_has_critical_images(question_text):
    """
    检查题干是否含有大量无法处理的图片标记，导致AI无法独立做题。
    返回 True 表示图片过多/关键信息缺失，应留给人工。
    """
    if not question_text:
        return True

    img_markers = len(re.findall(r'\[公式图片\]|\[图片\]|\[含公式图片\]', question_text))
    text_chars = len(re.sub(r'\s+', '', re.sub(
        r'\[公式图片\]|\[图片\]|\[含公式图片\]', '', question_text)))

    if img_markers == 0:
        return False

    # 如果图片标记 >= 3 且有效文字少于 50 字，视为无法解答
    if img_markers >= 3 and text_chars < 50:
        return True

    # 如果题干几乎全是图片（图片占比过高）
    if text_chars < 20 and img_markers >= 1:
        return True

    return False


class AIClient:
    """DeepSeek AI 接口"""

    def __init__(self):
        self.client = OpenAI(
            api_key=CONFIG["api_key"],
            base_url=CONFIG["base_url"],
        )
        self.model = CONFIG["model"]

    def _call_ai(self, prompt, temperature=0.1):
        """统一的 AI 调用方法，返回解析后的 JSON dict 或 None"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=temperature,
            )
            content = response.choices[0].message.content.strip()
            content = re.sub(r"^```json\s*", "", content)
            content = re.sub(r"\s*```$", "", content)
            return json.loads(content)
        except json.JSONDecodeError:
            try:
                return eval(content)
            except Exception:
                return None
        except Exception as e:
            print(f"  ⚠ AI调用失败: {e}")
            return None

    def solve_question_independently(self, question_text, question_type, position):
        """
        AI 独立做题，不给出期望答案，纯粹根据题干作答。
        返回: {"my_answer": str, "confidence": "high"/"medium"/"low", "reason": str}
        """
        type_hints = {
            "choice": "这是一道选择题，请给出你认为正确的选项字母（如A/B/C/D）。",
            "true_false": '这是一道判断题，请判断题目表述是否正确，回答"正确"或"错误"。',
            "fill_blank": "这是一道填空题，请给出你认为应该填写的内容。",
            "essay": "这是一道简答题，请给出你的答案要点。",
            "unknown": "请根据题目类型给出你的答案。",
        }
        type_hint = type_hints.get(question_type, type_hints["unknown"])

        prompt = f"""你是一名电子技术专业教师，请认真阅读以下试题并独立作答。

## 题目位置
{position}

## 试题内容
{question_text}

## 要求
{type_hint}

注意：
- [公式图片] 标记表示此处原有数学公式或图片，请尽量根据上下文推断内容
- 如果因为图片缺失导致关键信息不足，无法确定答案，请如实说明
- 请基于电子技术专业知识严谨作答

## 输出格式（严格JSON，不要其他内容）
{{
  "my_answer": "你的答案（选择题写字母如A，判断题写正确或错误，填空题写具体内容）",
  "confidence": "high/medium/low（你对答案的把握程度）",
  "reason": "简短解题思路（50字以内）"
}}"""

        result = self._call_ai(prompt)
        if result is None:
            return {"my_answer": "", "confidence": "low", "reason": "AI调用或解析失败"}
        return result

    def verify_answer_with_description(self, question_text, description, position):
        """
        AI 仔细阅读具体描述和题干，综合判断描述中指出的答案错误是否成立。
        用于 AI 独立做题结果与描述不一致时的二次确认，或无法独立做题时的兜底判断。
        返回: {"agree": bool, "my_answer": str, "reason": str}
        """
        prompt = f"""你是一名电子技术专业教师。请仔细阅读以下试题和质检描述，判断质检指出的答案错误是否正确。

## 题目位置
{position}

## 试题内容
{question_text}

## 质检报告的具体描述（请仔细阅读！）
{description}

## 你的任务
1. 先独立思考题目的正确答案
2. 再仔细阅读质检描述中给出的分析和理由
3. 综合判断：质检描述指出的问题是否合理

注意：
- [公式图片] 标记表示此处有公式/图片，请根据上下文推断
- 质检描述中的分析可能正确也可能有误，请独立判断
- 不确定时输出 agree=false

## 输出格式（严格JSON，不要其他内容）
{{
  "agree": true或false,
  "my_answer": "你认为的正确答案",
  "reason": "简短说明（30字以内）"
}}"""

        result = self._call_ai(prompt)
        if result is None:
            return {"agree": False, "my_answer": "", "reason": "AI调用失败"}
        return result

    def analyze_error(self, error_type, position, description, context_text):
        """
        AI 分析错误是否成立，给出修正方案。
        context_text 可以为空（当没有 docx 文件时仅根据描述判断）。
        """
        prompt = f"""你是一名专业的试卷校对员。请根据以下AI质检报错信息，判断该错误是否真实存在并需要修改。

## 质检报错信息
- 错误类型: {error_type}
- 具体位置: {position}
- 具体描述: {description}

"""
        if context_text:
            prompt += f"""## 文档原文（报错位置附近的内容）
{context_text}

"""

        prompt += """## 判断规则
1. 仔细阅读"具体描述"，理解质检AI指出的问题
2. 如果描述提到"内容缺失""空白""无内容"等，且位置处可能存在数学公式图片（标记为 [公式图片]），则判定为【误报】
3. 如果描述指出了具体的文字错误（错别字、标点、格式问题等），且原文确实有该问题，则判定为【需修改】，给出原文和修正
4. 如果描述内容不合理或原文没有问题，则判定为【误报】
5. 对于"知识性错误""表述不当"等需要专业判断的，如果你不确定则判定为【不确定】

## 输出格式（严格JSON，不要其他内容）
{
  "judgment": "需修改" 或 "误报" 或 "不确定",
  "original": "原文中需要修改的片段（仅 judgment 为 需修改 时填写，否则留空）",
  "fix": "修正后的文字（仅 judgment 为 需修改 时填写，否则留空）",
  "reason": "简短说明判断理由（15字以内）"
}"""

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            content = response.choices[0].message.content.strip()
            content = re.sub(r"^```json\s*", "", content)
            content = re.sub(r"\s*```$", "", content)
            return json.loads(content)
        except json.JSONDecodeError:
            try:
                return eval(content)
            except Exception:
                return {"judgment": "不确定", "original": "", "fix": "", "reason": "AI返回格式异常"}
        except Exception as e:
            return {"judgment": "不确定", "original": "", "fix": "", "reason": f"AI调用失败: {e}"}

    def analyze_text_error(self, description, context_text, position):
        """
        处理文本类错误：错字/少字/多字/表述不当/文本有误等。

        两种情况都能处理：
        1. 描述具体（如"55时基电路少了一个5"）→ 直接提取修改方案
        2. 描述模糊（如"此处文本有误"）→ AI独立阅读原文，判断是否有错并给出修改
        """
        prompt = f"""你是一名电子技术专业的资深校对员。你的任务是判断文档中是否存在文字错误，并给出精确的修改方案。

## 具体位置
{position}

## 质检报告的具体描述
{description}

"""
        if context_text:
            prompt += f"""## 文档原文（报错位置附近，请逐字仔细阅读！）
{context_text}

"""

        prompt += """## 工作流程

### 第一步：理解描述
仔细阅读"具体描述"，判断它属于哪种情况：
- A) 描述很具体：明确指出了哪个字/词有误，或者说了"少了""多了""写成了"等
- B) 描述较模糊：只笼统地说"文本有误""表述不当""不准确"等，没有明确指出错在哪里

### 第二步：根据情况处理

**如果是情况A（描述具体）：**
1. 从描述中提取错误文本和正确文本
2. 在原文中确认该错误确实存在
3. 给出精确的替换方案
4. 举例：描述"55时基电路少了一个5" → original="55时基电路", fix="555时基电路"
5. 举例：描述"把输入写成了输人" → original="输人", fix="输入"

**如果是情况B（描述模糊）：**
1. 认真逐字阅读文档原文，结合电子技术专业知识独立判断
2. 重点检查：错别字、形近字（如"人/入""已/己""未/末"）、漏字、多字、专业术语拼写
3. 如果你在原文中确实发现了文字问题，给出修改方案
4. 如果原文阅读后觉得没有明显问题，设 can_fix=false
5. 特别注意：报错位置如果有 [公式图片] 标记，说明该处是图片/公式，不是文字错误，应设 can_fix=false

### 第三步：输出

## 输出格式（严格JSON，不要其他内容）
{
  "can_fix": true或false,
  "original": "需要替换的原文片段（必须是文档原文中实际存在的精确文本）",
  "fix": "替换后的正确文本",
  "reason": "简短说明（20字以内）"
}

## 关键约束
- original 必须是能在文档中直接搜到的精确文本，不能自己编造
- 如果文档原文未提供（为空），而描述又不够具体，请设 can_fix=false
- [公式图片] 标记处的问题设 can_fix=false
- 不确定是否有错时，宁可设 can_fix=false 留给人工，不要强行修改"""

        result = self._call_ai(prompt)
        if result is None:
            return {"can_fix": False, "original": "", "fix": "", "reason": "AI调用失败"}
        return result


class DocxProcessor:
    """Docx 文档处理器"""

    def __init__(self, docx_path):
        self.path = docx_path
        self.doc = Document(docx_path)
        self.modified = False

    def _paragraph_has_image(self, para):
        """检查段落是否含图片/公式"""
        xml = para._element.xml
        indicators = ["w:drawing", "w:pict", "w:object",
                      "mc:AlternateContent", "m:oMath", "m:oMathPara"]
        return any(ind in xml for ind in indicators)

    def find_paragraph_by_position(self, position_text, description=""):
        """根据位置描述和具体描述中的文字片段，在文档中部分匹配定位段落"""

        # 1. 直接匹配位置文本
        for i, para in enumerate(self.doc.paragraphs):
            if position_text and position_text in para.text:
                return i

        # 2. 从描述中提取可能的原文片段进行匹配
        #    常见格式："原文为XXX" "原答案为X" "应改为" 等
        text_fragments = []

        # 提取引号内的文字
        quoted = re.findall(r'[""「」\'](.*?)[""「」\']', description)
        text_fragments.extend(quoted)

        # 提取"原文为/原答案为/原内容为"后面的内容
        originals = re.findall(r'原(?:文|答案|内容)?(?:为|是)[：:]?\s*(.+?)(?:[，,；;。]|$)', description)
        text_fragments.extend(originals)

        # 提取较长的中文片段（4字以上连续中文）
        cn_fragments = re.findall(r'[\u4e00-\u9fff]{4,}', description)
        text_fragments.extend(cn_fragments)

        # 按长度降序排列（越长越精确）
        text_fragments = sorted(set(text_fragments), key=len, reverse=True)

        for frag in text_fragments:
            if len(frag) < 2:
                continue
            for i, para in enumerate(self.doc.paragraphs):
                if frag in para.text:
                    return i

        # 3. 位置文本中的关键词匹配（如"第一大题""选择题"等）
        pos_keywords = re.findall(r'[\u4e00-\u9fff]+', position_text)
        for kw in pos_keywords:
            if len(kw) >= 2:
                for i, para in enumerate(self.doc.paragraphs):
                    if kw in para.text:
                        return i

        return None

    def get_context(self, para_index, window=3):
        """获取指定段落及上下文文本"""
        paras = self.doc.paragraphs
        start = max(0, para_index - window)
        end = min(len(paras), para_index + window + 1)

        lines = []
        for i in range(start, end):
            text = paras[i].text.strip()
            has_img = self._paragraph_has_image(paras[i])

            if has_img and not text:
                lines.append(f"[第{i+1}段] [公式图片]")
            elif has_img:
                lines.append(f"[第{i+1}段] {text} [含公式图片]")
            else:
                lines.append(f"[第{i+1}段] {text}")

            if i == para_index:
                lines[-1] += "  ← 报错位置"

        return "\n".join(lines)

    def get_question_context(self, para_index, window=15):
        """获取完整题目内容（更大范围，用于AI做题）"""
        paras = self.doc.paragraphs
        start = max(0, para_index - 5)
        end = min(len(paras), para_index + window)

        lines = []
        for i in range(start, end):
            text = paras[i].text.strip()
            has_img = self._paragraph_has_image(paras[i])

            if has_img and not text:
                lines.append("[公式图片]")
            elif has_img:
                lines.append(f"{text} [公式图片]")
            elif text:
                lines.append(text)

        return "\n".join(lines)

    def fix_answer_and_remove_analysis(self, para_index, old_answer, new_answer, description, question_type="choice"):
        """修改答案并删除对应的解析/详解，支持不同题型"""
        paras = self.doc.paragraphs
        answer_fixed = False
        search_start = max(0, para_index - 3)
        search_end = min(len(paras), para_index + 25)

        if question_type == "true_false":
            return self._fix_true_false_answer(
                para_index, old_answer, new_answer, search_start, search_end)

        if question_type == "fill_blank":
            return self._fix_fill_blank_answer(
                para_index, old_answer, new_answer, search_start, search_end)

        # ── 选择题答案替换 ──
        for i in range(search_start, search_end):
            text = paras[i].text.strip()
            if re.search(r'(答案|Answer)[：:\s】\]]*' + re.escape(old_answer), text, re.IGNORECASE):
                for run in paras[i].runs:
                    if old_answer in run.text:
                        run.text = run.text.replace(old_answer, new_answer, 1)
                        answer_fixed = True
                        self.modified = True
                        break
                if answer_fixed:
                    self._remove_analysis_after(i)
                    break

        if not answer_fixed:
            for i in range(search_start, search_end):
                text = paras[i].text
                if old_answer in text and any(kw in text for kw in ["答案", "Answer", "答"]):
                    for run in paras[i].runs:
                        if old_answer in run.text:
                            run.text = run.text.replace(old_answer, new_answer, 1)
                            answer_fixed = True
                            self.modified = True
                            break
                    if answer_fixed:
                        self._remove_analysis_after(i)
                        break

        return answer_fixed

    def _fix_true_false_answer(self, para_index, old_answer, new_answer, search_start, search_end):
        """修改判断题答案（正确↔错误, 对↔错, √↔×）"""
        paras = self.doc.paragraphs

        old_variants = []
        new_variant = new_answer
        if old_answer in ("正确", "对"):
            old_variants = ["正确", "对", "√", "✓"]
            new_variant = "错误" if new_answer in ("错误", "错") else new_answer
        elif old_answer in ("错误", "错"):
            old_variants = ["错误", "错", "×", "✗"]
            new_variant = "正确" if new_answer in ("正确", "对") else new_answer

        for i in range(search_start, search_end):
            text = paras[i].text.strip()
            if not any(kw in text for kw in ["答案", "Answer", "答"]):
                continue
            for old_v in old_variants:
                if old_v in text:
                    for run in paras[i].runs:
                        if old_v in run.text:
                            run.text = run.text.replace(old_v, new_variant, 1)
                            self.modified = True
                            self._remove_analysis_after(i)
                            return True
        return False

    def _fix_fill_blank_answer(self, para_index, old_answer, new_answer, search_start, search_end):
        """修改填空题答案（替换答案区域中的文本）"""
        paras = self.doc.paragraphs

        if not old_answer or not new_answer:
            return False

        for i in range(search_start, search_end):
            text = paras[i].text
            if old_answer in text and any(kw in text for kw in ["答案", "Answer", "答"]):
                for run in paras[i].runs:
                    if old_answer in run.text:
                        run.text = run.text.replace(old_answer, new_answer, 1)
                        self.modified = True
                        self._remove_analysis_after(i)
                        return True
                # 跨 run 的情况
                full_text = "".join(r.text for r in paras[i].runs)
                if old_answer in full_text:
                    new_text = full_text.replace(old_answer, new_answer, 1)
                    if paras[i].runs:
                        for r in paras[i].runs:
                            r.text = ""
                        paras[i].runs[0].text = new_text
                        self.modified = True
                        self._remove_analysis_after(i)
                        return True

        # 退而求其次：在附近段落中直接搜索旧答案文本
        for i in range(search_start, search_end):
            if old_answer in paras[i].text:
                for run in paras[i].runs:
                    if old_answer in run.text:
                        run.text = run.text.replace(old_answer, new_answer, 1)
                        self.modified = True
                        return True
        return False

    def _remove_analysis_after(self, answer_para_index):
        """删除答案行之后的解析/详解段落"""
        paras = self.doc.paragraphs
        analysis_keywords = ["解析", "详解", "解答", "分析", "解题"]

        i = answer_para_index + 1
        while i < len(paras) and i < answer_para_index + 10:
            text = paras[i].text.strip()

            # 如果遇到下一题的开始，停止删除
            if re.match(r'^\d+[\.\、\s]', text):
                break
            if any(kw in text for kw in ["答案", "【答案】"]) and i > answer_para_index + 1:
                break

            # 如果是解析相关内容，清空该段落
            if any(kw in text for kw in analysis_keywords) or (
                    i == answer_para_index + 1 and text and "答案" not in text):
                for run in paras[i].runs:
                    run.text = ""
                self.modified = True

            i += 1

    def check_image_at_position(self, position_text, description=""):
        """检查报错位置附近是否有图片/公式"""
        para_idx = self.find_paragraph_by_position(position_text, description)
        if para_idx is None:
            return False, None

        if self._paragraph_has_image(self.doc.paragraphs[para_idx]):
            return True, para_idx

        for offset in range(1, 4):
            for idx in [para_idx - offset, para_idx + offset]:
                if 0 <= idx < len(self.doc.paragraphs):
                    if self._paragraph_has_image(self.doc.paragraphs[idx]):
                        return True, para_idx

        return False, para_idx

    def replace_text(self, para_index, original, replacement):
        """替换段落中的文本"""
        if para_index is not None and para_index < len(self.doc.paragraphs):
            para = self.doc.paragraphs[para_index]
            if original in para.text:
                for run in para.runs:
                    if original in run.text:
                        run.text = run.text.replace(original, replacement, 1)
                        self.modified = True
                        return True

        for i, para in enumerate(self.doc.paragraphs):
            if original in para.text:
                for run in para.runs:
                    if original in run.text:
                        run.text = run.text.replace(original, replacement, 1)
                        self.modified = True
                        return True

                full_text = para.text
                new_text = full_text.replace(original, replacement, 1)
                if para.runs:
                    for run in para.runs:
                        run.text = ""
                    para.runs[0].text = new_text
                    self.modified = True
                    return True

        return False

    def save(self):
        """保存文档"""
        self.doc.save(self.path)


def find_docx_for_resource(resource_id, resource_name, docx_dir):
    """根据资料ID或名称找到对应的 docx 文件"""
    if resource_id in RESOURCE_FILE_MAP:
        path = os.path.join(docx_dir, RESOURCE_FILE_MAP[resource_id])
        if os.path.exists(path):
            return path

    docx_files = [f for f in os.listdir(docx_dir)
                  if f.endswith(".docx") and not f.startswith("~$")]

    # 匹配 [资料ID] 格式，如 [58106924]第52卷...docx
    for f in docx_files:
        if f"[{resource_id}]" in f:
            return os.path.join(docx_dir, f)

    # 匹配文件名中直接包含资料ID
    for f in docx_files:
        if resource_id in f:
            return os.path.join(docx_dir, f)

    # 匹配卷号/套号
    name_keywords = re.findall(r"第\d+卷|第\d+套|\d+卷|\d+套", resource_name)
    for kw in name_keywords:
        for f in docx_files:
            if kw in f:
                return os.path.join(docx_dir, f)

    return None


def is_formula_related(error_type, description):
    """判断是否可能是公式相关误报"""
    keywords = ["内容缺失", "空白", "无内容", "缺少内容", "内容为空",
                "公式", "缺失", "空", "无文本", "图片"]
    combined = f"{error_type} {description}"
    return any(kw in combined for kw in keywords)


def is_missing_content_error(description):
    """判断是否为"缺少参数/内容缺失"类型——可能是图片导致的误判"""
    keywords = ["缺少参数", "缺少内容", "内容缺失", "空白", "无内容",
                "内容为空", "无文本", "参数缺失", "为空", "缺少数据",
                "没有内容", "内容不完整"]
    return any(kw in description for kw in keywords)


def is_text_typo_error(error_type, description):
    """判断是否为文本类错误（错字/少字/多字/表述不当/文本有误等）。
    宽松匹配：只要不是答案错误、不是内容缺失，涉及文本/文字问题的都归入此类，
    交由 AI 读取题干后做独立判断。"""
    specific_kw = ["错字", "少字", "多字", "错别字", "笔误", "错写", "误写",
                   "漏写", "少了", "多了一个", "缺少一个", "写成了", "误用",
                   "写错", "应为", "应该是", "打错", "漏了", "缺一个"]
    broad_kw = ["文本错误", "文字错误", "文本有误", "文字有误", "内容有误",
                "表述不当", "表述有误", "表述错误", "用词不当", "用词错误",
                "语句不通", "不准确", "有误", "错误", "不正确",
                "需要修改", "需修改", "建议修改"]
    combined = f"{error_type} {description}"
    if any(kw in combined for kw in specific_kw):
        return True
    if any(kw in combined for kw in broad_kw):
        return True
    return False


def process_report():
    """主处理逻辑"""
    script_dir = os.environ.get("EXAM_TOOL_WORKDIR")
    if not script_dir or not os.path.isdir(script_dir):
        script_dir = os.path.dirname(os.path.abspath(__file__))
    report_path = os.path.join(script_dir, CONFIG["report_xlsx"])

    if not os.path.exists(report_path):
        report_path = input("请输入报错 Excel 文件路径: ").strip().strip('"')
        if not os.path.exists(report_path):
            print("文件不存在，退出。")
            return

    docx_dir = os.path.dirname(report_path) or script_dir

    print(f"\n报错表: {report_path}")
    print(f"文档目录: {docx_dir}")

    wb = openpyxl.load_workbook(report_path)
    ws = wb.active

    # 自动识别列
    headers = [ws.cell(row=1, column=j).value for j in range(1, ws.max_column + 1)]
    print(f"表格列: {headers}")

    col_map = {}
    for j, h in enumerate(headers):
        if h:
            col_map[h] = j + 1

    col_id = col_map.get("资料ID", col_map.get("试卷序号", 2))
    col_name = col_map.get("资料名称", 3)
    col_type = col_map.get("错误类型", col_map.get("问题类型", 4))
    col_pos = col_map.get("具体位置", 5)
    col_desc = col_map.get("具体描述", col_map.get("问题详细描述", 6))
    col_done = col_map.get("已修改", ws.max_column)

    ai = AIClient()
    docx_cache = {}
    stats = {"total": 0, "skipped_answer": 0, "fixed": 0, "false_positive": 0,
             "uncertain": 0, "no_docx": 0, "image_skip": 0}

    total_rows = ws.max_row - 1
    print(f"\n共 {total_rows} 条报错记录")
    print("处理策略:")
    print("  - 答案错误 → AI独立做题，与具体描述对比验证后修正")
    print("  - 图片题无法处理 → 留给人工")
    print("  - 缺少参数/内容缺失 → 检查该位置有无图片，有则误判打✔，无则人工")
    print("  - 错字/少字/多字 → AI精读描述，提取原文与修正文本，自动替换")
    print("  - 其他错误 → AI综合分析并自动修改 → 打 ✔")
    print("=" * 60)

    for row_idx in range(2, ws.max_row + 1):
        resource_id = str(ws.cell(row=row_idx, column=col_id).value or "")
        resource_name = str(ws.cell(row=row_idx, column=col_name).value or "")
        error_type = str(ws.cell(row=row_idx, column=col_type).value or "")
        position = str(ws.cell(row=row_idx, column=col_pos).value or "")
        description = str(ws.cell(row=row_idx, column=col_desc).value or "")
        done_val = ws.cell(row=row_idx, column=col_done).value

        stats["total"] += 1
        progress = f"[{stats['total']}/{total_rows}]"

        # 跳过已处理的
        if done_val == "✔":
            print(f"{progress} 跳过（已处理）")
            continue

        # ═══════════════ 处理"答案错误" ═══════════════
        if "答案错误" in error_type:
            print(f"\n{progress} ID:{resource_id} | 答案错误 | {position}")
            print(f"  描述: {description[:120]}")

            # ── 步骤1: 解析具体描述，提取答案信息 ──
            ans_info = parse_answer_info(position, description)
            q_type = ans_info["question_type"]
            old_answer = ans_info["old_answer"]
            new_answer_from_desc = ans_info["new_answer"]
            print(f"  解析: 题型={q_type} | 原答案={old_answer or '?'} | 描述建议={new_answer_from_desc or '?'}")

            # 简答题目前无法自动处理
            if q_type == "essay":
                print(f"  — 简答题，留给人工处理")
                stats["skipped_answer"] += 1
                continue

            # ── 步骤2: 找到 docx 并读取题干 ──
            docx_path = find_docx_for_resource(resource_id, resource_name, docx_dir)
            if not docx_path:
                stats["no_docx"] += 1
                print(f"  ⚠ 未找到对应 docx，留给人工处理")
                stats["skipped_answer"] += 1
                continue

            if docx_path not in docx_cache:
                if CONFIG["backup"]:
                    backup = docx_path.replace(".docx", "_备份.docx")
                    if not os.path.exists(backup):
                        shutil.copy2(docx_path, backup)
                docx_cache[docx_path] = DocxProcessor(docx_path)
            processor = docx_cache[docx_path]

            para_idx = processor.find_paragraph_by_position(position, description)
            if para_idx is None:
                print(f"  ⚠ 无法定位题目位置，留给人工处理")
                stats["skipped_answer"] += 1
                continue

            question_text = processor.get_question_context(para_idx)

            # ── 步骤3: 检查题目是否含有无法处理的图片 ──
            if question_has_critical_images(question_text):
                print(f"  — 题干含关键图片信息，AI无法处理，留给人工")
                stats["image_skip"] += 1
                stats["skipped_answer"] += 1
                continue

            # ── 步骤4: AI 独立做题（不给出期望答案） ──
            time.sleep(CONFIG["ai_delay"])
            solve_result = ai.solve_question_independently(question_text, q_type, position)
            ai_answer = str(solve_result.get("my_answer", "")).strip()
            confidence = solve_result.get("confidence", "low")
            reason = solve_result.get("reason", "")
            print(f"  AI独立做题: {ai_answer} (置信度:{confidence}) | {reason}")

            # ── 步骤5: 对比 AI 答案与具体描述 ──
            confirmed = False
            final_new_answer = ""

            if new_answer_from_desc:
                # 有明确的描述建议答案，直接对比
                ai_ans_normalized = ai_answer.upper().strip()
                desc_ans_normalized = new_answer_from_desc.upper().strip()

                if q_type == "true_false":
                    # 判断题：正确/错误/对/错 归一化
                    ai_norm = "正确" if ai_ans_normalized in ("正确", "对", "√", "TRUE", "是") else \
                              "错误" if ai_ans_normalized in ("错误", "错", "×", "FALSE", "否") else ai_ans_normalized
                    desc_norm = "正确" if desc_ans_normalized in ("正确", "对") else \
                                "错误" if desc_ans_normalized in ("错误", "错") else desc_ans_normalized
                    if ai_norm == desc_norm:
                        confirmed = True
                        final_new_answer = new_answer_from_desc
                else:
                    # 选择题：比较字母
                    if ai_ans_normalized and desc_ans_normalized and \
                       ai_ans_normalized[0] == desc_ans_normalized[0]:
                        confirmed = True
                        final_new_answer = new_answer_from_desc

            if not confirmed and confidence in ("high", "medium"):
                # AI 答案与描述不一致，但 AI 有一定把握
                # 做二次确认：让 AI 同时看题干和描述来综合判断
                print(f"  → AI答案与描述建议不一致，进行二次确认...")
                time.sleep(CONFIG["ai_delay"])
                verify_result = ai.verify_answer_with_description(
                    question_text, description, position)
                agree = verify_result.get("agree", False)
                v_answer = str(verify_result.get("my_answer", "")).strip()
                v_reason = verify_result.get("reason", "")
                print(f"  二次确认: agree={agree} | 答案={v_answer} | {v_reason}")

                if agree:
                    confirmed = True
                    # 优先使用描述中的建议答案，其次用二次确认的答案
                    final_new_answer = new_answer_from_desc or v_answer

            if not confirmed:
                # 如果还是无法确认且没有明确建议答案，尝试最后的兜底
                if not new_answer_from_desc and confidence == "high":
                    # AI 高置信度但描述中没有明确答案，尝试用描述做最终验证
                    time.sleep(CONFIG["ai_delay"])
                    verify_result = ai.verify_answer_with_description(
                        question_text, description, position)
                    if verify_result.get("agree", False):
                        confirmed = True
                        final_new_answer = str(verify_result.get("my_answer", "")).strip()
                        print(f"  兜底确认成功: {final_new_answer}")

            # ── 步骤6: 决定是否修改 ──
            if confirmed and final_new_answer:
                if not old_answer:
                    # 尝试从更多模式中提取原答案
                    if q_type == "choice":
                        # 从文档的答案行中尝试读取
                        for search_i in range(max(0, para_idx - 3), min(len(processor.doc.paragraphs), para_idx + 25)):
                            t = processor.doc.paragraphs[search_i].text.strip()
                            m = re.search(r'(?:答案|Answer)[：:\s】\]]*([A-Da-d])', t, re.IGNORECASE)
                            if m:
                                old_answer = m.group(1).upper()
                                break

                if old_answer:
                    success = processor.fix_answer_and_remove_analysis(
                        para_idx, old_answer, final_new_answer, description, q_type)
                    if success:
                        ws.cell(row=row_idx, column=col_done, value="✔")
                        stats["fixed"] += 1
                        print(f"  ✔ 答案已修正: {old_answer} → {final_new_answer}，解析已删除")
                    else:
                        print(f"  ⚠ 在文档中未找到答案位置，留给人工处理")
                        stats["skipped_answer"] += 1
                else:
                    print(f"  ⚠ 无法提取原答案，留给人工处理")
                    stats["skipped_answer"] += 1
            else:
                stats["skipped_answer"] += 1
                print(f"  — AI无法确认答案错误，留给人工处理")

            continue

        # ═══════════════ 处理非答案错误 ═══════════════
        print(f"\n{progress} ID:{resource_id} | {error_type} | {position}")
        print(f"  描述: {description[:120]}")

        # 尝试找到对应 docx
        context_text = ""
        docx_path = find_docx_for_resource(resource_id, resource_name, docx_dir)
        processor = None
        para_idx = None

        if docx_path:
            if docx_path not in docx_cache:
                if CONFIG["backup"]:
                    backup = docx_path.replace(".docx", "_备份.docx")
                    if not os.path.exists(backup):
                        shutil.copy2(docx_path, backup)
                docx_cache[docx_path] = DocxProcessor(docx_path)
            processor = docx_cache[docx_path]
            para_idx = processor.find_paragraph_by_position(position, description)
            if para_idx is not None:
                context_text = processor.get_context(para_idx)
        else:
            stats["no_docx"] += 1
            print(f"  (未找到对应 docx，仅根据描述判断)")

        # ── 分支A: "缺少参数/内容缺失"类 → 检查是否有图片导致误判 ──
        if is_missing_content_error(description) or is_formula_related(error_type, description):
            if processor:
                has_img, _ = processor.check_image_at_position(position, description)
                if has_img:
                    print(f"  ✔ 该位置存在图片/公式 → 质检误判，标记完成")
                    ws.cell(row=row_idx, column=col_done, value="✔")
                    stats["false_positive"] += 1
                    continue
                else:
                    print(f"  — 该位置无图片，描述为内容缺失类，留给人工处理")
                    stats["uncertain"] += 1
                    continue
            else:
                print(f"  — 无docx文件，无法检查图片，留给人工处理")
                stats["uncertain"] += 1
                continue

        # ── 分支B: 文本类错误 → AI阅读题干 + 描述后判断并修正 ──
        if is_text_typo_error(error_type, description):
            print(f"  → 识别为文本类错误，AI将阅读题干内容并判断...")

            # 使用更大的上下文窗口，让AI有足够内容做独立判断
            wide_context = ""
            if processor and para_idx is not None:
                wide_context = processor.get_question_context(para_idx, window=10)
            elif context_text:
                wide_context = context_text

            time.sleep(CONFIG["ai_delay"])
            result = ai.analyze_text_error(description, wide_context, position)

            can_fix = result.get("can_fix", False)
            original = result.get("original", "")
            fix = result.get("fix", "")
            reason = result.get("reason", "")
            print(f"  AI解析: can_fix={can_fix} | 「{original}」→「{fix}」 | {reason}")

            if can_fix and original and fix and original != fix:
                if processor:
                    success = processor.replace_text(para_idx, original, fix)
                    if success:
                        ws.cell(row=row_idx, column=col_done, value="✔")
                        stats["fixed"] += 1
                        print(f"  ✔ 已修改文档: 「{original}」→「{fix}」")
                    else:
                        print(f"  ⚠ 在文档中未找到原文「{original}」，留给人工处理")
                        stats["uncertain"] += 1
                else:
                    stats["uncertain"] += 1
                    print(f"  ⚠ 无docx文件，无法自动修改，留给人工处理")
            elif can_fix is False and reason:
                # AI明确判断不需要修改（可能是误报或图片问题）
                if "误" in reason or "图片" in reason or "公式" in reason:
                    ws.cell(row=row_idx, column=col_done, value="✔")
                    stats["false_positive"] += 1
                    print(f"  ✔ AI判断无需修改: {reason}")
                else:
                    stats["uncertain"] += 1
                    print(f"  — AI无法确定修改内容，留给人工处理")
            else:
                stats["uncertain"] += 1
                print(f"  — AI无法确定修改内容，留给人工处理")
            continue

        # ── 分支C: 其他错误类型 → 通用AI分析 ──
        time.sleep(CONFIG["ai_delay"])
        result = ai.analyze_error(error_type, position, description, context_text)

        judgment = result.get("judgment", "不确定")
        reason = result.get("reason", "")
        print(f"  AI判断: {judgment} | {reason}")

        if judgment == "误报":
            ws.cell(row=row_idx, column=col_done, value="✔")
            stats["false_positive"] += 1
            print(f"  ✔ 误报，标记完成")

        elif judgment == "需修改":
            original = result.get("original", "")
            fix = result.get("fix", "")
            print(f"  修改: 「{original}」→「{fix}」")

            if processor and original and fix:
                success = processor.replace_text(para_idx, original, fix)
                if success:
                    ws.cell(row=row_idx, column=col_done, value="✔")
                    stats["fixed"] += 1
                    print(f"  ✔ 已修改文档")
                else:
                    print(f"  ⚠ 在文档中未找到原文，无法自动修改")
                    stats["uncertain"] += 1
            elif not processor:
                ws.cell(row=row_idx, column=col_done, value="✔")
                stats["fixed"] += 1
                print(f"  ✔ AI确认需修改（无docx，已标记）")
            else:
                stats["uncertain"] += 1
                print(f"  ⚠ AI未给出具体修改内容")

        else:
            stats["uncertain"] += 1
            print(f"  — 不确定，保留待人工确认")

    # 保存修改过的 docx
    for path, proc in docx_cache.items():
        if proc.modified:
            proc.save()
            print(f"\n已保存文档: {path}")

    # 保存更新后的 Excel
    wb.save(report_path)
    print(f"\n已更新报错表: {report_path}")

    # 统计
    print(f"\n{'=' * 60}")
    print(f"处理完成！统计:")
    print(f"  总记录数: {stats['total']}")
    print(f"  AI已修正（答案+其他）: {stats['fixed']}")
    print(f"  质检误判（已标记✔）: {stats['false_positive']}")
    print(f"  答案错误-留给人工: {stats['skipped_answer']}")
    print(f"    其中-图片题无法处理: {stats['image_skip']}")
    print(f"  不确定（待人工）: {stats['uncertain']}")
    print(f"  无对应docx文件: {stats['no_docx']}")
    print(f"{'=' * 60}")


def main():
    print("=" * 60, flush=True)
    print("  AI 自动修正试卷报错工具", flush=True)
    print("=" * 60, flush=True)
    print(flush=True)
    print("规则：", flush=True)
    print("  - 答案错误 → AI重读题干独立做题，与具体描述对比后修正", flush=True)
    print("  - 图片题无法处理 → 留给人工", flush=True)
    print("  - 缺少参数/内容缺失 → 检查图片存在则误判打✔，否则人工", flush=True)
    print("  - 错字/少字/多字 → AI精读描述提取修改，自动替换", flush=True)
    print("  - 其他错误 → AI综合分析并自动修改 → 打 ✔", flush=True)
    print(flush=True)

    process_report()
    input("\n按 Enter 退出...")


if __name__ == "__main__":
    main()
