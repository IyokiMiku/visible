"""Format repaired Inner Mongolia fill-in-the-blank DOCX sections.

This local-only repair adjusts the second section formatting in the two affected
textbooks, regenerates original-version DOCX files, and refreshes ZIP packages.
"""
from __future__ import annotations

import json
import re
import shutil
import sys
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt

BASE_DIR = Path(__file__).resolve().parents[2]
if str(BASE_DIR / "01_工具脚本") not in sys.path:
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本"))

from 生成器.postprocess import _convert_to_blank  # noqa: E402

TARGET_DIRS = [
    BASE_DIR / "04_生成输出" / "生成结果" / "内蒙古自治区 机电类" / "电工电子技术与技能",
    BASE_DIR / "04_生成输出" / "生成结果" / "内蒙古自治区 机电类" / "电气测量技术",
]

SECTION_START_RE = re.compile(r"^二[、.．]\s*填空题")
SECTION_END_RE = re.compile(r"^三[、.．]\s*\S+")
ANSWER_LABELS = ("【答案】", "【解析】")


@dataclass
class FormatResult:
    analysis_path: Path
    status: str
    message: str
    original_path: str = ""
    zip_path: str = ""


def block_text(block) -> str:
    texts = [node.text or "" for node in block.iter() if node.tag.rsplit("}", 1)[-1] == "t"]
    return "".join(texts).strip()


def find_fill_section_bounds(doc: Document) -> tuple[int, int]:
    children = list(doc._body._element)
    start_idx = None
    for i, child in enumerate(children):
        if SECTION_START_RE.match(block_text(child)):
            start_idx = i
            break
    if start_idx is None:
        raise ValueError("找不到二、填空题")

    end_idx = len(children)
    for j in range(start_idx + 1, len(children)):
        if SECTION_END_RE.match(block_text(children[j])):
            end_idx = j
            break
    if end_idx == len(children):
        raise ValueError("找不到三、题型标题作为结束锚点")
    return start_idx, end_idx


def set_run_font(run, font_name: str, size=Pt(10.5), *, red=False, bold=False) -> None:
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if red:
        run.font.color.rgb = RGBColor(255, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def is_times_char(ch: str) -> bool:
    return ch.isascii() and (ch.isalnum() or ch in " .,:;+-=*/()[]{}_Ω%")


def add_mixed_text(paragraph, text: str, *, red=False, bold=False, size=Pt(10.5)) -> None:
    if not text:
        return
    buf = ""
    current = None
    for ch in text:
        kind = "Times New Roman" if is_times_char(ch) else "宋体"
        if current is None:
            current = kind
            buf = ch
        elif kind == current:
            buf += ch
        else:
            run = paragraph.add_run(buf)
            set_run_font(run, current, size, red=red, bold=bold)
            current = kind
            buf = ch
    if buf:
        run = paragraph.add_run(buf)
        set_run_font(run, current or "宋体", size, red=red, bold=bold)


def rebuild_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.line_spacing = 1.5

    if SECTION_START_RE.match(text):
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        set_run_font(run, "黑体", Pt(12), bold=True)
        return

    for label in ANSWER_LABELS:
        if text.startswith(label):
            label_run = paragraph.add_run(label)
            set_run_font(label_run, "黑体", Pt(10.5), red=True, bold=True)
            add_mixed_text(paragraph, text[len(label):], red=True, size=Pt(10.5))
            return

    add_mixed_text(paragraph, text, size=Pt(10.5))


def format_fill_section(analysis_path: Path) -> None:
    doc = Document(str(analysis_path))
    start_idx, end_idx = find_fill_section_bounds(doc)
    children = list(doc._body._element)
    target_elements = set(children[start_idx:end_idx])

    for paragraph in doc.paragraphs:
        if paragraph._element in target_elements:
            rebuild_paragraph(paragraph, paragraph.text.strip())

    doc.save(str(analysis_path))


def original_docx_path(analysis_path: Path) -> Path:
    return analysis_path.with_name(analysis_path.name.replace("解析版", "原卷版"))


def refresh_original_and_zip(analysis_path: Path) -> tuple[str, str]:
    original_path = original_docx_path(analysis_path)
    if original_path.exists():
        original_path.unlink()
    _convert_to_blank(str(analysis_path), str(original_path))

    zip_path = analysis_path.with_name(
        analysis_path.name.replace("（解析版）", "").replace("(解析版)", "").strip() + ".zip"
    )
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(analysis_path, analysis_path.name)
        zf.write(original_path, original_path.name)
    return str(original_path), str(zip_path)


def collect_analysis_docx() -> list[Path]:
    files: list[Path] = []
    for root in TARGET_DIRS:
        files.extend(
            sorted(
                [p for p in root.glob("第*练*解析版*.docx") if p.is_file() and not p.name.startswith("~$")],
                key=lambda p: int(re.match(r"第(\d+)练", p.name).group(1)),
            )
        )
    return files


def copy_backup(path: Path, backup_root: Path) -> None:
    rel = path.relative_to(BASE_DIR)
    dst = backup_root / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, dst)


def write_report(results: list[FormatResult], report_path: Path, backup_root: Path) -> None:
    counts: dict[str, int] = {}
    for result in results:
        counts[result.status] = counts.get(result.status, 0) + 1
    lines = [
        "# 内蒙古填空题格式修复报告",
        "",
        f"- 时间：{time.strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 备份目录：`{backup_root}`",
        f"- 统计：{json.dumps(counts, ensure_ascii=False)}",
        "",
        "| 状态 | 解析版DOCX | 说明 | 原卷版DOCX | ZIP |",
        "|---|---|---|---|---|",
    ]
    for result in results:
        analysis = result.analysis_path.relative_to(BASE_DIR)
        original = Path(result.original_path).relative_to(BASE_DIR) if result.original_path else ""
        zip_path = Path(result.zip_path).relative_to(BASE_DIR) if result.zip_path else ""
        lines.append(f"| {result.status} | `{analysis}` | {result.message} | `{original}` | `{zip_path}` |")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_root = BASE_DIR / "04_生成输出" / "生成结果" / f"_fill_blank_format_backup_{timestamp}"
    report_path = BASE_DIR / "04_生成输出" / "质检报告" / f"内蒙古填空题格式修复报告_{timestamp}.md"
    report_path.parent.mkdir(parents=True, exist_ok=True)

    results: list[FormatResult] = []
    files = collect_analysis_docx()
    print(f"准备处理 {len(files)} 份解析版 DOCX。")
    for idx, analysis_path in enumerate(files, 1):
        print(f"[{idx}/{len(files)}] {analysis_path.relative_to(BASE_DIR)}")
        try:
            original_path = original_docx_path(analysis_path)
            copy_backup(analysis_path, backup_root)
            if original_path.exists():
                copy_backup(original_path, backup_root)
            format_fill_section(analysis_path)
            original, zip_path = refresh_original_and_zip(analysis_path)
            print("  -> ok: 格式修复成功")
            results.append(FormatResult(analysis_path, "ok", "格式修复成功", original, zip_path))
        except Exception as exc:
            print(f"  -> failed: {exc}")
            results.append(FormatResult(analysis_path, "failed", str(exc)))

    write_report(results, report_path, backup_root)
    print(f"报告已写入: {report_path}")
    return 1 if any(r.status == "failed" for r in results) else 0


if __name__ == "__main__":
    raise SystemExit(main())
