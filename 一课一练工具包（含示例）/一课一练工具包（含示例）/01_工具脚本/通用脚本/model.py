"""
题目文档套用模板工具

功能：
  将含标题和题目的 .docx 文档内容套入模板，生成带完整版面的文档。
  - 模板首页保留蓝色虚线框编写说明、标题等（不修改）
  - 自动跳过输入文档开头的标题/学生信息等非题目段落
  - 原文档其余内容（含答案、解析等）原样保留放入模板
  - 各题型标题前自动插入宋体五号单倍行距空行
  - 后续页面自动继承模板背景（来自页眉）

用法：
  python model.py
  弹窗选择模板文件 → 选择一个或多个题目文档 → 自动生成

依赖：
  pip install python-docx
"""

import copy
import re
import os
import sys

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI
except ImportError:
    print("缺少依赖，请先执行: pip install python-docx")
    sys.exit(1)

A_BLIP = qn("a:blip")
R_EMBED = qn("r:embed")
R_ID = qn("r:id")
W_SECTPR = qn("w:sectPr")

NS_VML = "urn:schemas-microsoft-com:vml"
NS_OFFICE = "urn:schemas-microsoft-com:office:office"

HEADER_RE = re.compile(r"^[一二三四五六七八九十]+\s*[、.．]")
QNUM_RE = re.compile(r"^(\d+)\s*[.、．)\）]")


# ════════════════════ 段落筛选 ════════════════════


def _classify_paragraphs(doc):
    """
    分析文档段落，返回应保留的段落索引列表。
    跳过：开头非题目段落（标题/学生信息等）、末尾空段落。
    保留：从第一个题型标题或题号开始的所有内容（原样保留）。
    """
    paras = doc.paragraphs
    total = len(paras)

    first_header = None
    for i, p in enumerate(paras):
        if HEADER_RE.match(p.text.strip()):
            first_header = i
            break

    if first_header is None:
        for i, p in enumerate(paras):
            if QNUM_RE.match(p.text.strip()):
                first_header = i
                break

    if first_header is None:
        first_header = 0

    keep = list(range(first_header, total))

    while keep and not paras[keep[-1]].text.strip():
        keep.pop()

    return keep


# ════════════════════ 图片/OLE 复制 ════════════════════


def _make_unique_partname(dst_package, original_partname):
    """为目标包生成不重复的 partname"""
    name = str(original_partname)
    base, ext = os.path.splitext(name)

    existing = set()
    for part in dst_package.iter_parts():
        existing.add(str(part.partname))

    if name not in existing:
        return PackURI(name)

    counter = 1
    while True:
        candidate = f"{base}_c{counter}{ext}"
        if candidate not in existing:
            return PackURI(candidate)
        counter += 1


def _remap_images(element, src_doc, dst_doc, rid_cache):
    """
    将段落 XML 中引用的图片/OLE 关系从 src 复制到 dst，更新 rId。
    使用唯一 partname 避免同名文件覆盖。rid_cache 缓存已复制的映射。
    """
    src_part = src_doc.part
    dst_part = dst_doc.part
    dst_package = dst_part.package

    refs = []
    for blip in element.findall(".//" + A_BLIP):
        old_rid = blip.get(R_EMBED)
        if old_rid:
            refs.append((blip, R_EMBED, old_rid))
    for imgdata in element.findall(".//{" + NS_VML + "}imagedata"):
        old_rid = imgdata.get(R_ID)
        if old_rid:
            refs.append((imgdata, R_ID, old_rid))
    for ole in element.findall(".//{" + NS_OFFICE + "}OLEObject"):
        old_rid = ole.get(R_ID)
        if old_rid:
            refs.append((ole, R_ID, old_rid))

    for elem, attr, old_rid in refs:
        if old_rid in rid_cache:
            elem.set(attr, rid_cache[old_rid])
            continue

        if old_rid not in src_part.rels:
            continue

        rel = src_part.rels[old_rid]
        src_target = rel.target_part

        unique_name = _make_unique_partname(dst_package, src_target.partname)
        new_part = OpcPart(
            unique_name,
            src_target.content_type,
            src_target.blob,
            dst_package,
        )
        new_rid = dst_part.relate_to(new_part, rel.reltype)

        rid_cache[old_rid] = new_rid
        elem.set(attr, new_rid)


# ════════════════════ 空行生成 ════════════════════


def _make_spacer_para():
    """生成宋体五号单倍行距的空段落 XML 元素"""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "21")
    rPr.append(szCs)
    pPr.append(rPr)

    p.append(pPr)
    return p


# ════════════════════ 模板处理 ════════════════════


def _remove_trailing_empty(doc):
    """移除文档末尾的空段落"""
    paras = list(doc.paragraphs)
    for p in reversed(paras):
        if p.text.strip():
            break
        p._element.getparent().remove(p._element)


def merge(template_path, input_path, output_path):
    """将输入文档的题目内容套入模板并保存"""
    tpl_doc = Document(template_path)
    src_doc = Document(input_path)

    _remove_trailing_empty(tpl_doc)

    keep_indices = _classify_paragraphs(src_doc)
    paras = src_doc.paragraphs

    body = tpl_doc.element.body
    sect_pr = body.find(W_SECTPR)

    rid_cache = {}
    kept = 0

    for i in keep_indices:
        p = paras[i]

        if HEADER_RE.match(p.text.strip()):
            spacer = _make_spacer_para()
            if sect_pr is not None:
                sect_pr.addprevious(spacer)
            else:
                body.append(spacer)

        new_p = copy.deepcopy(p._element)
        _remap_images(new_p, src_doc, tpl_doc, rid_cache)

        if sect_pr is not None:
            sect_pr.addprevious(new_p)
        else:
            body.append(new_p)
        kept += 1

    tpl_doc.save(output_path)
    return kept


# ════════════════════ 弹窗 ════════════════════


def ask_template():
    """弹窗选择模板文件（单选）"""
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    path = filedialog.askopenfilename(
        title="第1步：选择模板文件",
        filetypes=[("Word 文档", "*.docx")],
    )
    root.destroy()

    if not path:
        print("未选择模板，程序退出。")
        sys.exit(0)
    return path


def ask_inputs():
    """弹窗选择题目文档（可多选）"""
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    paths = filedialog.askopenfilenames(
        title="第2步：选择题目文档（可多选）",
        filetypes=[("Word 文档", "*.docx")],
    )
    root.destroy()

    if not paths:
        print("未选择题目文档，程序退出。")
        sys.exit(0)
    return list(paths)


# ════════════════════ 主程序 ════════════════════


def main():
    out_dir = os.environ.get("EXAM_TOOL_WORKDIR")
    if not out_dir or not os.path.isdir(out_dir):
        if getattr(sys, "frozen", False):
            out_dir = os.path.dirname(sys.executable)
        else:
            out_dir = os.path.dirname(os.path.abspath(__file__)) or "."

    template_path = ask_template()
    print(f"模板: {os.path.basename(template_path)}")

    input_paths = ask_inputs()
    print(f"题目文档: {len(input_paths)} 个\n")

    for src in input_paths:
        name = os.path.basename(src)
        base, ext = os.path.splitext(name)

        out_name = f"{base}（套模板）{ext}"
        dst = os.path.join(os.path.dirname(src), out_name)

        print(f"处理: {name}")
        try:
            kept = merge(template_path, src, dst)
            print(f"  -> {out_name}  ({kept} 个段落)")
        except Exception as e:
            print(f"  失败: {e}")

    print(f"\n全部完成！共处理 {len(input_paths)} 个文件。")
    input("按 Enter 退出...")


if __name__ == "__main__":
    main()
