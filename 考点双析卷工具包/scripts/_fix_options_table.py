"""一键修复：选项对齐（制表符→表格 + 双列→智能单列）。

两步合一：
  第一步  将制表符分隔的选项段落替换为无边框表格（智能选单列/双列）
  第二步  扫描已有双列选项表格，超阈值或不平衡的切换为单列

输入格式：A. 选项A\\t\\tB. 选项B\\nC. 选项C\\t\\tD. 选项D
输出格式：题干段落 + 无边框选项表格（短→双列，长→单列，均对齐）

适用范围：所有一课一练目录下的（解析版）.docx 文件
"""

import os
import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
#  中英混排工具（内嵌，无需外部依赖）
# ============================================================
LATIN_FONT = 'Times New Roman'

_LATIN_PATTERN = re.compile(
    r'[a-zA-Z0-9\.\,\;\:\!\?\-\+\=\/\*\(\)\[\]\{\}\<\>\~\@\#\$\%\^\&\|\\\''
    r'\u03c0\u03b1-\u03c9\u0391-\u03a9\u2126\u00b0\u00d7\u00b7\u221a\u2264'
    r'\u2265\u2248\u2260\u221e\u03bc\u2103]+'
)

def _split_cjk_latin(text):
    """将文本拆分为中文/拉丁文交替的片段列表"""
    tokens = []
    current = ''
    current_type = None
    for ch in text:
        ch_type = 'latin' if _LATIN_PATTERN.match(ch) else 'cjk'
        if current_type is None:
            current_type, current = ch_type, ch
        elif ch_type == current_type:
            current += ch
        else:
            tokens.append((current, current_type))
            current, current_type = ch, ch_type
    if current:
        tokens.append((current, current_type))
    return tokens


# ============================================================
#  配置常量
# ============================================================
_SINGLE_COL_THRESHOLD = 16      # 任一选项超此字数 → 单列（半页宽≈16汉字不换行）
_ROW_BALANCE_RATIO = 1.6         # 同排两选项长度比超此值 → 单列（避免短选项下方大片空白）


# ============================================================
#  辅助函数
# ============================================================

def _should_use_single_col(options_texts):
    """判断是否应使用单列布局。

    返回 True 的条件（任一满足）：
      1. 任一选项长度 > _SINGLE_COL_THRESHOLD
      2. 同排两个选项长度悬殊（比例 > _ROW_BALANCE_RATIO）
    """
    n = len(options_texts)
    if n <= 1:
        return True
    if max(len(o) for o in options_texts) > _SINGLE_COL_THRESHOLD:
        return True
    for i in range(0, n - 1, 2):
        a = len(options_texts[i])
        b = len(options_texts[i + 1])
        if max(a, b) / max(min(a, b), 1) > _ROW_BALANCE_RATIO:
            return True
    return False


def _is_option_paragraph(text):
    """判断段落文本是否包含制表符分隔的选项"""
    return '\t' in text and bool(re.search(r'[A-D][\.\．\、]', text))


def _is_option_table(table):
    """判断 Table 对象是否选项表格（首行首列以 A. 或 A．开头）"""
    try:
        if not table.rows or not table.rows[0].cells:
            return False
        c0 = table.rows[0].cells[0].text.strip()
        return c0.startswith('A.') or c0.startswith('A．')
    except:
        return False


def _parse_question_text(text):
    """解析题目文本 → (stem, [option1, option2, ...])"""
    lines = text.split('\n')
    stem = lines[0].strip()
    options = []
    for line in lines[1:]:
        for p in re.split(r'\t+', line):
            p = p.strip()
            if p:
                options.append(p)
    return stem, options


# ============================================================
#  选项表格构建（lxml级别，单列/双列通用）
# ============================================================

def _build_options_table(options, font_name='宋体', font_size=10.5):
    """构建选项表格 lxml 元素。自动选择单列/双列布局。

    Returns:
        (tbl_element, rows, cols) 或 (None, 0, 0) 若无选项
    """
    if not options:
        return None, 0, 0

    n = len(options)
    if _should_use_single_col(options):
        rows, cols = n, 1
    elif n <= 2:
        rows, cols = 1, n
    else:
        rows = (n + 1) // 2
        cols = 2

    col_w = '9000' if cols == 1 else '4500'

    tbl = OxmlElement('w:tbl')

    # -- tblPr: 100%宽 + 无边框 --
    tblPr = OxmlElement('w:tblPr')
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), '5000')
    tw.set(qn('w:type'), 'pct')
    tblPr.append(tw)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), 'auto')
        borders.append(b)
    tblPr.append(borders)
    tbl.append(tblPr)

    # -- tblGrid --
    tg = OxmlElement('w:tblGrid')
    for _ in range(cols):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), col_w)
        tg.append(gc)
    tbl.append(tg)

    # -- 填充行 --
    for r in range(rows):
        tr = OxmlElement('w:tr')
        for c in range(cols):
            idx = r * cols + c
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), col_w); tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW); tc.append(tcPr)

            wp = _make_paragraph('')
            if idx < len(options):
                _fill_mixed_text(wp, options[idx], font_name, font_size)
            tc.append(wp); tr.append(tc)
        tbl.append(tr)

    return tbl, rows, cols


def _make_paragraph(text='', font_name='宋体', font_size=10.5):
    """创建带1.5倍行距的段落 lxml 元素"""
    wp = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:after'), '0'); sp.set(qn('w:before'), '0')
    sp.set(qn('w:line'), '360'); sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp); wp.append(pPr)
    if text:
        _fill_mixed_text(wp, text, font_name, font_size)
    return wp


def _fill_mixed_text(paragraph_elem, text, font_name='宋体', font_size=10.5):
    """向段落元素添加中英混排文本 runs"""
    for seg, seg_type in _split_cjk_latin(text):
        wr = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        font = LATIN_FONT if seg_type == 'latin' else font_name
        rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font)
        rf.set(qn('w:eastAsia'), font)
        rPr.append(rf)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz); wr.append(rPr)
        wt = OxmlElement('w:t')
        wt.set(qn('xml:space'), 'preserve'); wt.text = seg
        wr.append(wt); paragraph_elem.append(wr)


# ============================================================
#  核心修复逻辑（两步合一）
# ============================================================

def fix_file(filepath):
    """修复单个文件：
      第1步：制表符段落 → 选项表格
      第2步：已有双列选项表格 → 单列（如需要）
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        # 跳过无法打开的文件
        return 0, 0

    body = doc.element.body
    step1 = 0   # 制表符段落修复数
    step2 = 0   # 双列→单列切换数
    need_save = False

    # ---- 第1步：制表符段落 → 表格 ----
    para_items = []
    for para in doc.paragraphs:
        text = para.text
        if not _is_option_paragraph(text):
            continue
        elem = para._element
        for j, child in enumerate(list(body)):
            if child is elem:
                para_items.append((j, text, elem))
                break

    for idx, text, para_elem in reversed(para_items):
        stem, options = _parse_question_text(text)
        if not options:
            continue
        body.remove(para_elem)
        # 题干段落
        body.insert(idx, _make_paragraph(stem))
        # 选项表格
        tbl, _, _ = _build_options_table(options)
        if tbl is not None:
            body.insert(idx + 1, tbl)
        step1 += 1
        need_save = True

    # ---- 第2步：已有双列选项表格 → 智能切换单列 ----
    tbl_items = []
    for i, child in enumerate(list(body)):
        if child.tag != qn('w:tbl'):
            continue
        tbl_idx = sum(1 for c in list(body)[:i] if c.tag == qn('w:tbl'))
        if tbl_idx >= len(doc.tables):
            continue
        tbl = doc.tables[tbl_idx]
        if not _is_option_table(tbl):
            continue
        n_cols = len(tbl.rows[0].cells) if tbl.rows else 0
        if n_cols != 2:
            continue

        # 收集现有文本
        texts = []
        for row in tbl.rows:
            for cell in row.cells:
                texts.append(cell.text.strip())

        if _should_use_single_col(texts):
            tbl_items.append((i, child, texts))

    for idx, elem, texts in reversed(tbl_items):
        body.remove(elem)
        new_tbl, _, _ = _build_options_table(texts)
        if new_tbl is not None:
            body.insert(idx, new_tbl)
        step2 += 1
        need_save = True

    if need_save:
        doc.save(filepath)

    return step1, step2


# ============================================================
#  主入口
# ============================================================

if __name__ == '__main__':
    base = os.path.dirname(__file__)
    total_s1 = total_s2 = files_s1 = files_s2 = 0

    for root, dirs, files in os.walk(base):
        for fn in files:
            if '（解析版）.docx' not in fn or fn.startswith('~'):
                continue
            fp = os.path.join(root, fn)
            s1, s2 = fix_file(fp)
            if s1 or s2:
                parts = []
                if s1: parts.append(f'段落→表格 ×{s1}')
                if s2: parts.append(f'双列→单列 ×{s2}')
                print(f"  ✅ {os.path.basename(fp)}: {'  '.join(parts)}")
                total_s1 += s1; total_s2 += s2
                if s1: files_s1 += 1
                if s2: files_s2 += 1

    print(f"\n{'='*60}")
    if total_s1:
        print(f"  第1步：{files_s1} 个文件，{total_s1} 处段落→表格")
    if total_s2:
        print(f"  第2步：{files_s2} 个文件，{total_s2} 处双列→单列")
    if not total_s1 and not total_s2:
        print(f"  全部文件已是最优格式，无需修复 ✓")
    print(f"{'='*60}")
