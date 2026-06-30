"""
题目文档批量处理工具

用法:
    python process_exam.py

流程:
    1. 自动读取同目录下的 xlsx 规划表
    2. 从最后一列未标绿的首行开始处理
    3. 弹窗询问该行对应的原文档（docx）
    4. 弹窗输入各题型每小题分值（校验总分100）
    5. 生成 4 个文件，成功后该行标绿，继续下一行
    6. 失败则停止并报错

依赖:
    pip install python-docx openpyxl
"""

import re
import os
import sys
import shutil

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖，请先执行: pip install python-docx openpyxl")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import PatternFill
except ImportError:
    print("缺少依赖，请先执行: pip install openpyxl")
    sys.exit(1)

# ════════════════════ XML tag 常量 ════════════════════

W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_PPR = qn("w:pPr")
W_COLOR = qn("w:color")
W_VAL = qn("w:val")
W_SHD = qn("w:shd")
W_HL = qn("w:highlight")

GREEN_FILL = PatternFill(start_color="92D050", end_color="92D050", fill_type="solid")
RED_FILL = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")

# ════════════════════ 文本识别 ════════════════════

HEADER_RE = re.compile(r"^[一二三四五六七八九十]+\s*[、.．]")
QNUM_RE = re.compile(r"^(\d+)\s*[.、．)\）]")
LUE_RE = re.compile(r"【(?:详解|解析)】\s*略")


def is_header(text):
    return bool(text and HEADER_RE.match(text))


def get_qnum(text):
    m = QNUM_RE.match(text or "")
    return int(m.group(1)) if m else None


# ════════════════════ 文档结构分析 ════════════════════


def analyze(doc):
    out = []
    sec, cur_q, in_ans = -1, None, False

    for p in doc.paragraphs:
        t = p.text.strip()

        if is_header(t):
            sec += 1
            cur_q = None
            in_ans = False
            out.append(dict(tp="hdr", q=None, s=sec, lue=False))
            continue

        n = get_qnum(t)
        if n is not None and in_ans and cur_q is not None and n <= cur_q:
            n = None

        if n is not None:
            cur_q = n
            in_ans = False
            if "【答案】" in t:
                in_ans = True
                out.append(dict(tp="mix", q=cur_q, s=sec, lue=False))
            else:
                out.append(dict(tp="body", q=cur_q, s=sec, lue=False))
            continue

        if "【答案】" in t:
            in_ans = True
        if in_ans or "【详解】" in t or "【解析】" in t:
            in_ans = True
            out.append(
                dict(tp="ans", q=cur_q, s=sec, lue=bool(LUE_RE.search(t)))
            )
            continue

        if cur_q is not None:
            out.append(dict(tp="body", q=cur_q, s=sec, lue=False))
        else:
            out.append(
                dict(tp="sp" if sec >= 0 else "pre", q=None, s=sec, lue=False)
            )

    return out


# ════════════════════ 格式操作 ════════════════════


def _make_red(para):
    for r in para._element.findall(W_R):
        rPr = r.find(W_RPR)
        if rPr is None:
            rPr = r.makeelement(W_RPR, {})
            r.insert(0, rPr)
        c = rPr.find(W_COLOR)
        if c is None:
            c = rPr.makeelement(W_COLOR, {})
            rPr.append(c)
        c.set(W_VAL, "FF0000")


def _remove_shading(para):
    pPr = para._element.find(W_PPR)
    if pPr is not None:
        for e in pPr.findall(W_SHD):
            pPr.remove(e)
    for r in para._element.findall(W_R):
        rPr = r.find(W_RPR)
        if rPr is None:
            continue
        for tag in (W_SHD, W_HL):
            for e in rPr.findall(tag):
                rPr.remove(e)


def _trim_answer(para):
    idx = para.text.find("【答案】")
    if idx < 0:
        return
    pos = 0
    for run in para.runs:
        rt = run.text or ""
        rs = pos
        re_ = pos + len(rt)
        if rs >= idx:
            run.text = ""
        elif re_ > idx:
            run.text = rt[: idx - rs]
        pos = re_


def _renumber_para(para, old_num, new_num):
    if old_num == new_num:
        return
    m = re.match(r"^(\d+)", para.text)
    if not m or int(m.group(1)) != old_num:
        return

    num_end = m.end()
    new_str = str(new_num)
    pos = 0
    first = True

    for run in para.runs:
        rt = run.text or ""
        rs, re_ = pos, pos + len(rt)
        if re_ <= num_end:
            run.text = new_str if first else ""
            first = False
        elif rs < num_end:
            cut = num_end - rs
            run.text = (new_str if first else "") + rt[cut:]
            first = False
        pos = re_


def _make_blank_para():
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "21")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "21")
    rPr.append(szCs)

    pPr.append(rPr)
    p.append(pPr)
    return p


def _set_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


# ════════════════════ 编写说明修改 ════════════════════


def _replace_in_runs(para, old, new):
    """跨 run 查找替换（文本可能被拆分到多个 run 中），保留各 run 格式"""
    full = para.text
    idx = full.find(old)
    if idx < 0:
        return False
    end = idx + len(old)

    pos = 0
    first = True
    for run in para.runs:
        rt = run.text or ""
        rs, re_ = pos, pos + len(rt)

        if re_ <= idx or rs >= end:
            pass
        elif rs >= idx and re_ <= end:
            run.text = new if first else ""
            first = False
        elif rs < idx and re_ > end:
            run.text = rt[:idx - rs] + new + rt[end - rs:]
            first = False
        elif rs < idx:
            run.text = rt[:idx - rs] + (new if first else "")
            first = False
        elif re_ > end:
            run.text = (new if first else "") + rt[end - rs:]
            first = False

        pos = re_
    return True


def _replace_after_marker(para, marker, new_text):
    """将段落中 marker 之后的所有文本替换为 new_text，保留各 run 格式"""
    full = para.text
    idx = full.find(marker)
    if idx < 0:
        return
    cut_pos = idx + len(marker)
    pos = 0
    first_after = True
    for run in para.runs:
        rt = run.text or ""
        rs, re_ = pos, pos + len(rt)
        if re_ <= cut_pos:
            pass
        elif rs >= cut_pos:
            run.text = new_text if first_after else ""
            first_after = False
        else:
            keep = cut_pos - rs
            run.text = rt[:keep] + (new_text if first_after else "")
            first_after = False
        pos = re_


def _replace_before_marker(para, marker, new_text):
    """将段落中 marker 前面的名称文字替换为 new_text，空格原样保留"""
    full = para.text
    idx = full.find(marker)
    if idx < 0:
        return
    # 找到名称结束位置（空格之前）
    name_end = idx
    while name_end > 0 and full[name_end - 1] in (" ", "\u3000"):
        name_end -= 1
    # 只替换 0 ~ name_end 的文字部分，name_end ~ idx 的空格保留
    _replace_in_runs(para, full[:name_end], new_text)


def _modify_preamble(doc, vol_num, exam_name, paper_type):
    """
    修改编写说明（虚线框）+ 标题（红圈），保留原始 run 格式：
    编写说明：
      1. "本专辑第X卷" → 替换卷号数字
      2. "掌握内容：" 后 → 替换为试卷名称
    标题：
      3. "考点双析卷 第X卷" → 替换卷号数字
      4. 试卷名称行前半部分 → 替换为试卷名称
      5. "教师讲解卷"/"学生练习卷" → 对应类型
    """
    suffix = "教师讲解卷" if paper_type == "教师" else "学生练习卷"

    for p in doc.paragraphs:
        text = p.text

        # 编写说明："本专辑第X卷……精准对标"
        if "本专辑第" in text and "精准对标" in text:
            m = re.search(r"第(\d+)卷", text)
            if m:
                _replace_in_runs(p, f"第{m.group(1)}卷", f"第{vol_num}卷")

        # 编写说明："掌握内容：XXX"
        elif "掌握内容" in text and ("：" in text or ":" in text):
            marker = "掌握内容：" if "：" in text else "掌握内容:"
            _replace_after_marker(p, marker, exam_name)

        # 标题："考点双析卷 第X卷"
        elif "考点双析卷" in text and re.search(r"第\d+卷", text):
            m = re.search(r"第(\d+)卷", text)
            if m:
                _replace_in_runs(p, f"第{m.group(1)}卷", f"第{vol_num}卷")

        # 标题最后一行："XXX    教师讲解卷/学生练习卷"
        elif "讲解卷" in text or "练习卷" in text:
            old_suffix = "教师讲解卷" if "讲解卷" in text else "学生练习卷"
            # 替换试卷名称（old_suffix 之前的部分）
            _replace_before_marker(p, old_suffix, exam_name)
            # 替换卷类型
            _replace_in_runs(p, old_suffix, suffix)


# ════════════════════ 生成文档 ════════════════════


def build(src, dst, meta, qf, keep_ans, red,
          sec_scores=None, vol_num=None, exam_name=None, paper_type=None):
    sec_scores = sec_scores or {}
    shutil.copy2(src, dst)
    doc = Document(dst)
    paras = list(doc.paragraphs)

    if len(paras) != len(meta):
        raise RuntimeError(
            f"段落数不一致（{len(paras)} vs {len(meta)}）")

    active = {
        m["s"]
        for m in meta
        if m["q"] is not None and qf(m["q"]) and m["s"] >= 0
    }

    rm = []

    for i, (p, m) in enumerate(zip(paras, meta)):
        tp = m["tp"]
        keep = True

        if tp in ("hdr", "sp"):
            keep = m["s"] in active
        elif tp == "body":
            keep = m["q"] is not None and qf(m["q"])
        elif tp == "ans":
            keep = keep_ans and not m["lue"] and m["q"] is not None and qf(m["q"])
            if keep and red:
                _make_red(p)
                _remove_shading(p)
        elif tp == "mix":
            if m["q"] is None or not qf(m["q"]):
                keep = False
            elif not keep_ans:
                _trim_answer(p)
            elif red:
                _make_red(p)
                _remove_shading(p)

        if not keep:
            rm.append(i)

    rm_set = set(rm)
    kept_qs = sorted({
        m["q"] for i, m in enumerate(meta)
        if m["q"] is not None and qf(m["q"])
    })
    renum = {old: new for new, old in enumerate(kept_qs, 1)}

    for i, (p, m) in enumerate(zip(paras, meta)):
        if i in rm_set or m["q"] is None or m["q"] not in renum:
            continue
        if get_qnum(p.text.strip()) == m["q"]:
            _renumber_para(p, m["q"], renum[m["q"]])

    hdr_items = []
    for i, (p, m) in enumerate(zip(paras, meta)):
        if m["tp"] == "hdr" and m["s"] in active and i not in rm_set:
            sec = m["s"]
            count = len({
                m2["q"] for m2 in meta
                if m2["q"] is not None and qf(m2["q"]) and m2["s"] == sec
            })
            score = sec_scores.get(sec)
            hdr_items.append((p, count, score))

    for i in reversed(rm):
        paras[i]._element.getparent().remove(paras[i]._element)

    for idx, (p, count, score) in enumerate(hdr_items):
        if idx > 0:
            p._element.addprevious(_make_blank_para())

        hdr_text = p.text.strip()
        if "本大题共" in hdr_text:
            continue

        fmt_bold, fmt_size, fmt_color = None, None, None
        if p.runs:
            r0 = p.runs[0]
            fmt_bold = r0.bold
            fmt_size = r0.font.size
            if r0.font.color and r0.font.color.rgb:
                fmt_color = r0.font.color.rgb

        total = count * score
        info = f"（本大题共{count}小题，每小题{score}分，共{total}分）"
        run = p.add_run(info)
        run.bold = fmt_bold
        run.font.size = fmt_size
        if fmt_color:
            run.font.color.rgb = fmt_color

    if vol_num is not None:
        _modify_preamble(doc, vol_num, exam_name, paper_type)

    # 全文档去除所有文字底纹
    for p in doc.paragraphs:
        _remove_shading(p)

    doc.save(dst)


# ════════════════════ Excel 操作 ════════════════════


def _is_green(cell):
    """检查单元格是否已标绿"""
    fill = cell.fill
    if fill.patternType != "solid":
        return False
    fg = fill.fgColor
    if fg is None:
        return False
    # fgColor 可能是 theme/indexed 类型而非 rgb
    rgb_val = fg.rgb
    if rgb_val is None or str(rgb_val) == "00000000":
        return False
    rgb = str(rgb_val).upper()
    # openpyxl 返回的 rgb 可能是 8 位（含 alpha）如 "0092D050"
    # 或 6 位如 "92D050"，统一取后 6 位
    if len(rgb) == 8:
        hex_rgb = rgb[2:]  # 去掉 alpha 前缀
    elif len(rgb) == 6:
        hex_rgb = rgb
    else:
        return False
    try:
        r = int(hex_rgb[0:2], 16)
        g = int(hex_rgb[2:4], 16)
        b = int(hex_rgb[4:6], 16)
        return g > 150 and g > r and g > b
    except ValueError:
        return False


def _find_header_row(ws):
    """找到含'序号'的表头行号"""
    for row_idx in range(1, ws.max_row + 1):
        val = ws.cell(row=row_idx, column=1).value
        if val is not None and str(val).strip() == "序号":
            return row_idx
    return None


def _find_mark_col(ws, header_row):
    """找到标记列：表头行中最后一个有内容的列的下一列，或'卷数'列的下一列"""
    last_content_col = 1
    for col in range(1, ws.max_column + 1):
        val = ws.cell(row=header_row, column=col).value
        if val is not None and str(val).strip():
            last_content_col = col
    # 标记列 = 有内容的最后列 + 1
    mark_col = last_content_col + 1
    # 但如果该列已经有绿色/红色标记，就用它
    # 也检查 last_content_col 本身是否就是标记列（无表头文字但有颜色）
    for col in [last_content_col + 1, last_content_col]:
        for row in range(header_row + 1, min(header_row + 5, ws.max_row + 1)):
            cell = ws.cell(row=row, column=col)
            if cell.fill.patternType == "solid":
                fg = cell.fill.fgColor
                if fg and fg.rgb and str(fg.rgb) != "00000000":
                    return col
    return mark_col


def _get_pending_rows(ws, header_row):
    """
    返回待处理的行列表：[(excel行号, 序号, 试卷名称), ...]
    跳过标记列已标绿的行。
    """
    last_col = _find_mark_col(ws, header_row)
    rows = []
    for row_idx in range(header_row + 1, ws.max_row + 1):
        seq_cell = ws.cell(row=row_idx, column=1)
        if seq_cell.value is None or not str(seq_cell.value).strip().isdigit():
            continue
        mark_cell = ws.cell(row=row_idx, column=last_col)
        if _is_green(mark_cell):
            continue
        seq = int(seq_cell.value)
        name_val = ws.cell(row=row_idx, column=3).value
        exam_name = str(name_val).strip() if name_val else f"序号{seq}"
        rows.append((row_idx, seq, exam_name))
    return rows


# ════════════════════ 弹窗 ════════════════════


def _ask_xlsx_file(out_dir):
    """弹窗让用户选择 xlsx 规划表文件"""
    import tkinter as tk
    from tkinter import messagebox

    xlsx_files = sorted([
        f for f in os.listdir(out_dir)
        if f.endswith(".xlsx") and not f.startswith("~")
    ])
    if not xlsx_files:
        print("脚本目录下没有 .xlsx 文件")
        sys.exit(1)

    result = {}

    def submit():
        sel = listbox.curselection()
        if not sel:
            messagebox.showwarning("提示", "请选择一个文件")
            return
        result["file"] = xlsx_files[sel[0]]
        root.destroy()

    root = tk.Tk()
    root.title("选择规划表")
    root.resizable(False, False)

    tk.Label(root, text="请选择规划表文件：",
             font=("宋体", 11)).pack(padx=20, pady=(15, 5))

    frame = tk.Frame(root)
    frame.pack(padx=20, pady=5)

    scrollbar = tk.Scrollbar(frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    listbox = tk.Listbox(frame, width=55, height=min(len(xlsx_files), 10),
                         font=("宋体", 10), yscrollcommand=scrollbar.set)
    for f in xlsx_files:
        listbox.insert(tk.END, f"  {f}")
    listbox.pack(side=tk.LEFT)
    scrollbar.config(command=listbox.yview)

    if xlsx_files:
        listbox.selection_set(0)

    tk.Button(root, text="确 定", font=("宋体", 11), width=10,
              command=submit).pack(pady=(5, 15))

    root.eval('tk::PlaceWindow . center')
    root.mainloop()

    if "file" not in result:
        print("用户取消，程序退出。")
        sys.exit(0)

    return os.path.join(out_dir, result["file"]), result["file"]


def _ask_input_file(out_dir, exam_name):
    """弹窗让用户选择当前行对应的原文档"""
    import tkinter as tk
    from tkinter import messagebox

    docx_files = sorted([
        f for f in os.listdir(out_dir)
        if f.endswith(".docx") and not f.startswith("~")
    ])
    if not docx_files:
        print("脚本目录下没有 .docx 文件")
        sys.exit(1)

    result = {}

    def submit():
        sel = listbox.curselection()
        if not sel:
            messagebox.showwarning("提示", "请选择一个文件")
            return
        result["file"] = docx_files[sel[0]]
        root.destroy()

    root = tk.Tk()
    root.title("选择原文档")
    root.resizable(False, False)

    tk.Label(root, text=f"正在处理: {exam_name}\n请选择对应的原文档：",
             font=("宋体", 11), justify="left").pack(padx=20, pady=(15, 5))

    frame = tk.Frame(root)
    frame.pack(padx=20, pady=5)

    scrollbar = tk.Scrollbar(frame)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    listbox = tk.Listbox(frame, width=55, height=min(len(docx_files), 12),
                         font=("宋体", 10), yscrollcommand=scrollbar.set)
    for f in docx_files:
        listbox.insert(tk.END, f"  {f}")
    listbox.pack(side=tk.LEFT)
    scrollbar.config(command=listbox.yview)

    if docx_files:
        listbox.selection_set(0)

    tk.Button(root, text="确 定", font=("宋体", 11), width=10,
              command=submit).pack(pady=(5, 15))

    root.eval('tk::PlaceWindow . center')
    root.mainloop()

    if "file" not in result:
        print("用户取消，程序退出。")
        sys.exit(0)

    return os.path.join(out_dir, result["file"])


def _ask_name_suffix():
    """弹窗询问输出文件名中试卷名称后面的后缀文本"""
    import tkinter as tk

    result = {}

    def submit():
        val = entry.get().strip()
        if not val:
            from tkinter import messagebox
            messagebox.showerror("输入错误", "后缀文本不能为空，请输入内容。")
            return
        result["suffix"] = val
        root.destroy()

    root = tk.Tk()
    root.title("设置输出文件名后缀")
    root.resizable(False, False)

    tk.Label(root, text="请输入输出文件名中的后缀文本：",
             font=("宋体", 11)).pack(padx=20, pady=(15, 5))
    tk.Label(root, text="格式: 第?卷 试卷名称 【此处文本】（解析版/原卷版）.docx",
             font=("宋体", 9), fg="gray").pack(padx=20)

    entry = tk.Entry(root, width=50, font=("宋体", 10))
    entry.pack(padx=20, pady=10)

    tk.Button(root, text="确 定", font=("宋体", 11), width=10,
              command=submit).pack(pady=(0, 15))

    entry.focus_set()
    root.eval('tk::PlaceWindow . center')
    root.mainloop()

    if "suffix" not in result:
        print("未输入后缀文本，程序退出。")
        sys.exit(1)

    return result["suffix"]


def _ask_scores(sec_info):
    """弹窗输入每题型分值，校验总分100。返回 {分区序号: 分值}"""
    import tkinter as tk
    from tkinter import messagebox

    result = {}
    secs = sorted(sec_info.keys())

    def submit():
        scores = {}
        total = 0
        for sec in secs:
            val = entries[sec].get().strip()
            if not val.isdigit() or int(val) <= 0:
                name = sec_info[sec][0]
                messagebox.showerror("输入错误", f"【{name}】的分值必须为正整数")
                return
            s = int(val)
            scores[sec] = s
            total += s * sec_info[sec][1]

        if total != 100:
            messagebox.showwarning(
                "总分不是100",
                f"当前总分为 {total} 分（需要100分）\n请重新调整各题型分值。"
            )
            return

        result.update(scores)
        root.destroy()

    root = tk.Tk()
    root.title("设置各题型分值")
    root.resizable(False, False)

    tk.Label(root, text="请输入每个题型的 每小题 分值：\n",
             font=("宋体", 11)).grid(row=0, column=0, columnspan=3, padx=15, pady=(15, 5))

    entries = {}
    for row, sec in enumerate(secs, start=1):
        name, count = sec_info[sec]
        tk.Label(root, text=f"{name}（共{count}题）",
                 font=("宋体", 10), anchor="w").grid(
            row=row, column=0, sticky="w", padx=(20, 5), pady=3)
        e = tk.Entry(root, width=6, font=("宋体", 10), justify="center")
        e.grid(row=row, column=1, padx=5, pady=3)
        tk.Label(root, text="分/题", font=("宋体", 10)).grid(
            row=row, column=2, sticky="w", padx=(0, 20), pady=3)
        entries[sec] = e

    tk.Button(root, text="确 定", font=("宋体", 11), width=10,
              command=submit).grid(
        row=len(secs) + 1, column=0, columnspan=3, pady=(10, 15))

    entries[secs[0]].focus_set()
    root.eval('tk::PlaceWindow . center')
    root.mainloop()

    if not result:
        print("用户取消，程序退出。")
        sys.exit(0)

    return result


# ════════════════════ 单行处理 ════════════════════


def process_row(src, out_dir, seq, exam_name, name_suffix):
    """处理一行：分析文档 → 问分值 → 生成4文件。成功返回 True。"""
    doc = Document(src)
    meta = analyze(doc)

    qs = sorted({m["q"] for m in meta if m["q"] is not None})
    ns = len({m["s"] for m in meta if m["s"] >= 0})
    print(f"  文档: {ns} 个题型, {len(qs)} 道题")

    # 统计拆分后每套的题目数
    sec_info = {}
    for i, (p, m) in enumerate(zip(doc.paragraphs, meta)):
        if m["tp"] == "hdr":
            sec = m["s"]
            total_in_sec = len({
                m2["q"] for m2 in meta
                if m2["q"] is not None and m2["s"] == sec
            })
            sec_info[sec] = (p.text.strip(), total_in_sec // 2)

    sec_scores = _ask_scores(sec_info)

    teacher_vol = seq * 2 - 1
    student_vol = seq * 2

    odd = lambda n: n % 2 == 1
    even = lambda n: n % 2 == 0

    jobs = [
        (f"第{teacher_vol}卷 {exam_name} {name_suffix}（解析版）.docx",
         odd, True, True, teacher_vol, "教师"),
        (f"第{student_vol}卷 {exam_name} {name_suffix}（解析版）.docx",
         even, True, True, student_vol, "学生"),
        (f"第{teacher_vol}卷 {exam_name} {name_suffix}（原卷版）.docx",
         odd, False, False, teacher_vol, "教师"),
        (f"第{student_vol}卷 {exam_name} {name_suffix}（原卷版）.docx",
         even, False, False, student_vol, "学生"),
    ]

    for name, qf, ans, red, vol, ptype in jobs:
        build(src, os.path.join(out_dir, name), meta, qf, ans, red,
              sec_scores, vol, exam_name, ptype)
        print(f"    -> {name}")

    return True


# ════════════════════ 主程序 ════════════════════


def main():
    out_dir = os.environ.get("EXAM_TOOL_WORKDIR")
    if not out_dir or not os.path.isdir(out_dir):
        if getattr(sys, 'frozen', False):
            out_dir = os.path.dirname(sys.executable)
        else:
            out_dir = os.path.dirname(os.path.abspath(__file__)) or "."

    # 弹窗选择 xlsx 规划表
    xlsx_path, xlsx_name = _ask_xlsx_file(out_dir)
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active
    print(f"已连接规划表: {xlsx_name}")

    header_row = _find_header_row(ws)
    if header_row is None:
        print("规划表中未找到表头行（'序号'列）")
        sys.exit(1)

    # 弹窗询问输出文件名中的后缀文本
    name_suffix = _ask_name_suffix()
    print(f"文件名后缀: {name_suffix}")

    pending = _get_pending_rows(ws, header_row)

    if not pending:
        print("所有序号均已处理完毕（全部标绿）。")
        return

    print(f"待处理: {len(pending)} 行 "
          f"(序号 {', '.join(str(r[1]) for r in pending)})\n")

    mark_col = _find_mark_col(ws, header_row)

    for row_idx, seq, exam_name in pending:
        print(f"{'='*50}")
        print(f"序号 {seq}: {exam_name}")
        print(f"{'='*50}")

        # 弹窗选择原文档
        src = _ask_input_file(out_dir, exam_name)
        print(f"  原文档: {os.path.basename(src)}")

        try:
            process_row(src, out_dir, seq, exam_name, name_suffix)

            # 成功 → 标绿并保存
            ws.cell(row=row_idx, column=mark_col).fill = GREEN_FILL
            wb.save(xlsx_path)
            print(f"  序号 {seq} 已完成，已标绿。\n")

        except Exception as e:
            # 失败 → 标红并保存
            ws.cell(row=row_idx, column=mark_col).fill = RED_FILL
            wb.save(xlsx_path)
            print(f"\n[错误] 序号 {seq} 处理失败，已标红: {e}")
            print("程序停止。请修复问题后重新运行，已完成的序号（绿色）会自动跳过。")
            sys.exit(1)

    print("\n全部完成!")


if __name__ == "__main__":
    main()
