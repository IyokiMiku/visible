"""
解析版 → 原卷版 转换工具

功能：
  将解析版 .docx（含【答案】【详解】【解析】等）转换为原卷版 .docx
  - 删除所有【答案】【详解】【解析】段落
  - 混合段落（题干+【答案】同行）只保留题干部分
  - 去除全文底纹和高亮

用法：
  python answer2none.py
  弹窗选择一个或多个解析版 docx → 自动在同目录生成对应原卷版

依赖：
  pip install python-docx
"""

import re
import os
import sys
import shutil

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖，请先执行: pip install python-docx")
    sys.exit(1)

# ════════════════════ XML tag 常量 ════════════════════

W_R = qn("w:r")
W_RPR = qn("w:rPr")
W_PPR = qn("w:pPr")
W_SHD = qn("w:shd")
W_HL = qn("w:highlight")

# ════════════════════ 文本识别 ════════════════════

HEADER_RE = re.compile(r"^[一二三四五六七八九十]+\s*[、.．]")
QNUM_RE = re.compile(r"^(\d+)\s*[.、．)\）]")
LUE_RE = re.compile(r"【(?:详解|解析)】\s*略")


def is_header(text):
    return bool(text and HEADER_RE.match(text))


def get_qnum(text):
    m = QNUM_RE.match(text or "")
    return int(m.group(1)) if m else None


# ════════════════════ 文档结构分析 ════════════════════


def analyze(doc):
    out = []
    sec, cur_q, in_ans = -1, None, False

    for p in doc.paragraphs:
        t = p.text.strip()

        if is_header(t):
            sec += 1
            cur_q = None
            in_ans = False
            out.append(dict(tp="hdr", q=None, s=sec, lue=False))
            continue

        n = get_qnum(t)
        if n is not None and in_ans and cur_q is not None and n <= cur_q:
            n = None

        if n is not None:
            cur_q = n
            in_ans = False
            if "【答案】" in t:
                in_ans = True
                out.append(dict(tp="mix", q=cur_q, s=sec, lue=False))
            else:
                out.append(dict(tp="body", q=cur_q, s=sec, lue=False))
            continue

        if "【答案】" in t:
            in_ans = True
        if in_ans or "【详解】" in t or "【解析】" in t:
            in_ans = True
            out.append(
                dict(tp="ans", q=cur_q, s=sec, lue=bool(LUE_RE.search(t)))
            )
            continue

        if cur_q is not None:
            out.append(dict(tp="body", q=cur_q, s=sec, lue=False))
        else:
            out.append(
                dict(tp="sp" if sec >= 0 else "pre", q=None, s=sec, lue=False)
            )

    return out


# ════════════════════ 格式操作 ════════════════════


def _remove_shading(para):
    pPr = para._element.find(W_PPR)
    if pPr is not None:
        for e in pPr.findall(W_SHD):
            pPr.remove(e)
    for r in para._element.findall(W_R):
        rPr = r.find(W_RPR)
        if rPr is None:
            continue
        for tag in (W_SHD, W_HL):
            for e in rPr.findall(tag):
                rPr.remove(e)


def _trim_answer(para):
    """截掉段落中【答案】及其后面的全部内容"""
    idx = para.text.find("【答案】")
    if idx < 0:
        return
    pos = 0
    for run in para.runs:
        rt = run.text or ""
        rs = pos
        re_ = pos + len(rt)
        if rs >= idx:
            run.text = ""
        elif re_ > idx:
            run.text = rt[: idx - rs]
        pos = re_


# ════════════════════ 核心转换 ════════════════════


def convert(src_path, dst_path):
    """将解析版 docx 转换为原卷版 docx"""
    shutil.copy2(src_path, dst_path)
    doc = Document(dst_path)
    paras = list(doc.paragraphs)
    meta = analyze(doc)

    if len(paras) != len(meta):
        raise RuntimeError(
            f"段落数不一致（{len(paras)} vs {len(meta)}），文档结构异常")

    rm_indices = []

    for i, (p, m) in enumerate(zip(paras, meta)):
        tp = m["tp"]

        if tp == "ans":
            rm_indices.append(i)
        elif tp == "mix":
            _trim_answer(p)

    for i in reversed(rm_indices):
        paras[i]._element.getparent().remove(paras[i]._element)

    for p in doc.paragraphs:
        _remove_shading(p)

    doc.save(dst_path)


# ════════════════════ 弹窗选择文件 ════════════════════


def ask_files():
    """弹窗让用户选择一个或多个解析版 docx 文件"""
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()

    files = filedialog.askopenfilenames(
        title="选择解析版 docx 文件（可多选）",
        filetypes=[("Word 文档", "*.docx")],
    )
    root.destroy()

    if not files:
        print("未选择文件，程序退出。")
        sys.exit(0)

    return list(files)


# ════════════════════ 主程序 ════════════════════


def main():
    files = ask_files()

    print(f"共选择 {len(files)} 个文件\n")

    for src in files:
        name = os.path.basename(src)
        directory = os.path.dirname(src)

        # 生成输出文件名：将"解析版"替换为"原卷版"，否则追加"（原卷版）"
        if "解析版" in name:
            out_name = name.replace("解析版", "原卷版")
        else:
            base, ext = os.path.splitext(name)
            out_name = f"{base}（原卷版）{ext}"

        dst = os.path.join(directory, out_name)

        print(f"处理: {name}")
        try:
            convert(src, dst)
            print(f"  -> {out_name}")
        except Exception as e:
            print(f"  失败: {e}")

    print(f"\n全部完成！共处理 {len(files)} 个文件。")
    input("按 Enter 退出...")


if __name__ == "__main__":
    main()
