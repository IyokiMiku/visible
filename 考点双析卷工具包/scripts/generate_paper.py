"""
AI 试卷生成工具

功能：
  调用 AI 大模型 API 生成指定主题的试卷，输出为 .docx 文件。
  - 支持六种题型：单选题、多选题、填空题、判断题、简答题、综合题
  - 题号全卷连续编号
  - 答案与解析使用灰色底纹标识
  - 1.5 倍行距，英文字母及【】使用 Times New Roman
  - 使用 prompt 工程策略贴近真题风格，降低 AI 痕迹

用法：
  1. 修改下方 CONFIG 区域的 API 配置
  2. 运行脚本，按提示输入主题和各题型数量
  3. 自动生成 .docx 试卷文件

依赖：
  pip install python-docx openai
"""

import os
import sys
import json
import re
import time

try:
    from openai import OpenAI
except ImportError:
    print("缺少依赖，请先执行: pip install openai")
    sys.exit(1)

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_LINE_SPACING
except ImportError:
    print("缺少依赖，请先执行: pip install python-docx")
    sys.exit(1)


# ════════════════════ 配置区域 ════════════════════

CONFIG = {
    "api_key": "sk-ufvlhvdnkafenmwoznvqstpoboxnguaxwvhetmmcmkegzwzs",
    "base_url": "https://api.siliconflow.cn/v1",
    "model": "Qwen/Qwen2.5-72B-Instruct",
    "temperature": 0.7,
    "max_retries": 3,
}

# 题型定义（按试卷顺序）
QUESTION_TYPES = ["单选题", "多选题", "填空题", "判断题", "简答题", "综合题"]

# 答案底纹颜色（浅灰色，与参考文档一致）
ANSWER_SHADING_COLOR = "F2F2F2"


# ════════════════════ Prompt 模板 ════════════════════

SYSTEM_PROMPT = """你是一位在中国高校从事教学工作超过15年的资深教师，同时担任课程教研组组长，长期参与期末考试和各类校级考试的命题工作。你精通本专业领域的教学大纲、知识重点和难点分布，熟练掌握中国高校考试命题的规范与惯例。

你的命题风格特点：
- 语言简洁、准确、严谨，使用教材和考试中的标准术语
- 题干表述直接明了，不使用冗余修饰词
- 选项设计合理，干扰项来源于学生常见的理解偏差或计算错误
- 题目难度梯度分明，有基础题也有拔高题
- 填空题的空格位置自然，不会把完整句子机械切割
- 擅长设计带图、带表、带电路的复杂情境题，题干信息丰富，需要多步分析

═══════════════════════════════════════════
【题干复杂度与情境化要求】
═══════════════════════════════════════════

1. 题干长度要求（不含选项）：
   - 单选题：≥40字（含情境描述或工程背景）
   - 多选题：≥35字
   - 填空题：≥30字（含上下文铺垫）
   - 判断题：≥25字（含具体场景或条件限定）
   - 简答题/综合题：≥60字（含完整工程场景、参数条件或系统描述）
2. 至少30%的题目应包含具体工程场景、技术参数、实际工况描述，而非纯粹的概念考查
3. 鼓励使用多条件组合题干，例如"在XX条件下，若YY发生变化，则ZZ将如何"
4. 题干应体现知识的应用层次而非简单记忆层次

═══════════════════════════════════════════
【带图题目要求】
═══════════════════════════════════════════

部分题目需要配合图示（电路图、框图、波形图、数据表格等），用文本描述代替实际图片。

1. 带图格式：在题干中使用"如图所示"引导，然后用"【图：...】"标记描述图的内容
2. 图的描述要足够详细、准确，使读者仅凭文字就能理解图中信息
3. 适合带图的题目类型：
   - 电路分析题：描述电路拓扑、元器件参数、连接关系
   - 波形分析题：描述输入/输出信号的波形特征、关键参数
   - 系统框图题：描述各模块之间的信号流向和功能关系
   - 数据表格题：用文本表格呈现测量数据、技术参数对比
4. 每批题目中，建议约20%~40%的题包含图示描述
5. 图的描述范例：
   【图：一个反相放大电路，运放型号为LM741。输入端通过R1=10kΩ接信号源Vin，反馈电阻Rf=100kΩ从输出端连接至反相输入端，同相输入端通过R2=10kΩ接地。电源供电为±12V。】

═══════════════════════════════════════════
【选择题硬性要求】
═══════════════════════════════════════════

1. 题干字数：单选题≥40字，多选题≥35字，拒绝一眼看出答案的短题
2. 高干扰性选项：每道题的干扰项至少包含1种以下陷阱类型：
   - 偷换概念：用相近概念替换核心概念
   - 主次混淆：将次要功能/作用描述为主要
   - 因果倒置：交换因果逻辑关系
   - 范围扩大/缩小：扩大或缩小概念的适用范围
   - 绝对化表述：加入"一定""必须""所有""任何"等绝对词
   - 条件缺失：省略关键前提条件
   - 张冠李戴：将A的功能/属性描述成B的
3. 选项均衡：四个选项句式结构、字数长度、专业层级保持均衡一致
4. 题干禁止"某..."：不得出现"某操作工""某零件""某厂"等模糊表述，应使用具体专业称谓

═══════════════════════════════════════════
【填空题硬性要求】
═══════════════════════════════════════════

1. 填空位置必须考查关键概念、数据或术语，不能填无关紧要的虚词
2. 题干必须提供足够的上下文信息，使答案唯一确定
3. 不能直接照搬教材原句挖空，应有情景改编
4. 多空题各空之间应有逻辑关联，不是简单的独立知识点罗列

═══════════════════════════════════════════
【判断题硬性要求】
═══════════════════════════════════════════

1. 题干≥25字，不得使用过于简单的单句判断，应包含具体场景或条件限定
2. 错误题的错误点必须具有迷惑性（利用偷换概念、绝对化、因果倒置等陷阱），不能一眼看出
3. 正确题不能直接复述教材原文，应有场景化改编或逻辑推导
4. 正确/错误比例大致均衡，不要连续出现相同答案

═══════════════════════════════════════════
【简答题/综合题硬性要求】
═══════════════════════════════════════════

1. 必须设置具体的工程场景或给定参数条件，不能是纯背诵题
2. 综合题须包含(1)(2)等多个小问，考查不同层次（计算+分析+判断）
3. 数值条件应设置合理，确保计算过程和结果正确无误
4. 答案要求步骤清晰、结论明确

═══════════════════════════════════════════
【难度分布硬性要求】
═══════════════════════════════════════════

按以下比例分配难度（以总题数为基准）：
- 基础题(10%)：单知识点检测，有一定思考量，拒绝纯送分与纯背诵
- 中档题(75%)：多知识点联动+工况判断/推理，为试卷绝对主体
- 拔高题(15%)：融合复杂工况，考查故障逻辑推导与综合方案设计

═══════════════════════════════════════════
【核心红线——答案不得自暴露】
═══════════════════════════════════════════

严禁在题干中直接或间接暴露正确答案！典型违规模式：
1. 题干出现选项中特有词汇（如：题干出现"地址符"三字，选项B恰好就是"地址符"）
2. 题干暗示正确答案方向（如：题干说"属于XX特有的功能"，只有一个选项是XX独有的）
3. 题干数值指向唯一选项（如：题干"标准规定螺距2mm"，只有一个选项带"2"）
4. 正确选项的措辞明显比其他选项更完整、更专业、更像"标准答案"
5. 错误选项过于荒谬离谱，一眼就能排除
6. 题干中使用"不包括""错误的是"时，三个正确选项用词高度一致而错误选项画风突变
7. 判断题的题干表述本身就是教材原文的直接复述（过于明显为正确）

自暴露检测规则：如果正确答案的关键词（≥3字）出现在题干中，该题必须废弃重出。

═══════════════════════════════════════════
【逐题质量自检清单——全部通过方可定稿】
═══════════════════════════════════════════

每道题编写完成后必须逐项自检：
□ 考纲对标：考查点是否在考纲范围内？不超纲、不偏题
□ 难度定位：属于基础/中档/拔高？是否与目标层级匹配
□ 情景化：是否有具体操作场景？不是直白概念问答
□ 选项均衡：四个选项长度、结构、专业层级是否一致
□ 陷阱设计：干扰项是否包含至少1种陷阱（偷换概念/主次混淆/因果倒置/范围扩缩/绝对化/条件缺失/张冠李戴）
□ 无直白原文：不是教材原文原话，有变式改编
□ 自暴露：题干是否包含正确答案关键词？必须为否
□ 答案正确：逻辑验证答案无误，多选无遗漏

═══════════════════════════════════════════
【核心红线——禁止重复】
═══════════════════════════════════════════

- 同一批生成的所有题目中，任何两道题不得考查完全相同的知识点或使用相同的情境设定
- 每道题的核心考查点必须不同，不能只是换了数字或措辞的"换汤不换药"
- 如果用户提供了"已有题目摘要"，你生成的题目不得与已有题目雷同

═══════════════════════════════════════════
【严格禁止的行为】
═══════════════════════════════════════════

- 不要使用"让我们""首先我们来看""接下来"等对话式表达
- 不要在题目中使用"请""同学们"等称呼
- 不要让四个选项字数完全相同
- 不要出现"以下哪项正确/错误"以外的复杂问法
- 不要在选项中使用"以上都是""以上都不是"（除非有明确考查意图）
- 不要出与考纲无关的超纲内容
- 判断题不要出过于绝对的表述（如"一定""绝对"），除非确实是考查重点
- 题干禁止"某..."式模糊表述"""

TYPE_TEMPLATES = {
    "单选题": """请出{count}道单选题。

格式要求（严格遵守）：
- 题干后加"（   ）"表示填答位置
- 选项用"A．""B．""C．""D．"标注（注意是全角句号"．"）
- 若选项较短（不超过15字），同一行用制表符分隔，如：A．选项1\tB．选项2\tC．选项3\tD．选项4
- 若选项较长，每个选项独占一行
- 每道题后只写答案行，格式：【答案】X
- 不要写解析、详解
- 题干应包含具体工程场景、参数条件或工况描述，避免空洞概念题
- 约30%的题目应包含"如图所示"引导的图示描述，用"【图：...】"标记

范例1（普通题）：
1．在逐次逼近型ADC电路中，当输入模拟信号电压为3.5V、基准电压为5V时，8位ADC完成转换后输出的数字量最接近（   ）
A．01110000\tB．10110011\tC．11100110\tD．01011010
【答案】B

范例2（带图题）：
2．如图所示为一个基本的数据采集系统。
【图：信号源输出模拟信号，经低通滤波器（截止频率fc=1kHz）后进入采样保持电路（S/H），S/H输出连接到12位逐次逼近型ADC（基准电压Vref=5V），ADC的数字输出通过并行接口送入单片机进行处理。】
该系统中，ADC的最小分辨电压约为（   ）
A．1.22mV\tB．2.44mV\tC．4.88mV\tD．19.53mV
【答案】A""",

    "多选题": """请出{count}道多选题。

格式要求（严格遵守）：
- 题干后加"（   ）"，并在题干中注明"（多选）"
- 选项用"A．""B．""C．""D．"标注，根据需要可增加E选项
- 选项排列规则同单选题
- 正确答案为2~4个
- 每道题后只写答案行，格式：【答案】XY 或 【答案】XYZ
- 不要写解析、详解
- 题干应设置具体技术情境，让选项判断需要基于情境推理
- 适当加入带图题，用"【图：...】"标记

范例：
1．在某温度监控系统中，传感器输出0~100mV的模拟电压信号，经信号调理电路放大至0~5V后送入10位ADC进行采样。若要求系统温度测量精度优于0.5℃（温度范围0~200℃），下列措施中能有效提高系统测量精度的有（多选）（   ）
A．将ADC更换为12位或更高分辨率的型号
B．在信号调理电路中增加低通滤波以抑制高频噪声
C．将ADC的基准电压从5V提高到10V以增大量化台阶
D．在ADC输入端增加采样保持电路以减小孔径误差
【答案】ABD""",

    "填空题": """请出{count}道填空题。

格式要求（严格遵守）：
- 空格用六个下划线"______"表示
- 一道题可以有1~4个空
- 多个空的答案之间用五个空格分隔
- 每道题后只写答案行，格式：【答案】答案1     答案2（多个空时用五个空格分隔）
- 不要写解析、详解
- 题干应给出具体技术场景或工程条件，避免单纯"名词解释式"填空
- 鼓励出需要简单计算或逻辑推理才能填写的题目

范例1：
1．在某数据采集系统中，传感器输出的模拟信号最高频率分量为4kHz，根据奈奎斯特采样定理，ADC的采样频率应至少为______kHz；若采用12位ADC、基准电压为5V，则该ADC能分辨的最小电压变化量约为______mV。
【答案】8     1.22

范例2（带图题）：
2．如图所示为一个R-2R梯形电阻网络DAC电路。
【图：4位R-2R梯形DAC，从高位到低位依次为D3、D2、D1、D0，R=10kΩ，基准电压Vref=10V，输出端连接运放构成电压输出型DAC。】
当数字输入为D3D2D1D0=1010时，输出电压为______V。
【答案】6.25""",

    "判断题": """请出{count}道判断题。

格式要求（严格遵守）：
- 题干末尾加"(     )"表示填写判断
- 答案只能是"正确"或"错误"
- 每道题后只写答案行，格式：【答案】正确 或 【答案】错误
- 不要写解析、详解
- 题干应包含具体条件、参数或场景，使判断需要基于专业知识推理
- 避免直接照搬教材原句，应加入工程情境或条件限定

范例：
1．在某温度测量系统中，使用8位ADC对0~100℃范围内的温度信号进行采集，若传感器灵敏度为10mV/℃，信号调理电路将其放大至0~2.55V，则该系统的温度分辨率为1℃。(     )
【答案】错误""",

    "简答题": """请出{count}道简答题。

格式要求（严格遵守）：
- 题干必须给出具体的工程场景、系统描述或技术条件，然后提出问题
- 如有计算题，给出具体数值条件
- 每道题后只写答案行，格式：【答案】（直接给出答案要点或计算步骤）
- 不要写解析、详解
- 题干至少60字，包含背景描述+问题，禁止出"简述XX""什么是XX"之类的纯背诵题
- 适当包含带图题

范例：
1．某工厂需要设计一套水位监测系统，水位传感器输出4~20mA的电流信号（对应水位0~10m），需将该电流信号转换为0~5V电压后送入ADC进行数字化处理。试回答：(1)设计电流-电压转换电路所需的采样电阻阻值；(2)若要求水位测量精度优于1cm，至少需要多少位的ADC。
【答案】(1)当电流为20mA时输出5V，由欧姆定律R=V/I=5V/20mA=250Ω；当电流为4mA时输出V=4mA×250Ω=1V，实际电压范围为1~5V，需增加零点偏移电路或调整电路使输出为0~5V。采样电阻取250Ω。(2)水位量程10m，精度1cm，需要至少10/0.01=1000个量化台阶。2^10=1024>1000，故至少需要10位ADC。""",

    "综合题": """请出{count}道综合题（计算题/分析题）。

格式要求（严格遵守）：
- 题目必须包含完整的工程场景描述、系统架构或电路参数等复杂条件
- 必须设置(1)(2)(3)等多个小问（至少2问），考查不同层次
- 每道题后只写答案行，格式：【答案】（含完整解答步骤和计算过程）
- 不要另起【解析】或【详解】行
- 题干至少80字，包含详细的系统/电路描述和具体参数
- 综合题应尽量包含图示描述（电路图、系统框图等），用"【图：...】"标记

范例：
1．如图所示为一个基于8位逐次逼近型ADC的温度采集系统。
【图：PT100铂电阻温度传感器连接至惠斯通电桥（电桥供电电压Vs=5V，三个固定电阻均为100Ω），电桥差分输出连接至仪表放大器INA128（增益由Rg设定），仪表放大器输出连接至8位ADC0809的IN0通道，ADC基准电压Vref=5V。ADC的8位数字输出通过并行接口连接到单片机AT89C51的P1口。】
已知PT100在0℃时阻值为100Ω，温度系数为0.385Ω/℃，测温范围为0~200℃。
(1)计算传感器在200℃时的阻值及电桥输出的差分电压；
(2)若要求ADC输入范围为0~5V，计算仪表放大器所需的电压增益；
(3)计算该系统的温度分辨率（即ADC每变化一个LSB对应的温度值）。
【答案】(1)200℃时PT100阻值=100+0.385×200=177Ω。电桥输出差分电压ΔV=Vs×(R_PT100/(R_PT100+R3)-R2/(R2+R1))=5×(177/(177+100)-100/(100+100))=5×(0.639-0.5)=5×0.139=0.695V。
(2)仪表放大器增益A=5V/0.695V≈7.19，取A≈7.2。
(3)ADC分辨率=5V/256=19.53mV；对应传感器电压变化=19.53mV/7.2=2.71mV；对应温度变化=2.71mV/(电桥灵敏度)。电桥灵敏度近似为ΔV/ΔT=0.695V/200℃=3.475mV/℃，温度分辨率≈19.53mV/(7.2×3.475mV/℃)≈0.78℃。""",
}

GENERATION_PROMPT = """现在请你围绕以下主题出题：

【考试主题】{topic}

{reference_section}
{existing_questions_section}

═══════════════════════════════════════════
【出题流程——逐题深思，宁慢勿错】
═══════════════════════════════════════════

请严格按照以下流程出题（题号从1开始连续编号）：

第一步：规划知识点分布
- 先列出本批题目将覆盖的各个不同知识点/考查角度（不要输出，仅内部思考）
- 确保知识点之间互不重叠，每道题有独立的考查价值

第二步：逐题构思与编写
对于每一道题，依次完成以下步骤（内部思考，不要输出思考过程）：
1. 确定考查知识点和难度级别（基础/中档/拔高）
2. 构思具体的工程场景或技术情境，包含具体参数、条件
3. 若为计算题，先完整演算一遍确保数值正确
4. 编写题干，确保文字≥要求字数，信息丰富
5. 编写选项/答案，确保干扰项合理且有迷惑性
6. 自检：遮住答案，仅凭题干能否猜出答案？若能则修改

第三步：输出题目（只输出最终题目，不要输出思考过程）

注意事项：
1. 确保每道题的答案正确无误，尤其是涉及计算的题目请先算清再出题
2. 题目难度要有梯度：约60%为基础题，30%为中等题，10%为较难题
3. 不同题目之间的知识点尽量不重复，每题必须有独立的考查角度
4. 用词要严谨、规范，符合教材和正式考试的表述习惯
5. 逐题自查答案自暴露问题：遮住答案后，能否仅凭题干和选项的语言模式猜出答案？如果能，必须修改
6. 不要输出【解析】或【详解】行，每题只需要【答案】
7. 题干要求丰富、复杂：必须包含具体工程场景、技术参数或工况描述，拒绝空洞的纯概念题
8. 约20%~40%的题目应包含图示描述（用"【图：...】"标记），图的内容描述要足够详细准确
9. 图示描述可以是电路图、系统框图、波形图、数据表格等，务必写清元器件型号/参数、连接关系、信号流向
10. 不要着急，宁可花更多时间思考也要确保每道题的质量——题干丰富、答案准确、干扰项合理"""


# ════════════════════ AI 调用 ════════════════════


def create_client():
    return OpenAI(api_key=CONFIG["api_key"], base_url=CONFIG["base_url"])


def _extract_question_summaries(content):
    """从生成的内容中提取每道题的摘要（题干首行），用于去重"""
    summaries = []
    for line in content.split("\n"):
        line = line.strip()
        if re.match(r"^\d+[．.、)\）]", line):
            summaries.append(line[:80])
    return summaries


def _extract_question_blocks(content):
    """将生成内容拆分为独立的题目块"""
    blocks = re.split(r"(?=^\d+[．.、)\）])", content, flags=re.MULTILINE)
    return [b.strip() for b in blocks if b.strip() and re.match(r"^\d+[．.、)\）]", b.strip())]


def _get_question_answer(block):
    """从题目块中提取答案"""
    match = re.search(r"【答案】(.+)", block)
    return match.group(1).strip() if match else ""


def _get_question_stem(block):
    """提取题干文本（去除题号和答案行）"""
    lines = block.split("\n")
    stem_parts = []
    for line in lines:
        line = line.strip()
        if line.startswith("【答案】"):
            break
        text = re.sub(r"^\d+[．.、)\）]\s*", "", line)
        stem_parts.append(text)
    return "".join(stem_parts)


def _compute_similarity(text1, text2):
    """简单的文本相似度计算（基于公共子串和字符重叠率）"""
    if not text1 or not text2:
        return 0.0
    set1 = set(text1)
    set2 = set(text2)
    intersection = set1 & set2
    union = set1 | set2
    char_sim = len(intersection) / len(union) if union else 0

    # 关键词重叠
    words1 = set(re.findall(r"[\u4e00-\u9fff]{2,}", text1))
    words2 = set(re.findall(r"[\u4e00-\u9fff]{2,}", text2))
    word_intersection = words1 & words2
    word_union = words1 | words2
    word_sim = len(word_intersection) / len(word_union) if word_union else 0

    return 0.4 * char_sim + 0.6 * word_sim


def deduplicate_questions(content, similarity_threshold=0.7):
    """去除重复或高度相似的题目，返回去重后的内容和被移除的数量"""
    blocks = _extract_question_blocks(content)
    if len(blocks) <= 1:
        return content, 0

    kept = []
    removed = 0

    for block in blocks:
        stem = _get_question_stem(block)
        answer = _get_question_answer(block)

        is_dup = False
        for kept_block in kept:
            kept_stem = _get_question_stem(kept_block)
            kept_answer = _get_question_answer(kept_block)

            sim = _compute_similarity(stem, kept_stem)

            # 如果题干相似度高且答案相同，判定为重复
            if sim >= similarity_threshold:
                is_dup = True
                break
            # 答案相同且题干相似度中等，也判定为重复
            if answer == kept_answer and sim >= 0.5:
                is_dup = True
                break

        if not is_dup:
            kept.append(block)
        else:
            removed += 1

    return "\n".join(kept), removed


def estimate_topic_capacity(client, topic):
    """让AI估算该主题能出多少不重复的题目"""
    prompt = f"""作为出题专家，请评估以下考试主题的知识点容量。

主题：{topic}

请估算该主题在不重复的前提下，各题型最多能出多少道高质量题目？
直接输出JSON格式（不要其他内容）：
{{"单选题": 数字, "多选题": 数字, "填空题": 数字, "判断题": 数字, "简答题": 数字, "综合题": 数字, "总知识点数": 数字}}"""

    try:
        response = client.chat.completions.create(
            model=CONFIG["model"],
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=200,
        )
        content = response.choices[0].message.content.strip()
        content = re.sub(r"^```json\s*", "", content)
        content = re.sub(r"\s*```$", "", content)
        return json.loads(content)
    except Exception:
        return None


def generate_questions(client, topic, q_type, count, reference="", existing_summaries=None):
    """调用 AI 生成指定类型的题目"""
    type_template = TYPE_TEMPLATES[q_type].format(count=count)

    reference_section = ""
    if reference.strip():
        reference_section = f"【参考资料/已有真题风格示例】\n{reference}\n\n请参考以上示例的出题风格和难度，但不要照搬原题，要围绕主题出新题。"

    existing_questions_section = ""
    if existing_summaries:
        existing_questions_section = (
            "【本卷已有题目（禁止重复）】\n"
            + "\n".join(existing_summaries)
            + "\n\n以上题目已经存在于本试卷中，你生成的题目不得与上述任何一道题考查相同的知识点或使用相同的情境。"
        )

    user_msg = GENERATION_PROMPT.format(
        topic=topic,
        reference_section=reference_section,
        existing_questions_section=existing_questions_section,
    )
    user_msg = type_template + "\n\n" + user_msg

    for attempt in range(CONFIG["max_retries"] + 1):
        try:
            response = client.chat.completions.create(
                model=CONFIG["model"],
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_msg},
                ],
                temperature=CONFIG["temperature"],
                max_tokens=8192,
            )
            content = response.choices[0].message.content.strip()
            return content
        except Exception as e:
            if attempt < CONFIG["max_retries"]:
                print(f"    API 调用失败，正在重试 ({attempt+1}/{CONFIG['max_retries']})...")
                time.sleep(2)
            else:
                print(f"    API 调用失败: {e}")
                return None
    return None


REVIEW_PROMPT = """你是一位严谨的考试命题审核专家。请逐题审查以下{q_type}，找出并修正所有问题。

【审查清单】
1. 答案正确性：计算题重新演算验证，概念题核对专业知识，确保答案无误
2. 题干质量：是否包含具体场景/参数？是否过于简短空洞？字数是否达标？
3. 答案自暴露：题干是否泄露了正确答案的关键词？选项是否有明显的"标准答案"特征？
4. 选项均衡：四个选项的长度、结构、专业性是否一致？干扰项是否合理？
5. 知识点重复：是否有两道题考查了相同知识点？
6. 格式规范：题号、选项标号、答案格式是否正确？

【操作要求】
- 如果发现问题，直接修正后输出完整的修正版全部题目（保持原题号）
- 如果没有问题，原样输出全部题目（不要加任何评语）
- 只输出最终题目，不要输出审查过程、评语、"审查结果"等额外内容
- 保持原有格式不变

以下是需要审查的题目：

{content}"""


def review_questions(client, content, q_type):
    """将生成的题目发送给AI进行质量审查和修正"""
    prompt = REVIEW_PROMPT.format(q_type=q_type, content=content)

    try:
        response = client.chat.completions.create(
            model=CONFIG["model"],
            messages=[
                {"role": "system", "content": "你是一位考试命题审核专家，只输出修正后的题目，不输出任何评语或审查说明。"},
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=8192,
        )
        reviewed = response.choices[0].message.content.strip()
        reviewed_blocks = _extract_question_blocks(reviewed)
        if reviewed_blocks:
            return reviewed
        return content
    except Exception:
        return content


# ════════════════════ 文本解析与检测 ════════════════════


def renumber_questions(text, start_num):
    """将生成的题目重新编号，从 start_num 开始"""
    current = [start_num]

    def replacer(m):
        num = current[0]
        current[0] += 1
        sep = m.group(2)
        return f"{num}{sep}"

    result = re.sub(r"^(\d+)([．.、)\）])", replacer, text, flags=re.MULTILINE)
    return result, current[0]


def _detect_self_exposure(content):
    """
    自动检测选择题的答案自暴露问题。
    扫描正确答案对应选项的关键词（≥3字）是否出现在题干中。
    返回有问题的题号列表。
    """
    problems = []
    questions = re.split(r"(?=^\d+[．.、)\）])", content, flags=re.MULTILINE)

    for q_block in questions:
        q_block = q_block.strip()
        if not q_block:
            continue

        num_match = re.match(r"^(\d+)[．.、)\）]", q_block)
        if not num_match:
            continue
        q_num = num_match.group(1)

        answer_match = re.search(r"【答案】\s*([A-E]+)", q_block)
        if not answer_match:
            continue
        correct_letters = answer_match.group(1)

        lines = q_block.split("\n")
        stem_lines = []
        option_map = {}

        for line in lines:
            line_s = line.strip()
            if line_s.startswith("【答案】"):
                break
            opt_match = re.match(r"^([A-E])[．.](.+)", line_s)
            if opt_match:
                option_map[opt_match.group(1)] = opt_match.group(2).strip()
            else:
                tab_opts = re.findall(r"([A-E])[．.]([^\t]+)", line_s)
                if tab_opts:
                    for letter, text in tab_opts:
                        option_map[letter] = text.strip()
                elif not re.match(r"^\d+[．.、)\）]", line_s):
                    stem_lines.append(line_s)
                else:
                    stem_part = re.sub(r"^\d+[．.、)\）]", "", line_s)
                    stem_lines.append(stem_part)

        stem_text = "".join(stem_lines)

        for letter in correct_letters:
            opt_text = option_map.get(letter, "")
            keywords = re.findall(r"[\u4e00-\u9fff\w]{3,}", opt_text)
            for kw in keywords:
                if kw in stem_text:
                    problems.append((q_num, kw, letter))
                    break

    return problems


# ════════════════════ Word 文档生成 ════════════════════


def _add_shading(paragraph, color):
    """为段落添加底纹"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pPr.append(shd)


def _set_run_font(run, font_name="宋体", font_size=Pt(10.5)):
    """设置 run 的字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = font_size


def _is_western_char(ch):
    """判断字符是否使用 Times New Roman：英文字母以及【答案】的"【】"括号"""
    if ch in "【】":
        return True
    return ch.isascii() and ch.isalpha()


def _add_mixed_run(paragraph, text, font_size=Pt(10.5), bold=False):
    """按字符类型混排字体：英文字母与"【】"使用 Times New Roman，其余使用宋体。"""
    if not text:
        return

    segments = []
    buf = ""
    buf_kind = None
    for ch in text:
        kind = _is_western_char(ch)
        if buf_kind is None:
            buf_kind = kind
            buf = ch
        elif kind == buf_kind:
            buf += ch
        else:
            segments.append((buf, buf_kind))
            buf = ch
            buf_kind = kind
    if buf:
        segments.append((buf, buf_kind))

    for seg_text, is_west in segments:
        run = paragraph.add_run(seg_text)
        run.font.size = font_size
        if bold:
            run.bold = True
        rFonts = run._element.rPr.rFonts
        if is_west:
            run.font.name = "Times New Roman"
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:eastAsia"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")
        else:
            run.font.name = "宋体"
            rFonts.set(qn("w:eastAsia"), "宋体")


def _set_line_spacing_15(paragraph):
    """设置段落为 1.5 倍行距"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.line_spacing = 1.5


def _add_para_border(paragraph, color="A6A6A6"):
    """为段落添加四边框线（用于图示描述占位框）"""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        border = OxmlElement(side)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)
        pBdr.append(border)
    pPr.append(pBdr)


def _add_figure_para(doc, fig_text):
    """添加一个带边框的图示描述段落，楷体、灰色字、带边框"""
    para = doc.add_paragraph()
    _add_para_border(para)
    _set_line_spacing_15(para)
    para.paragraph_format.left_indent = Cm(1)
    para.paragraph_format.right_indent = Cm(1)
    run = para.add_run(fig_text)
    run.font.size = Pt(10)
    run.font.name = "楷体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True


def build_docx(sections, output_path):
    """
    将各题型内容构建为 Word 文档。
    sections: [(题型名, 文本内容), ...]
    """
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    section_index = 0
    section_labels = ["一", "二", "三", "四", "五", "六"]

    for idx, (q_type, content) in enumerate(sections):
        if not content:
            continue

        label = section_labels[section_index] if section_index < len(section_labels) else str(section_index + 1)
        header_text = f"{label}、{q_type}"
        header_para = doc.add_paragraph()
        _add_mixed_run(header_para, header_text, bold=True)
        _set_line_spacing_15(header_para)
        section_index += 1

        lines = content.split("\n")
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if re.match(r"^[一二三四五六七八九十]+\s*[、.．]", line_stripped):
                continue
            if line_stripped.startswith("【解析】") or line_stripped.startswith("【详解】"):
                continue

            fig_match = re.match(r"^【图[：:](.+?)】$", line_stripped)
            if fig_match:
                _add_figure_para(doc, f"（图示）{fig_match.group(1).strip()}")
                continue

            if re.search(r"【图[：:](.+?)】", line_stripped):
                parts = re.split(r"(【图[：:].+?】)", line_stripped)
                for part in parts:
                    part = part.strip()
                    if not part:
                        continue
                    fig_inner = re.match(r"^【图[：:](.+?)】$", part)
                    if fig_inner:
                        _add_figure_para(doc, f"（图示）{fig_inner.group(1).strip()}")
                    else:
                        para = doc.add_paragraph()
                        _add_mixed_run(para, part)
                        _set_line_spacing_15(para)
                        if part.startswith("【答案】"):
                            _add_shading(para, ANSWER_SHADING_COLOR)
                continue

            para = doc.add_paragraph()
            _add_mixed_run(para, line)
            _set_line_spacing_15(para)

            if line_stripped.startswith("【答案】"):
                _add_shading(para, ANSWER_SHADING_COLOR)

    doc.save(output_path)
    print(f"\n试卷已保存: {output_path}")


# ════════════════════ 用户交互 ════════════════════


def get_user_input():
    """获取用户输入的主题和各题型数量"""
    print("=" * 50)
    print("        AI 试卷生成工具")
    print("=" * 50)
    print()

    topic = input("请输入试卷主题（如：模拟电子技术-运算放大器）：").strip()
    if not topic:
        print("主题不能为空，程序退出。")
        sys.exit(0)

    print(f"\n请输入各题型的数量（输入0表示不需要该题型）：")
    counts = {}
    for qt in QUESTION_TYPES:
        while True:
            try:
                n = int(input(f"  {qt}数量: "))
                if n < 0:
                    print("    数量不能为负数，请重新输入")
                    continue
                counts[qt] = n
                break
            except ValueError:
                print("    请输入有效的整数")

    total = sum(counts.values())
    if total == 0:
        print("总题数为0，程序退出。")
        sys.exit(0)

    print(f"\n是否提供参考资料/真题示例？（可选，直接回车跳过）")
    print("  提供参考资料可帮助 AI 更好地模仿出题风格和把控难度。")
    print("  输入文件路径(.txt)或直接粘贴内容，输入空行结束：")
    ref_lines = []
    first_line = input("  > ").strip()
    if first_line:
        if os.path.isfile(first_line):
            try:
                with open(first_line, "r", encoding="utf-8") as f:
                    ref_lines = [f.read()]
            except Exception:
                ref_lines = [first_line]
        else:
            ref_lines.append(first_line)
            while True:
                line = input("  > ")
                if not line.strip():
                    break
                ref_lines.append(line)
    reference = "\n".join(ref_lines)

    return topic, counts, reference


# ════════════════════ 主程序 ════════════════════


BATCH_SIZE = 5  # 每批最多生成的题目数（小批量以确保质量）


def _generate_in_batches(client, topic, q_type, count, reference, all_summaries):
    """
    分批生成题目（每批最多 BATCH_SIZE 道），每批传入已生成的题目摘要避免重复。
    生成后自动去重，不够则补生成。
    """
    all_content_blocks = []
    remaining = count
    max_attempts = 5  # 补生成最多尝试次数

    for attempt in range(max_attempts):
        if remaining <= 0:
            break

        batch_size = min(remaining, BATCH_SIZE)
        batch_summaries = all_summaries + _extract_question_summaries("\n".join(all_content_blocks))

        print(f"    批次 {attempt + 1}: 生成 {batch_size} 题（已有摘要 {len(batch_summaries)} 条）...", flush=True)

        content = generate_questions(
            client, topic, q_type, batch_size, reference,
            existing_summaries=batch_summaries if batch_summaries else None,
        )

        if not content:
            print(f"    批次生成失败", flush=True)
            break

        # 去重
        combined = "\n".join(all_content_blocks + [content]) if all_content_blocks else content
        deduped, removed = deduplicate_questions(combined)

        if removed > 0:
            print(f"    去重: 移除 {removed} 道重复题目", flush=True)

        all_content_blocks = [deduped]
        actual_count = len(_extract_question_blocks(deduped))
        remaining = count - actual_count

        if remaining <= 0:
            break

        time.sleep(1)

    final_content = "\n".join(all_content_blocks)
    final_blocks = _extract_question_blocks(final_content)

    # 如果去重后题目不够，提示用户
    if len(final_blocks) < count:
        print(f"    ⚠ 去重后仅保留 {len(final_blocks)}/{count} 道不重复题目（主题知识点有限）", flush=True)

    return "\n".join(final_blocks[:count])


def main():
    topic, counts, reference = get_user_input()

    print(f"\n{'=' * 50}")
    print(f"主题: {topic}")
    print(f"题型分布: {', '.join(f'{k}{v}题' for k, v in counts.items() if v > 0)}")
    print(f"{'=' * 50}\n")

    client = create_client()

    # 方案1：窄主题检测，估算容量并提醒用户
    print("正在评估主题知识点容量...", flush=True)
    capacity = estimate_topic_capacity(client, topic)

    if capacity:
        print(f"  预估知识点数: {capacity.get('总知识点数', '未知')}")
        warnings = []
        for qt, cnt in counts.items():
            if cnt == 0:
                continue
            max_cap = capacity.get(qt, 999)
            if isinstance(max_cap, int) and cnt > max_cap:
                warnings.append(f"  ⚠ {qt}: 你要求 {cnt} 道，但该主题建议最多 {max_cap} 道")

        if warnings:
            print("\n主题容量警告（题目数可能超出知识点范围）：")
            for w in warnings:
                print(w)
            print()
            choice = input("是否调整数量？(y=调整 / 回车=继续生成，超出部分将自动去重): ").strip().lower()
            if choice == "y":
                print("\n请重新输入各题型数量：")
                for qt in QUESTION_TYPES:
                    if counts.get(qt, 0) > 0:
                        max_cap = capacity.get(qt, counts[qt])
                        while True:
                            try:
                                n = int(input(f"  {qt}（建议≤{max_cap}）: "))
                                if n < 0:
                                    continue
                                counts[qt] = n
                                break
                            except ValueError:
                                print("    请输入整数")
                print()
    else:
        print("  容量评估跳过（API未响应），继续生成...\n")

    sections = []
    current_num = 1
    all_summaries = []

    for q_type in QUESTION_TYPES:
        count = counts.get(q_type, 0)
        if count == 0:
            continue

        print(f"正在生成: {q_type} × {count} ...", flush=True)

        # 方案3：分批生成（大于 BATCH_SIZE 时自动分批）
        if count > BATCH_SIZE:
            content = _generate_in_batches(client, topic, q_type, count, reference, all_summaries)
        else:
            content = generate_questions(
                client, topic, q_type, count, reference,
                existing_summaries=all_summaries if all_summaries else None,
            )

        if content:
            # 方案2：自动去重
            content, removed = deduplicate_questions(content)
            if removed > 0:
                print(f"  去重: 移除 {removed} 道重复题目", flush=True)

            # AI 质量审查（答案验证、题干改进、格式修正）
            print(f"  正在进行质量审查...", flush=True)
            content = review_questions(client, content, q_type)
            print(f"  质量审查完成 ✓", flush=True)

            all_summaries.extend(_extract_question_summaries(content))
            content, current_num = renumber_questions(content, current_num)

            if q_type in ("单选题", "多选题"):
                exposure_issues = _detect_self_exposure(content)
                if exposure_issues:
                    print(f"  ⚠ 检测到自暴露问题：")
                    for q_num, kw, letter in exposure_issues:
                        print(f'    第{q_num}题：题干含正确答案({letter})关键词"{kw}"')
                    print(f"  正在重新生成以修复...")
                    content2 = generate_questions(
                        client, topic, q_type, count, reference,
                        existing_summaries=all_summaries,
                    )
                    if content2:
                        content2, _ = deduplicate_questions(content2)
                        actual_count = len(_extract_question_blocks(content))
                        content2, current_num_new = renumber_questions(content2, current_num - actual_count)
                        issues2 = _detect_self_exposure(content2)
                        if len(issues2) < len(exposure_issues):
                            content = content2
                            current_num = current_num_new
                            if issues2:
                                print(f"  重新生成后仍有{len(issues2)}题存在自暴露，请人工复查")
                            else:
                                print(f"  重新生成成功，自暴露问题已修复 ✓")
                        else:
                            print(f"  重新生成未改善，保留原版本，请人工复查标记题目")

            sections.append((q_type, content))
            print(f"  完成 ✓ （{len(_extract_question_blocks(content))} 道）")
        else:
            print(f"  生成失败，跳过该题型")

    if not sections:
        print("\n所有题型均生成失败，请检查 API 配置。")
        sys.exit(1)

    out_dir = os.environ.get("EXAM_TOOL_WORKDIR")
    if not out_dir or not os.path.isdir(out_dir):
        out_dir = os.path.dirname(os.path.abspath(__file__)) if not getattr(sys, "frozen", False) else os.path.dirname(sys.executable)
    safe_topic = re.sub(r'[\\/:*?"<>|]', "_", topic)
    output_path = os.path.join(out_dir, f"{safe_topic}_试卷.docx")

    build_docx(sections, output_path)

    print(f"\n全部完成！")
    input("按 Enter 退出...")


if __name__ == "__main__":
    main()
