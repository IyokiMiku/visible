"""批量替换已生成 DOCX 文档中的考试类型文字。

用途：
  将已生成 Word 文档正文、表格、页眉、页脚中的“对口招生”替换为“高职分类考试”。

默认行为：
  - 扫描 ``04_生成输出/生成结果`` 下所有 .docx 文件；
  - 处理前为每个将被修改的文档生成 ``.examtypebak.docx`` 备份；
  - 递归处理正文段落、表格单元格、页眉、页脚中的段落；
  - 跳过 Word 临时文件 ``~$*.docx`` 和备份文件 ``*.examtypebak.docx``。

示例：

    python 01_工具脚本/文档处理/batch_replace_exam_type_docx.py
    python 01_工具脚本/文档处理/batch_replace_exam_type_docx.py --path "04_生成输出/生成结果/重庆市 机械加工类"
    python 01_工具脚本/文档处理/batch_replace_exam_type_docx.py --path "某个文件.docx"
    python 01_工具脚本/文档处理/batch_replace_exam_type_docx.py --dry-run
    python 01_工具脚本/文档处理/batch_replace_exam_type_docx.py --no-backup
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_TARGET = PROJECT_ROOT / "04_生成输出" / "生成结果"
DEFAULT_OLD = "对口招生"
DEFAULT_NEW = "高职分类考试"
BACKUP_SUFFIX = ".examtypebak.docx"


def _iter_table_paragraphs(tables) -> Iterable:
    """递归遍历表格单元格中的段落。"""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def _iter_story_paragraphs(doc: Document) -> Iterable:
    """遍历正文、表格、页眉、页脚中的段落。"""
    yield from doc.paragraphs
    yield from _iter_table_paragraphs(doc.tables)

    for section in doc.sections:
        header = section.header
        yield from header.paragraphs
        yield from _iter_table_paragraphs(header.tables)

        first_header = section.first_page_header
        yield from first_header.paragraphs
        yield from _iter_table_paragraphs(first_header.tables)

        even_header = section.even_page_header
        yield from even_header.paragraphs
        yield from _iter_table_paragraphs(even_header.tables)

        footer = section.footer
        yield from footer.paragraphs
        yield from _iter_table_paragraphs(footer.tables)

        first_footer = section.first_page_footer
        yield from first_footer.paragraphs
        yield from _iter_table_paragraphs(first_footer.tables)

        even_footer = section.even_page_footer
        yield from even_footer.paragraphs
        yield from _iter_table_paragraphs(even_footer.tables)


def _copy_run_style(src, dst) -> None:
    """尽量保留原 run 的常见样式。"""
    dst.bold = src.bold
    dst.italic = src.italic
    dst.underline = src.underline
    dst.style = src.style
    dst.font.name = src.font.name
    dst.font.size = src.font.size
    dst.font.bold = src.font.bold
    dst.font.italic = src.font.italic
    dst.font.underline = src.font.underline
    if src.font.color and src.font.color.rgb:
        dst.font.color.rgb = src.font.color.rgb


def _replace_in_paragraph(paragraph, old: str, new: str) -> int:
    """替换段落文字；优先逐 run 替换，跨 run 时重建段落。"""
    if old not in (paragraph.text or ""):
        return 0

    count = 0
    for run in paragraph.runs:
        if old in run.text:
            count += run.text.count(old)
            run.text = run.text.replace(old, new)

    if count:
        return count

    # 极少数情况下关键词被 Word 拆成多个 run。此时重建段落文本，保留第一个 run 的基本样式。
    original_text = paragraph.text
    count = original_text.count(old)
    replacement = original_text.replace(old, new)
    first_run = paragraph.runs[0] if paragraph.runs else None
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    new_run = paragraph.add_run(replacement)
    if first_run is not None:
        _copy_run_style(first_run, new_run)
    return count


def _docx_paths(target: Path) -> list[Path]:
    if target.is_file():
        return [target] if target.suffix.lower() == ".docx" else []
    if not target.exists():
        return []
    return sorted(
        p for p in target.rglob("*.docx")
        if p.is_file()
        and not p.name.startswith("~$")
        and not p.name.endswith(BACKUP_SUFFIX)
    )


def process_docx(path: Path, old: str, new: str, dry_run: bool, backup: bool) -> int:
    doc = Document(str(path))
    count = 0
    for paragraph in _iter_story_paragraphs(doc):
        count += _replace_in_paragraph(paragraph, old, new)

    if count and not dry_run:
        if backup:
            backup_path = path.with_name(path.stem + BACKUP_SUFFIX)
            if not backup_path.exists():
                shutil.copy2(path, backup_path)
        doc.save(str(path))
    return count


def main() -> None:
    parser = argparse.ArgumentParser(description="批量替换 DOCX 中的考试类型文字")
    parser.add_argument("--path", default=str(DEFAULT_TARGET), help="要处理的 docx 文件或目录，默认 04_生成输出/生成结果")
    parser.add_argument("--old", default=DEFAULT_OLD, help="要替换的旧文本，默认：对口招生")
    parser.add_argument("--new", default=DEFAULT_NEW, help="替换后的新文本，默认：高职分类考试")
    parser.add_argument("--dry-run", action="store_true", help="只预览命中文档，不实际写入")
    parser.add_argument("--no-backup", action="store_true", help="不生成 .examtypebak.docx 备份")
    args = parser.parse_args()

    target = Path(args.path)
    if not target.is_absolute():
        target = PROJECT_ROOT / target

    paths = _docx_paths(target)
    if not paths:
        print(f"未找到 DOCX：{target}")
        return

    changed_files = 0
    total_replacements = 0
    for path in paths:
        try:
            count = process_docx(
                path,
                old=args.old,
                new=args.new,
                dry_run=args.dry_run,
                backup=not args.no_backup,
            )
        except Exception as exc:  # noqa: BLE001 - 批量处理时不中断后续文件
            print(f"[失败] {path}: {exc}")
            continue

        if count:
            changed_files += 1
            total_replacements += count
            action = "将替换" if args.dry_run else "已替换"
            print(f"[{action}] {path} ({count}处)")

    mode = "预览完成" if args.dry_run else "处理完成"
    print(f"{mode}：命中文档 {changed_files}/{len(paths)} 个，替换 {total_replacements} 处。")
    if changed_files and not args.dry_run and not args.no_backup:
        print(f"已为修改过的文档生成 *{BACKUP_SUFFIX} 备份。")


if __name__ == "__main__":
    main()
