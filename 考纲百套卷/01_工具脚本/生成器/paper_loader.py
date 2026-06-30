"""人工组卷试卷读取与拆题。

从 ``04_生成输出/组卷待质检`` 读取人工组好的 DOCX/TXT，拆分为统一题目结构，
供后续质检与按需修复使用。本模块不负责生成新题，只做读取、解析和规范化。
"""
from __future__ import annotations

from dataclasses import asdict, dataclass, field
import hashlib
import json
from pathlib import Path
import re
from typing import Any, Iterable

from .paths import MANUAL_PAPER_DIR
from .planning import PaperPlan

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


QUESTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+\s*[、.．]\s*(.+)")
QUESTION_NO_RE = re.compile(r"^(\d+)\s*[.、．)）]\s*(.*)")
DUPLICATE_QUESTION_NO_RE = re.compile(r"^(\d+)\s*[.、．)）]\s*\1\s*[.、．)）]\s*(.*)")
OPTION_RE = re.compile(r"^([A-H])\s*[.、．)]\s*(.+)", re.I)
INLINE_OPTION_RE = re.compile(r"(?<=[^\w])([A-H])\s*[.、．)]\s*", re.I)
ANSWER_RE = re.compile(r"^【?\s*(答案|参考答案)\s*】?[:：]?\s*(.*)")
ANALYSIS_RE = re.compile(r"^【?\s*(解析|答案解析|解题思路|详解|试题解析|题目解析|分析)\s*】?[:：]?\s*(.*)")

QUESTION_TYPE_KEYWORDS = [
    ("单项选择", "单选题"),
    ("单选", "单选题"),
    ("多项选择", "多选题"),
    ("多选", "多选题"),
    ("判断", "判断题"),
    ("填空", "填空题"),
    ("简答", "简答题"),
    ("计算", "计算题"),
    ("综合", "综合题"),
    ("分析", "分析题"),
    ("作图", "作图题"),
    ("识图", "识图题"),
]


@dataclass
class LoadedPaper:
    """人工组卷试卷及拆题结果。"""

    path: Path
    paper_label: str = ""
    title: str = ""
    questions: list[dict[str, Any]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass
class _DocxImageRef:
    """DOCX 段落中的真实图片引用，用于跨卷图片素材查重。"""

    paragraph_index: int
    rel_id: str
    target_ref: str = ""
    sha256: str = ""
    byte_size: int = 0
    width_emu: int | None = None
    height_emu: int | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class _DocxLineRecord:
    """DOCX 可见行及其原始段落/表格位置信息。"""

    text: str
    paragraph_index: int | str | None = None
    has_image: bool = False
    image_placeholder: bool = False
    image_refs: list[_DocxImageRef] = field(default_factory=list)
    is_table: bool = False
    table_index: int | None = None
    body_index: int | None = None
    has_protected_content: bool = False


@dataclass
class _QuestionDraft:
    question_no: int
    question_type: str = ""
    heading: str = ""
    source_docx_path: str = ""
    stem_lines: list[str] = field(default_factory=list)
    options: list[str] = field(default_factory=list)
    answer_lines: list[str] = field(default_factory=list)
    analysis_lines: list[str] = field(default_factory=list)
    raw_lines: list[str] = field(default_factory=list)
    source_paragraph_indices: list[int] = field(default_factory=list)
    source_body_indices: list[int] = field(default_factory=list)
    stem_paragraph_indices: list[int] = field(default_factory=list)
    answer_paragraph_indices: list[int] = field(default_factory=list)
    analysis_paragraph_indices: list[int] = field(default_factory=list)
    option_table_indices: list[int] = field(default_factory=list)
    has_stem_image: bool = False
    has_answer_image: bool = False
    has_analysis_image: bool = False
    has_protected_content: bool = False
    stem_image_refs: list[dict[str, Any]] = field(default_factory=list)
    answer_image_refs: list[dict[str, Any]] = field(default_factory=list)
    analysis_image_refs: list[dict[str, Any]] = field(default_factory=list)

    def _unique_indices(self, values: list[int]) -> list[int]:
        return list(dict.fromkeys(values))

    def to_question(self) -> dict[str, Any]:
        protected_reasons = []
        if self.has_stem_image:
            protected_reasons.append("stem_visual_object")
        if self.has_answer_image:
            protected_reasons.append("answer_visual_object")
        if self.has_analysis_image:
            protected_reasons.append("analysis_visual_object")
        if self.has_protected_content and not protected_reasons:
            protected_reasons.append("body_visual_object")
        protected = bool(self.source_docx_path and protected_reasons)
        image_refs = {
            "stem": self.stem_image_refs,
            "answer": self.answer_image_refs,
            "analysis": self.analysis_image_refs,
        }
        original_docx_image_refs = [
            ref
            for refs in image_refs.values()
            for ref in refs
        ]
        return {
            "question_no": self.question_no,
            "question_type": self.question_type,
            "heading": self.heading,
            "stem": _join_lines(self.stem_lines),
            "options": self.options,
            "answer": _join_lines(self.answer_lines),
            "analysis": _join_lines(self.analysis_lines),
            "knowledge_points": [],
            "difficulty": "",
            "source_path": "",
            "status": "loaded",
            "issues": [],
            "raw_text": _join_lines(self.raw_lines),
            "source_docx_path": self.source_docx_path,
            "source_paragraph_indices": self._unique_indices(self.source_paragraph_indices),
            "source_body_indices": self._unique_indices(self.source_body_indices),
            "original_docx_body_indices": self._unique_indices(self.source_body_indices),
            "original_docx_body_range": (
                {
                    "start": min(self.source_body_indices),
                    "end": max(self.source_body_indices),
                }
                if self.source_body_indices else None
            ),
            "stem_paragraph_indices": self._unique_indices(self.stem_paragraph_indices),
            "answer_paragraph_indices": self._unique_indices(self.answer_paragraph_indices),
            "analysis_paragraph_indices": self._unique_indices(self.analysis_paragraph_indices),
            "option_table_indices": self._unique_indices(self.option_table_indices),
            "image_flags": {
                "stem": self.has_stem_image,
                "answer": self.has_answer_image,
                "analysis": self.has_analysis_image,
            },
            "image_refs": image_refs,
            "has_original_docx_images": bool(original_docx_image_refs),
            "original_docx_image_refs": original_docx_image_refs,
            "protected_original_docx_block": protected,
            "protection_reason": ",".join(protected_reasons) if protected else "",
        }


def _clean_line(text: str) -> str:
    text = str(text or "").strip()
    text = re.sub(r"(\d)．(\d)", r"\1.\2", text)  # 全角小数点 → 半角
    text = re.sub(r"[ \t　]+", " ", text)
    duplicate = DUPLICATE_QUESTION_NO_RE.match(text)
    if duplicate:
        text = f"{duplicate.group(1)}. {duplicate.group(2).strip()}".rstrip()
    return text


def _join_lines(lines: Iterable[str]) -> str:
    return "\n".join(line for line in (_clean_line(line) for line in lines) if line).strip()


def _strip_analysis_label(text: str) -> str:
    """Remove duplicated analysis labels; DOCX generation adds the final label."""
    value = str(text or "").strip()
    while True:
        cleaned = re.sub(
            r"^(?:【\s*(?:解析|详解|分析|答案解析|试题解析|题目解析|解题思路)\s*】|(?:解析|详解|分析)\s*[:：])\s*",
            "",
            value,
        )
        if cleaned == value:
            return cleaned.strip()
        value = cleaned.strip()


def _element_has_protected_content(element) -> bool:
    """检测任意 DOCX body XML 块中是否包含图片/公式/OLE 等受保护对象。"""
    xml = getattr(element, "xml", "") or ""
    return any(
        tag in xml
        for tag in (
            "<wp:inline",
            "<wp:anchor",
            "<pic:pic",
            "<a:blip",
            "nvPicPr",
            "<w:pict",
            "<v:imagedata",
            "<w:object",
            "OLEObject",
            "Equation.DSMT4",
            "<m:oMath",
            "<m:oMathPara",
            "AlternateContent",
        )
    )


def _paragraph_has_image(paragraph) -> bool:
    """检测段落 XML 中是否包含受保护视觉对象。"""
    return _element_has_protected_content(paragraph._element)


def _paragraph_image_refs(paragraph, paragraph_index: int) -> list[_DocxImageRef]:
    """提取段落中嵌入图片/对象的关系 ID、字节 hash 与尺寸。"""
    refs: list[_DocxImageRef] = []
    seen_rel_ids: set[str] = set()

    def add_ref(rel_id: str, target_ref: str = "", blob: bytes = b"", width_emu: int | None = None, height_emu: int | None = None) -> None:
        if not rel_id or rel_id in seen_rel_ids:
            return
        seen_rel_ids.add(rel_id)
        refs.append(_DocxImageRef(
            paragraph_index=paragraph_index,
            rel_id=rel_id,
            target_ref=target_ref,
            sha256=hashlib.sha256(blob).hexdigest() if blob else "",
            byte_size=len(blob),
            width_emu=width_emu,
            height_emu=height_emu,
        ))

    try:
        blips = paragraph._element.xpath('.//*[local-name()="blip"]')
    except Exception:
        blips = []
    for blip in blips:
        rel_id = ""
        for key, value in getattr(blip, "attrib", {}).items():
            if key.endswith("}embed") or key.endswith("}link") or key in {"r:embed", "r:link"}:
                rel_id = value
                break
        if not rel_id:
            continue
        try:
            part = paragraph.part.related_parts.get(rel_id)
        except Exception:
            part = None
        blob = getattr(part, "blob", b"") or b""
        target_ref = str(getattr(part, "partname", "") or getattr(part, "target_ref", ""))
        width_emu: int | None = None
        height_emu: int | None = None
        try:
            parents = blip.xpath("ancestor::*[local-name()='inline' or local-name()='anchor']")
            if parents:
                extents = parents[0].xpath('.//*[local-name()="extent"]')
                if extents:
                    cx = extents[0].get("cx")
                    cy = extents[0].get("cy")
                    width_emu = int(cx) if cx else None
                    height_emu = int(cy) if cy else None
        except Exception:
            width_emu = None
            height_emu = None
        add_ref(rel_id, target_ref, blob, width_emu, height_emu)

    try:
        v_imagedata = paragraph._element.xpath('.//*[local-name()="imagedata"]')
    except Exception:
        v_imagedata = []
    for img in v_imagedata:
        rel_id = img.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") or img.get("r:id") or img.get("id") or ""
        if not rel_id:
            continue
        try:
            part = paragraph.part.related_parts.get(rel_id)
        except Exception:
            part = None
        blob = getattr(part, "blob", b"") or b""
        target_ref = str(getattr(part, "partname", "") or getattr(part, "target_ref", ""))
        add_ref(rel_id, target_ref, blob)

    try:
        ole_objects = paragraph._element.xpath('.//*[local-name()="OLEObject"]')
    except Exception:
        ole_objects = []
    for ole in ole_objects:
        rel_id = ole.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id") or ole.get("r:id") or ole.get("id") or ""
        if not rel_id:
            continue
        try:
            part = paragraph.part.related_parts.get(rel_id)
        except Exception:
            part = None
        blob = getattr(part, "blob", b"") or b""
        target_ref = str(getattr(part, "partname", "") or getattr(part, "target_ref", ""))
        add_ref(rel_id, target_ref, blob)

    return refs


def _paragraph_records_from_docx(path: Path) -> list[_DocxLineRecord]:
    if Document is None:
        raise RuntimeError("缺少 python-docx，请先安装 DOCX 读取依赖。")

    doc = Document(str(path))
    records: list[_DocxLineRecord] = []
    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    paragraph_by_element = {id(paragraph._element): (idx, paragraph) for idx, paragraph in enumerate(paragraphs)}
    table_by_element = {id(table._element): (idx, table) for idx, table in enumerate(tables)}

    for body_index, element in enumerate(doc.element.body.iterchildren()):
        paragraph_entry = paragraph_by_element.get(id(element))
        if paragraph_entry is not None:
            idx, paragraph = paragraph_entry
            text = _clean_line(paragraph.text)
            image_refs = _paragraph_image_refs(paragraph, idx)
            has_image = bool(image_refs) or _paragraph_has_image(paragraph)
            has_protected = has_image or _element_has_protected_content(element)
            if text:
                records.append(_DocxLineRecord(
                    text=text,
                    paragraph_index=idx,
                    has_image=has_image,
                    image_refs=image_refs,
                    body_index=body_index,
                    has_protected_content=has_protected,
                ))
            if has_image:
                records.append(_DocxLineRecord(
                    text="[图片]",
                    paragraph_index=idx,
                    has_image=True,
                    image_placeholder=True,
                    image_refs=image_refs,
                    body_index=body_index,
                    has_protected_content=has_protected,
                ))
            continue

        table_entry = table_by_element.get(id(element))
        if table_entry is not None:
            table_idx, table = table_entry
            table_protected = _element_has_protected_content(element)
            emitted = False
            for row in table.rows:
                cells = [_clean_line(cell.text) for cell in row.cells]
                text = " ".join(cell for cell in cells if cell)
                if text:
                    emitted = True
                    records.append(_DocxLineRecord(
                        text=text,
                        is_table=True,
                        table_index=table_idx,
                        body_index=body_index,
                        has_image=table_protected,
                        has_protected_content=table_protected,
                    ))
            if table_protected and not emitted:
                records.append(_DocxLineRecord(
                    text="[图片]",
                    has_image=True,
                    image_placeholder=True,
                    is_table=True,
                    table_index=table_idx,
                    body_index=body_index,
                    has_protected_content=True,
                ))
    return records


def _paragraph_lines_from_docx(path: Path) -> list[str]:
    return [record.text for record in _paragraph_records_from_docx(path)]


def _lines_from_txt(path: Path) -> list[str]:
    return [_clean_line(line) for line in path.read_text(encoding="utf-8").splitlines() if _clean_line(line)]


def read_paper_lines(path: str | Path) -> list[str]:
    """读取人工组卷文件的可见文本行。"""
    paper_path = Path(path)
    suffix = paper_path.suffix.lower()
    if suffix == ".docx":
        return _paragraph_lines_from_docx(paper_path)
    if suffix in {".txt", ".md"}:
        return _lines_from_txt(paper_path)
    raise ValueError(f"暂不支持的人工组卷文件类型：{paper_path.suffix}")


def infer_question_type(text: str) -> str:
    """从大题标题推断统一题型名称。"""
    compact = re.sub(r"\s+", "", text or "")
    for keyword, question_type in QUESTION_TYPE_KEYWORDS:
        if keyword in compact:
            return question_type
    return ""


def _extract_title(lines: list[str]) -> str:
    for line in lines[:10]:
        if not QUESTION_HEADING_RE.match(line) and not QUESTION_NO_RE.match(line):
            return line
    return ""


def _append_to_current_part(question: _QuestionDraft, part: str, line: str) -> str:
    answer_match = ANSWER_RE.match(line)
    if answer_match:
        content = answer_match.group(2).strip()
        if content:
            question.answer_lines.append(content)
        return "answer"

    analysis_match = ANALYSIS_RE.match(line)
    if analysis_match:
        content = _strip_analysis_label(analysis_match.group(2).strip())
        if content:
            question.analysis_lines.append(content)
        return "analysis"

    option_match = OPTION_RE.match(line)
    if option_match and part not in {"answer", "analysis"}:
        label = option_match.group(1).upper()
        content = option_match.group(2).strip()
        # 检查同一行内是否还有后续选项（如 "A. xxx B．xxx C．xxx D．xxx"）
        subsequent = [(m.group(1).upper(), m.start(), m.end()) for m in INLINE_OPTION_RE.finditer(content)]
        if subsequent and subsequent[0][1] > 0:
            pieces: list[str] = []
            next_start = len(content)
            for sub_label, sub_start, sub_end in reversed(subsequent):
                piece = content[sub_end:next_start].strip()
                if piece:
                    pieces.append(f"{sub_label}. {piece}")
                next_start = sub_start
            if next_start > 0:
                first_piece = content[:next_start].strip()
                if first_piece:
                    pieces.append(f"{label}. {first_piece}")
            pieces.reverse()
            question.options.extend(pieces)
        else:
            question.options.append(f"{label}. {content}")
        return "stem"

    if part == "answer":
        question.answer_lines.append(line)
    elif part == "analysis":
        question.analysis_lines.append(line)
    else:
        question.stem_lines.append(line)
    return part


def _record_question_source(question: _QuestionDraft, part: str, record: _DocxLineRecord | None) -> None:
    if record is None:
        return
    if record.body_index is not None:
        question.source_body_indices.append(record.body_index)
    if record.has_protected_content:
        question.has_protected_content = True
    if record.is_table:
        if part == "stem" and record.table_index is not None:
            question.option_table_indices.append(record.table_index)
        return
    if record.paragraph_index is None:
        return
    idx = record.paragraph_index
    question.source_paragraph_indices.append(idx)
    image_refs = [ref.to_dict() for ref in record.image_refs]
    if part == "answer":
        question.answer_paragraph_indices.append(idx)
        if record.has_image:
            question.has_answer_image = True
            question.answer_image_refs.extend(image_refs)
    elif part == "analysis":
        question.analysis_paragraph_indices.append(idx)
        if record.has_image:
            question.has_analysis_image = True
            question.analysis_image_refs.extend(image_refs)
    else:
        question.stem_paragraph_indices.append(idx)
        if record.has_image:
            question.has_stem_image = True
            question.stem_image_refs.extend(image_refs)


def parse_questions_from_records(records: list[_DocxLineRecord], source_path: str | Path = "") -> tuple[list[dict[str, Any]], list[str]]:
    """将 DOCX 可见行拆成题目，并记录图片所在题目分区。"""
    questions: list[dict[str, Any]] = []
    warnings: list[str] = []
    current_heading = ""
    current_type = ""
    current: _QuestionDraft | None = None
    current_part = "stem"
    source_docx_path = str(source_path) if source_path else ""

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return
        question = current.to_question()
        question["source_path"] = source_docx_path
        if not question["stem"] and question["options"]:
            warnings.append(f"第{question['question_no']}题未解析到题干。")
        questions.append(question)
        current = None

    for record in records:
        line = _clean_line(record.text)
        if not line:
            continue

        heading_match = QUESTION_HEADING_RE.match(line)
        if heading_match:
            inferred_type = infer_question_type(heading_match.group(1))
            if inferred_type:
                flush_current()
                current_heading = line
                current_type = inferred_type
                current_part = "stem"
                continue

        question_match = QUESTION_NO_RE.match(line)
        if question_match:
            flush_current()
            current = _QuestionDraft(
                question_no=int(question_match.group(1)),
                question_type=current_type,
                heading=current_heading,
                source_docx_path=source_docx_path,
                raw_lines=[line],
            )
            current_part = "stem"
            _record_question_source(current, current_part, record)
            remainder = question_match.group(2).strip()
            if remainder:
                current_part = _append_to_current_part(current, current_part, remainder)
            continue

        if current is None:
            continue

        current.raw_lines.append(line)
        before_part = current_part
        current_part = _append_to_current_part(current, current_part, line)
        label_part = current_part if (ANSWER_RE.match(line) or ANALYSIS_RE.match(line)) else before_part
        _record_question_source(current, label_part, record)

    flush_current()

    if not questions:
        warnings.append("未解析到任何题目，请检查题号是否使用 1.、1、 或 1）等格式。")

    for idx, question in enumerate(questions, 1):
        if question["question_no"] != idx:
            warnings.append(f"题号不连续或顺序异常：解析到第{question['question_no']}题，位置为第{idx}题。")

    return questions, warnings


def parse_questions_from_lines(lines: list[str], source_path: str | Path = "") -> tuple[list[dict[str, Any]], list[str]]:
    """将试卷文本行拆成题目字典列表。"""
    questions: list[dict[str, Any]] = []
    warnings: list[str] = []
    current_heading = ""
    current_type = ""
    current: _QuestionDraft | None = None
    current_part = "stem"

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return
        question = current.to_question()
        question["source_path"] = str(source_path) if source_path else ""
        if not question["stem"] and question["options"]:
            warnings.append(f"第{question['question_no']}题未解析到题干。")
        questions.append(question)
        current = None

    for raw_line in lines:
        line = _clean_line(raw_line)
        if not line:
            continue

        heading_match = QUESTION_HEADING_RE.match(line)
        if heading_match:
            inferred_type = infer_question_type(heading_match.group(1))
            if inferred_type:
                flush_current()
                current_heading = line
                current_type = inferred_type
                current_part = "stem"
                continue

        question_match = QUESTION_NO_RE.match(line)
        if question_match:
            flush_current()
            current = _QuestionDraft(
                question_no=int(question_match.group(1)),
                question_type=current_type,
                heading=current_heading,
                raw_lines=[line],
            )
            remainder = question_match.group(2).strip()
            current_part = "stem"
            if remainder:
                current_part = _append_to_current_part(current, current_part, remainder)
            continue

        if current is None:
            continue

        current.raw_lines.append(line)
        current_part = _append_to_current_part(current, current_part, line)

    flush_current()

    if not questions:
        warnings.append("未解析到任何题目，请检查题号是否使用 1.、1、 或 1）等格式。")

    for idx, question in enumerate(questions, 1):
        if question["question_no"] != idx:
            warnings.append(f"题号不连续或顺序异常：解析到第{question['question_no']}题，位置为第{idx}题。")

    return questions, warnings


def load_manual_paper(path: str | Path, paper: PaperPlan | None = None) -> LoadedPaper:
    """读取一个人工组卷文件并拆题。"""
    paper_path = Path(path)
    if paper_path.suffix.lower() == ".docx":
        records = _paragraph_records_from_docx(paper_path)
        lines = [record.text for record in records]
        questions, warnings = parse_questions_from_records(records, paper_path)
    else:
        lines = read_paper_lines(paper_path)
        questions, warnings = parse_questions_from_lines(lines, paper_path)
    paper_label = paper.paper_label if paper else _paper_label_from_filename(paper_path)
    return LoadedPaper(
        path=paper_path,
        paper_label=paper_label,
        title=_extract_title(lines),
        questions=questions,
        warnings=warnings,
    )


def _paper_label_from_filename(path: Path) -> str:
    match = re.search(r"第\s*(\d+)\s*卷", path.stem)
    return f"第{int(match.group(1))}卷" if match else ""


def _candidate_score(path: Path, paper: PaperPlan) -> int:
    name = path.stem
    score = 0
    if paper.paper_label in name:
        score += 100
    if f"第{paper.paper_no}卷" in name or f"第 {paper.paper_no} 卷" in name:
        score += 80
    if paper.paper_type and paper.paper_type in name:
        score += 20
    if paper.topic and paper.topic in name:
        score += 10
    if paper.point_name and paper.point_name in name:
        score += 10
    if paper.module and paper.module in name:
        score += 5
    if paper.meta:
        if paper.meta.province and paper.meta.province in name:
            score += 5
        if paper.meta.exam_category and paper.meta.exam_category in str(path):
            score += 5
    suffix = path.suffix.lower()
    if "_repaired" in name:
        score += 250
    if suffix == ".txt" and "_repaired" in name:
        # 兼容旧版本文本修复产物；DOCX 修复产物同样应优先作为最终排版来源。
        score += 200
    elif suffix == ".txt":
        score += 150
    return score


def _is_allowed_manual_source(path: Path) -> bool:
    """Return whether a pending manual paper file can be treated as source material."""
    name = path.name
    if name.startswith("~$"):
        return False
    artifact_markers = (
        "questions_repaired",
        "（待人工审核）",
        "(待人工审核)",
        "考纲百套卷（解析版）",
        "考纲百套卷（原卷版）",
        "考纲百套卷(解析版)",
        "考纲百套卷(原卷版)",
    )
    return not any(marker in name for marker in artifact_markers)


def find_manual_paper_for_plan(
    paper: PaperPlan,
    search_dir: str | Path = MANUAL_PAPER_DIR,
    supported_suffixes: set[str] | None = None,
) -> Path | None:
    """在组卷待质检目录中按卷号查找人工组卷文件。"""
    base_dir = Path(search_dir)
    if not base_dir.exists():
        return None

    supported_suffixes = supported_suffixes or {".docx", ".txt", ".md"}
    candidates = [
        path
        for path in base_dir.rglob("*")
        if path.is_file()
        and path.suffix.lower() in supported_suffixes
        and _is_allowed_manual_source(path)
        and paper.paper_label in path.stem
    ]
    if not candidates:
        candidates = [
            path
            for path in base_dir.rglob("*")
            if path.is_file()
            and path.suffix.lower() in supported_suffixes
            and _is_allowed_manual_source(path)
            and re.search(rf"第\s*{paper.paper_no}\s*卷", path.stem)
        ]
    if not candidates:
        return None

    candidates.sort(key=lambda path: (_candidate_score(path, paper), -len(str(path))), reverse=True)
    return candidates[0]


def load_manual_paper_for_plan(
    paper: PaperPlan,
    search_dir: str | Path = MANUAL_PAPER_DIR,
) -> LoadedPaper | None:
    """查找并读取某卷对应的人工组卷文件。"""
    path = find_manual_paper_for_plan(paper, search_dir)
    if path is None:
        return None
    return load_manual_paper(path, paper)


def _strip_json_fence(text: str) -> str:
    text = str(text or "").strip()
    fence = re.match(r"^```(?:json)?\s*(.*?)\s*```$", text, flags=re.S | re.I)
    return fence.group(1).strip() if fence else text


def parse_generated_paper_output(output: dict[str, Any] | str | list[Any]) -> dict[str, Any]:
    """从模型整卷输出中提取试卷 JSON 对象。"""
    if isinstance(output, dict):
        data = output
    elif isinstance(output, list):
        data = {"title": "", "questions": output}
    else:
        text = _strip_json_fence(str(output or ""))
        candidates = [text]
        start, end = text.find("{"), text.rfind("}")
        if start >= 0 and end > start:
            candidates.append(text[start : end + 1])

        data = None
        for candidate in candidates:
            try:
                data = json.loads(candidate)
                break
            except json.JSONDecodeError:
                continue
        if data is None:
            raise ValueError("AI 生成结果不是可解析的 JSON 对象")

    if isinstance(data, list):
        data = {"title": "", "questions": data}
    if not isinstance(data, dict):
        raise ValueError("AI 生成结果必须是 JSON 对象")
    if isinstance(data.get("paper"), dict) and not data.get("questions"):
        data = data["paper"]
    if not isinstance(data.get("questions"), list) or not data["questions"]:
        raise ValueError("AI 生成结果缺少非空 questions 列表")
    return data


def _normalize_generated_options(options: Any) -> list[str]:
    if isinstance(options, dict):
        result = []
        for label in sorted(options):
            value = str(options[label] or "").strip()
            if value:
                result.append(f"{str(label).upper()[:1]}. {value}")
        return result

    result = []
    for index, option in enumerate(options or []):
        text = str(option or "").strip()
        if not text:
            continue
        if OPTION_RE.match(text):
            result.append(text)
        else:
            label = chr(ord("A") + index)
            result.append(f"{label}. {text}")
    return result


def _normalize_generated_question(raw: dict[str, Any], index: int, source_path: str | Path = "") -> dict[str, Any]:
    question_no = raw.get("question_no") or raw.get("no") or raw.get("number") or index
    try:
        question_no = int(question_no)
    except Exception as exc:
        raise ValueError(f"第{index}题题号不是数字：{question_no}") from exc

    stem = str(raw.get("stem") or raw.get("question") or raw.get("content") or "").strip()
    answer = str(raw.get("answer") or raw.get("答案") or "").strip()
    analysis = str(raw.get("analysis") or raw.get("解析") or raw.get("explanation") or "").strip()
    if not stem:
        raise ValueError(f"第{question_no}题缺少题干")
    if not answer:
        raise ValueError(f"第{question_no}题缺少答案")
    if not analysis:
        raise ValueError(f"第{question_no}题缺少解析")

    knowledge_points = raw.get("knowledge_points") or raw.get("knowledge_point") or []
    if isinstance(knowledge_points, str):
        knowledge_points = [item.strip() for item in re.split(r"[、,，;；\n]+", knowledge_points) if item.strip()]

    return {
        "question_no": question_no,
        "question_type": str(raw.get("question_type") or raw.get("type") or "").strip(),
        "heading": str(raw.get("heading") or "").strip(),
        "stem": stem,
        "options": _normalize_generated_options(raw.get("options")),
        "answer": answer,
        "analysis": analysis,
        "knowledge_points": knowledge_points if isinstance(knowledge_points, list) else [],
        "difficulty": str(raw.get("difficulty") or "").strip(),
        "source_path": str(source_path) if source_path else "",
        "status": "ai_generated",
        "issues": [],
        "raw_text": str(raw.get("raw_text") or "").strip(),
    }


def load_generated_paper_json(
    output: dict[str, Any] | str | list[Any],
    paper: PaperPlan,
    source_path: str | Path = "",
) -> LoadedPaper:
    """将 AI 直接生成的整卷 JSON 转成 LoadedPaper。"""
    data = parse_generated_paper_output(output)
    questions = [
        _normalize_generated_question(raw, index, source_path)
        for index, raw in enumerate(data["questions"], 1)
        if isinstance(raw, dict)
    ]
    if len(questions) != len(data["questions"]):
        raise ValueError("AI 生成结果 questions 中存在非对象题目")
    if not questions:
        raise ValueError("AI 生成结果未包含有效题目")

    warnings: list[str] = []
    for idx, question in enumerate(questions, 1):
        if question["question_no"] != idx:
            warnings.append(f"题号不连续或顺序异常：解析到第{question['question_no']}题，位置为第{idx}题。")

    blueprint_numbers = [row.question_no for row in paper.blueprint_rows if row.question_no is not None]
    if blueprint_numbers and [question["question_no"] for question in questions] != blueprint_numbers:
        warnings.append("AI 生成题号与细目表题号不完全一致，请重点复核。")

    return LoadedPaper(
        path=Path(source_path) if source_path else Path("AI直接生成"),
        paper_label=paper.paper_label,
        title=str(data.get("title") or "").strip(),
        questions=questions,
        warnings=warnings,
    )


def loaded_paper_to_dict(loaded_paper: LoadedPaper) -> dict[str, Any]:
    """将 LoadedPaper 转成可 JSON 序列化字典。"""
    data = asdict(loaded_paper)
    data["path"] = str(loaded_paper.path)
    return data
