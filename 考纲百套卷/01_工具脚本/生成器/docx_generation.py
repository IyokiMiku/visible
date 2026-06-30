"""DOCX 生成。"""
import copy
import hashlib
import os
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from .paths import BASE_DIR, TEMPLATE_PATH
from .text_processing import _clean_paper_text
from .paper_loader import _strip_analysis_label

COMMON_SCRIPT_NAMES = ("class.py", "answer2none.py", "zip.py", "fix.py")
RESIDUAL_FORMULA_RE = re.compile(r"\{\{?math:|\\\(|\\\)|\$[^$]+\$|\\(?:frac|times|Omega|omega|rho|Phi|Delta|mu|sqrt)\b")
A_BLIP = qn("a:blip")
R_EMBED = qn("r:embed")
R_ID = qn("r:id")
W_SECTPR = qn("w:sectPr")
NS_VML = "urn:schemas-microsoft-com:vml"
NS_OFFICE = "urn:schemas-microsoft-com:office:office"


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def _find_residual_formula_text(doc):
    """扫描 DOCX 可见文本中的公式标记/LaTeX 残留，防止乱码进入成品。"""
    hits = []
    paragraphs = list(doc.paragraphs) + list(_iter_table_paragraphs(doc.tables))
    for paragraph in paragraphs:
        text = paragraph.text or ""
        if RESIDUAL_FORMULA_RE.search(text):
            hits.append(text.strip()[:120])
    return hits


def _make_unique_partname(dst_package, original_partname):
    name = str(original_partname)
    base, ext = os.path.splitext(name)
    existing = {str(part.partname) for part in dst_package.iter_parts()}
    if name not in existing:
        return PackURI(name)
    counter = 1
    while True:
        candidate = f"{base}_c{counter}{ext}"
        if candidate not in existing:
            return PackURI(candidate)
        counter += 1


EMU_PER_CM = 360000
_OPTION_IMAGE_MAX_W_CM = 5.0
_OPTION_IMAGE_MAX_H_CM = 4.0
_STEM_IMAGE_MAX_W_CM = 12.0
_STEM_IMAGE_MAX_H_CM = 7.0


def _shape_extents(paragraph_elem) -> list[tuple[int, int]]:
    extents = []
    try:
        extent_elems = paragraph_elem.xpath('.//*[local-name()="extent"]')
    except Exception:
        extent_elems = []
    for extent in extent_elems:
        try:
            cx = int(extent.get('cx') or 0)
            cy = int(extent.get('cy') or 0)
        except (TypeError, ValueError):
            continue
        if cx > 0 and cy > 0:
            extents.append((cx, cy))
    return extents


def _looks_like_image_option_paragraph(paragraph_elem) -> bool:
    text = re.sub(r"\s+", "", _paragraph_visible_text(paragraph_elem) or "")
    if not re.match(r"^[A-H][.．、)]?$", text, flags=re.I):
        return False
    extents = _shape_extents(paragraph_elem)
    if not extents:
        return False
    return len(extents) == 1 or any(cx / EMU_PER_CM >= 5.5 or cy / EMU_PER_CM >= 4.5 for cx, cy in extents)


def _image_cap_for_copied_paragraph(paragraph_elem, default_w_cm: float, default_h_cm: float) -> tuple[float, float]:
    if _looks_like_image_option_paragraph(paragraph_elem):
        return _OPTION_IMAGE_MAX_W_CM, _OPTION_IMAGE_MAX_H_CM
    return default_w_cm, default_h_cm


def _remap_embedded_relationships(element, src_doc, dst_doc, rid_cache):
    src_part = src_doc.part
    dst_part = dst_doc.part
    dst_package = dst_part.package
    refs = []
    try:
        blips = element.xpath('.//*[local-name()="blip"]')
    except Exception:
        blips = []
    for blip in blips:
        old_rid = blip.get(R_EMBED)
        if old_rid:
            refs.append((blip, R_EMBED, old_rid))
    try:
        imagedata_elems = element.xpath('.//*[local-name()="imagedata"]')
    except Exception:
        imagedata_elems = []
    for imgdata in imagedata_elems:
        old_rid = imgdata.get(R_ID)
        if old_rid:
            refs.append((imgdata, R_ID, old_rid))
    try:
        ole_elems = element.xpath('.//*[local-name()="OLEObject"]')
    except Exception:
        ole_elems = []
    for ole in ole_elems:
        old_rid = ole.get(R_ID)
        if old_rid:
            refs.append((ole, R_ID, old_rid))

    for elem, attr, old_rid in refs:
        if old_rid in rid_cache:
            elem.set(attr, rid_cache[old_rid])
            continue
        if old_rid not in src_part.rels:
            continue
        rel = src_part.rels[old_rid]
        src_target = rel.target_part
        unique_name = _make_unique_partname(dst_package, src_target.partname)
        new_part = Part(unique_name, src_target.content_type, src_target.blob, dst_package)
        new_rid = dst_part.relate_to(new_part, rel.reltype)
        rid_cache[old_rid] = new_rid
        elem.set(attr, new_rid)


def _append_body_element(doc, element):
    body = doc.element.body
    sect_pr = body.find(W_SECTPR)
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)


# -- 受保护视觉/数学标记，含这些东西的段落应保持原样 --
# 注意：mc:(AlternateContent) 的命名空间前缀未在 python-docx nsmap 中注册，
# 需手动拼 Clark 名: {urn}AlternateContent
_NS_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_PROTECTED_CONTENT_TAGS = (
    qn('w:drawing'), qn('w:pict'), qn('w:object'),
    qn('m:oMath'), qn('m:oMathPara'),
    f"{_NS_MC}AlternateContent",
)


def _paragraph_has_protected_content(element) -> bool:
    """检测段落 XML 元素中是否含图片/公式/OLE/备用内容等不可覆写的视觉对象。"""
    for tag in _PROTECTED_CONTENT_TAGS:
        if element.find(f'.//{tag}') is not None:
            return True
    return False


def _set_run_text_style(r_elem, font_name: str, color_hex: str | None = None, *, bold: bool = False, size_half_points: str = '21') -> None:
    """Set run font/color/size directly on a copied XML run."""
    from lxml import etree

    rpr = r_elem.find(qn('w:rPr'))
    if rpr is None:
        rpr = etree.Element(qn('w:rPr'))
        r_elem.insert(0, rpr)

    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn('w:rFonts'))
    western = 'Times New Roman' if font_name == '宋体' else font_name
    rfonts.set(qn('w:ascii'), western)
    rfonts.set(qn('w:hAnsi'), western)
    rfonts.set(qn('w:eastAsia'), font_name)

    for tag in (qn('w:b'), qn('w:bCs')):
        for elem in list(rpr.findall(tag)):
            rpr.remove(elem)
    if bold:
        etree.SubElement(rpr, qn('w:b'))

    for tag in (qn('w:sz'), qn('w:szCs')):
        elem = rpr.find(tag)
        if elem is None:
            elem = etree.SubElement(rpr, tag)
        elem.set(qn('w:val'), size_half_points)

    if color_hex is not None:
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = etree.SubElement(rpr, qn('w:color'))
        color_el.set(qn('w:val'), color_hex)


def _split_label_run_if_needed(r_elem, label_tags: list[str]) -> list:
    """Split a single text run like '【答案】A' so label/content can use different fonts."""
    from copy import deepcopy

    text_nodes = r_elem.findall(qn('w:t'))
    if len(text_nodes) != 1:
        return [r_elem]
    text = text_nodes[0].text or ''
    label = next((tag for tag in label_tags if tag and tag in text), '')
    if not label or text == label:
        return [r_elem]

    before, after_label = text.split(label, 1)
    new_runs = []
    anchor = r_elem

    if before:
        text_nodes[0].text = before
        new_runs.append(r_elem)
        label_run = deepcopy(r_elem)
        label_run.find(qn('w:t')).text = label
        anchor.addnext(label_run)
        anchor = label_run
        new_runs.append(label_run)
    else:
        text_nodes[0].text = label
        new_runs.append(r_elem)

    if after_label:
        content_run = deepcopy(r_elem)
        content_run.find(qn('w:t')).text = after_label
        anchor.addnext(content_run)
        new_runs.append(content_run)

    return new_runs


def _apply_answer_analysis_format(new_element, color_hex: str, label_tags: list[str]) -> None:
    """Format copied answer/analysis paragraphs: label red 黑体五号不加粗, content red 宋体五号."""
    runs = list(new_element.findall(qn('w:r')))
    expanded_runs = []
    for r_elem in runs:
        if (r_elem.find(qn('w:drawing')) is not None
                or r_elem.find(qn('w:pict')) is not None
                or r_elem.find(qn('w:object')) is not None):
            continue
        expanded_runs.extend(_split_label_run_if_needed(r_elem, label_tags))

    for r_elem in expanded_runs:
        text_content = ''.join((t.text or '') for t in r_elem.findall(qn('w:t')))
        is_label = text_content in label_tags
        _set_run_text_style(
            r_elem,
            '黑体' if is_label else '宋体',
            color_hex,
            bold=is_label,
            size_half_points='21',
        )


def _apply_text_run_format(new_element, color_hex: str, bold_tags: list[str] | None = None) -> None:
    """Apply red answer/analysis formatting to copied source paragraphs."""
    tags = bold_tags or []
    if tags:
        _apply_answer_analysis_format(new_element, color_hex, tags)
        return
    for r_elem in new_element.findall(qn('w:r')):
        if (r_elem.find(qn('w:drawing')) is not None
                or r_elem.find(qn('w:pict')) is not None
                or r_elem.find(qn('w:object')) is not None):
            continue
        _set_run_text_style(r_elem, '宋体', color_hex, bold=False, size_half_points='21')

def _cap_inline_image_extents(element, max_w_cm: float = 12.0, max_h_cm: float = 8.0) -> None:
    """Clamp copied DOCX image display sizes while preserving aspect ratio."""
    max_w_cm, max_h_cm = _image_cap_for_copied_paragraph(element, max_w_cm, max_h_cm)
    max_w = int(max_w_cm * EMU_PER_CM)
    max_h = int(max_h_cm * EMU_PER_CM)
    try:
        extent_elems = element.xpath('.//*[local-name()="extent"]')
    except Exception:
        extent_elems = []
    for extent in extent_elems:
        try:
            cx = int(extent.get('cx') or 0)
            cy = int(extent.get('cy') or 0)
        except (TypeError, ValueError):
            continue
        if cx <= 0 or cy <= 0:
            continue
        scale = min(1.0, max_w / cx if cx > max_w else 1.0, max_h / cy if cy > max_h else 1.0)
        if scale >= 1.0:
            continue
        new_cx = max(1, int(cx * scale))
        new_cy = max(1, int(cy * scale))
        extent.set('cx', str(new_cx))
        extent.set('cy', str(new_cy))
        inline_or_anchor = extent.getparent()
        if inline_or_anchor is not None:
            try:
                graphic_exts = inline_or_anchor.xpath('.//*[local-name()="xfrm"]/*[local-name()="ext"]')
            except Exception:
                graphic_exts = []
            if graphic_exts:
                graphic_exts[0].set('cx', str(new_cx))
                graphic_exts[0].set('cy', str(new_cy))


def _format_copied_element(new_element, doc_rid_cache, dst_doc, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> None:
    """对刚复制到目标文档的 w:p 元素施加字体/字号/行距。
    含图片/公式/OLE 对象的段落只保留颜色+加粗覆写，不改变行距和字体防止渲染劣化。
    直接操作 lxml 元素，不依赖 python-docx 的 paragraph 缓存。

    text_color: None=不改变颜色; (R,G,B)=设为指定颜色（如 (255,0,0) 红色）。
    bold_tags: 【答案】/【解析】等需加粗的 tag 列表。
    """
    from lxml import etree

    # --- 清掉所有段落底纹（w:shd），防止源文档背景色污染输出 ---
    for scope in (new_element, *new_element.findall(qn('w:r'))):
        for parent_tag in (qn('w:pPr'), qn('w:rPr')):
            parent = scope.find(parent_tag)
            if parent is not None:
                shd = parent.find(qn('w:shd'))
                if shd is not None:
                    parent.remove(shd)

    color_hex = ''.join(f'{value:02X}' for value in text_color) if text_color is not None else None

    # 含图片/公式/OLE 对象的段落也清洗段落属性和文字 run，但不移动图片/公式对象。

    # --- 段落属性：1.5 倍行距，段前0段后2pt ---
    ppr = new_element.find(qn('w:pPr'))
    if ppr is None:
        ppr = etree.SubElement(new_element, qn('w:pPr'))
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(ppr, qn('w:spacing'))
    spacing.set(qn('w:line'), '360')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '40')  # 2pt = 40 twips

    for r_elem in new_element.findall(qn('w:r')):
        # 跳过含 drawing / pict 的图片 run
        if (r_elem.find(qn('w:drawing')) is not None
                or r_elem.find(qn('w:pict')) is not None
                or r_elem.find(qn('w:object')) is not None):
            continue

        # 字体/字号：宋体 10.5pt
        rpr = r_elem.find(qn('w:rPr'))
        if rpr is None:
            rpr = etree.SubElement(r_elem, qn('w:rPr'))

        # 西文字体
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), '宋体')

        # 字号 10.5pt = 21 half-points
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = etree.SubElement(rpr, qn('w:sz'))
        sz.set(qn('w:val'), '21')
        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = etree.SubElement(rpr, qn('w:szCs'))
        sz_cs.set(qn('w:val'), '21')

        # 颜色；标签/正文的加粗差异在循环后统一拆分处理。
        if color_hex is not None:
            color_el = rpr.find(qn('w:color'))
            if color_el is None:
                color_el = etree.SubElement(rpr, qn('w:color'))
            color_el.set(qn('w:val'), color_hex)

    if color_hex is not None and bold_tags:
        _apply_answer_analysis_format(new_element, color_hex, bold_tags)


def _body_range_from_question(question: dict[str, Any], body_count: int) -> tuple[int, int] | None:
    body_range = question.get("original_docx_body_range") or {}
    if isinstance(body_range, dict):
        start = body_range.get("start")
        end = body_range.get("end")
        if (isinstance(start, int) or str(start).isdigit()) and (isinstance(end, int) or str(end).isdigit()):
            start_i = int(start)
            end_i = int(end)
            if 0 <= start_i <= end_i < body_count:
                return start_i, end_i
    indices = [int(value) for value in question.get("original_docx_body_indices") or question.get("source_body_indices") or [] if isinstance(value, int) or str(value).isdigit()]
    if indices:
        start_i = min(indices)
        end_i = max(indices)
        if 0 <= start_i <= end_i < body_count:
            return start_i, end_i
    return None


def _copy_source_body_range(doc, source_path: str, question: dict[str, Any], rid_cache: dict | None = None) -> int:
    """Copy a protected question's complete original DOCX body range."""
    if not source_path:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        print(f"  → 警告：源 DOCX 不存在，无法复制原始题块：{src_path}")
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        body_elements = list(src_doc.element.body.iterchildren())
        body_range = _body_range_from_question(question, len(body_elements))
        if body_range is None:
            return 0
        copied = 0
        for element in body_elements[body_range[0]:body_range[1] + 1]:
            if element.tag == W_SECTPR:
                continue
            if element.tag not in (qn("w:p"), qn("w:tbl")):
                continue
            new_element = copy.deepcopy(element)
            _remove_placeholder_text(new_element)
            if new_element.tag == qn("w:p") and not _paragraph_has_protected_content(new_element) and not _paragraph_visible_text(new_element).strip():
                continue
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _cap_inline_image_extents(new_element, _STEM_IMAGE_MAX_W_CM, _STEM_IMAGE_MAX_H_CM)
            if new_element.tag == qn("w:p"):
                _format_copied_element(new_element, rid_cache, doc, text_color=None, bold_tags=None)
            elif new_element.tag == qn("w:tbl"):
                _format_copied_table(new_element)
            _append_body_element(doc, new_element)
            copied += 1
        return copied
    except Exception as exc:
        print(f"  → 警告：复制原始题块失败：{exc}")
        return 0


# -- 按段落索引从源 DOCX 复制段落的缓存 --
_SRC_DOC_CACHE: dict[str, Any] = {}


def _get_cached_src_doc(source_path: str):
    if source_path not in _SRC_DOC_CACHE:
        _SRC_DOC_CACHE[source_path] = Document(source_path)
    return _SRC_DOC_CACHE[source_path]


def _copy_source_paragraphs(doc, source_path: str, indices: list[int], rid_cache: dict | None = None, text_color: tuple | None = None, bold_tags: list[str] | None = None, visual_only: bool = False) -> int:
    """从源 DOCX 复制指定段落到目标文档（含图片/公式 remap）。
    共享 rid_cache 避免同一次生成中重复注入相同图片资源。
    text_color: None=不改变颜色; (R,G,B)=设为指定颜色。
    bold_tags: 需加粗的标签文本列表，如 ["【答案】"]。
    visual_only=True 时只保留图片/公式/OLE 对象，丢弃源段落文字和旧题号。
    返回复制的段落数。
    """
    if not source_path or not indices:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        print(f"  → 警告：源 DOCX 不存在，无法复制段落：{src_path}")
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        paragraphs = list(src_doc.paragraphs)
        copied = 0
        sorted_indices = sorted(set(int(i) for i in indices if isinstance(i, int) or str(i).isdigit()))
        for idx in sorted_indices:
            if idx < 0 or idx >= len(paragraphs):
                continue
            new_element = copy.deepcopy(paragraphs[idx]._element)
            if visual_only and not _keep_only_visual_content(new_element):
                continue
            _remove_placeholder_text(new_element)
            if not visual_only and not _paragraph_has_protected_content(new_element) and not _paragraph_visible_text(new_element).strip():
                continue
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _cap_inline_image_extents(new_element, _STEM_IMAGE_MAX_W_CM, _STEM_IMAGE_MAX_H_CM)
            copied += _append_copied_paragraph_latex_safe(doc, new_element, text_color=text_color, bold_tags=bold_tags)
        return copied
    except Exception as exc:
        print(f"  → 警告：复制源段落失败：{exc}")
        return 0


def _paragraph_body_index(paragraphs, paragraph_index: int, body_elements: list[Any]) -> int | None:
    if paragraph_index < 0 or paragraph_index >= len(paragraphs):
        return None
    target = paragraphs[paragraph_index]._element
    for body_index, element in enumerate(body_elements):
        if element is target:
            return body_index
    return None


def _copy_source_tables_between_paragraphs(doc, source_path: str, start_indices: list[int], end_indices: list[int], rid_cache: dict | None = None) -> int:
    """Copy option tables located between copied stem paragraphs and answer paragraphs."""
    return _copy_source_option_tables(
        doc,
        source_path,
        start_indices=start_indices,
        end_indices=end_indices,
        rid_cache=rid_cache,
    )


def _copy_source_option_tables(
    doc,
    source_path: str,
    start_indices: list[int] | None = None,
    end_indices: list[int] | None = None,
    table_indices: list[int] | None = None,
    rid_cache: dict | None = None,
) -> int:
    """Copy source DOCX option tables either by explicit table ids or paragraph bounds."""
    if not source_path:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        paragraphs = list(src_doc.paragraphs)
        tables = list(src_doc.tables)
        table_elements: list[Any] = []

        explicit_ids = [int(value) for value in (table_indices or []) if isinstance(value, int) or str(value).isdigit()]
        for table_idx in dict.fromkeys(explicit_ids):
            if 0 <= table_idx < len(tables):
                table_elements.append(tables[table_idx]._element)

        if not table_elements and start_indices and end_indices:
            body_elements = list(src_doc.element.body.iterchildren())
            stem_body_indices = [
                idx for idx in (
                    _paragraph_body_index(paragraphs, int(value), body_elements)
                    for value in start_indices
                    if isinstance(value, int) or str(value).isdigit()
                )
                if idx is not None
            ]
            answer_body_indices = [
                idx for idx in (
                    _paragraph_body_index(paragraphs, int(value), body_elements)
                    for value in end_indices
                    if isinstance(value, int) or str(value).isdigit()
                )
                if idx is not None
            ]
            if not stem_body_indices or not answer_body_indices:
                return 0
            start = max(stem_body_indices)
            following_answers = [idx for idx in answer_body_indices if idx > start]
            if not following_answers:
                return 0
            end = min(following_answers)
            table_elements = [element for element in body_elements[start + 1:end] if element.tag == qn("w:tbl")]

        copied = 0
        seen: set[int] = set()
        for element in table_elements:
            if id(element) in seen:
                continue
            seen.add(id(element))
            new_element = copy.deepcopy(element)
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _cap_inline_image_extents(new_element, _OPTION_IMAGE_MAX_W_CM, _OPTION_IMAGE_MAX_H_CM)
            if _looks_like_options_table(new_element) and _append_normalized_options_table(doc, new_element):
                copied += 1
                continue
            if _looks_like_options_table(new_element):
                _style_copied_options_table(new_element)
            else:
                _format_copied_table(new_element)
            _append_body_element(doc, new_element)
            copied += 1
        return copied
    except Exception as exc:
        print(f"  → 警告：复制源选项表格失败：{exc}")
        return 0


def _replace_first_question_number(paragraph_elem, display_no: int) -> None:
    """Replace the leading question number in a copied source paragraph."""
    remaining = []
    seen_text = ""
    for run_elem in paragraph_elem.findall(qn('w:r')):
        for text_elem in run_elem.findall(qn('w:t')):
            remaining.append(text_elem)
            seen_text += text_elem.text or ""
            if len(seen_text) >= 20:
                break
        if len(seen_text) >= 20:
            break
    match = re.match(r"^\s*\d+\s*[\.．、)]\s*", seen_text)
    if not match:
        return
    replacement = f"{display_no}. "
    chars_to_remove = match.end()
    first = True
    for text_elem in remaining:
        value = text_elem.text or ""
        if first:
            take = min(chars_to_remove, len(value))
            text_elem.text = replacement + value[take:]
            chars_to_remove -= take
            first = False
        elif chars_to_remove > 0:
            take = min(chars_to_remove, len(value))
            text_elem.text = value[take:]
            chars_to_remove -= take
        if chars_to_remove <= 0:
            break


def _paragraph_visible_text(paragraph_elem) -> str:
    return "".join(t.text or "" for t in paragraph_elem.findall(f".//{qn('w:t')}"))


def _add_latex_safe_text_paragraph(doc, text: str, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> None:
    """Write copied visible text through the normal rich-text path so LaTeX becomes Word math."""
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本" / "核心脚本"))
    from docx_utils1 import _add_rich_text_run  # noqa: WPS433 - internal shared DOCX writer

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.5

    remaining = text or ""
    label = next((tag for tag in (bold_tags or []) if remaining.startswith(tag)), "")
    if label:
        label_run = para.add_run(label)
        label_run.font.name = "黑体"
        label_run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
        label_run.font.size = Pt(10.5)
        label_run.bold = True
        if text_color:
            label_run.font.color.rgb = RGBColor(*text_color)
        remaining = remaining[len(label):]

    _add_rich_text_run(para, remaining, "宋体", 10.5, False, text_color)


def _append_copied_paragraph_latex_safe(doc, new_element, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> int:
    """Append a copied paragraph, converting LaTeX text while preserving visual objects when present."""
    visible = _paragraph_visible_text(new_element).strip()
    if not visible or not RESIDUAL_FORMULA_RE.search(visible):
        _format_copied_element(new_element, {}, doc, text_color=text_color, bold_tags=bold_tags)
        _append_body_element(doc, new_element)
        return 1

    if _paragraph_has_protected_content(new_element):
        _add_latex_safe_text_paragraph(doc, visible, text_color=text_color, bold_tags=bold_tags)
        visual_element = copy.deepcopy(new_element)
        if _keep_only_visual_content(visual_element):
            _format_copied_element(visual_element, {}, doc, text_color=text_color, bold_tags=None)
            _append_body_element(doc, visual_element)
        return 1

    _add_latex_safe_text_paragraph(doc, visible, text_color=text_color, bold_tags=bold_tags)
    return 1


def _looks_like_plain_option_paragraph(paragraph_elem) -> bool:
    """Detect option-only source paragraphs so standard option tables do not duplicate them."""
    if _paragraph_has_protected_content(paragraph_elem):
        return False
    text = re.sub(r"\s+", " ", _paragraph_visible_text(paragraph_elem)).strip()
    if not text:
        return False
    markers = re.findall(r"(?:^|\s)([A-D])\s*[.．、)]", text, flags=re.I)
    return len(markers) >= 2 or bool(re.match(r"^[A-D]\s*[.．、)]", text, flags=re.I))


def _set_cell_text_style(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run.font.size = Pt(10.5)


def _clear_table_borders(table) -> None:
    from docx.oxml import OxmlElement
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _add_options_table_only(doc, option_lines: list[str]) -> None:
    options = []
    for line in option_lines:
        for part in re.split(r"\t+", line):
            part = part.strip()
            if part:
                options.append(part)
    if not options:
        return
    rows = (len(options) + 1) // 2 if len(options) > 2 else 1
    cols = 2 if len(options) > 1 else 1
    table = doc.add_table(rows=rows, cols=cols)
    _clear_table_borders(table)
    for i, opt in enumerate(options):
        cell = table.cell(i // cols, i % cols)
        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        _set_cell_text_style(paragraph, opt)



def _run_has_protected_content(run_elem) -> bool:
    return any(run_elem.find(f'.//{tag}') is not None for tag in _PROTECTED_CONTENT_TAGS)


def _keep_only_visual_content(paragraph_elem) -> bool:
    """Remove text from a copied paragraph while retaining pictures/formulas/OLE."""
    if not _paragraph_has_protected_content(paragraph_elem):
        return False
    for run_elem in list(paragraph_elem.findall(qn('w:r'))):
        if _run_has_protected_content(run_elem):
            for text_elem in list(run_elem.findall(qn('w:t'))):
                run_elem.remove(text_elem)
            continue
        parent = run_elem.getparent()
        if parent is not None:
            parent.remove(run_elem)
    return _paragraph_has_protected_content(paragraph_elem)


def _remove_placeholder_text(paragraph_elem) -> None:
    """Remove literal [图片] placeholders while preserving real image objects."""
    for text_elem in paragraph_elem.findall(f".//{qn('w:t')}"):
        value = text_elem.text or ""
        value = re.sub(r"\s*(?:\[图片\]|【图片】)\s*", "", value)
        value = re.sub(r"\s*【(?:解析|详解|分析)】\s*(?:略|无|—|－－|--|\.|。)?\s*$", "", value)
        text_elem.text = value


def _strip_image_placeholders(text: str) -> str:
    """Remove image placeholders from regenerated text; real visuals are copied from DOCX."""
    return re.sub(r"\s*(?:\[图片\]|【图片】)\s*", "", str(text or "")).strip()


def _copy_labeled_paragraphs_reflowed(doc, source_path: str, indices: list[int], rid_cache: dict | None = None, text_color: tuple | None = None, label_tags: list[str] | None = None, drop_stub_analysis: bool = False) -> int:
    """Copy answer/analysis paragraphs in-place, preserving image layout and cleaning placeholders."""
    if not source_path or not indices:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        paragraphs = list(src_doc.paragraphs)
        copied = 0
        for idx in sorted(set(int(i) for i in indices if isinstance(i, int) or str(i).isdigit())):
            if idx < 0 or idx >= len(paragraphs):
                continue
            new_element = copy.deepcopy(paragraphs[idx]._element)
            _remove_placeholder_text(new_element)
            visible = _paragraph_visible_text(new_element).strip()
            if drop_stub_analysis and not _paragraph_has_protected_content(new_element) and not _should_emit_analysis(None, re.sub(r"^【(?:解析|详解|分析)】\s*", "", visible)):
                continue
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _cap_inline_image_extents(new_element, _STEM_IMAGE_MAX_W_CM, _STEM_IMAGE_MAX_H_CM)
            copied += _append_copied_paragraph_latex_safe(doc, new_element, text_color=text_color, bold_tags=label_tags)
        return copied
    except Exception as exc:
        print(f"  → 警告：复制答案/解析段落失败：{exc}")
        return 0

def _copy_stem_paragraphs_reflowed(doc, source_path: str, indices: list[int], display_no: int, rid_cache: dict | None = None, skip_plain_options: bool = False) -> int:
    """Copy stem paragraphs in-place, preserving image layout while normalizing text and question number."""
    if not source_path or not indices:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        print(f"  → 警告：源 DOCX 不存在，无法复制带图题题干：{src_path}")
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        paragraphs = list(src_doc.paragraphs)
        copied = 0
        for pos, idx in enumerate(sorted(set(int(i) for i in indices if isinstance(i, int) or str(i).isdigit()))):
            if idx < 0 or idx >= len(paragraphs):
                continue
            new_element = copy.deepcopy(paragraphs[idx]._element)
            if skip_plain_options and _looks_like_plain_option_paragraph(new_element):
                continue
            if pos == 0:
                _replace_first_question_number(new_element, display_no)
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _cap_inline_image_extents(new_element, _STEM_IMAGE_MAX_W_CM, _STEM_IMAGE_MAX_H_CM)
            copied += _append_copied_paragraph_latex_safe(doc, new_element, text_color=None, bold_tags=None)
        return copied
    except Exception as exc:
        print(f"  → 警告：复制带图题题干失败：{exc}")
        return 0


def _normalize_option_text(option: Any) -> str:
    if isinstance(option, dict):
        label = str(option.get('label') or '').strip().rstrip('.．、)）')
        value = str(option.get('text') or option.get('content') or '').strip()
        return f"{label}. {value}".strip() if label else value
    return str(option or '').strip()


def _option_lines_for_table(options: Any) -> list[str]:
    if isinstance(options, dict):
        options = [f"{label}. {options[label]}" for label in sorted(options) if str(options[label]).strip()]
    normalized = []
    for option in options or []:
        text = _normalize_option_text(option)
        if not text:
            continue
        match = re.match(r"^([A-H])\s*[.、．)]\s*(.+)", text, re.I)
        normalized.append(f"{match.group(1).upper()}. {match.group(2).strip()}" if match else text)
    if len(normalized) == 4 and all(len(opt) <= 5 for opt in normalized):
        return ["\t\t".join(normalized)]
    return [
        f"{normalized[i]}\t\t{normalized[i + 1]}" if i + 1 < len(normalized) else normalized[i]
        for i in range(0, len(normalized), 2)
    ]


def _image_ref_indices(question: dict[str, Any], part: str) -> list[int]:
    refs = (question.get("image_refs") or {}).get(part) or []
    indices: list[int] = []
    for ref in refs:
        if not isinstance(ref, dict):
            continue
        value = ref.get("paragraph_index")
        if isinstance(value, int) or str(value).isdigit():
            indices.append(int(value))
    # Some formulas are stored in ordinary paragraphs and do not appear in
    # image_refs. Copy only the current section's original paragraphs so stem,
    # answer, and analysis visuals do not repeat into each other.
    for value in question.get(f"{part}_paragraph_indices") or []:
        if isinstance(value, int) or str(value).isdigit():
            indices.append(int(value))
    return list(dict.fromkeys(indices))


def _copy_visual_paragraphs_for_part(doc, question: dict[str, Any], part: str, rid_cache: dict, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> int:
    source_path = question.get("source_docx_path") or question.get("source_path")
    if not source_path:
        return 0
    return _copy_source_paragraphs(
        doc,
        source_path,
        _image_ref_indices(question, part),
        rid_cache,
        text_color=text_color,
        bold_tags=bold_tags,
        visual_only=True,
    )


def _question_has_original_docx_images(question: dict[str, Any] | None) -> bool:
    if not question:
        return False
    if question.get("has_original_docx_images"):
        return True
    refs = question.get("image_refs") or {}
    return any(refs.get(part) for part in ("stem", "answer", "analysis"))


def _append_protected_question_reflowed(doc, question: dict[str, Any], display_no: int, add_question_options_table, add_paragraph_with_style, add_labeled_text) -> bool:
    """Render text with standard formatting while copying original visual paragraphs."""
    rid_cache: dict = {}
    stem = _strip_question_number(question.get("stem") or question.get("raw_text") or "")
    stem = _strip_image_placeholders(stem)
    if stem:
        add_paragraph_with_style(doc, f"{display_no}. {stem}".rstrip(), font_name="宋体", font_size=10.5, space_after=2)
    else:
        add_paragraph_with_style(doc, f"{display_no}.".rstrip(), font_name="宋体", font_size=10.5, space_after=2)

    copied_stem = _copy_visual_paragraphs_for_part(doc, question, "stem", rid_cache)

    options = question.get("options") or []
    has_option_images = any(isinstance(opt, dict) and (opt.get("image") or opt.get("images")) for opt in options)
    source_path = question.get("source_docx_path") or question.get("source_path")
    source_table_copied = 0
    if source_path:
        source_table_copied = _copy_source_option_tables(
            doc,
            source_path,
            start_indices=question.get("stem_paragraph_indices") or [],
            end_indices=question.get("answer_paragraph_indices") or [],
            table_indices=question.get("option_table_indices") or [],
            rid_cache=rid_cache,
        )
    if options and not has_option_images and not source_table_copied:
        _add_options_table_only(doc, _option_lines_for_table(options))

    answer = _strip_image_placeholders(str(question.get("answer") or "")).strip()
    if answer:
        add_labeled_text(
            doc, "【答案】", answer,
            label_font="黑体", label_size=10.5,
            content_font="宋体", content_size=10.5,
            color=_ANSWER_RED, bold_label=True, bold_content=False,
        )
    answer_copied = _copy_visual_paragraphs_for_part(
        doc, question, "answer", rid_cache,
        text_color=_ANSWER_RED, bold_tags=["【答案】"],
    )

    analysis = _strip_image_placeholders(str(question.get("analysis") or question.get("explanation") or "")).strip()
    analysis = _strip_analysis_label(analysis)
    emit_analysis = _should_emit_analysis(question, analysis)
    if emit_analysis:
        add_labeled_text(
            doc, "【解析】", analysis,
            label_font="黑体", label_size=10.5,
            content_font="宋体", content_size=10.5,
            color=_ANSWER_RED, bold_label=True, bold_content=False,
        )
    analysis_copied = _copy_visual_paragraphs_for_part(
        doc, question, "analysis", rid_cache,
        text_color=_ANSWER_RED, bold_tags=["【解析】", "【详解】", "【分析】"],
    )

    return bool(stem or copied_stem or options or answer or answer_copied or analysis_copied or (analysis and emit_analysis))


def validate_docx_preserves_images(output_path: str | Path, questions: list[dict[str, Any]] | None) -> None:
    """Raise if source DOCX image hashes referenced by questions are missing in output."""
    expected: set[str] = set()
    fallback_count = 0
    for question in questions or []:
        refs = question.get("original_docx_image_refs") or []
        if not refs:
            image_refs = question.get("image_refs") or {}
            refs = [ref for part_refs in image_refs.values() for ref in (part_refs or [])]
        for ref in refs:
            if not isinstance(ref, dict):
                continue
            sha = str(ref.get("sha256") or "").strip()
            if sha:
                expected.add(sha)
            elif ref.get("byte_size"):
                fallback_count += 1
    if not expected and not fallback_count:
        return

    output = Path(output_path)
    actual: set[str] = set()
    actual_count = 0
    try:
        with zipfile.ZipFile(output) as zf:
            for name in zf.namelist():
                if not name.startswith("word/media/") or name.endswith("/"):
                    continue
                data = zf.read(name)
                actual_count += 1
                actual.add(hashlib.sha256(data).hexdigest())
    except Exception as exc:
        raise RuntimeError(f"输出 DOCX 图片校验失败：无法读取 {output}：{exc}") from exc

    missing = expected - actual
    expected_count = len(expected) + fallback_count
    if missing or (fallback_count and actual_count < expected_count):
        raise RuntimeError(
            "输出 DOCX 图片校验失败：原始含图题图片缺失或数量减少 "
            f"(expected_hashes={len(expected)}, missing_hashes={len(missing)}, "
            f"expected_count={expected_count}, actual_count={actual_count})"
        )


# 答案/解析颜色
_ANSWER_RED = (255, 0, 0)


def _append_original_question_block(doc, question: dict[str, Any]) -> int:
    """复制带图题的原始段落到目标文档，并按分区上色。
    题干 → 默认黑色；答案 → 红色；解析 → 红色。
    返回复制的段落数。
    """
    source_path = question.get("source_docx_path") or question.get("source_path")
    if not source_path:
        return 0
    rid_cache: dict = {}
    total = _copy_source_body_range(doc, source_path, question, rid_cache)
    if total:
        return total
    if question.get("original_docx_body_range") or question.get("source_body_indices"):
        print(f"  → 警告：第{question.get('question_no')}题原始题块范围复制失败，回退为段落索引复制。")
    else:
        print(f"  → 警告：第{question.get('question_no')}题缺少原始题块范围，回退为段落索引复制。")

    total = 0

    # 题干段落：黑色（text_color=None 即不覆写颜色）
    stem_indices = question.get("stem_paragraph_indices") or []
    total += _copy_source_paragraphs(doc, source_path, stem_indices, rid_cache, text_color=None)

    # DOCX 选择题的选项常以表格存放在题干图和答案之间，也要原样带回。
    answer_indices = question.get("answer_paragraph_indices") or []
    total += _copy_source_tables_between_paragraphs(doc, source_path, stem_indices, answer_indices, rid_cache)

    # 答案段落：标签黑体五号不加粗，正文宋体五号，均为红色
    total += _copy_source_paragraphs(doc, source_path, answer_indices, rid_cache, text_color=_ANSWER_RED, bold_tags=["【答案】"])

    # 解析段落：标签黑体五号不加粗，正文宋体五号，均为红色
    analysis_indices = question.get("analysis_paragraph_indices") or []
    total += _copy_source_paragraphs(doc, source_path, analysis_indices, rid_cache, text_color=_ANSWER_RED, bold_tags=["【解析】", "【详解】", "【分析】"])

    # 兜底：如果有 source_paragraph_indices 但分区索引缺失，统一用黑色
    all_indices = question.get("source_paragraph_indices") or []
    covered = set(stem_indices + answer_indices + analysis_indices)
    fallback = [i for i in all_indices if i not in covered]
    if fallback:
        total += _copy_source_paragraphs(doc, source_path, fallback, rid_cache, text_color=None)

    return total



def _question_is_protected(question: dict[str, Any] | None) -> bool:
    return bool(question and (question.get("protected_original_docx_block") or _question_has_original_docx_images(question)))




def _is_true_false_question(question: dict[str, Any] | None, q_text: str = "") -> bool:
    """判断题不应输出 A/B 正确/错误选项。"""
    qtype = str((question or {}).get("question_type") or (question or {}).get("type") or "").strip()
    return "判断" in qtype or bool(re.match(r"^\d+[.．、].*（\s*）\s*$", str(q_text or "").splitlines()[0].strip() if q_text else ""))


def _strip_true_false_options(q_text: str) -> str:
    lines = []
    for line in str(q_text or "").splitlines():
        compact = re.sub(r"\s+", "", line)
        if re.fullmatch(r"A[.．、)]正确B[.．、)]错误", compact, flags=re.I):
            continue
        if re.fullmatch(r"[AB][.．、)](?:正确|错误)", compact, flags=re.I):
            continue
        lines.append(line)
    return "\n".join(lines).strip()

def _question_has_structured_images(question: dict[str, Any] | None) -> bool:
    """判断题目结构中是否已有可排版的题干图/选项图字段。
    兼容 stem_images/option_images（API 丰富来源）和 image_refs（DOCX 解析来源）。
    """
    if not isinstance(question, dict):
        return False
    # API 丰富来源
    if question.get("stem_images") or question.get("option_images"):
        return True
    # DOCX 解析来源：image_refs.stem 存在即说明题干含图
    image_refs = question.get("image_refs") or {}
    if image_refs.get("stem"):
        return True
    options = question.get("options") or question.get("option") or []
    if isinstance(options, list):
        return any(isinstance(opt, dict) and opt.get("images") for opt in options)
    return False


def _strip_question_number(text: str) -> str:
    return re.sub(r"^\s*\d+[\.．、]\s*", "", str(text or "").strip())


_STUB_ANALYSIS_RE = re.compile(r"^(?:(?:略|无)\s*[。.]?|—|－－|--|[.。])?\s*$")


def _should_emit_analysis(question: dict[str, Any] | None, analysis: str) -> bool:
    text = str(analysis or "").strip()
    if not text or _STUB_ANALYSIS_RE.match(text):
        return False
    if _question_is_protected(question) and len(text) <= 10:
        return False
    return True


def _parse_text_options_from_fallback(fallback_text: str) -> list[dict[str, Any]]:
    """从解析版文本兜底提取 A-D 选项，供图片字段缺少文字时补齐。"""
    lines = [line.strip() for line in str(fallback_text or "").splitlines() if line.strip()]
    text = "\n".join(lines[1:]) if len(lines) > 1 else ""
    pattern = re.compile(r"([A-D])[\.．]\s*(.*?)(?=(?:\s+[A-D][\.．])|$)", re.S)
    options = []
    for match in pattern.finditer(text.replace("\t", " ")):
        content = re.sub(r"\s+", " ", match.group(2)).strip()
        options.append({"label": match.group(1), "text": content, "images": []})
    return options


def _normalize_structured_choice_question(question: dict[str, Any], fallback_text: str) -> dict[str, Any] | None:
    """把不同来源的图片选择题字段归一化为 docx_utils1 可消费的结构。"""
    if not isinstance(question, dict):
        return None
    fallback_lines = [line.strip() for line in str(fallback_text or "").splitlines() if line.strip()]
    stem_text = str(
        question.get("stem")
        or question.get("stem_text")
        or question.get("question")
        or question.get("content")
        or (fallback_lines[0] if fallback_lines else "")
        or ""
    ).strip()
    if stem_text and not re.match(r"^\d+[\.．、]", stem_text) and fallback_lines and re.match(r"^\d+[\.．、]", fallback_lines[0]):
        # 保留解析版文本中的题号，避免结构化 stem 只有裸题干时丢失编号。
        stem_text = re.match(r"^\s*\d+[\.．、]", fallback_lines[0]).group(0) + stem_text
    elif stem_text:
        stem_text = _strip_question_number(stem_text)
        if fallback_lines and re.match(r"^\d+[\.．、]", fallback_lines[0]):
            stem_text = re.match(r"^\s*\d+[\.．、]", fallback_lines[0]).group(0) + stem_text

    stem_images = question.get("stem_images") or question.get("images") or []
    raw_options = question.get("options") or question.get("option") or []
    text_options = _parse_text_options_from_fallback(fallback_text)
    text_by_label = {opt["label"]: opt.get("text", "") for opt in text_options}

    options = []
    if isinstance(raw_options, list):
        for idx, opt in enumerate(raw_options):
            if isinstance(opt, dict):
                label = str(opt.get("label") or chr(ord("A") + idx)).strip().rstrip(".．")
                text = str(opt.get("text") or opt.get("content") or text_by_label.get(label, "") or "").strip()
                images = opt.get("images") or opt.get("image") or []
                if isinstance(images, (str, dict)):
                    images = [images]
                options.append({"label": label, "text": text, "images": images})
            elif idx < len(text_options):
                options.append(text_options[idx])
    elif isinstance(raw_options, dict):
        for idx, label in enumerate("ABCD"):
            opt = raw_options.get(label) or raw_options.get(label.lower())
            if isinstance(opt, dict):
                images = opt.get("images") or opt.get("image") or []
                if isinstance(images, (str, dict)):
                    images = [images]
                options.append({
                    "label": label,
                    "text": str(opt.get("text") or opt.get("content") or text_by_label.get(label, "") or "").strip(),
                    "images": images,
                })
    if not options:
        options = text_options

    option_images = question.get("option_images") or {}
    if isinstance(option_images, dict):
        by_label = {opt["label"]: opt for opt in options}
        for label, images in option_images.items():
            key = str(label).strip().rstrip(".．")
            if isinstance(images, (str, dict)):
                images = [images]
            by_label.setdefault(key, {"label": key, "text": text_by_label.get(key, ""), "images": []})["images"].extend(images or [])
        options = [by_label[label] for label in sorted(by_label.keys())]

    if not stem_images and not any(opt.get("images") for opt in options):
        return None
    return {"stem_text": stem_text, "stem_images": stem_images, "options": options}


def _copy_common_scripts_to_output(sub_dir):
    """将常用后处理脚本复制到当前考类/教材输出目录，便于直接双击使用。"""
    source_dir = BASE_DIR / "01_工具脚本" / "通用脚本"
    for script_name in COMMON_SCRIPT_NAMES:
        src = source_dir / script_name
        if not src.exists():
            print(f"  → 警告：未找到通用脚本 {src}")
            continue
        dst = sub_dir / script_name
        try:
            shutil.copy2(src, dst)
        except Exception as exc:
            print(f"  → 警告：复制通用脚本失败 {script_name}: {exc}")

def _get_value(source, *names, default=""):
    """兼容 dict 和 dataclass/object 的安全取值。"""
    if source is None:
        return default
    for name in names:
        if isinstance(source, dict):
            value = source.get(name)
        else:
            value = getattr(source, name, None)
        if value not in (None, ""):
            return value
    return default


def _normalize_province_name(name):
    """规范省市名称，自治区全称保持不变。"""
    text = str(name or "").strip()
    if not text:
        return ""
    direct = {
        "北京": "北京市", "天津": "天津市", "上海": "上海市", "重庆": "重庆市",
        "河北": "河北省", "山西": "山西省", "辽宁": "辽宁省", "吉林": "吉林省",
        "黑龙江": "黑龙江省", "江苏": "江苏省", "浙江": "浙江省", "安徽": "安徽省",
        "福建": "福建省", "江西": "江西省", "山东": "山东省", "河南": "河南省",
        "湖北": "湖北省", "湖南": "湖南省", "广东": "广东省", "海南": "海南省",
        "四川": "四川省", "贵州": "贵州省", "云南": "云南省", "陕西": "陕西省",
        "甘肃": "甘肃省", "青海": "青海省", "台湾": "台湾省",
        "内蒙古": "内蒙古自治区", "广西": "广西壮族自治区", "西藏": "西藏自治区",
        "宁夏": "宁夏回族自治区", "新疆": "新疆维吾尔自治区",
    }
    if text in direct:
        return direct[text]
    if text.endswith(("省", "市", "自治区")):
        return text
    return direct.get(text, text)


def _safe_path_part(text):
    """清理 Windows 文件名/目录名非法字符。"""
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", str(text or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or "未命名"


def _parse_edition(raw):
    """从简短版次说明中兜底提取出版社和版次。"""
    text = str(raw or "").strip()
    if not text:
        return "高教版", "第一版"
    if "·" in text:
        publisher, edition = text.split("·", 1)
        return publisher.strip() or "高教版", edition.strip() or "第一版"
    if "版" in text and len(text) <= 6:
        return text, "第一版"
    return "高教版", text or "第一版"


def _extract_textbook_detail(meta, topic):
    """提取教材/课程显示所需信息。"""
    detail = None
    course = _get_value(topic, "course", "module")
    course_map = _get_value(meta, "course_textbook_map", default={}) or {}
    if course_map and course in course_map:
        detail = course_map[course]
    elif _get_value(meta, "textbook_details", default=None):
        detail = _get_value(meta, "textbook_details")[0]

    if detail:
        if isinstance(detail, dict):
            name = detail.get("name") or detail.get("textbook_name") or ""
            publisher = detail.get("publisher") or "高教版"
            edition = detail.get("edition") or "第一版"
            display = detail.get("display") or f"{publisher}·{edition}"
        else:
            name = getattr(detail, "name", "")
            publisher = getattr(detail, "publisher", "高教版")
            edition = getattr(detail, "edition", "第一版")
            display = getattr(detail, "display", f"{publisher}·{edition}")
        return {"name": name or "教材", "publisher": publisher, "edition": edition, "display": display}

    textbook_list = _get_value(meta, "textbook_list", default=None) or []
    if textbook_list:
        first = textbook_list[0]
        if isinstance(first, (list, tuple)) and len(first) >= 2:
            publisher, edition = _parse_edition(first[1])
            return {"name": first[0] or "教材", "publisher": publisher, "edition": edition, "display": f"{publisher}·{edition}"}
        return {"name": str(first), "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}

    textbooks = _get_value(meta, "textbooks")
    match = re.search(r"《(.+?)》", str(textbooks or ""))
    if match:
        return {"name": match.group(1), "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}

    fallback = _get_value(meta, "subject") or _get_value(topic, "course", "module") or "教材"
    return {"name": fallback, "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}


def _format_paper_type(topic):
    text = str(_get_value(topic, "paper_type", "paper_kind", "type", "卷型") or "").strip()
    aliases = {
        "考点": "考点训练卷", "考点卷": "考点训练卷", "point": "考点训练卷",
        "专题": "专题训练卷", "专题卷": "专题训练卷", "topic": "专题训练卷",
        "综合": "课程综合卷", "课程综合": "课程综合卷", "综合卷": "课程综合卷", "course": "课程综合卷",
    }
    if text in {"考点训练卷", "专题训练卷", "课程综合卷"}:
        return text
    lowered = text.lower()
    if lowered in aliases:
        return aliases[lowered]
    return aliases.get(text, text or "考点训练卷")


def _format_exam_category(meta):
    return str(_get_value(meta, "category", "exam_category", "major_category") or "").strip()


def _format_paper_name(topic):
    paper_type = _format_paper_type(topic)
    if paper_type == "考点训练卷":
        name = _get_value(topic, "point_name", "theme", "topic", "subject")
    elif paper_type == "专题训练卷":
        name = _get_value(topic, "topic", "theme", "point_name", "subject")
    else:
        name = _get_value(topic, "theme", "topic", "module", "subject") or "课程综合"
    return str(name or "未命名试卷").strip()


def _format_course_name(meta, topic, textbook_name=None):
    return str(
        _get_value(topic, "course", "module")
        or _get_value(meta, "subject", "course")
        or textbook_name
        or "课程"
    ).strip()


def _unique_nonempty(values):
    result = []
    seen = set()
    for value in values:
        text = str(value or "").strip()
        if text and text not in seen:
            result.append(text)
            seen.add(text)
    return result


def _iter_topic_rows(topic):
    rows = _get_value(topic, "rows", default=[]) or []
    return rows if isinstance(rows, (list, tuple)) else []


def _row_value(row, *names):
    for name in names:
        value = _get_value(row, name, default="")
        if value:
            return value
        raw = _get_value(row, "raw", default={}) or {}
        if isinstance(raw, dict):
            for key in names:
                value = raw.get(key)
                if value:
                    return value
    return ""


def _format_all_courses(meta, topic, course_name):
    """按规划总表收集的全考类课程列表生成专辑课程信息和计数。"""
    # 优先使用 meta.all_courses（从全部行收集的课程列表）
    courses = getattr(meta, "all_courses", None) or []
    if not courses:
        # 兜底：从传入的 rows 或 topic 中提取
        rows = _get_value(meta, "rows", default=None) or _get_value(topic, "all_rows", default=None) or _iter_topic_rows(topic)
        courses = _unique_nonempty(_row_value(row, "module", "知识模块") for row in rows)
    if not courses:
        explicit = _get_value(meta, "course_names", "courses", default=None)
        if explicit:
            courses = _unique_nonempty(explicit if isinstance(explicit, (list, tuple)) else re.split(r"[、,，]\s*", str(explicit)))
    if not courses:
        courses = [course_name]
    return "、".join(courses), len(courses)



def _strip_and_reindex_content(raw_text: str) -> str:
    """去掉原始序号（如 "2."、"（2）"、"一、"），从 1 重新编号，多条用分号分隔。"""
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return raw_text
    stripped = []
    for line in lines:
        # 先按行切分；行内若含多个序号则再切
        parts = re.split(r'(?:^|[（(])\s*\d+\s*[）).、．]\s*|^[一二三四五六七八九十]+\s*[、．.]\s*', line)
        for part in parts:
            cleaned = re.sub(r'^[、．.,，；;]+\s*', '', part).strip()
            if cleaned:
                stripped.append(cleaned)
    if not stripped:
        return raw_text
    if len(stripped) == 1:
        return f'1. {stripped[0]}'
    return '；'.join(f'{i}. {item}' for i, item in enumerate(stripped, 1))


def _format_specific_content(topic):
    """按卷型映射“具体内容为”：考点卷用考点内容，专题/综合卷用覆盖名称。"""
    paper_type = _format_paper_type(topic)
    rows = _iter_topic_rows(topic)
    if paper_type == "考点训练卷":
        content = _get_value(topic, "point_content", "knowledge", "content")
        if not content and rows:
            content = _row_value(rows[0], "point_content", "掌握/理解/了解考点内容", "knowledge", "content")
        raw = str(content or _format_paper_name(topic)).strip()
        return _strip_and_reindex_content(raw)

    if paper_type == "专题训练卷":
        names = _unique_nonempty(
            _row_value(row, "point_name", "考点名称（考点训练卷名称）", "knowledge_point")
            for row in rows
        )
        if not names:
            names = [_format_paper_name(topic)]
        return "、".join(names)

    names = _unique_nonempty(_row_value(row, "topic", "考纲一级标题（专题名称）") for row in rows)
    if not names:
        names = _unique_nonempty(
            _row_value(row, "point_name", "考点名称（考点训练卷名称）")
            for row in rows
        )
    if not names:
        names = [_format_paper_name(topic)]
    return "、".join(names)


def _format_exam_table_title(meta, province):
    year = _get_value(meta, "exam_year", "year", default="2025")
    return _get_value(
        meta,
        "exam_table_title",
        default=f"{province}{year}年中等职业学校毕业生进入普通高校学习专业基础课和专业课考试科目表",
    )


def _get_topic_output_base(meta, topic, output_dir):
    """按"省份 考类"组织输出目录（单级目录，省/市与考类合并）。"""
    province = _safe_path_part(_normalize_province_name(_get_value(meta, "province")))
    category = _safe_path_part(_format_exam_category(meta))
    combined = " ".join(part for part in (province, category) if part and part != "未命名").strip()
    sub_dir = Path(output_dir)
    if combined:
        sub_dir = sub_dir / combined
    return sub_dir


def _format_exam_type(meta):
    """生成文件名和编写说明使用的考试类型。"""
    return str(_get_value(meta, "exam_type_name", "exam_type", default="高职分类考试") or "高职分类考试").strip()


def _build_docx_filename(meta, topic, set_idx=1, variant="解析版"):
    """按考纲百套卷命名规范生成 DOCX 文件名。"""
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    category = _format_exam_category(meta)
    textbook = _extract_textbook_detail(meta, topic)
    course_name = _format_course_name(meta, topic, textbook["name"])
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    sets = int(_get_value(topic, "sets", default=1) or 1)
    set_suffix = f"-{set_idx}" if sets > 1 else ""
    filename = (
        f"第{seq}卷{set_suffix} {paper_name} {paper_type}《{course_name}》"
        f"{province}（{exam_type}）{category} 考纲百套卷（{variant}）.docx"
    )
    return _safe_path_part(filename)


def _build_docx_titles(meta, topic):
    """生成考纲百套卷三行标题，与模板脚本的标题区域保持一致。"""
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    category = _format_exam_category(meta)
    textbook = _extract_textbook_detail(meta, topic)
    course_name = _format_course_name(meta, topic, textbook["name"])
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    category_part = f"《{category}考纲百套卷》" if category else "《考纲百套卷》"
    return (
        f"{province}（{exam_type}）{category_part}第{seq}卷",
        f"《{course_name}》",
        f"{paper_name} {paper_type}",
    )


def _paper_has_blank_before_first_heading(paper_text: str) -> bool:
    """判断第一类大题前是否已经保留空行。"""
    lines = paper_text.split("\n")
    for idx, line in enumerate(lines):
        if re.match(r"^\s*[一二三四五六七八九十][、.．]", line):
            return idx > 0 and not lines[idx - 1].strip()
    return False


def _add_exam_info_header(doc, add_para, add_separator_blank=True) -> None:
    """在标题下方添加考试时间、满分和考生信息。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    line1 = "考试时间：60分钟        满分：100分"
    line2 = "班级：________ 姓名：________ 学号：________ 成绩：________"

    add_para(doc, line1, font_name="宋体", font_size=10.5,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, line2, font_name="宋体", font_size=10.5,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    if add_separator_blank:
        add_para(doc, " ", font_name="宋体", font_size=10.5,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


def _source_docx_from_questions(questions: list[dict[str, Any]] | None) -> Path | None:
    """Return the common source DOCX when all questions came from the same DOCX."""
    paths = []
    for question in questions or []:
        source = question.get("source_docx_path") or question.get("source_path")
        if source:
            paths.append(Path(source))
    if not paths:
        return None
    first = paths[0]
    if all(path == first for path in paths) and first.suffix.lower() == ".docx" and first.exists():
        return first
    return None


def _is_repaired_layout_source(path: Path | None) -> bool:
    return bool(path and "_repaired" in path.stem.lower())


def _format_copied_table(table_elem) -> None:
    """Normalize table text formatting without changing cell/order layout."""
    for paragraph_elem in table_elem.findall(f".//{qn('w:p')}"):
        _format_copied_element(paragraph_elem, {}, None, text_color=None, bold_tags=None)


def _append_repaired_docx_body(doc, source_path: Path) -> int:
    """Copy the repaired DOCX body as the authoritative layout, preserving order exactly."""
    src_doc = _get_cached_src_doc(str(source_path))
    rid_cache: dict = {}
    copied = 0
    for element in src_doc.element.body.iterchildren():
        if element.tag == W_SECTPR:
            continue
        new_element = copy.deepcopy(element)
        _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
        if new_element.tag == qn("w:p"):
            text = _paragraph_visible_text(new_element).strip()
            if text.startswith(("【答案】", "【解析】", "【详解】", "【分析】")):
                _normalize_answer_analysis_label(new_element)
                _format_copied_element(new_element, rid_cache, doc, text_color=_ANSWER_RED, bold_tags=["【答案】", "【解析】", "【详解】", "【分析】"])
            else:
                _format_copied_element(new_element, rid_cache, doc, text_color=None, bold_tags=None)
        elif new_element.tag == qn("w:tbl"):
            _format_copied_table(new_element)
        _append_body_element(doc, new_element)
        copied += 1
    return copied


def _set_paragraph_spacing(paragraph_elem, *, line: str = '360', before: str = '0', after: str = '40') -> None:
    """Apply the standard generated-paper paragraph spacing to a copied paragraph."""
    from lxml import etree
    ppr = paragraph_elem.find(qn('w:pPr'))
    if ppr is None:
        ppr = etree.SubElement(paragraph_elem, qn('w:pPr'))
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(ppr, qn('w:spacing'))
    spacing.set(qn('w:line'), line)
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), before)
    spacing.set(qn('w:after'), after)


def _is_option_text(text: str) -> bool:
    return bool(re.match(r"^\s*[A-H]\s*[.．、)]", str(text or ""), flags=re.I))


def _is_option_paragraph_text(text: str) -> bool:
    text = str(text or "").strip()
    if not text:
        return False
    parts = [part.strip() for part in re.split(r"\t+", text) if part.strip()]
    if len(parts) >= 2:
        return all(_is_option_text(part) for part in parts)
    return _is_option_text(text)


def _split_option_texts(text: str) -> list[str]:
    """Split compact option text like 'A. 6 WB. 10 W' into aligned options."""
    source = re.sub(r"\s+", " ", str(text or "")).strip()
    if not source:
        return []
    matches = list(re.finditer(r"([A-H])\s*[.．、)]", source, flags=re.I))
    if not matches:
        return []
    options = []
    for pos, match in enumerate(matches):
        end = matches[pos + 1].start() if pos + 1 < len(matches) else len(source)
        content = source[match.end():end].strip()
        options.append(f"{match.group(1).upper()}. {content}".strip())
    return options if all(_is_option_text(option) for option in options) else []


def _append_options_table_from_lines(doc, option_lines: list[str]) -> bool:
    options: list[str] = []
    for line in option_lines:
        for part in re.split(r"\t+", str(line or "")):
            part = part.strip()
            if part:
                split = _split_option_texts(part)
                options.extend(split or [part])
    if not options or not all(_is_option_text(opt) for opt in options):
        return False
    _add_options_table_only(doc, options)
    return True


def _is_question_stem_text(text: str) -> bool:
    return bool(re.match(r"^\s*\d+\s*[.．、]", str(text or "")))


def _is_answer_or_analysis_text(text: str) -> bool:
    return str(text or "").strip().startswith(("【答案】", "【解析】", "【详解】", "【分析】"))



def _trim_trailing_option_block(text: str) -> str:
    """Remove option text accidentally carried into answer/analysis paragraphs."""
    source = str(text or "")
    matches = list(re.finditer(r"([A-H])\s*[.．、)]", source, flags=re.I))
    if len(matches) < 2:
        return source
    for pos, match in enumerate(matches[:-1]):
        next_label = matches[pos + 1].group(1).upper()
        if match.group(1).upper() == "A" and next_label == "B" and match.start() > 20:
            return source[:match.start()].rstrip()
    return source


def _trim_paragraph_text_after(paragraph_elem, keep_text: str) -> None:
    remaining = len(keep_text)
    for text_elem in paragraph_elem.findall(f".//{qn('w:t')}"):
        value = text_elem.text or ""
        if remaining >= len(value):
            remaining -= len(value)
            continue
        if remaining > 0:
            text_elem.text = value[:remaining]
            remaining = 0
        else:
            text_elem.text = ""

def _replace_visible_prefix(paragraph_elem, old_len: int, replacement: str) -> None:
    """Replace only the leading visible text, preserving later runs and inline math positions."""
    if old_len <= 0:
        return
    text_nodes = paragraph_elem.findall(f".//{qn('w:t')}")
    remaining = old_len
    inserted = False
    for node in text_nodes:
        value = node.text or ""
        if remaining <= 0:
            break
        if remaining >= len(value):
            node.text = replacement if not inserted else ""
            inserted = True
            remaining -= len(value)
            continue
        suffix = value[remaining:]
        node.text = (replacement if not inserted else "") + suffix
        inserted = True
        remaining = 0
        break


def _normalize_answer_analysis_label(paragraph_elem) -> None:
    """Normalize copied labels while preserving inline formula/image positions."""
    text_nodes = paragraph_elem.findall(f".//{qn('w:t')}")
    if not text_nodes:
        return
    visible = "".join(node.text or "" for node in text_nodes)
    stripped = visible.lstrip()
    leading_ws = visible[: len(visible) - len(stripped)]

    label = ""
    match = re.match(r"^(?:【\s*答案\s*】)+", stripped)
    if match:
        label = "【答案】"
    else:
        match = re.match(r"^(?:【\s*(?:解析|详解|分析)\s*】|(?:解析|详解|分析)\s*[:：])+\s*", stripped)
        if match:
            label = "【解析】"
    if not match:
        return

    replacement = leading_ws + label
    old_len = len(leading_ws) + match.end()
    _replace_visible_prefix(paragraph_elem, old_len, replacement)

    # Only trim accidental trailing option text when no protected inline object can
    # be displaced. Formula/image paragraphs must keep their run order untouched.
    if _paragraph_has_protected_content(paragraph_elem):
        return
    visible = "".join(node.text or "" for node in paragraph_elem.findall(f".//{qn('w:t')}"))
    trimmed = _trim_trailing_option_block(visible)
    if trimmed != visible:
        _trim_paragraph_text_after(paragraph_elem, trimmed)


def _style_table_cell_paragraph(paragraph_elem) -> None:
    _set_paragraph_spacing(paragraph_elem, line='360', before='0', after='0')
    for r_elem in paragraph_elem.findall(qn('w:r')):
        if _run_has_protected_content(r_elem):
            continue
        _set_run_text_style(r_elem, '宋体', None, bold=False, size_half_points='21')


def _table_text_lines(table_elem) -> list[str]:
    lines = []
    for text_elem in table_elem.findall(f".//{qn('w:t')}"):
        value = text_elem.text or ""
        if value.strip():
            lines.append(value.strip())
    return lines


def _table_option_texts(table_elem) -> list[str]:
    options: list[str] = []
    for line in _table_text_lines(table_elem):
        split = _split_option_texts(line)
        if split:
            options.extend(split)
        elif _is_option_text(line):
            options.append(line)
    return options


def _looks_like_options_table(table_elem) -> bool:
    options = _table_option_texts(table_elem)
    return bool(options) and len(options) >= len(_table_text_lines(table_elem))


def _table_has_protected_content(table_elem) -> bool:
    return any(table_elem.find(f'.//{tag}') is not None for tag in _PROTECTED_CONTENT_TAGS)


def _append_normalized_options_table(doc, table_elem) -> bool:
    if _table_has_protected_content(table_elem):
        return False
    return _append_options_table_from_lines(doc, _table_option_texts(table_elem))


def _is_blank_paragraph_element(paragraph_elem) -> bool:
    return not _paragraph_visible_text(paragraph_elem).strip() and not _paragraph_has_protected_content(paragraph_elem)


def _is_structural_blank_before_heading(body: list[Any], idx: int) -> bool:
    if idx <= 0 or idx + 1 >= len(body):
        return False
    prev_el = body[idx - 1]
    next_el = body[idx + 1]
    if prev_el.tag not in (qn('w:p'), qn('w:tbl')) or next_el.tag != qn('w:p'):
        return False
    next_text = _paragraph_visible_text(next_el).strip()
    return bool(re.match(r"^\s*[一二三四五六七八九十][、.．]", next_text))


def _style_copied_options_table(table_elem) -> None:
    """Keep repaired table placement but apply final no-border option styling."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    tbl_pr = table_elem.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table_elem.insert(0, tbl_pr)

    borders = tbl_pr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = borders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            borders.append(border)
        border.set(_qn('w:val'), 'nil')

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), '5000')
    tbl_w.set(qn('w:type'), 'pct')

    for paragraph_elem in table_elem.findall(f".//{qn('w:p')}"):
        _style_table_cell_paragraph(paragraph_elem)


def _copy_paragraph_with_standard_style(doc, paragraph_elem, *, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> None:
    _format_copied_element(paragraph_elem, {}, doc, text_color=text_color, bold_tags=bold_tags)
    _append_body_element(doc, paragraph_elem)


def _copy_repaired_docx_body_formatted(doc, source_path: Path) -> int:
    """Copy repaired DOCX order exactly while applying generated-paper spacing/options style."""
    src_doc = _get_cached_src_doc(str(source_path))
    rid_cache: dict = {}
    body = list(src_doc.element.body.iterchildren())
    copied = 0
    idx = 0
    while idx < len(body):
        element = body[idx]
        if element.tag == W_SECTPR:
            idx += 1
            continue

        if element.tag == qn('w:p'):
            if _is_blank_paragraph_element(element) and not _is_structural_blank_before_heading(body, idx):
                idx += 1
                continue

            text = _paragraph_visible_text(element).strip()
            if _is_question_stem_text(text):
                group = [copy.deepcopy(element)]
                j = idx + 1
                while j < len(body):
                    next_el = body[j]
                    if next_el.tag == W_SECTPR:
                        break
                    if next_el.tag == qn('w:p'):
                        next_text = _paragraph_visible_text(next_el).strip()
                        if _is_answer_or_analysis_text(next_text):
                            break
                        if _is_question_stem_text(next_text):
                            break
                        group.append(copy.deepcopy(next_el))
                        j += 1
                        continue
                    if next_el.tag == qn('w:tbl'):
                        group.append(copy.deepcopy(next_el))
                        j += 1
                        continue
                    break

                option_buffer: list[str] = []
                for new_el in group:
                    _remap_embedded_relationships(new_el, src_doc, doc, rid_cache)
                    if new_el.tag == qn('w:p'):
                        if _is_blank_paragraph_element(new_el):
                            copied += 1
                            continue
                        group_text = _paragraph_visible_text(new_el).strip()
                        if _is_option_paragraph_text(group_text) and not _paragraph_has_protected_content(new_el):
                            option_buffer.append(group_text)
                            copied += 1
                            continue
                        if option_buffer:
                            _append_options_table_from_lines(doc, option_buffer)
                            option_buffer = []
                        _copy_paragraph_with_standard_style(doc, new_el)
                    elif new_el.tag == qn('w:tbl'):
                        if option_buffer:
                            _append_options_table_from_lines(doc, option_buffer)
                            option_buffer = []
                        if _looks_like_options_table(new_el) and _append_normalized_options_table(doc, new_el):
                            pass
                        elif _looks_like_options_table(new_el):
                            _style_copied_options_table(new_el)
                            _append_body_element(doc, new_el)
                        else:
                            _format_copied_table(new_el)
                            _append_body_element(doc, new_el)
                    copied += 1
                if option_buffer:
                    _append_options_table_from_lines(doc, option_buffer)
                idx = j
                continue

            new_element = copy.deepcopy(element)
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            if _is_answer_or_analysis_text(text):
                _normalize_answer_analysis_label(new_element)
                _copy_paragraph_with_standard_style(
                    doc, new_element,
                    text_color=_ANSWER_RED,
                    bold_tags=["【答案】", "【解析】", "【详解】", "【分析】"],
                )
            elif re.match(r"^\s*[一二三四五六七八九十][、.．]", text):
                _set_paragraph_spacing(new_element, line='360', before='0', after='120')
                for r_elem in new_element.findall(qn('w:r')):
                    if not _run_has_protected_content(r_elem):
                        _set_run_text_style(r_elem, '黑体', None, bold=True, size_half_points='24')
                _append_body_element(doc, new_element)
            else:
                _copy_paragraph_with_standard_style(doc, new_element)
            copied += 1
            idx += 1
            continue

        new_element = copy.deepcopy(element)
        _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
        if new_element.tag == qn('w:tbl'):
            if _looks_like_options_table(new_element) and _append_normalized_options_table(doc, new_element):
                copied += 1
                idx += 1
                continue
            if _looks_like_options_table(new_element):
                _style_copied_options_table(new_element)
            else:
                _format_copied_table(new_element)
        _append_body_element(doc, new_element)
        copied += 1
        idx += 1
    return copied


def generate_docx(meta, topic, set_idx, paper_text, output_dir, needs_manual_review=False, questions=None):
    """将试卷文本生成为格式化的 DOCX 文件"""
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本" / "核心脚本"))
    from docx_utils1 import (
        copy_template, add_editorial_note,
        add_paragraph_with_style, add_question_options_table,
        add_structured_choice_question,
        add_labeled_text, save_docx
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 清理 AI 返回文本
    paper_text = _clean_paper_text(paper_text)

    # 确定文件名和标题元数据
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    textbook = _extract_textbook_detail(meta, topic)
    textbook_name = textbook["name"]
    publisher = textbook["publisher"]
    edition = textbook["edition"]
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    course_name = _format_course_name(meta, topic, textbook_name)

    filename = _build_docx_filename(meta, topic, set_idx)
    if needs_manual_review:
        filename = f"（待人工审核）{filename}"

    # 按"省份/考类/课程"组织子目录，与规划表位置保持一致
    sub_dir = _get_topic_output_base(meta, topic, output_dir)
    os.makedirs(sub_dir, exist_ok=True)
    _copy_common_scripts_to_output(sub_dir)
    output_path = sub_dir / filename

    # 使用 02_配置资源/模板和资源/template.docx 作为唯一模板来源，保留其页眉、页脚和样式。
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"未找到 DOCX 模板：{TEMPLATE_PATH}")
    doc = copy_template(str(TEMPLATE_PATH), str(output_path))

    # 不强制覆盖模板页边距，避免破坏 template.docx 中已配置的页眉页脚版式。

    # 编写说明；新版考纲百套卷模板不再插入 separator 图片。
    sep_img = None
    all_courses, course_count = _format_all_courses(meta, topic, course_name)
    specific_content = _format_specific_content(topic)
    add_editorial_note(
        doc,
        textbook_name=textbook_name,
        edition=edition,
        chapter_seq=seq,
        knowledge_scope=_get_value(topic, "knowledge", "point_content", "topic", default=paper_name),
        province=province,
        publisher=publisher,
        separator_image=sep_img,
        exam_type=exam_type,
        series_name="考纲百套卷",
        volume_label=f"第{seq}卷",
        paper_type=paper_type,
        paper_name=paper_name,
        course_name=course_name,
        category=_format_exam_category(meta),
        all_courses=all_courses,
        course_count=course_count,
        exam_table_title=_format_exam_table_title(meta, province),
        specific_content=specific_content,
    )

    # 标题
    title1, title2, title3 = _build_docx_titles(meta, topic)

    add_paragraph_with_style(doc, title1, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title2, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title3, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 考试信息行
    _add_exam_info_header(
        doc,
        add_paragraph_with_style,
        add_separator_blank=not _paper_has_blank_before_first_heading(paper_text),
    )

    repaired_source_docx = _source_docx_from_questions(questions)
    if _is_repaired_layout_source(repaired_source_docx):
        _copy_repaired_docx_body_formatted(doc, repaired_source_docx)
        save_docx(doc, str(output_path))
        validate_docx_preserves_images(output_path, questions)
        residuals = _find_residual_formula_text(doc)
        if residuals:
            print(f"  [!] 公式转换后仍发现 {len(residuals)} 处疑似 LaTeX/math 残留，请检查 DOCX：{output_path}")
            for sample in residuals[:3]:
                print(f"      - {sample}")
        return str(output_path)

    # 答案/解析使用红色字体
    RED_COLOR = (255, 0, 0)

    # 逐行解析并输出
    lines = paper_text.split("\n")
    protected_questions = questions or []
    question_by_no = {
        int(question.get("question_no")): question
        for question in protected_questions
        if isinstance(question.get("question_no"), int) or str(question.get("question_no", "")).isdigit()
    }
    question_cursor = 0
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 大题标题（加粗，黑体）
        if re.match(r"^[一二三四五六七八九十][、.．]", line):
            add_paragraph_with_style(doc, line, font_name="黑体", font_size=12,
                                     bold=True, space_after=6)
            i += 1
            continue

        # 题目开头（数字编号）
        if re.match(r"^\d+[\.．、]", line):
            no_match = re.match(r"^(\d+)[\.．、]", line)
            current_question = None
            if no_match:
                current_question = question_by_no.get(int(no_match.group(1)))
            if current_question is None:
                current_question = protected_questions[question_cursor] if question_cursor < len(protected_questions) else None
            question_cursor += 1
            if _question_is_protected(current_question):
                rendered = _append_protected_question_reflowed(
                    doc, current_question, question_cursor,
                    add_question_options_table, add_paragraph_with_style, add_labeled_text,
                )
                if rendered:
                    j = i + 1
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if re.match(r"^\d+[\.．、]", next_line):
                            break
                        if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                            break
                        j += 1
                    i = j
                    continue
                print(f"  → 警告：第{current_question.get('question_no')}题带图题重排失败，回退为文本生成。")

            # 收集题目完整文本（直到下一个【答案】或下一题）
            q_lines = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if next_line.startswith(("【答案】", "【解析】", "【详解】", "【分析】")):
                    break
                if not next_line:
                    # 防止题目文本和【答案】之间的偶发空行进入 DOCX。
                    lookahead = j + 1
                    while lookahead < len(lines) and not lines[lookahead].strip():
                        lookahead += 1
                    if lookahead < len(lines) and lines[lookahead].strip().startswith(("【答案】", "【解析】", "【详解】", "【分析】")):
                        j = lookahead
                        continue
                    j += 1
                    continue
                if re.match(r"^\d+[\.．、]", next_line):
                    break
                if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                    break
                q_lines.append(lines[j])
                j += 1

            q_text = "\n".join(q_lines)

            # 结构化图片题：题干文字 → 题干图 → 图片选项（protected 原始题块优先级更高）
            if _question_has_structured_images(current_question):
                normalized = _normalize_structured_choice_question(current_question, q_text)
                if normalized:
                    rendered = add_structured_choice_question(
                        doc,
                        normalized["stem_text"],
                        normalized["options"],
                        stem_images=normalized["stem_images"],
                    )
                    if rendered:
                        i = j
                        continue
                    q_no = current_question.get("question_no") if current_question else ""
                    print(f"  → 警告：第{q_no}题结构化图片排版失败，回退为文本生成。")

            # 判断题不输出 A/B 正确/错误选项；选择题才排选项表。
            if _is_true_false_question(current_question, q_text):
                add_paragraph_with_style(doc, _strip_true_false_options(q_text), font_size=10.5, space_after=2)
            elif re.search(r"[A-D][\.．]\s*\S", q_text):
                add_question_options_table(doc, q_text)
            else:
                add_paragraph_with_style(doc, q_text, font_size=10.5, space_after=2)

            i = j
            continue

        # 【答案】— 红色字体；多行答案继续归入同一红色段落
        if line.startswith("【答案】"):
            answer_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith(("【解析】", "【详解】", "【分析】", "【答案】"))
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    answer_lines.append(next_line)
                j += 1
            answer_content = "\n".join(answer_lines)

            # 答案图片：从 source DOCX 按段落索引复制原段（与题干图片处理方式一致）
            answer_copied = False
            if current_question and not _question_is_protected(current_question):
                answer_refs = (current_question.get("image_refs") or {}).get("answer") or []
                answer_indices = current_question.get("answer_paragraph_indices") or []
                if answer_refs and answer_indices:
                    source_path = current_question.get("source_docx_path") or current_question.get("source_path")
                    if source_path:
                        copied = _copy_source_paragraphs(doc, source_path, answer_indices, text_color=_ANSWER_RED, bold_tags=["【答案】"])
                        answer_copied = copied > 0

            if answer_copied:
                # 源段落已包含完整答案（标签+图片+文字），只需追加可能的纯文本后续行
                if answer_content:
                    add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            else:
                add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            i = j
            continue

        # 【解析】/【详解】/【分析】— 红色字体，统一显示为【解析】
        if line.startswith(("【解析】", "【详解】", "【分析】")):
            label = "【解析】"
            explanation_lines = [line[line.index("】") + 1:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith(("【答案】", "【解析】", "【详解】"))
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    explanation_lines.append(next_line)
                j += 1
            explanation_content = "\n".join(explanation_lines)
            explanation_content = _strip_analysis_label(explanation_content)
            if not _should_emit_analysis(current_question, explanation_content):
                i = j
                continue

            # 解析图片：从 source DOCX 按段落索引复制原段（与题干图片处理方式一致）
            analysis_copied = False
            if current_question and not _question_is_protected(current_question):
                analysis_refs = (current_question.get("image_refs") or {}).get("analysis") or []
                analysis_indices = current_question.get("analysis_paragraph_indices") or []
                if analysis_refs and analysis_indices:
                    source_path = current_question.get("source_docx_path") or current_question.get("source_path")
                    if source_path:
                        copied = _copy_source_paragraphs(doc, source_path, analysis_indices, text_color=_ANSWER_RED, bold_tags=["【解析】", "【详解】", "【分析】"])
                        analysis_copied = copied > 0

            if analysis_copied:
                # 源段落已包含完整解析（标签+图片+文字），只需追加可能的纯文本后续行
                if explanation_content:
                    add_labeled_text(doc, label, explanation_content, color=RED_COLOR)
            else:
                add_labeled_text(doc, label, explanation_content, color=RED_COLOR)
            i = j
            continue

        # 普通文本行
        if line:
            add_paragraph_with_style(doc, line, font_size=10.5, space_after=2)

        i += 1

    save_docx(doc, str(output_path))
    validate_docx_preserves_images(output_path, questions)
    residuals = _find_residual_formula_text(doc)
    if residuals:
        print(f"  [!] 公式转换后仍发现 {len(residuals)} 处疑似 LaTeX/math 残留，请检查 DOCX：{output_path}")
        for sample in residuals[:3]:
            print(f"      - {sample}")
    return str(output_path)
