# -*- coding: utf-8 -*-
"""
通用细目表生成器 — 从考点规划总表 xlsx 自动生成全部专题/课程综合卷细目表

用法：
  python generate_mesh.py "04_生成输出/生产规划/重庆市 电子信息类/"

行为：
  1. 在指定目录下查找 *考点规划总表.xlsx
  2. 解析规划表数据（含合并单元格）
  3. 提取省份/考类信息
  4. 为每个专题生成「专题训练卷细目表」
  5. 为每门课程生成「课程综合卷细目表」
  6. 全部输出到同一目录
"""

from __future__ import annotations

import argparse
import os
import re
import sys
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 1. 规划表解析
# ============================================================

def find_plan_xlsx(search_dir: str) -> str:
    """在指定目录中查找 *考点规划总表.xlsx 文件，返回第一个匹配的完整路径。"""
    d = Path(search_dir)
    if not d.is_dir():
        raise FileNotFoundError(f"目录不存在: {search_dir}")
    candidates = list(d.glob("*考点规划总表.xlsx"))
    if not candidates:
        raise FileNotFoundError(f"在 {search_dir} 中未找到 *考点规划总表.xlsx")
    if len(candidates) > 1:
        print(f"[警告] 找到多个规划表文件，使用第一个: {candidates[0].name}")
    return str(candidates[0])


def _forward_fill(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """对 A/B/G/H/I/J 列做向下填充，解决合并单元格读取为 None 的问题。"""
    for key in ("A", "B", "G", "H", "I", "J"):
        carry = None
        for r in rows:
            if r.get(key) is not None and str(r[key]).strip():
                carry = r[key]
            else:
                r[key] = carry
    return rows


def parse_plan(filepath: str) -> dict[str, Any]:
    """从规划表 xlsx 解析出课程、专题、考点结构。"""
    wb = openpyxl.load_workbook(filepath)
    ws = wb.active

    # ---- 读取标题，推断省份/考类 ----
    title = str(ws["A1"].value or "")
    prov_cat_match = re.search(r"《(.+?)》", title)
    prov_cat = prov_cat_match.group(1) if prov_cat_match else ""
    province = prov_cat  # 先整体使用
    # 如果标题中无空格，从目录名拆分
    if " " not in prov_cat:
        dir_name = Path(filepath).parent.name
        if " " in dir_name:
            prov_cat = dir_name
    # 拆分为省份和考类
    parts = prov_cat.rsplit(" ", 1)
    province = parts[0] if len(parts) > 1 else prov_cat

    # ---- 读取题型结构（底部汇总表区域） ----
    qtype_map = _parse_question_types(ws)

    # ---- 读取数据行（仅读取 E 列为 1 的行：真正的考点训练卷数据行） ----
    raw_rows: list[dict[str, Any]] = []
    data_start = 7
    for row_idx in range(data_start, ws.max_row + 1):
        e_val = ws.cell(row=row_idx, column=5).value
        # E列=1 才是考点训练卷行；E列="" 或其他值则跳过（题型汇总表、合计行等）
        if e_val is None or str(e_val).strip() != "1":
            continue
        row_data = {}
        for col, key in enumerate(("A","B","C","D","E","F","G","H","I","J"), 1):
            v = ws.cell(row=row_idx, column=col).value
            row_data[key] = str(v).strip() if v is not None else None
        raw_rows.append(row_data)

    if not raw_rows:
        raise ValueError("未解析到任何数据行")

    _forward_fill(raw_rows)

    # ---- 构建结构化数据 ----
    courses: dict[str, dict] = {}
    for r in raw_rows:
        course = r["A"]
        topic = r["B"]
        if not course or not topic:
            continue

        if course not in courses:
            comp_vols = _parse_vol_range(r["J"])
            courses[course] = {
                "topics": {},
                "comp_count": int(r["I"]) if r["I"] and r["I"].isdigit() else len(comp_vols),
                "comp_vols": comp_vols,
            }

        if topic not in courses[course]["topics"]:
            topic_vol = _parse_vol_number(r["H"])
            courses[course]["topics"][topic] = {
                "vol": topic_vol,
                "checkpoints": [],
            }

        d_text = r["D"] or ""
        # 为每个知识点评定考纲要求（从段首动词提取）
        lines = [l.strip() for l in d_text.split("\n") if l.strip()]
        req_str = _extract_requirement(lines)

        courses[course]["topics"][topic]["checkpoints"].append({
            "C": r["C"],
            "D": d_text,
            "req": req_str,
            "F": r["F"],
        })

    return {
        "province": province,
        "province_category": prov_cat,
        "courses": courses,
        "qtype_map": qtype_map,
    }


def _parse_vol_range(j_text: str) -> list[int]:
    """解析 '第34-36卷' → [34, 35, 36] 或 '第93-94卷' → [93, 94]"""
    if not j_text:
        return []
    nums = re.findall(r"\d+", j_text)
    if len(nums) >= 2:
        return list(range(int(nums[0]), int(nums[-1]) + 1))
    elif len(nums) == 1:
        return [int(nums[0])]
    return []


def _parse_vol_number(h_text: str) -> int:
    """解析 '第29卷' → 29"""
    if not h_text:
        return 0
    nums = re.findall(r"\d+", h_text)
    return int(nums[0]) if nums else 0


def _extract_requirement(lines: list[str]) -> str:
    """从知识点文本中提取考纲要求关键词。"""
    reqs = set()
    for line in lines:
        for kw in ["掌握", "理解", "了解"]:
            if kw in line:
                reqs.add(kw)
    if not reqs:
        return "了解"
    # 优先级：掌握 > 理解 > 了解
    for kw in ["掌握", "理解", "了解"]:
        if kw in reqs:
            return kw
    return "了解"


def _parse_question_types(ws) -> dict[str, list[dict]]:
    """从规划表底部题型汇总表区域读取每门课程的题型结构。"""
    qtype_map: dict[str, list[dict]] = {}
    current_course = None

    # 从第80行附近开始扫描题型区域
    for row_idx in range(max(ws.max_row - 30, 80), ws.max_row + 1):
        a_val = str(ws.cell(row=row_idx, column=1).value or "")
        b_val = str(ws.cell(row=row_idx, column=2).value or "")

        # 课程标题行
        if "题型" in b_val and a_val:
            # 提取课程名
            course_name = re.sub(r"[（(].*?[)）]", "", a_val).strip()
            if course_name:
                current_course = course_name
                qtype_map[current_course] = []
            continue

        # 题型数据行
        if current_course and b_val and b_val != "题型" and "各课程题型" not in a_val:
            b_cell = ws.cell(row=row_idx, column=2).value
            c_cell = ws.cell(row=row_idx, column=3).value
            d_cell = ws.cell(row=row_idx, column=4).value
            e_cell = ws.cell(row=row_idx, column=5).value
            if b_cell and c_cell:
                qtype_map[current_course].append({
                    "type": str(b_cell).strip(),
                    "count": str(c_cell).strip(),
                    "score_per": str(d_cell).strip() if d_cell else "",
                    "subtotal": str(e_cell).strip() if e_cell else "",
                })

    return qtype_map


# ============================================================
# 2. 知识点权重与题量分配
# ============================================================

def _kp_weight(kp_line: str, req: str) -> str:
    """判断单个知识点的权重：high / medium / low"""
    text = kp_line.strip()

    # 考纲要求驱动
    if "掌握" in req:
        return "high"
    if "理解" in req:
        return "medium"
    if "了解" in req:
        return "low"

    # 文本内容驱动
    high_kw = ["计算", "应用", "定律", "编程", "公式", "方法", "原理"]
    med_kw = ["分析", "判别", "特点", "功能", "关系", "结构"]

    if any(kw in text for kw in high_kw):
        return "high"
    if any(kw in text for kw in med_kw):
        return "medium"
    return "low"


def _flatten_checkpoints(checkpoints: list[dict]) -> list[dict]:
    """将考点列表展平为知识点列表。"""
    all_kps = []
    for cp in checkpoints:
        lines = [l.strip() for l in cp["D"].split("\n") if l.strip()]
        for line in lines:
            all_kps.append({
                "cp_name": cp["C"],
                "kp_text": line,
                "req": cp["req"],
                "weight": _kp_weight(line, cp["req"]),
            })
    return all_kps


def allocate_topic_questions(checkpoints: list[dict]) -> list[dict]:
    """为专题训练卷分配44题。"""
    TOTAL_Q = 44
    all_kps = _flatten_checkpoints(checkpoints)
    if not all_kps:
        return []

    high = [k for k in all_kps if k["weight"] == "high"]
    med = [k for k in all_kps if k["weight"] == "medium"]
    low = [k for k in all_kps if k["weight"] == "low"]
    all_sorted = high + med + low

    def _circular(seq, idx):
        return seq[idx % len(seq)] if seq else all_sorted[idx % len(all_sorted)]

    questions = []
    q = 0
    idx = 0

    # 单选16题 (1-16)：简单为主，末尾少数适中
    for i in range(16):
        q += 1
        kp = _circular(all_sorted, idx); idx += 1
        diff = "适中" if (kp["weight"] == "high" and i >= 12) else "简单"
        questions.append(_make_q(q, "单项选择题", diff, kp, 1))

    # 判断16题 (17-32)：全部简单
    for i in range(16):
        q += 1
        kp = _circular(all_sorted, idx); idx += 1
        questions.append(_make_q(q, "判断题", "简单", kp, 1))

    # 填空8题 (33-40)：重要知识点适中，其余简单
    fill_pool = (high * 2 + med + low)[:16] if high else all_sorted * 2
    fi = 0
    for i in range(8):
        q += 1
        kp = fill_pool[fi % len(fill_pool)]; fi += 1
        diff = "适中" if kp["weight"] == "high" else "简单"
        questions.append(_make_q(q, "填空题", diff, kp, 1))

    # 综合4题 (41-44)：2困难+2适中，全部来自高权重知识点
    comp_pool = high[:4] if len(high) >= 4 else (high + med)[:4]
    if len(comp_pool) < 4:
        comp_pool = (comp_pool + low)[:4]
    for i, kp in enumerate(comp_pool[:4]):
        q += 1
        diff = "困难" if i < 2 else "适中"
        questions.append(_make_q(q, "综合题", diff, kp, 1))

    return questions


def allocate_comp_questions(course_checkpoints: list[dict], vol_idx: int, total_vols: int) -> list[dict]:
    """为课程综合卷之一分配44题。vol_idx 从0开始。"""
    TOTAL_Q = 44
    all_kps = _flatten_checkpoints(course_checkpoints)
    total = len(all_kps)
    kps_per = total // total_vols
    rem = total % total_vols

    start = vol_idx * kps_per + min(vol_idx, rem)
    end = start + kps_per + (1 if vol_idx < rem else 0)
    my_kps = all_kps[start:end]

    high = [k for k in my_kps if k["weight"] == "high"]
    med = [k for k in my_kps if k["weight"] == "medium"]
    low = [k for k in my_kps if k["weight"] == "low"]
    all_local = high + med + low

    def _safe_get(seq, idx):
        return seq[idx % len(seq)] if seq else my_kps[idx % len(my_kps)] if my_kps else {"cp_name":"综合","kp_text":"综合考查","req":"综合","weight":"medium"}

    questions = []
    q = 0
    idx = 0

    for i in range(16):
        q += 1; kp = _safe_get(all_local, idx); idx += 1
        diff = "适中" if kp["weight"] == "high" and i >= 12 else "简单"
        questions.append(_make_q(q, "单项选择题", diff, kp, 1))

    for i in range(16):
        q += 1; kp = _safe_get(all_local, idx); idx += 1
        questions.append(_make_q(q, "判断题", "简单", kp, 1))

    fill_pool = (high * 2 + med) if high else all_local * 2
    fi = 0
    for i in range(8):
        q += 1; kp = _safe_get(fill_pool, fi); fi += 1
        diff = "适中" if kp["weight"] == "high" else "简单"
        questions.append(_make_q(q, "填空题", diff, kp, 1))

    comp_pool = high[:4] if len(high) >= 4 else (high + med)[:4]
    if len(comp_pool) < 4:
        comp_pool = (comp_pool + low)[:4]
    for i, kp in enumerate(comp_pool[:4]):
        q += 1
        diff = "困难" if i < 2 else "适中"
        questions.append(_make_q(q, "综合题", diff, kp, 1))

    return questions


def _make_q(q_num, q_type, difficulty, kp, _unused=1):
    intent_map = {
        "单项选择题": {
            "简单": f"考查{kp['cp_name']}中的基本概念。",
            "适中": f"考查{kp['cp_name']}中的理解与应用。",
            "困难": f"深入考查{kp['cp_name']}的辨析。",
        },
        "判断题": {
            "简单": f"考查{kp['cp_name']}中的结论判断。",
            "适中": f"考查{kp['cp_name']}中的分析判别。",
            "困难": f"考查{kp['cp_name']}的深层理解。",
        },
        "填空题": {
            "简单": f"考查{kp['cp_name']}的基础内容填空。",
            "适中": f"考查{kp['cp_name']}中关键参数或公式。",
            "困难": f"考查{kp['cp_name']}的综合填空。",
        },
        "综合题": {
            "简单": f"综合考查{kp['cp_name']}的基础应用。",
            "适中": f"综合考查{kp['cp_name']}的理解应用。",
            "困难": f"综合考查{kp['cp_name']}，涉及计算分析。",
        },
    }
    intent = intent_map.get(q_type, {}).get(difficulty, f"考查{kp['cp_name']}。")
    return {
        "q_num": q_num, "q_type": q_type, "difficulty": difficulty,
        "cp_name": kp["cp_name"], "kp_text": kp["kp_text"],
        "req": kp["req"], "intent": intent,
    }


# ============================================================
# 3. Docx 生成
# ============================================================

HEADERS = ["题号", "题型", "难度", "考查内容", "对应考点", "对应知识点", "考纲要求", "出题意图/覆盖说明"]
COL_WIDTHS = [Cm(1.0), Cm(2.0), Cm(1.2), Cm(6.0), Cm(3.5), Cm(5.0), Cm(3.0), Cm(5.0)]


def safe_name(s: str) -> str:
    for ch in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
        s = s.replace(ch, '-')
    return s


def _add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_cell(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.font.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)


def _set_header_cell(cell, text):
    _set_cell(cell, text, bold=True, size=9)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E1F2" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _build_table(doc, questions):
    table = doc.add_table(rows=1 + len(questions), cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, w in enumerate(COL_WIDTHS):
        for row in table.rows:
            row.cells[i].width = w
    for j, h in enumerate(HEADERS):
        _set_header_cell(table.rows[0].cells[j], h)
    for i, q in enumerate(questions):
        row = table.rows[i + 1]
        vals = [str(q["q_num"]), q["q_type"], q["difficulty"],
                q["kp_text"][:120], q["cp_name"], q["kp_text"],
                q["req"], q["intent"]]
        for j, v in enumerate(vals):
            align = WD_ALIGN_PARAGRAPH.LEFT if j >= 3 else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell(row.cells[j], v, size=8, align=align)
        row.height = Cm(0.55)
    _add_borders(table)
    return table


def _add_cover_paragraph(doc, text_lines, bold_indices=None):
    if bold_indices is None:
        bold_indices = set()
    for i, (text, is_bold) in enumerate(text_lines):
        if i == 0:
            p = doc.add_paragraph()
        else:
            p.add_run(" ")
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.bold = is_bold
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def generate_topic_docx(province, course_name, topic_name, vol_num, checkpoints, output_dir):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"专题训练卷 第{vol_num}卷 细目表")
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14); run.font.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"地区：{province}    |    课程：{course_name}    |    专题：{topic_name}    |    卷号：第{vol_num}卷")
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    doc.add_paragraph()

    questions = allocate_topic_questions(checkpoints)
    _build_table(doc, questions)

    doc.add_paragraph()
    stats = {}
    for q in questions:
        stats[q["difficulty"]] = stats.get(q["difficulty"], 0) + 1
    _add_cover_paragraph(doc, [
        (f"【覆盖说明】本卷共44题（单选16+判断16+填空8+综合4），覆盖「{topic_name}」专题下全部{len(checkpoints)}个考点。", True),
        (f"难度分布：简单{stats.get('简单',0)}题 / 适中{stats.get('适中',0)}题 / 困难{stats.get('困难',0)}题。", False),
    ])

    os.makedirs(output_dir, exist_ok=True)
    fname = f"专题训练卷_第{vol_num}卷_{province}_{safe_name(topic_name)}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    print(f"  [OK] {fname}")
    return fpath


def generate_comp_docx(province, course_name, vol_num, vol_idx, total_vols, all_cps, output_dir):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)

    label = f"（第{vol_idx+1}套）" if total_vols > 1 else ""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"课程综合卷 第{vol_num}卷 细目表 {label}")
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14); run.font.bold = True

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"地区：{province}    |    课程：{course_name}    |    卷号：第{vol_num}卷    |    共{total_vols}套综合卷")
    run.font.name = '宋体'; run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

    doc.add_paragraph()

    questions = allocate_comp_questions(all_cps, vol_idx, total_vols)
    _build_table(doc, questions)

    doc.add_paragraph()
    stats = {}
    for q in questions:
        stats[q["difficulty"]] = stats.get(q["difficulty"], 0) + 1
    _add_cover_paragraph(doc, [
        (f"【覆盖说明】本卷为{total_vols}套课程综合卷中的第{vol_idx+1}套，共44题。{total_vols}套合起来必须覆盖「{course_name}」课程的全部考点。", True),
        (f"难度分布：简单{stats.get('简单',0)}题 / 适中{stats.get('适中',0)}题 / 困难{stats.get('困难',0)}题。", False),
    ])

    os.makedirs(output_dir, exist_ok=True)
    fname = f"课程综合卷_第{vol_num}卷_{province}_{safe_name(course_name)}.docx"
    fpath = os.path.join(output_dir, fname)
    doc.save(fpath)
    print(f"  [OK] {fname}")
    return fpath


# ============================================================
# 4. 主流程
# ============================================================

def run(target: str) -> int:
    """target 可以是目录或 xlsx 文件路径。若是文件则以其所在目录作为工作目录。"""
    target_path = Path(target)
    if target_path.is_file():
        plan_path = str(target_path)
        output_dir = str(target_path.parent)
    else:
        plan_path = find_plan_xlsx(target)
        output_dir = target

    print(f"规划表: {plan_path}")
    print(f"输出目录: {output_dir}")

    data = parse_plan(plan_path)
    province = data["province"]
    courses = data["courses"]
    qtype_map = data["qtype_map"]

    print(f"地区/考类: {province}")
    print(f"课程数: {len(courses)}")
    print(f"题型结构: {qtype_map if qtype_map else '使用默认 单选16+判断16+填空8+综合4'}")

    topic_count = 0
    comp_count = 0

    for course_name, cdata in courses.items():
        print(f"\n{'='*60}")
        print(f"课程: {course_name}")

        # 专题训练卷
        for topic_name, tdata in cdata["topics"].items():
            vol_num = tdata["vol"]
            checkpoints = tdata["checkpoints"]
            print(f"  专题训练卷 第{vol_num}卷: {topic_name} ({len(checkpoints)}考点)")
            generate_topic_docx(province, course_name, topic_name, vol_num, checkpoints, output_dir)
            topic_count += 1

        # 课程综合卷
        comp_vols = cdata["comp_vols"]
        all_cps = []
        for tdata in cdata["topics"].values():
            all_cps.extend(tdata["checkpoints"])
        print(f"  课程综合卷 ({len(comp_vols)}套): {len(all_cps)}个考点")
        for i, vol_num in enumerate(comp_vols):
            generate_comp_docx(province, course_name, vol_num, i, len(comp_vols), all_cps, output_dir)
            comp_count += 1

    print(f"\n{'='*60}")
    print(f"生成完毕：专题{ topic_count }个 + 课程综合{ comp_count }个 = 共{ topic_count + comp_count }个细目表")
    return 0


PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _search_all_plans() -> list[Path]:
    """在 04_生成输出/生产规划/ 下递归搜索所有 *考点规划总表.xlsx。"""
    search_root = PROJECT_ROOT / "04_生成输出" / "生产规划"
    if not search_root.is_dir():
        return []
    return sorted(search_root.rglob("*考点规划总表.xlsx"))


def _pick_plan(plans: list[Path]) -> Path | None:
    """列出候选文件让用户选择，返回选中的路径。"""
    print("\n找到以下考点规划总表：\n")
    for i, p in enumerate(plans, 1):
        rel = p.relative_to(PROJECT_ROOT)
        print(f"  [{i}] {rel}")
    print(f"\n  输入编号选择（1-{len(plans)}），输入 q 退出")
    while True:
        try:
            choice = input("> ").strip()
        except (EOFError, KeyboardInterrupt):
            print()
            return None
        if choice.lower() in ("q", "quit", "exit"):
            return None
        if choice.isdigit():
            idx = int(choice)
            if 1 <= idx <= len(plans):
                return plans[idx - 1]
        print(f"  请输入 1-{len(plans)} 之间的数字，或 q 退出")


def main():
    parser = argparse.ArgumentParser(
        description="通用细目表生成器 — 从考点规划总表生成专题/综合卷细目表"
    )
    parser.add_argument(
        "plan", nargs="?",
        help="考点规划总表 xlsx 路径或其所在目录。省略则列出全部候选供选择。"
    )
    args = parser.parse_args()

    if args.plan:
        raise SystemExit(run(args.plan))

    # 无参数：搜索并让用户选择
    plans = _search_all_plans()
    if not plans:
        print("未找到任何 *考点规划总表.xlsx，请检查 04_生成输出/生产规划/ 目录。")
        raise SystemExit(1)

    chosen = _pick_plan(plans)
    if chosen is None:
        print("已取消。")
        raise SystemExit(0)

    print(f"\n已选择: {chosen.relative_to(PROJECT_ROOT)}")
    raise SystemExit(run(str(chosen)))


if __name__ == "__main__":
    main()
