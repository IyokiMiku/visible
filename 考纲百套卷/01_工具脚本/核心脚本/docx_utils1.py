"""Word 文档操作公共工具（基于 python-docx）

统一管理考纲百套卷 DOCX 的排版样式，新对话中生成试卷只需调用这里的功能即可保持一致格式。

核心功能：
  copy_template(tpl, out)        — 复制模板 docx，返回 Document
  set_margins(doc, t,b,l,r)      — 设页边距（默认 2.54/2.54/3.18/3.18 cm）
  add_editorial_note(doc, ...)    — 编写说明蓝框+分隔图（核心排版入口）
  add_paragraph_with_style(doc, ...) — 段落（中英自动分字体）
  add_heading_with_style(doc, ...)   — 标题（兼容中文模板缺样式）
  add_labeled_text(doc, ...)         — 【答案】/【解析】标签段落
  add_table_with_style(doc, ...)     — 带样式表格
  save_docx(doc, out)            — 保存文档

中英混排：所有文字输出函数通过 _add_mixed_run() 自动将拉丁字符设为
Times New Roman，中文字符保持指定字体（宋体/楷体/黑体）。

编写说明格式：
  - 蓝色虚线单格表格 (#4472C4, 1.5pt)
  - 楷体 10.5pt 加粗，字体颜色 #4472C4
  - 单倍行距，段前段后 4pt，首行缩进 2 字符
  - 可选分隔图：4.60cm 宽，左对齐，紧贴后续标题
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import copy
import os
import re
from pathlib import Path

try:
    from lxml import etree
except ImportError:  # 允许没有依赖时降级为普通文本公式
    etree = None

try:
    import latex2mathml.converter
except ImportError:  # 允许没有依赖时降级为普通文本公式
    latex2mathml = None

# -- 字体规则：中文保持指定字体，数字/字母统一 Times New Roman --
LATIN_FONT = 'Times New Roman'

def _has_latin(text):
    return bool(re.search(r'[a-zA-Z0-9]', text))

def _split_cjk_latin(text):
    """将文本拆分为中文/拉丁文交替的片段列表"""
    tokens = []
    current = ''
    current_type = None
    latin_pattern = re.compile(r'[a-zA-Z0-9\.\,\;\:\!\?\-\+\=\/\*\(\)\[\]\{\}\<\>\~\@\#\$\%\^\&\|\\\'\u03c0\u03b1-\u03c9\u0391-\u03a9\u2126\u00b0\u00d7\u00b7\u221a\u2264\u2265\u2248\u2260\u221e\u03bc\u2103]+')
    
    for ch in text:
        ch_type = 'latin' if latin_pattern.match(ch) else 'cjk'
        if current_type is None:
            current_type = ch_type
            current = ch
        elif ch_type == current_type:
            current += ch
        else:
            tokens.append((current, current_type))
            current = ch
            current_type = ch_type
    if current:
        tokens.append((current, current_type))
    return tokens

def _add_formula_run(para, formula_text, font_size, bold=False, font_color=None):
    """添加简易 LaTeX 公式片段：字母/数字 Times New Roman，并处理常见上下标。"""
    text = _latex_to_plain_text(formula_text)
    if "\n" in text:
        for idx, part in enumerate(text.split("\n")):
            if idx > 0:
                para.add_run().add_break()
            if part:
                _add_formula_run(para, part, font_size, bold, font_color)
        return

    i = 0
    while i < len(text):
        ch = text[i]
        if ch in ('_', '^') and i + 1 < len(text):
            is_sub = ch == '_'
            j = i + 1
            if text[j] == '{':
                k = text.find('}', j + 1)
                if k != -1:
                    payload = text[j + 1:k]
                    i = k + 1
                else:
                    payload = text[j + 1:]
                    i = len(text)
            else:
                payload = text[j]
                i = j + 1
            run = para.add_run(payload)
            run.font.name = LATIN_FONT
            run._element.rPr.rFonts.set(qn('w:eastAsia'), LATIN_FONT)
            run.font.size = Pt(font_size)
            run.bold = bold
            run.font.italic = True
            run.font.subscript = is_sub
            run.font.superscript = not is_sub
            if font_color:
                run.font.color.rgb = RGBColor(*font_color)
            continue
        run = para.add_run(ch)
        run.font.name = LATIN_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), LATIN_FONT)
        run.font.size = Pt(font_size)
        run.bold = bold
        run.font.italic = True
        if font_color:
            run.font.color.rgb = RGBColor(*font_color)
        i += 1


_MATH_WARNED = False
_MML2OMML_XSL = None


def _warn_math_fallback(reason):
    """公式转换不可用时只提示一次，并降级为普通文本。"""
    global _MATH_WARNED
    if not _MATH_WARNED:
        print(f"  [!] Word 原生公式转换不可用，公式将降级为普通文本：{reason}")
        _MATH_WARNED = True


def _split_math_markers(text):
    """拆分 {math:...}/{{math:...}} 公式标记和普通文本片段。

    公式内容本身经常含有 LaTeX 花括号，如 X_{L甲}、\frac{U}{R}。
    不能用简单的非贪婪正则匹配到第一个 } 就结束，否则会把公式截断。
    """
    source = text or ''
    parts = []
    pos = 0
    length = len(source)

    while pos < length:
        starts = []
        double_start = source.find('{{math:', pos)
        single_start = source.find('{math:', pos)
        if double_start != -1:
            starts.append((double_start, '{{math:', '}}'))
        if single_start != -1:
            starts.append((single_start, '{math:', '}'))
        if not starts:
            parts.append(('text', source[pos:]))
            break

        start, opener, closer = min(starts, key=lambda item: item[0])
        if start > pos:
            parts.append(('text', source[pos:start]))

        content_start = start + len(opener)
        i = content_start
        depth = 0
        end = -1
        if closer == '}}':
            # {{math:...}}：公式内部允许单个 { } 成对出现；只有 depth=0 的 }} 才结束。
            while i < length:
                if source.startswith('}}', i) and depth == 0:
                    end = i
                    break
                if source[i] == '{':
                    depth += 1
                elif source[i] == '}' and depth > 0:
                    depth -= 1
                i += 1
            marker_end = end + 2 if end != -1 else -1
        else:
            # {math:...}：公式内部允许 LaTeX 单花括号成对出现；depth=0 的 } 才结束。
            while i < length:
                ch = source[i]
                if ch == '{':
                    depth += 1
                elif ch == '}':
                    if depth == 0:
                        end = i
                        break
                    depth -= 1
                i += 1
            marker_end = end + 1 if end != -1 else -1

        if end == -1:
            # 标记不完整时保守当普通文本处理，避免误删内容。
            parts.append(('text', source[start:]))
            break

        parts.append(('math', source[content_start:end].strip()))
        pos = marker_end

    return [(kind, chunk) for kind, chunk in parts if chunk]


def _find_mml2omml_xsl():
    """查找 Microsoft Office 自带的 MathML→OMML 转换样式表。"""
    global _MML2OMML_XSL
    if _MML2OMML_XSL is not None:
        return _MML2OMML_XSL

    candidates = []
    program_roots = [
        os.environ.get('ProgramFiles'),
        os.environ.get('ProgramFiles(x86)'),
    ]
    for root in program_roots:
        if not root:
            continue
        root_path = Path(root) / 'Microsoft Office'
        candidates.extend(root_path.glob('**/MML2OMML.XSL'))

    for path in candidates:
        if path.exists():
            _MML2OMML_XSL = path
            return path

    _MML2OMML_XSL = False
    return None


def _normalize_latex_for_mathml(latex):
    """将少量中文命题常用写法规范为 latex2mathml 更易识别的形式。"""
    import html

    text = html.unescape(latex or '').strip()
    text = text.strip('$')
    text = re.sub(r'\\(?:displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b\s*', '', text)
    text = re.sub(r'\\(?:newline|linebreak)\b', r'\\', text)
    text = text.replace(r'\mathrm{（}', r'\left(')
    text = text.replace(r'\mathrm{）}', r'\right)')
    text = text.replace(r'\mathrm{(}', r'\left(')
    text = text.replace(r'\mathrm{)}', r'\right)')
    text = text.replace(r'\mathrm{\prime}', r'\prime')
    text = text.replace(r"\mathrm{'}", r'\prime')
    text = text.replace(r"\mathrm{’}", r'\prime')
    latex_command_replacements = {
        r'\Updelta': r'\Delta ',
        r'\updelta': r'\delta ',
        r'\Uppsi': r'\Psi ',
        r'\uppsi': r'\psi ',
    }
    for old, new in latex_command_replacements.items():
        text = text.replace(old, new)
    replacements = {
        '×': r'\times ',
        '·': r'\cdot ',
        '÷': r'\div ',
        '≈': r'\approx ',
        '≤': r'\le ',
        '≥': r'\ge ',
        '≠': r'\ne ',
        'ρ': r'\rho ',
        'Ω': r'\Omega ',
        'Φ': r'\Phi ',
        'φ': r'\phi ',
        'Δ': r'\Delta ',
        'μ': r'\mu ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def _clean_latex_residue_text(text):
    """Clean malformed LaTeX residues that are not wrapped as formula markers."""
    text = text or ''
    text = re.sub(r'\\\s*(?:Up)?[Oo]\s*m\s*e\s*g\s*a', 'Ω', text)
    text = re.sub(r'(?<![A-Za-z])(?:Up)?[Oo]mega\b', 'Ω', text)
    greek_command_replacements = {
        r'\Updelta': 'Δ',
        r'\updelta': 'δ',
        r'\Uppsi': 'Ψ',
        r'\uppsi': 'ψ',
    }
    for old, new in greek_command_replacements.items():
        text = text.replace(old, new)

    open_brace = re.escape('{')
    close_brace = re.escape('}')
    backslash = re.escape(chr(92))
    no_backslash_prefix = f'(^|[^{backslash}])'

    def _replace_bare_frac(match):
        prefix, numerator, denominator = match.groups()
        numerator = numerator.strip()
        denominator = denominator.strip()
        if re.search(r'\s|[+\-−×*/÷]', denominator) and not (denominator.startswith('(') and denominator.endswith(')')):
            denominator = f'({denominator})'
        return f'{prefix}{numerator}/{denominator}'

    text = re.sub(no_backslash_prefix + r'(?:mathrm|text|operatorname|mathbf|mathit|mathsf|boldsymbol|overline|underline)' + open_brace + r'([^{}]+)' + close_brace, lambda m: m.group(1) + m.group(2), text)
    text = re.sub(r'\\(?:displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b\s*', '', text)
    text = re.sub(r'\\(?:quad|qquad)\b\s*', ' ', text)
    text = re.sub(r'\\[,;:!]\s*', ' ', text)
    text = re.sub(r'\\(?:newline|linebreak)\b|\\\\', '\n', text)
    text = re.sub(r'\\(?:left|right)\s*', '', text)
    bare_frac_pattern = no_backslash_prefix + r'frac\s*' + open_brace + r'([^{}]+)' + close_brace + r'\s*' + open_brace + r'([^{}]+)' + close_brace
    text = re.sub(bare_frac_pattern, _replace_bare_frac, text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])_\{([^{}]+)\}', lambda m: m.group(1) + '_' + m.group(2), text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])\^\{([^{}]+)\}', lambda m: m.group(1) + '^' + m.group(2), text)
    text = re.sub(bare_frac_pattern, _replace_bare_frac, text)
    return text

def _latex_to_plain_text(latex):
    """公式转换失败时，把常见 LaTeX 降级成可读纯文本，避免 DOCX 出现反斜杠乱码。"""
    text = _clean_latex_residue_text((latex or '').strip().strip('$'))
    text = re.sub(r'\\dfrac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'\1/\2', text)
    text = re.sub(r'\\tfrac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'\1/\2', text)
    text = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'\1/\2', text)
    text = re.sub(r'\\sqrt\s*\{([^{}]+)\}', r'√\1', text)
    text = re.sub(r'\\sqrt\s*\[([^\]]+)\]\s*\{([^{}]+)\}', r'\2^(1/\1)', text)
    text = re.sub(r'\\(?:mathrm|text|operatorname|mathbf|mathit|mathsf|boldsymbol|overline|underline)\{([^{}]+)\}', r'\1', text)
    text = re.sub(r'\\(?:displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b\s*', '', text)
    replacements = {
        r'\newline': '\n',
        r'\linebreak': '\n',
        r'\times': '×',
        r'\cdot': '·',
        r'\div': '÷',
        r'\le': '≤',
        r'\ge': '≥',
        r'\neq': '≠',
        r'\ne': '≠',
        r'\approx': '≈',
        r'\pi': 'π',
        r'\rho': 'ρ',
        r'\Omega': 'Ω',
        r'\omega': 'Ω',
        r'\Phi': 'Φ',
        r'\phi': 'φ',
        r'\rightarrow': '→',
        r'\to': '→',
        r'\leftarrow': '←',
        r'\leftrightarrow': '↔',
        r'\infty': '∞',
        r'\pm': '±',
        r'\mp': '∓',
        r'\angle': '∠',
        r'\parallel': '∥',
        r'\perp': '⊥',
        r'\circ': '°',
        r'\quad': ' ',
        r'\qquad': ' ',
        r'\,': ' ',
        r'\;': ' ',
        r'\:': ' ',
        r'\!': '',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r'\\(?:left|right)\s*', '', text)
    text = re.sub(r'\\(?:begin|end)\s*\{(?:aligned|align|array|matrix|cases|gathered)\}', '\n', text)
    text = re.sub(r'&', '', text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])_\{([^{}]+)\}', r'\1_\2', text)
    text = re.sub(r'([A-Za-z0-9α-ωΑ-Ω])\^\{([^{}]+)\}', r'\1^\2', text)
    text = text.replace('{', '').replace('}', '')
    text = re.sub(r'\\[A-Za-z]+\b\s*', '', text)
    text = re.sub(r'\\+', '', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r' *\n *', '\n', text)
    return text.strip()


def _latex_to_omml(latex):
    """把 LaTeX/线性公式转换为 Word 原生 OMML XML 元素。失败返回 None。"""
    if etree is None:
        _warn_math_fallback('缺少 lxml，请先安装 pip install lxml latex2mathml')
        return None
    if latex2mathml is None:
        _warn_math_fallback('缺少 latex2mathml，请先安装 pip install latex2mathml')
        return None

    xsl_path = _find_mml2omml_xsl()
    if not xsl_path:
        _warn_math_fallback('未找到 Microsoft Office 的 MML2OMML.XSL')
        return None

    try:
        mathml = latex2mathml.converter.convert(_normalize_latex_for_mathml(latex))
        mathml_root = etree.fromstring(mathml.encode('utf-8'))
        transform = etree.XSLT(etree.parse(str(xsl_path)))
        omml_tree = transform(mathml_root)
        omml_root = omml_tree.getroot()
        # MML2OMML.XSL 常返回 m:oMathPara；段落内混排时只插入其中的 m:oMath。
        if omml_root.tag == qn('m:oMathPara'):
            inline_math = omml_root.find(qn('m:oMath'))
            if inline_math is not None:
                omml_root = inline_math
        return parse_xml(etree.tostring(omml_root, encoding='unicode'))
    except Exception as exc:
        _warn_math_fallback(f'公式“{latex}”转换失败：{exc}')
        return None


def _set_omml_color(omml_element, font_color):
    """给 OMML 内部 run 设置字体颜色。"""
    if not font_color:
        return
    color = ''.join(f'{value:02X}' for value in font_color)
    # python-docx 的 BaseOxmlElement.xpath() 不接收 namespaces 参数；直接按
    # Clark name 遍历，避免 lxml 抛出 Undefined namespace prefix。
    for run in omml_element.iter(qn('m:r')):
        rpr = run.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run.insert(0, rpr)
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = OxmlElement('w:color')
            rpr.append(color_el)
        color_el.set(qn('w:val'), color)


def _set_omml_italic(omml_element):
    """将 OMML 公式内部文本标记为斜体。"""
    for run in omml_element.iter(qn('m:r')):
        rpr = run.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run.insert(0, rpr)
        italic_el = rpr.find(qn('w:i'))
        if italic_el is None:
            italic_el = OxmlElement('w:i')
            rpr.append(italic_el)
        italic_el.set(qn('w:val'), '1')


def _append_omml_to_paragraph(para, omml_element, font_color=None):
    """向段落追加 Word 原生公式节点。"""
    _set_omml_color(omml_element, font_color)
    _set_omml_italic(omml_element)
    para._p.append(omml_element)


def _add_math_marker_run(para, formula_text, cjk_font, font_size, bold, font_color):
    """添加 {{math:...}} 公式；不可转换时降级为清洗后的普通文本。"""
    formula_parts = re.split(r'\\(?:newline|linebreak)\b', formula_text or '')
    for idx, formula_part in enumerate(formula_parts):
        if idx > 0:
            para.add_run().add_break()
        formula_part = formula_part.strip()
        if not formula_part:
            continue
        omml = _latex_to_omml(formula_part)
        if omml is not None:
            _append_omml_to_paragraph(para, omml, font_color)
        else:
            _add_formula_run(para, _latex_to_plain_text(formula_part), font_size, bold, font_color)


def _add_plain_text_with_scripts(para, text, cjk_font, font_size, bold, font_color, latin_font=LATIN_FONT):
    """Render plain-text residues like U_5/R_ab/P^2 as formula script runs."""
    pattern = re.compile(r'([A-Za-z0-9α-ωΑ-Ω])([_^])([A-Za-z0-9α-ωΑ-Ω一-鿿]+)')
    pos = 0
    for match in pattern.finditer(text or ''):
        if match.start() > pos:
            _add_mixed_run(para, text[pos:match.start()], cjk_font, font_size, bold, font_color, latin_font)
        base, marker, payload = match.groups()
        _add_formula_run(para, f"{base}{marker}{{{payload}}}", font_size, bold, font_color)
        pos = match.end()
    if pos < len(text or ''):
        _add_mixed_run(para, text[pos:], cjk_font, font_size, bold, font_color, latin_font)


def _add_rich_text_run(para, text, cjk_font, font_size, bold, font_color, latin_font=LATIN_FONT):
    """添加文本；{{math:...}}、$...$、\\(...\\) 片段优先转为 Word 原生公式，失败时降级为清洗后的普通文本。"""
    text = _clean_latex_residue_text(text or '')
    text = re.sub(r'([A-Za-z0-9\)）])\s*乘以?\s*([A-Za-z0-9\(（])', r'\1×\2', text)
    for kind, chunk in _split_math_markers(text):
        if not chunk:
            continue
        if kind == 'math':
            _add_math_marker_run(para, chunk, cjk_font, font_size, bold, font_color)
            continue
        parts = re.split(r'(\\\(.+?\\\)|\$[^$]+\$)', chunk)
        for part in parts:
            if not part:
                continue
            if len(part) >= 2 and part.startswith('$') and part.endswith('$'):
                _add_math_marker_run(para, part[1:-1], cjk_font, font_size, bold, font_color)
            elif len(part) >= 4 and part.startswith('\\(') and part.endswith('\\)'):
                _add_math_marker_run(para, part[2:-2], cjk_font, font_size, bold, font_color)
            else:
                _add_plain_text_with_scripts(para, part, cjk_font, font_size, bold, font_color, latin_font)


def _add_mixed_run(para, text, cjk_font, font_size, bold, font_color, latin_font=LATIN_FONT):
    """添加段落文本，自动分中英文设置字体"""
    lines = str(text or '').split('\n')
    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            para.add_run().add_break()
        tokens = _split_cjk_latin(line)
        for seg_text, seg_type in tokens:
            run = para.add_run(seg_text)
            font = latin_font if seg_type == 'latin' else cjk_font
            run.font.name = font
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
            run.font.size = Pt(font_size)
            run.bold = bold
            if font_color:
                run.font.color.rgb = RGBColor(*font_color)


def copy_template(template_path: str, output_path: str) -> Document:
    """复制 Word 模板到输出路径，返回 Document 对象（保留页眉页脚）"""
    doc = Document(template_path)
    doc.save(output_path)
    return Document(output_path)


def set_margins(doc: Document, top: float = 2.54, bottom: float = 2.54,
                left: float = 3.18, right: float = 3.18):
    """设置页边距（单位：厘米）"""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_heading_with_style(doc: Document, text: str, level: int = 1,
                           font_name: str = "黑体", font_size: int = 16):
    """添加带样式的标题（兼容中文Word模板缺少Heading样式的场景）"""
    try:
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)
        return heading
    except KeyError:
        # 中文Word模板没有"Heading N"样式，退化为加粗段落
        return add_paragraph_with_style(
            doc, text, font_name=font_name, font_size=font_size,
            bold=True, space_after=12
        )


def add_paragraph_with_style(doc: Document, text: str, font_name: str = "宋体",
                             font_size: float = 12, bold: bool = False,
                             font_color: tuple = None,
                             alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after: float = 6):
    """添加带样式的段落，中英文自动分字体（中文=font_name，数字字母=Times New Roman）

    Args:
        font_color: RGB 颜色元组，如 (255,0,0) 为红色，None 为默认黑色
    """
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.5
    _add_rich_text_run(para, text, font_name, font_size, bold, font_color)
    return para


def add_table_with_style(doc: Document, headers: list, rows: list,
                         style: str = "Table Grid"):
    """添加带样式的表格，headers 为表头列表，rows 为数据行列表（每行一个 list）"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn('w:eastAsia'), "黑体")
                run.font.size = Pt(10)
                run.bold = True

    # 数据行
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
                    run.font.size = Pt(10)

    return table


def append_docx_after_template(template_path: str, content_path: str) -> str:
    """将内容文档追加到模板后面，返回合并后的文档路径"""
    template_doc = Document(template_path)
    content_doc = Document(content_path)

    for element in content_doc.element.body:
        template_doc.element.body.append(copy.deepcopy(element))

    merged_path = content_path.replace(".docx", "_merged.docx")
    template_doc.save(merged_path)
    return merged_path


def save_docx(doc: Document, output_path: str) -> str:
    """保存 Document 到文件"""
    doc.save(output_path)
    return output_path


def add_labeled_text(doc: Document, label: str, content: str,
                     label_font: str = "黑体", label_size: float = 10.5,
                     content_font: str = "宋体", content_size: float = 10.5,
                     color: tuple = None, bold_label: bool = True,
                     bold_content: bool = False):
    """添加标签+正文混合样式段落（标签和正文可分别设字体/颜色）

    用于【答案】xxx 和【解析】xxx 等场景

    Args:
        label: 标签文字（如"【答案】"）
        content: 正文内容
        label_font: 标签字体
        label_size: 标签字号（pt）
        content_font: 正文字体
        content_size: 正文字号（pt）
        color: RGB 颜色元组，如 (255,0,0)，None 为黑色
        bold_label: 标签是否加粗
        bold_content: 正文是否加粗
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.5

    # 标签 run
    label_run = para.add_run(label)
    label_run.font.name = label_font
    label_run._element.rPr.rFonts.set(qn('w:eastAsia'), label_font)
    label_run.font.size = Pt(label_size)
    label_run.bold = bold_label
    if color:
        label_run.font.color.rgb = RGBColor(*color)

    # 正文 run —— 自动分中英文设置字体；$...$、\\(...\\) 片段优先转为 Word 原生公式，失败时清洗为普通文本
    _add_rich_text_run(para, content, content_font, content_size, bold_content, color)

    return para


def add_editorial_note(doc: Document, textbook_name: str, edition: str,
                       chapter_seq: int, knowledge_scope: str,
                       province: str = "重庆市",
                       publisher: str = "高教版",
                       separator_image: str = None,
                       exam_type: str = "高职分类考试",
                       series_name: str = "考纲百套卷",
                       volume_label: str = None,
                       paper_type: str = "",
                       paper_name: str = "",
                       course_name: str = "",
                       category: str = "",
                       all_courses: str = "",
                       course_count: int = 0,
                       exam_table_title: str = "",
                       specific_content: str = ""):
    """添加编写说明表格 — 蓝色虚线边框表格，单倍行距，段前段后4pt，首行缩进2字符

    Args:
        province: 省份名，如"重庆市"；自治区应使用全称，如"内蒙古自治区"、"新疆维吾尔自治区"、"西藏自治区"
        publisher: 出版社简称，如"高教版"、"机工版"，默认"高教版"
        separator_image: 可选的分隔图片路径，插入在表格与标题之间
        exam_type: 考试类型，如"高职分类考试"、"对口招生"
        series_name: 系列名称，默认使用“考纲百套卷”
    """
    if series_name == "考纲百套卷":
        course_display = course_name or textbook_name
        volume_display = volume_label or f"第{chapter_seq}卷"
        subject = f'{province}《{category}考纲百套卷》' if category else f'{province}《考纲百套卷》'
        exam_table_display = exam_table_title or f'{province}2025年中等职业学校毕业生进入普通高校学习专业基础课和专业课考试科目表'
        courses_display = all_courses or course_display
        count_display = course_count or len([item for item in courses_display.replace('，', '、').split('、') if item.strip()])
        content_display = specific_content or knowledge_scope
        prefix = (
            f'{subject}，依据《{exam_table_display}》编写。'
            f'本专辑涵盖{courses_display}共{count_display}个课程，'
            f'且每个课程均采用三阶递进式训练体系：基础层（具象化支架）拆解考点为微目标，'
            f'紧扣考纲“掌握”“理解”要求编写考点训练卷；巩固层（关联性支架）强化知识交叉与场景关联，'
            f'按考纲专题编写专题训练卷；应用层（引导性支架）聚焦综合提升，'
            f'结合知识模块与教材编写课程综合卷。'
        )
        info = (
            f'本试卷是{volume_display}{paper_type or "训练卷"}，按《{course_display}》中的{paper_name or knowledge_scope}范围和要求编写。'
            f'具体内容为：{content_display}。'
        )
    else:
        prefix = (
            f'编写说明：依据{province}（{exam_type}）考试要求，'
            f'围绕《{textbook_name}》（{publisher}·{edition}）相关课程内容编制{series_name}。'
            f'专辑里的每一份试卷，都与考纲知识点紧密相关，题目围绕应掌握的知识和能力呈现。'
            f'目的在于帮助学生扎实掌握课程的基本概念与基本方法，'
            f'为后续复习提升奠定坚实基础。'
        )
        info = (
            f'本试卷是{province}（{exam_type}）《{textbook_name}》（{publisher}·{edition}）{series_name}的第{chapter_seq}卷，'
            f'内容涵盖{knowledge_scope}。'
        )

    # 创建 1行1列表格
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True

    # 去除模板自带的顶部空行（表格前第一个空白段落）
    body = doc.element.body
    first_p = body.find(qn('w:p'))
    if first_p is not None:
        first_t = first_p.find(qn('w:r'))
        if first_t is None or not (first_t.text or '').strip():
            body.remove(first_p)

    # 设置表格蓝色虚线边框 1.5pt
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'dashed')
        border.set(qn('w:sz'), '12')  # 1.5pt = 12/8 pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        borders.append(border)
    tblPr.append(borders)

    # 设置表格宽度为页面宽度
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    cell = table.rows[0].cells[0]

    # 第一段
    p1 = cell.paragraphs[0]
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    p1.paragraph_format.first_line_indent = Cm(0.74)
    _add_mixed_run(p1, prefix, "楷体", 10.5, True, (0x44, 0x72, 0xC4))

    # 第二段
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.first_line_indent = Cm(0.74)
    _add_mixed_run(p2, info, "楷体", 10.5, True, (0x44, 0x72, 0xC4))

    # -- 可选分隔图片（表格与标题之间，左对齐，紧贴标题）--
    if separator_image and os.path.exists(separator_image):
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        img_para.paragraph_format.space_before = Pt(0)
        img_para.paragraph_format.space_after = Pt(0)
        run = img_para.add_run()
        run.add_picture(separator_image, width=Cm(4.60))
    else:
        doc.add_paragraph()  # 无图片时表格与标题之间空行


# -- 选项布局自动适配阈值 --
_OPTION_SINGLE_COL_THRESHOLD = 12    # 任一选项超此字数 → 单列
_OPTION_ROW_BALANCE_RATIO = 1.6       # 同排两选项长度比超此值 → 单列（避免短选项下方大量空白）

# -- 图片选项排布参数（移植自 kaogangbaitao-v2 的带图选择题规则）--
LANDSCAPE_RATIO = 1.2                 # 宽高比 >= 1.2 → 横图
PORTRAIT_RATIO = 0.8                  # 宽高比 <= 0.8 → 竖图
_OPTION_IMAGE_CONSTRAINTS = {
    'landscape': {'max_w_cm': 5.0, 'max_h_cm': 4.0, 'layout': '2x2'},
    'portrait': {'max_w_cm': 3.5, 'max_h_cm': 4.5, 'layout': '1x4'},
    'square': {'max_w_cm': 5.0, 'max_h_cm': 5.0, 'layout': '2x2'},
    'mixed': {'max_w_cm': 3.5, 'max_h_cm': 4.5, 'layout': '2x2'},
}
_STEM_WITH_OPTION_IMAGE_MAX_CM = 5.0
_PX_PER_INCH = 96
_CM_PER_INCH = 2.54
_MIN_READABLE_IMAGE_CM = 1.2


def _apply_table_full_width_no_border(table):
    """将表格设为 100% 宽、无边框，复用选择题选项表格样式。"""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    tblPr.append(borders)


def _resolve_image_path(image_item):
    """从图片条目中取本地路径；URL 下载由数据层负责。"""
    if isinstance(image_item, (str, Path)):
        text = str(image_item).strip()
    elif isinstance(image_item, dict):
        text = str(
            image_item.get('local_path')
            or image_item.get('path')
            or image_item.get('file')
            or image_item.get('filename')
            or ''
        ).strip()
    else:
        return None
    if not text or re.match(r'^https?://', text, flags=re.IGNORECASE):
        return None
    return Path(text)


def _numeric_dim(value):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return 0
    return int(number) if number > 0 else 0


def _image_size_px(image_item):
    """读取图片真实像素尺寸；HTML 显示尺寸另行处理，避免误当原图尺寸。"""
    if isinstance(image_item, dict):
        width = _numeric_dim(image_item.get('pixel_width') or image_item.get('natural_width'))
        height = _numeric_dim(image_item.get('pixel_height') or image_item.get('natural_height'))
        if width and height:
            return width, height

    path = _resolve_image_path(image_item)
    if not path or not path.exists():
        return 0, 0
    try:
        from PIL import Image
        with Image.open(path) as img:
            return img.size
    except Exception:
        return 0, 0


def _display_size_px(image_item):
    if not isinstance(image_item, dict):
        return 0, 0
    width = _numeric_dim(image_item.get('display_w') or image_item.get('width'))
    height = _numeric_dim(image_item.get('display_h') or image_item.get('height'))
    return width, height


def _classify_image_orientation(width_px, height_px):
    if width_px <= 0 or height_px <= 0:
        return 'square'
    ratio = width_px / height_px
    if ratio >= LANDSCAPE_RATIO:
        return 'landscape'
    if ratio <= PORTRAIT_RATIO:
        return 'portrait'
    return 'square'


def _classify_option_images(options):
    """按所有选项图方向决定整体布局：全竖 1×4，否则 2×2。"""
    orientations = []
    for opt in options or []:
        if not isinstance(opt, dict):
            continue
        for image_item in opt.get('images') or []:
            w_px, h_px = _image_size_px(image_item)
            orientations.append(_classify_image_orientation(w_px, h_px))

    if not orientations:
        overall = 'square'
    else:
        unique = set(orientations)
        overall = unique.pop() if len(unique) == 1 else 'mixed'
    params = _OPTION_IMAGE_CONSTRAINTS[overall]
    return {
        'orientations': orientations,
        'overall': overall,
        'layout': params['layout'],
        'max_w_cm': params['max_w_cm'],
        'max_h_cm': params['max_h_cm'],
    }


def _fit_image_size_cm(image_item, max_w_cm, max_h_cm):
    """按网页显示尺寸/真实像素自适应，等比缩小且避免无依据放大。"""
    display_w, display_h = _display_size_px(image_item)
    width_px, height_px = (display_w, display_h) if display_w and display_h else _image_size_px(image_item)
    if width_px <= 0 or height_px <= 0:
        return max_w_cm, max_h_cm

    width_cm = width_px / _PX_PER_INCH * _CM_PER_INCH
    height_cm = height_px / _PX_PER_INCH * _CM_PER_INCH
    scale = min(1.0, max_w_cm / width_cm if max_w_cm and width_cm > max_w_cm else 1.0)
    if max_h_cm and height_cm * scale > max_h_cm:
        scale = min(scale, max_h_cm / height_cm)
    width_cm *= scale
    height_cm *= scale

    # 只对过小图做最低可读放大，且仍受当前场景的最大框约束。
    if max(width_cm, height_cm) < _MIN_READABLE_IMAGE_CM:
        readable_scale = _MIN_READABLE_IMAGE_CM / max(width_cm, height_cm, 0.01)
        readable_scale = min(
            readable_scale,
            max_w_cm / width_cm if max_w_cm else readable_scale,
            max_h_cm / height_cm if max_h_cm else readable_scale,
        )
        width_cm *= readable_scale
        height_cm *= readable_scale
    return width_cm, height_cm


def _add_image_run(paragraph, image_item, max_w_cm, max_h_cm, warn_label='图片'):
    """在段落中插入本地图片，成功返回 True；失败时不中断整卷生成。"""
    path = _resolve_image_path(image_item)
    if not path or not path.exists():
        print(f"  → 警告：{warn_label}缺少本地图片路径，已跳过。")
        return False
    width_cm, height_cm = _fit_image_size_cm(image_item, max_w_cm, max_h_cm)
    try:
        run = paragraph.add_run()
        run.add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
        return True
    except Exception as exc:
        print(f"  → 警告：插入{warn_label}失败：{exc}")
        return False


def _option_label(option, index):
    if isinstance(option, dict):
        label = str(option.get('label') or '').strip().rstrip('.．')
        if label:
            return label
    return chr(ord('A') + index)


def _option_text(option):
    if isinstance(option, dict):
        return str(option.get('text') or option.get('content') or '').strip()
    return str(option or '').strip()


def add_structured_choice_question(doc: Document, stem_text: str, options,
                                   stem_images=None, font_name: str = "宋体",
                                   font_size: float = 10.5):
    """输出含图片的结构化选择题。

    排布规则：题干文字 → 题干图 → 选项；选项图全竖向时 1×4，否则 2×2。
    """
    options = [opt for opt in (options or []) if isinstance(opt, dict)]
    stem_images = stem_images or []
    has_option_images = any(opt.get('images') for opt in options)
    if not stem_text and not stem_images and not has_option_images:
        return False

    if stem_text:
        add_paragraph_with_style(doc, stem_text, font_name, font_size, space_after=0)

    if stem_images:
        stem_max_cm = _STEM_WITH_OPTION_IMAGE_MAX_CM if has_option_images else 12.0
        for idx, image_item in enumerate(stem_images, 1):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            _add_image_run(para, image_item, stem_max_cm, stem_max_cm, f'题干图{idx}')

    if not options:
        return bool(stem_text or stem_images)

    if not has_option_images:
        option_lines = []
        for idx, opt in enumerate(options):
            label = _option_label(opt, idx)
            text = _option_text(opt)
            option_lines.append(f"{label}. {text}".strip())
        if option_lines:
            rows, cols = (len(option_lines), 1) if _should_use_single_col(option_lines) else ((len(option_lines) + 1) // 2, 2)
            table = doc.add_table(rows=rows, cols=cols)
            _apply_table_full_width_no_border(table)
            for i, opt_text in enumerate(option_lines):
                row_idx, col_idx = (i, 0) if cols == 1 else (i // cols, i % cols)
                cell = table.cell(row_idx, col_idx)
                for p in list(cell.paragraphs):
                    p._element.getparent().remove(p._element)
                p = cell.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                _add_rich_text_run(p, opt_text, font_name, font_size, False, None)
        return True

    layout_info = _classify_option_images(options)
    layout = layout_info['layout']
    if layout == '1x4':
        rows, cols = 1, min(max(len(options), 1), 4)
    else:
        cols = 2
        rows = (len(options) + 1) // 2
    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_full_width_no_border(table)

    for i, opt in enumerate(options):
        row_idx = 0 if layout == '1x4' else i // cols
        col_idx = i if layout == '1x4' else i % cols
        if row_idx >= rows or col_idx >= cols:
            continue
        cell = table.cell(row_idx, col_idx)
        # 彻底删除模板残留段落后新建，避免 cell 间 XML 引用串扰（如 C/D 文字互相污染）
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        _add_rich_text_run(p, f"{_option_label(opt, i)}. ", font_name, font_size, False, None)
        for img_idx, image_item in enumerate(opt.get('images') or [], 1):
            _add_image_run(
                p,
                image_item,
                layout_info['max_w_cm'],
                layout_info['max_h_cm'],
                f"选项{_option_label(opt, i)}图片{img_idx}",
            )
        text = _option_text(opt)
        if text:
            _add_rich_text_run(p, text, font_name, font_size, False, None)
    return True


def _should_use_single_col(options):
    """判断是否应该使用单列布局（非双列）。

    触发条件（任一满足即单列）：
      1. 任一选项长度 > _OPTION_SINGLE_COL_THRESHOLD
      2. 同排两个选项长度悬殊（比例 > _OPTION_ROW_BALANCE_RATIO），
         双列会导致短选项下方留大片空白
    """
    n = len(options)
    if n <= 1:
        return True

    # 条件1：长选项
    if max(len(o) for o in options) > _OPTION_SINGLE_COL_THRESHOLD:
        return True

    # 条件2：同排选项长度不平衡
    # 对于A/B/C/D四选项，配对为 (A,B) 和 (C,D)
    for i in range(0, n - 1, 2):
        a_len = len(options[i])
        b_len = len(options[i + 1])
        shorter = max(min(a_len, b_len), 1)
        if max(a_len, b_len) / shorter > _OPTION_ROW_BALANCE_RATIO:
            return True

    return False


def add_question_options_table(doc: Document, q_text: str,
                               font_name: str = "宋体", font_size: float = 10.5):
    """解析题目文本，用无边框表格输出选项，确保列对齐。

    输入格式示例：
      '1. 题干文本\\nA. 选项A\\t\\tB. 选项B\\nC. 选项C\\t\\tD. 选项D'

    自动适配布局：
      - 短且均衡 → 双列表格（紧凑省空间）
      - 长选项 或 同排长度悬殊 → 单列表格（避免挤窄/空白）
      - 非选项文本 → 普通段落
    """
    # -- 不包含制表符分隔的选项 → 普通段落 --
    if '\t' not in q_text:
        add_paragraph_with_style(doc, q_text, font_name, font_size, space_after=0)
        return

    # -- 按行分割 --
    lines = q_text.split('\n')
    stem = lines[0].strip()

    # -- 从剩余行提取选项（按制表符分割） --
    options = []
    for line in lines[1:]:
        parts = re.split(r'\t+', line)
        for p in parts:
            p = p.strip()
            if p:
                options.append(p)

    if not options:
        add_paragraph_with_style(doc, q_text, font_name, font_size, space_after=0)
        return

    # -- 题干段落 --
    add_paragraph_with_style(doc, stem, font_name, font_size, space_after=0)

    # -- 自动选择布局 --
    n = len(options)
    if _should_use_single_col(options):
        # 单列布局：每行一个选项
        rows, cols = n, 1
    elif n <= 2:
        rows, cols = 1, n
    else:
        # 双列布局：每行两个选项
        rows = (n + 1) // 2
        cols = 2

    table = doc.add_table(rows=rows, cols=cols)
    _apply_table_full_width_no_border(table)

    # -- 填充选项 --
    for i, opt in enumerate(options):
        if cols == 1:
            row_idx, col_idx = i, 0
        else:
            row_idx = i // cols
            col_idx = i % cols
        cell = table.cell(row_idx, col_idx)
        # 清空默认段落
        for p in cell.paragraphs:
            p.clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        _add_rich_text_run(p, opt, font_name, font_size, False, None)

    return table
