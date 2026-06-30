"""真题风格库生成工具 — 将少量真题样本蒸馏为可复用的命题风格说明

使用方法：
  # 生成可人工填写的风格库模板（不调用 API）
  python 01_工具脚本/真题风格/extract_exam_style.py --no-api --province "重庆市" --exam-category "机械加工类"

  # 从人工摘录/校对后的样题文本中蒸馏风格库（默认输出到 03_项目数据/参考资料/真题风格/省份/考类/风格总则.txt）
  python 01_工具脚本/真题风格/extract_exam_style.py --sample "真题样本.txt" --province "重庆市" --exam-category "机械加工类"

说明：
  本工具的目标不是全量 OCR 图片型 PDF，而是把少量代表题整理成“风格库”。
  后续 create.py 会读取 03_项目数据/参考资料/真题风格/ 下的 txt 文件，用于模仿真题口吻，
  但知识准确性仍以考纲、教材和当前主题为准。
"""

import argparse
import json
import os
import re
import shutil
import sys
import time
from pathlib import Path

from openai import OpenAI

BASE_DIR = Path(__file__).resolve().parents[2]
CONFIG_PATH = BASE_DIR / "02_配置资源" / "config.json"
DEFAULT_STYLE_DIR = BASE_DIR / "03_项目数据" / "参考资料" / "真题风格"
DEFAULT_TRUTH_TEXT_DIR = BASE_DIR / "03_项目数据" / "参考资料" / "真题"
STYLE_FILES = [
    ("风格总则.txt", "总体风格"),
    ("单选题风格.txt", "单项选择题风格"),
    ("多选题风格.txt", "多项选择题风格"),
    ("判断题风格.txt", "判断题风格"),
    ("填空题风格.txt", "填空题风格"),
    ("简答题风格.txt", "简答题风格"),
    ("计算题风格.txt", "计算题风格"),
    ("综合题风格.txt", "综合题风格"),
    ("代表样题.txt", "代表样题摘要"),
]
MAX_SAMPLE_CHARS = 60000  # API 蒸馏时保留的最大样本文本字符数，避免超长文档导致请求失败
OCR_MIN_TEXT_CHARS = 200
OCR_GARBAGE_TOKEN_RE = re.compile(r"\{#\{[^\n]{12,}\}#\}|\}#\}[^\n]{12,}\{#\{")


def find_tesseract():
    """自动查找 Tesseract 可执行文件路径。"""
    candidates = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
    ]
    for path in candidates:
        if os.path.isfile(path):
            return path
    return shutil.which("tesseract")


def check_tessdata(tessdata_dir=None, lang="chi_sim"):
    """检查 Tesseract 语言包是否可用。"""
    if tessdata_dir:
        return os.path.isfile(os.path.join(tessdata_dir, f"{lang}.traineddata"))

    default_dirs = [
        r"C:\Program Files\Tesseract-OCR\tessdata",
        r"C:\Program Files (x86)\Tesseract-OCR\tessdata",
        "/usr/share/tesseract-ocr/4.00/tessdata",
        "/usr/share/tesseract-ocr/5/tessdata",
        "/usr/share/tessdata",
    ]
    return any(os.path.isfile(os.path.join(d, f"{lang}.traineddata")) for d in default_dirs)


def _text_quality_score(text):
    """粗略评估 PDF 文本层是否可用，分数越高越像正常中文试题文本。"""
    if not text:
        return 0
    cjk = len(re.findall(r"[一-龥]", text))
    exam_terms = len(re.findall(r"选择题|判断题|填空题|简答题|计算题|试卷|答案|下列|正确|错误|电路|机械|机电", text))
    garbage = len(OCR_GARBAGE_TOKEN_RE.findall(text))
    return cjk + exam_terms * 20 - garbage * 120


def text_needs_ocr(text):
    """判断普通 PDF 文本提取结果是否过短或明显乱码，需要 OCR 兜底。"""
    stripped = (text or "").strip()
    if len(stripped) < OCR_MIN_TEXT_CHARS:
        return True
    garbage = len(OCR_GARBAGE_TOKEN_RE.findall(stripped))
    cjk = len(re.findall(r"[一-龥]", stripped))
    if garbage >= 2 and cjk < 300:
        return True
    if garbage >= 4 and garbage * 80 > max(cjk, 1):
        return True
    return False


if sys.stdout and hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")


def load_config():
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def call_api(client, model, system_prompt, user_prompt, max_tokens=6000, temperature=0.2):
    """调用 OpenAI 兼容 API，返回文本。"""
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=max_tokens,
                temperature=temperature,
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"[!] API 调用失败 (第{attempt + 1}次): {e}")
            if attempt < 2:
                wait = (attempt + 1) * 8
                print(f"    等待 {wait}s 后重试...")
                time.sleep(wait)
            else:
                raise


def _safe_path_part(text):
    """清理目录/文件名中的非法字符。"""
    text = (text or "").strip()
    return re.sub(r'[\\/:*?"<>|\s]+', "_", text).strip("_") or "未命名"


def _split_province_category(category_text):
    """兼容旧 --category 写法，如“重庆市机械加工类”。"""
    text = (category_text or "").strip()
    m = re.match(r"^([一-龥]+(?:省|市))(.+类)$", text)
    if m:
        return m.group(1), m.group(2)
    return "", text


def _resolve_output_path(args):
    """根据 --output 或 省份/考类 参数确定输出路径。"""
    if args.output:
        return Path(args.output)

    province = args.province
    category = args.exam_category
    if not province or not category:
        parsed_province, parsed_category = _split_province_category(args.category)
        province = province or parsed_province
        category = category or parsed_category

    if province and category:
        return DEFAULT_STYLE_DIR / _safe_path_part(province) / _safe_path_part(category) / "风格总则.txt"

    safe_name = _safe_path_part(args.category or category or "真题风格")
    return DEFAULT_STYLE_DIR / f"{safe_name}_风格总结.txt"


def _display_category(args):
    """用于提示词中的类别显示。"""
    province = args.province
    category = args.exam_category
    if not province or not category:
        parsed_province, parsed_category = _split_province_category(args.category)
        province = province or parsed_province
        category = category or parsed_category
    return f"{province}{category}" if province and category else (args.category or category or "未命名类别")


def parse_pages_arg(pages_str):
    """解析页码范围，如 '1,3,5-8' → [1,3,5,6,7,8]。"""
    if not pages_str:
        return []
    pages = []
    for part in pages_str.split(","):
        part = part.strip()
        if not part:
            continue
        m = re.match(r"^(\d+)\s*[-~]\s*(\d+)$", part)
        if m:
            start, end = int(m.group(1)), int(m.group(2))
            pages.extend(range(min(start, end), max(start, end) + 1))
        else:
            try:
                pages.append(int(part))
            except ValueError:
                pass
    return sorted(set(p for p in pages if p > 0))


def extract_text_from_docx(docx_path):
    """从 DOCX 提取段落和表格文本。"""
    try:
        from docx import Document
    except ImportError:
        print("错误：读取 docx 需要安装 python-docx：pip install python-docx")
        sys.exit(1)

    doc = Document(str(docx_path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def extract_text_from_pdf(pdf_path):
    """从文字型 PDF 提取文本；图片型 PDF 可能返回空或乱码。"""
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(pdf_path))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except ImportError:
            print("警告：缺少 pdfplumber/PyPDF2，跳过 PDF 文本提取。可安装：pip install pdfplumber")
    return text.strip()


def extract_text_from_pdf_tesseract(pdf_path, pages=None, tessdata_dir=None, lang="chi_sim", dpi=2.5, preprocess=False):
    """将 PDF 页面渲染为图片，并用 Tesseract OCR 提取文本。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("    OCR 跳过：缺少 PyMuPDF，可安装：pip install pymupdf")
        return ""
    try:
        import pytesseract
        from PIL import Image, ImageEnhance
    except ImportError:
        print("    OCR 跳过：缺少 pytesseract/Pillow，可安装：pip install pytesseract Pillow")
        return ""

    tesseract_cmd = find_tesseract()
    if not tesseract_cmd:
        print("    OCR 跳过：未找到 Tesseract，可安装：winget install --id UB-Mannheim.TesseractOCR")
        return ""
    if not check_tessdata(tessdata_dir, lang):
        print(f"    OCR 跳过：未找到 {lang}.traineddata，可通过 --tessdata 指定语言包目录。")
        return ""

    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    pdf_path = Path(pdf_path)
    doc = fitz.open(str(pdf_path))
    total_pages = len(doc)
    if pages:
        page_nums = [p for p in pages if 1 <= p <= total_pages]
    else:
        page_nums = list(range(1, total_pages + 1))

    env_backup = os.environ.get("TESSDATA_PREFIX")
    if tessdata_dir:
        os.environ["TESSDATA_PREFIX"] = tessdata_dir

    parts = []
    try:
        matrix = fitz.Matrix(dpi, dpi)
        config = "--oem 3 --psm 6"
        for page_num in page_nums:
            try:
                page = doc[page_num - 1]
                pix = page.get_pixmap(matrix=matrix, alpha=False)
                mode = "RGB" if pix.n < 4 else "RGBA"
                img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
                if preprocess:
                    img = img.convert("L")
                    img = ImageEnhance.Contrast(img).enhance(1.6)
                    img = img.point(lambda p: 255 if p > 180 else 0)
                text = pytesseract.image_to_string(img, lang=lang, config=config).strip()
                if text:
                    parts.append(f"【OCR第{page_num}页】\n{text}")
                print(f"    OCR 第 {page_num}/{total_pages} 页：{len(text)} 字")
            except Exception as e:
                print(f"    OCR 第 {page_num}/{total_pages} 页失败，已跳过：{e}")
    finally:
        doc.close()
        if tessdata_dir:
            if env_backup is None:
                os.environ.pop("TESSDATA_PREFIX", None)
            else:
                os.environ["TESSDATA_PREFIX"] = env_backup

    return "\n\n".join(parts).strip()


def load_text_from_file(path, args=None):
    """按文件类型读取真题文本，必要时对 PDF 启用 Tesseract OCR。"""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".pdf":
        text = extract_text_from_pdf(path)
        use_ocr = bool(args and getattr(args, "ocr_pdf", False))
        if use_ocr and text_needs_ocr(text):
            print("    文本层为空或疑似乱码，启用 Tesseract OCR...")
            ocr_text = extract_text_from_pdf_tesseract(
                path,
                pages=parse_pages_arg(getattr(args, "ocr_pages", "") or ""),
                tessdata_dir=getattr(args, "tessdata", None),
                lang=getattr(args, "ocr_lang", "chi_sim"),
                dpi=getattr(args, "ocr_dpi", 2.5),
                preprocess=getattr(args, "ocr_preprocess", False),
            )
            if ocr_text.strip() and _text_quality_score(ocr_text) > _text_quality_score(text):
                print(f"    已采用 OCR 文本：{len(ocr_text)} 字")
                return ocr_text
            if ocr_text.strip():
                print("    OCR 文本质量未明显优于文本层，保留原文本层。")
            else:
                print("    OCR 未得到可用文本，保留原文本层。")
        return text
    return ""


def collect_source_texts(source_dir, args=None):
    """读取某个考类目录下的 txt/docx/pdf 真题文本。"""
    source_dir = Path(source_dir)
    chunks = []
    skipped_pdfs = []
    garbage_pdfs = []
    for path in sorted(source_dir.rglob("*")):
        if not path.is_file() or path.name.startswith("~"):
            continue
        if path.suffix.lower() not in (".txt", ".docx", ".pdf"):
            continue
        print(f"  读取: {path.name}")
        text = load_text_from_file(path, args=args)
        if text.strip():
            chunks.append(f"\n\n【来源文件：{path.name}】\n{text.strip()}")
            if path.suffix.lower() == ".pdf" and text_needs_ocr(text):
                garbage_pdfs.append(path.name)
        elif path.suffix.lower() == ".pdf":
            skipped_pdfs.append(path.name)

    if skipped_pdfs:
        print("  警告：以下 PDF 未提取到文字，可能是图片型 PDF，需启用 --ocr-pdf 或先用 OCR/ocr_pdf.py 生成文本后再加入样本：")
        for name in skipped_pdfs:
            print(f"    - {name}")
    if garbage_pdfs and not (args and getattr(args, "ocr_pdf", False)):
        print("  警告：以下 PDF 文本层疑似乱码；可加 --ocr-pdf 使用 Tesseract 重新识别：")
        for name in garbage_pdfs:
            print(f"    - {name}")
    return "\n".join(chunks).strip()


def export_pdf_pages(pdf_path, pages, output_dir):
    """把 PDF 指定页导出为 PNG，供人工挑选/摘录。"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("错误：导出 PDF 页面图片需要安装 PyMuPDF")
        print("  pip install pymupdf")
        sys.exit(1)

    pdf_path = Path(pdf_path)
    output_dir = Path(output_dir)
    os.makedirs(output_dir, exist_ok=True)

    doc = fitz.open(str(pdf_path))
    total = len(doc)
    if not pages:
        pages = list(range(1, total + 1))

    exported = []
    for page_num in pages:
        if page_num < 1 or page_num > total:
            print(f"跳过无效页码: {page_num}")
            continue
        page = doc[page_num - 1]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        out_path = output_dir / f"{pdf_path.stem}_p{page_num:03d}.png"
        pix.save(str(out_path))
        exported.append(out_path)
        print(f"  导出: {out_path}")

    doc.close()
    print(f"共导出 {len(exported)} 页图片。请人工查看后摘录代表题到 txt，再用 --sample 蒸馏风格库。")


def build_template(category):
    category = category or "未命名类别"
    return f"""【资料定位】
本文件用于模仿{category}真题命题风格，不作为知识依据；知识准确性以教材、考纲和当前主题为准。

【总体风格】
- （填写该类别真题的整体口吻：简洁/情境化/重基础/重应用等）
- （填写常见考查角度：概念识记、故障判断、操作流程、计算应用等）

【单项选择题风格】
- 常见设问方式：
  - “下列关于____的说法正确的是（ ）”
  - “____的主要作用是（ ）”
  - “出现____现象时，可能的原因是（ ）”
- 选项风格：四个选项应为同类短语，长度接近，避免明显废选项。
- 干扰项来源：来自同系统、相邻概念、常见错误操作或相近部件。

【判断题风格】
- 题干应为完整判断句，避免过长。
- 判断点应明确，不使用模棱两可的限定词。

【填空题风格】
- 空格处应考查核心术语、参数或步骤名称。
- 不使用括号表示空格，统一用______。

【综合题风格】
- 多结合实际场景、故障现象、操作步骤或简单计算。
- 问题应分层清晰，答案简洁明确。

【代表样题摘要】
1. （填写代表样题的题干摘要与风格备注）
2. （填写代表样题的题干摘要与风格备注）
3. （填写代表样题的题干摘要与风格备注）

【禁止事项】
- 不得照搬样题题干、选项、情境或数值。
- 不得只替换年份、数字、部件名称来伪装新题。
- 真题风格仅用于口吻、设问方式、选项结构和解析简洁程度；知识点必须服从当前考纲和教材。
"""


def build_style_prompt(category, sample_text, max_examples, focus="完整风格库"):
    sample_text = sample_text[:MAX_SAMPLE_CHARS]
    system_prompt = """你是一位高考/高职分类考试真题风格分析专家。你的任务不是照抄真题，而是从少量样题中提炼可迁移的命题风格规则。

请严格区分：
- 真题风格：设问口吻、题干长度、选项结构、干扰项来源、解析表达。
- 知识依据：不得从样题外推新知识，后续出题仍以教材、考纲和当前主题为准。

输出应简洁、可复用，适合直接注入出题 prompt。"""

    if focus == "完整风格库":
        required = """风格库应包含以下栏目：
   - 【资料定位】
   - 【总体风格】
   - 【单项选择题风格】
   - 【判断题风格】
   - 【填空题风格】
   - 【综合题风格】
   - 【代表样题摘要】
   - 【禁止事项】"""
    elif focus == "代表样题摘要":
        required = f"""请只输出【代表样题】文件内容：
- 从样本中选取最多 {max_examples} 道最能代表该考类风格的题。
- 可保留题干、选项、答案和简短风格备注。
- 不要输出过多原题全文；代表样题用于学习口吻，不用于照搬。"""
    else:
        required = f"""请只输出【{focus}】文件内容：
- 聚焦该题型/维度的设问口吻、题干长度、选项或答案结构、干扰项来源、解析表达。
- 如果样本中该题型不足，请给出可后续补充的保守风格模板，并明确“当前样本不足”。
- 不要输出其他栏目。"""

    user_prompt = f"""请根据以下{category}真题样本，生成“{focus}”。

【输出要求】
1. 使用中文，直接输出正文。
2. 不要逐字 OCR 校对，不要追求还原所有题目；重点提炼风格。
3. 明确遵守：不得照搬样题题干、选项、情境或数值。
4. {required}

【真题样本】
{sample_text}
"""
    return system_prompt, user_prompt


def _style_output_dir(args):
    """返回结构化风格库输出目录。"""
    output_path = _resolve_output_path(args)
    if output_path.suffix.lower() == ".txt":
        return output_path.parent
    return output_path


def write_split_templates(output_dir, category):
    """生成多文件模板。"""
    os.makedirs(output_dir, exist_ok=True)
    templates = {
        "风格总则.txt": build_template(category),
        "单选题风格.txt": "【单项选择题风格】\n\n- （填写常见设问方式）\n- （填写选项长度、干扰项来源、题干场景特点）\n",
        "多选题风格.txt": "【多项选择题风格】\n\n- （填写多选题常见设问方式、正确项数量特点、选项层级和干扰项来源）\n",
        "判断题风格.txt": "【判断题风格】\n\n- （填写判断题常见考查点、表述方式、正误设置特点）\n",
        "填空题风格.txt": "【填空题风格】\n\n- （填写填空题常考术语、答案长度、题干语境特点）\n",
        "简答题风格.txt": "【简答题风格】\n\n- （填写简答题常见设问、答案分条方式、解析表达特点）\n",
        "计算题风格.txt": "【计算题风格】\n\n- （填写计算题常见公式类型、已知条件给法、单位和解析步骤要求）\n",
        "综合题风格.txt": "【综合题风格】\n\n- （填写综合题情境、设问层次、答案组织和解析特点）\n",
        "代表样题.txt": "【代表样题】\n\n请粘贴少量代表样题，并为每题补充【风格备注】。\n",
    }
    for filename, content in templates.items():
        path = Path(output_dir) / filename
        path.write_text(content.strip() + "\n", encoding="utf-8")
        print(f"  已生成: {path}")


def generate_split_files(client, model, category, sample_text, output_dir, max_examples, max_tokens):
    """调用 API 生成多文件风格库。"""
    os.makedirs(output_dir, exist_ok=True)
    for filename, focus in STYLE_FILES:
        print(f"  正在生成 {filename}...")
        sys_prompt, user_prompt = build_style_prompt(category, sample_text, max_examples, focus=focus)
        text = call_api(client, model, sys_prompt, user_prompt, max_tokens=max_tokens, temperature=0.2)
        path = Path(output_dir) / filename
        path.write_text(text.strip() + "\n", encoding="utf-8")
        print(f"    → {path} ({len(text)} 字)")


def _infer_category_from_dir(source_dir, province):
    """从题库目录推断考类。"""
    source_dir = Path(source_dir)
    if source_dir.name.endswith("类"):
        return source_dir.name
    m = re.search(r"([一-龥]+类)", source_dir.name)
    if m:
        return m.group(1)
    return source_dir.name.replace(province, "") or source_dir.name


def sync_truth_reference_text(province, exam_category, text, source_dir=None):
    """将真题汇总文本同步沉淀到 03_项目数据/参考资料/真题，作为可复用资料层。"""
    output_dir = DEFAULT_TRUTH_TEXT_DIR / _safe_path_part(province) / _safe_path_part(exam_category)
    os.makedirs(output_dir, exist_ok=True)
    output_path = output_dir / f"{_safe_path_part(province)}{_safe_path_part(exam_category)}_真题OCR汇总.txt"
    header = (
        f"{province}{exam_category}_真题OCR汇总\n"
        f"来源：{source_dir or '真题题库目录'}\n"
        "说明：本文件由真题文本提取/OCR 结果同步沉淀到 03_项目数据/参考资料/真题，"
        "作为可复用真题文本资料；OCR 文本需人工抽查，不作为知识准确性唯一依据。\n\n"
        "============================================================\n\n"
    )
    output_path.write_text(header + text.strip() + "\n", encoding="utf-8")
    return output_path


def process_source_dir(args, client=None, config=None):
    """处理单个考类题库目录，输出对应风格库。"""
    source_dir = Path(args.source_dir)
    if not source_dir.exists():
        print(f"错误：题库目录不存在 {source_dir}")
        return False
    province = args.province or source_dir.parent.name
    exam_category = args.exam_category or _infer_category_from_dir(source_dir, province)
    display = f"{province}{exam_category}"
    output_dir = DEFAULT_STYLE_DIR / _safe_path_part(province) / _safe_path_part(exam_category)

    print(f"\n处理题库目录：{source_dir}")
    print(f"  省份/考类：{display}")
    print(f"  输出目录：{output_dir}")
    sample_text = collect_source_texts(source_dir, args=args)
    if not sample_text:
        print("  未提取到可用文本，跳过。")
        return False
    original_len = len(sample_text)
    reference_text_path = sync_truth_reference_text(province, exam_category, sample_text, source_dir=source_dir)
    print(f"  已同步真题文本资料：{reference_text_path}")
    clipped_text = sample_text[:MAX_SAMPLE_CHARS]
    sample_path = output_dir / "_自动汇总样本.txt"
    os.makedirs(output_dir, exist_ok=True)
    sample_path.write_text(clipped_text + "\n", encoding="utf-8")
    if original_len > MAX_SAMPLE_CHARS:
        print(f"  已保存汇总样本：{sample_path}（原始 {original_len} 字，已截断为 {MAX_SAMPLE_CHARS} 字）")
    else:
        print(f"  已保存汇总样本：{sample_path}（{original_len} 字）")

    if args.sample_only:
        return True

    if args.no_api:
        write_split_templates(output_dir, display)
    else:
        if client is None:
            config = load_config()
            client = OpenAI(api_key=config["api_key"], base_url=config["base_url"])
        generate_split_files(
            client, config["model"], display, sample_text, output_dir,
            args.max_examples, config.get("max_tokens", 8000)
        )
    return True


def process_source_root(args):
    """批量处理省份题库根目录下各考类目录。"""
    root = Path(args.source_root)
    if not root.exists():
        print(f"错误：题库根目录不存在 {root}")
        sys.exit(1)
    province = args.province or root.name
    category_dirs = [p for p in sorted(root.iterdir()) if p.is_dir()]
    if not category_dirs:
        print(f"错误：题库根目录下没有考类子目录：{root}")
        sys.exit(1)

    config = None
    client = None
    if not args.no_api:
        config = load_config()
        client = OpenAI(api_key=config["api_key"], base_url=config["base_url"])

    ok = 0
    for category_dir in category_dirs:
        args.source_dir = str(category_dir)
        args.province = province
        args.exam_category = category_dir.name
        if process_source_dir(args, client=client, config=config):
            ok += 1
    print(f"\n批量处理完成：成功 {ok}/{len(category_dirs)} 个考类")


def main():
    parser = argparse.ArgumentParser(description="生成真题风格库 txt")
    parser.add_argument("--sample", help="少量真题样本 txt（人工摘录或 OCR 后校对）")
    parser.add_argument("--source-dir", help="单个考类题库目录，支持读取其中的 txt/docx/文字型pdf")
    parser.add_argument("--source-root", help="省份题库根目录，自动批量处理其下各考类子目录")
    parser.add_argument("--split-files", action="store_true", help="生成多文件风格库：风格总则/各题型风格/代表样题")
    parser.add_argument("--category", default="", help="兼容旧写法：省份+考类，如 重庆市机械加工类")
    parser.add_argument("--province", default="", help="省份，如 重庆市")
    parser.add_argument("--exam-category", default="", help="考类，如 机械加工类")
    parser.add_argument("--output", "-o", help="输出风格库 txt 路径")
    parser.add_argument("--max-examples", type=int, default=8, help="代表样题摘要最多保留数量")
    parser.add_argument("--no-api", action="store_true", help="不调用 API，只生成可人工填写的风格库模板")
    parser.add_argument("--sample-only", action="store_true", help="处理 --source-dir/--source-root 时只生成 _自动汇总样本.txt，不生成风格模板且不调用 API")
    parser.add_argument("--pdf", help="图片型 PDF 路径：仅导出指定页图片，供人工摘录")
    parser.add_argument("--pages", help="配合 --pdf 使用，页码如 1,3,5-8")
    parser.add_argument("--image-dir", help="PDF 页面图片输出目录")
    parser.add_argument("--ocr-pdf", action="store_true", help="处理 --source-dir/--source-root 时，对空文本层或疑似乱码 PDF 启用 Tesseract OCR")
    parser.add_argument("--ocr-pages", help="仅 OCR 指定页码，如 1,3,5-8；默认 OCR 全部页")
    parser.add_argument("--ocr-dpi", type=float, default=2.5, help="Tesseract OCR 前的 PDF 渲染倍率，默认 2.5；识别差可试 3.0")
    parser.add_argument("--ocr-lang", default="chi_sim", help="Tesseract OCR 语言代码，默认 chi_sim")
    parser.add_argument("--tessdata", help="Tesseract tessdata 目录（含 chi_sim.traineddata）")
    parser.add_argument("--ocr-preprocess", action="store_true", help="OCR 前做灰度/对比度/二值化增强")
    args = parser.parse_args()

    if args.source_root:
        process_source_root(args)
        return

    if args.source_dir:
        process_source_dir(args)
        return

    category = _display_category(args)
    output_path = _resolve_output_path(args)

    if args.pdf:
        image_dir = Path(args.image_dir) if args.image_dir else BASE_DIR / "_真题页面图片" / Path(args.pdf).stem
        export_pdf_pages(args.pdf, parse_pages_arg(args.pages), image_dir)
        if not args.sample and not args.no_api:
            print("\n提示：页面图片已导出。请人工摘录代表题为 txt 后，再运行：")
            print(f"  python 01_工具脚本/真题风格/extract_exam_style.py --sample 样题.txt --category \"{category}\" --output \"{output_path}\"")
            return

    if args.no_api:
        if args.split_files:
            write_split_templates(_style_output_dir(args), category)
            return
        text = build_template(category)
    else:
        if not args.sample:
            print("错误：调用 API 蒸馏风格库需要提供 --sample；如只生成模板，请加 --no-api")
            sys.exit(1)
        sample_path = Path(args.sample)
        if not sample_path.exists():
            print(f"错误：样本文件不存在 {sample_path}")
            sys.exit(1)
        sample_text = sample_path.read_text(encoding="utf-8")
        if not sample_text.strip():
            print("错误：样本文件为空")
            sys.exit(1)

        config = load_config()
        client = OpenAI(api_key=config["api_key"], base_url=config["base_url"])
        if args.split_files:
            generate_split_files(
                client, config["model"], category, sample_text,
                _style_output_dir(args), args.max_examples,
                config.get("max_tokens", 8000)
            )
            return
        sys_prompt, user_prompt = build_style_prompt(category, sample_text, args.max_examples)
        print(f"正在调用 API 蒸馏真题风格库：{category}")
        text = call_api(
            client, config["model"], sys_prompt, user_prompt,
            max_tokens=config.get("max_tokens", 8000),
            temperature=0.2,
        )

    os.makedirs(output_path.parent, exist_ok=True)
    output_path.write_text(text.strip() + "\n", encoding="utf-8")
    print(f"已保存: {output_path} ({len(text)} 字)")


if __name__ == "__main__":
    main()
