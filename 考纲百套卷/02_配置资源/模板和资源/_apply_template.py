# -*- coding: utf-8 -*-
"""
文档模板套用工具 v4.0 — 考纲百套卷模板套用。

用法: python _apply_template.py

说明：
- 当前模板文件为同目录下的 template.docx。
- 新版考纲百套卷模板不需要标识/分隔图片；separator_image 可以留空。
- 处理目标文件名格式：
  第x卷 试卷名称 卷型《课程名称》地名（考试类型）考类 考纲百套卷（解析版）.docx
"""

import copy, re, os, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ═════════ XML 常量 ═════════
W_R = qn('w:r'); W_T = qn('w:t'); W_PPR = qn('w:pPr'); W_RPR = qn('w:rPr')

# ═════════ 正则 ═════════
PAPER_TYPES = ('考点训练卷', '专题训练卷', '课程综合卷')
FILENAME_RE = re.compile(
    r'^(?:[（(]待人工审核[）)])?第(\d+)卷(?:-\d+)?\s+'
    r'(.+?)\s+'
    r'(考点训练卷|专题训练卷|课程综合卷)'
    r'《(.+?)》'
    r'(.+?)[（(](.+?)[）)]'
    r'(.*?)\s*考纲百套卷[（(](解析版|原卷版)[）)]\.docx$'
)
CN_NUM = {'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'十':10}
SECTION_RE = re.compile(r'^[（(]?\s*([一二三四五六七八九十]+)\s*[）)、．.]?\s*(.+)$')


# ═════════ 通用文本处理 ═════════
def parse_section(text):
    text = str(text or '').strip()
    m = SECTION_RE.match(text)
    if m: return CN_NUM.get(m.group(1), 0), m.group(2).strip()
    return None, text if text else ''


def normalize_topic(s):
    s = str(s or '')
    s = re.sub(r'[（(][一二三四五六七八九十]+[)）]', '', s)
    s = re.sub(r'[、/·\s]', '', s).strip()
    s = s.replace('等速万向节', '等角速万向节')
    return s


def _fuzzy_key(s):
    """超模糊查找键：去所有连接词和符号。"""
    return re.sub(r'[与和及、/·\s]', '', str(s or ''))


def _clean_course_name(text):
    text = str(text or '').strip()
    return re.sub(r'^课程\s*[:：]\s*', '', text).strip()


def _safe_int(value, default=1):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


# ═════════ 规划表 ═════════
def load_planning_table(xlsx_path):
    """加载生产规划表，返回卷号兜底映射、主题映射和全部行。"""
    from openpyxl import load_workbook
    wb = load_workbook(xlsx_path)
    ws = wb[wb.sheetnames[0]]
    rows, cur_course, cur_sec_num, cur_sec_name = [], '', None, ''
    for row in ws.iter_rows(min_row=6, max_row=ws.max_row, values_only=True):
        vals = [str(v).strip() if v is not None else '' for v in row[:8]]
        v0 = vals[0]
        if v0.startswith('课程') and '：' in v0:
            cur_course = _clean_course_name(v0)
            continue
        try:
            int(v0)
        except ValueError:
            if v0 and v0 != '序号' and '考纲' not in v0:
                cur_sec_num, cur_sec_name = parse_section(v0)
            continue
        rows.append({
            '序号': int(v0),
            '考纲知识点': vals[1],
            '试卷主题': vals[2],
            '级别': vals[3],
            '题型': vals[4],
            '难度': vals[5],
            '套数': _safe_int(vals[6], 1),
            '考纲标号': vals[7],
            '节号': cur_sec_num,
            '节主题': cur_sec_name,
            '课程': cur_course,
        })

    volume_map, volume_no = {}, 1
    for r in rows:
        for _ in range(r['套数']):
            volume_map[volume_no] = r
            volume_no += 1

    topic_map = {}
    for r in rows:
        keys = [r.get('试卷主题'), r.get('考纲知识点')]
        if r.get('课程') and r.get('试卷主题'):
            keys.append(f"{r['课程']}::{r['试卷主题']}")
        for key in keys:
            k = normalize_topic(key)
            if k and k not in topic_map:
                topic_map[k] = r
    return volume_map, topic_map, rows


# ═════════ 查表 ═════════
def lookup_plan(seq, paper_name, paper_type, course_name, volume_map, topic_map, allow_seq_fallback=True):
    """查规划表：试卷名称/课程组合匹配优先，可选卷号兜底。"""
    candidates = [paper_name, f'{course_name}::{paper_name}']
    if paper_type == '考点训练卷':
        candidates.append(paper_name.replace('基础', '').strip())

    row = None
    for candidate in candidates:
        k = normalize_topic(candidate)
        if not k:
            continue
        row = topic_map.get(k)
        if row is not None:
            break
        fk = _fuzzy_key(k)
        for tk, tv in topic_map.items():
            if k in tk or tk in k:
                row = tv; break
        if row is not None:
            break
        for tk, tv in topic_map.items():
            nk = _fuzzy_key(k); ntk = _fuzzy_key(tk)
            if nk and ntk and (nk in ntk or ntk in nk):
                row = tv; break
        if row is not None:
            break
        for tk, tv in topic_map.items():
            tfk = _fuzzy_key(tk)
            if fk and tfk and (fk in tfk or tfk in fk):
                row = tv; break
        if row is not None:
            break

    if row is None and allow_seq_fallback:
        row = volume_map.get(seq)

    if row:
        sn, nm = row['节号'], row['节主题']
        section_title = f'第{sn}节 {nm}' if sn is not None and nm else (nm or '')
        return row, section_title, row.get('考纲知识点') or paper_name
    return None, '', paper_name or '相关知识点'


# ═════════ 文件名解析 ═════════
def parse_filename(fp):
    fn = os.path.basename(fp)
    m = FILENAME_RE.match(fn)
    if not m:
        raise ValueError(f'无法解析: {fn}')
    return {
        'seq': int(m.group(1)),
        'paper_name': m.group(2).strip(),
        'paper_type': m.group(3).strip(),
        'course_name': m.group(4).strip(),
        'province': m.group(5).strip(),
        'exam_type': m.group(6).strip(),
        'category': m.group(7).strip(),
        'variant': m.group(8).strip(),
    }


# ═════════ XML工具 ═════════
def _set_run_text(re, txt):
    ts = re.findall(W_T)
    if ts:
        ts[0].text = txt; ts[0].set(qn('xml:space'),'preserve')
        for e in ts[1:]: re.remove(e)
    else:
        t = OxmlElement('w:t'); t.text = txt; t.set(qn('xml:space'),'preserve'); re.append(t)


def _clear_runs(p):
    rs = p.findall(W_R)
    if not rs:
        r = OxmlElement('w:r'); p.append(r); return r
    for e in rs[1:]: p.remove(e)
    return rs[0]


def _replace_para_text(p, text):
    r = _clear_runs(p)
    _set_run_text(r, text)
    return r


def _insert_blank_paragraph_after(p):
    new_p = OxmlElement('w:p')
    p.addnext(new_p)
    return new_p


# ═════════ 编写说明文本 ═════════
def _ed_p1(province, exam_type, category, course_name, exam_table_title):
    subject = f'{province}《{category}考纲百套卷》' if category else f'{province}《考纲百套卷》'
    table_title = exam_table_title or f'{province}中等职业学校毕业生进入普通高校学习专业基础课和专业课考试科目表'
    return (
        f'{subject}，依据《{table_title}》编写。'
        f'本专辑围绕{province}（{exam_type}）{category}考试要求，按课程和考纲内容组织训练；'
        f'每个课程采用三阶递进式训练体系：基础层拆解考点为微目标，'
        f'紧扣考纲“掌握”“理解”要求编写考点训练卷；巩固层强化知识交叉与场景关联，'
        f'按考纲专题编写专题训练卷；应用层聚焦综合提升，结合知识模块与教材编写课程综合卷。'
    )


def _ed_p2(seq, paper_name, paper_type, course_name, exam_points):
    return (
        f'本试卷是第{seq}卷{paper_type}，按《{course_name}》中的{paper_name}范围和要求编写。'
        f'具体内容为：{exam_points}。'
    )


# ═════════ 模板更新 ═════════
def _update_table(doc, info, exam_points, exam_table_title=''):
    if not doc.tables:
        return
    cell = doc.tables[0].rows[0].cells[0]
    paras = cell.paragraphs
    p1 = _ed_p1(info['province'], info['exam_type'], info['category'], info['course_name'], exam_table_title)
    if paras:
        pe = paras[0]._element; _clear_runs(pe); _set_run_text(pe.find(W_R), p1)
    p2 = _ed_p2(info['seq'], info['paper_name'], info['paper_type'], info['course_name'], exam_points)
    if len(paras) >= 2:
        pe = paras[1]._element; _clear_runs(pe); _set_run_text(pe.find(W_R), p2)
    elif paras:
        p1e = paras[0]._element; rpr = p1e.find('.//'+W_RPR); ppr = p1e.find(W_PPR)
        p2e = OxmlElement('w:p')
        if ppr is not None: p2e.append(copy.deepcopy(ppr))
        nr = OxmlElement('w:r')
        if rpr is not None: nr.append(copy.deepcopy(rpr))
        _set_run_text(nr, p2); p2e.append(nr); cell._element.append(p2e)


def _find_title_anchor(body):
    ch = list(body)
    for i, c in enumerate(ch):
        if etree.QName(c.tag).localname != 'p':
            continue
        text = ''.join(list(c.itertext()))
        if '考纲百套卷' in text or re.search(r'第\d+卷', text):
            return i
    for i, c in enumerate(ch):
        if etree.QName(c.tag).localname == 'p':
            return i
    return None


def _update_titles(body, info):
    ch = list(body)
    p1i = _find_title_anchor(body)
    if p1i is None:
        return
    title1 = f"{info['province']}（{info['exam_type']}）《{info['category']}考纲百套卷》第{info['seq']}卷" if info['category'] else f"{info['province']}（{info['exam_type']}）《考纲百套卷》第{info['seq']}卷"
    title2 = f"《{info['course_name']}》"
    title3 = f"{info['paper_name']} {info['paper_type']}"

    p = ch[p1i]
    _replace_para_text(p, title1)
    while p1i + 1 >= len(ch) or etree.QName(ch[p1i + 1].tag).localname != 'p':
        _insert_blank_paragraph_after(p)
        ch = list(body)
    _replace_para_text(ch[p1i + 1], title2)
    while p1i + 2 >= len(ch) or etree.QName(ch[p1i + 2].tag).localname != 'p':
        _insert_blank_paragraph_after(ch[p1i + 1])
        ch = list(body)
    _replace_para_text(ch[p1i + 2], title3)


def _insert_separator(doc, img_path):
    if not os.path.exists(img_path): return
    body = doc.element.body; ch = list(body)
    ti = next((i for i,c in enumerate(ch) if etree.QName(c.tag).localname == 'tbl'), None)
    if ti is None: return
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ip.paragraph_format.space_before = Pt(0); ip.paragraph_format.space_after = Pt(0)
    ip.add_run().add_picture(img_path, width=Cm(4.60))
    ie = ip._element; body.remove(ie); list(body)[ti].addnext(ie)
    ch = list(body)
    for i,c in enumerate(ch):
        if etree.QName(c.tag).localname == 'p':
            if list(c.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')):
                while i+1 < len(body):
                    nx = list(body)[i+1]
                    if etree.QName(nx.tag).localname == 'p' and not ''.join(list(nx.itertext())).strip():
                        body.remove(nx)
                    else: break
                break


# ═════════ 内容追加 ═════════
HEADER_RE = re.compile(r'^[一二三四五六七八九十]+\s*[、.．]')
QNUM_RE = re.compile(r'^(\d+)\s*[.、．)）]')


def _classify_paragraphs(doc):
    ps = doc.paragraphs; first = None
    for i,p in enumerate(ps):
        if HEADER_RE.match(p.text.strip()): first = i; break
    if first is None:
        for i,p in enumerate(ps):
            if QNUM_RE.match(p.text.strip()): first = i; break
    if first is None: first = 0
    keep = list(range(first, len(ps)))
    while keep and not ps[keep[-1]].text.strip(): keep.pop()
    return keep


def _remap_images(elem, sd, dd, rc):
    from docx.opc.part import Part; from docx.opc.packuri import PackURI
    AB = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    RE = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
    sp,dp = sd.part, dd.part; dkg = dp.package
    for blip,attr,rid in [(b,RE,b.get(RE)) for b in elem.findall('.//'+AB) if b.get(RE)]:
        if rid in rc: blip.set(attr, rc[rid]); continue
        if rid not in sp.rels: continue
        r = sp.rels[rid]; st = r.target_part
        nm = str(st.partname); ex = {str(p.partname) for p in dkg.iter_parts()}
        c,cd = 1,nm
        while cd in ex: c+=1; cd = f'{os.path.splitext(nm)[0]}_c{c}{os.path.splitext(nm)[1]}'
        np = Part(PackURI(cd), st.content_type, st.blob, dkg)
        nrid = dp.relate_to(np, r.reltype); rc[rid] = nrid; blip.set(attr, nrid)


def _append_content(td, sd):
    ki = _classify_paragraphs(sd)
    if not ki: return 0
    body = td.element.body; sp = body.find(qn('w:sectPr')); rc = {}; k = 0
    for i in ki:
        np = copy.deepcopy(sd.paragraphs[i]._element)
        _remap_images(np, sd, td, rc)
        (sp.addprevious(np) if sp is not None else body.append(np)); k += 1
    return k


# ═════════ 单文件 ═════════
def process_one(inp, out, tpl, sep_img, volume_map, topic_map, allow_seq_fallback=True, exam_table_title=''):
    try:
        info = parse_filename(inp)
    except ValueError as e:
        print(f'  跳过: {e}'); return False
    row, section_title, exam_points = lookup_plan(
        info['seq'], info['paper_name'], info['paper_type'], info['course_name'],
        volume_map, topic_map, allow_seq_fallback
    )
    print(
        f"  第{info['seq']}卷 | {info['province']} | {info['exam_type']} | "
        f"{info['category']} | 《{info['course_name']}》 | {info['paper_name']} | {info['paper_type']}"
    )
    if row:
        print(f'    → 规划: {row["级别"]} | 节:{section_title} | {exam_points[:60]}...')
    else:
        print('    → 规划: 未匹配，使用文件名信息')

    td = Document(tpl); sd = Document(inp)
    _update_table(td, info, exam_points, exam_table_title)
    _update_titles(td.element.body, info)
    if sep_img:
        _insert_separator(td, sep_img)
    n = _append_content(td, sd)
    print(f'  追加{n}段 → {os.path.basename(out)}')
    td.save(out); return True


# ═════════ 主程序 ═════════
def run(cfg):
    """cfg = {template, separator_image, planning_table, input_dir, output_dir, exam_table_title}"""
    tpl = cfg['template']; sep = cfg.get('separator_image', '')
    base = cfg['input_dir']; out_dir = cfg.get('output_dir', base)

    if not os.path.exists(tpl): print(f'错误: 模板 {tpl}'); return
    if not os.path.exists(base): print(f'错误: 目录 {base}'); return

    pt = cfg.get('planning_table', '')
    if pt and os.path.exists(pt):
        volume_map, topic_map, _ = load_planning_table(pt)
    else:
        if pt: print(f'警告: 规划表 {pt} 不存在')
        volume_map, topic_map = {}, {}

    files = []
    for root, dirs, fns in os.walk(base):
        for fn in fns:
            if fn.endswith('.docx') and fn.startswith(('第', '（待人工审核）第')) and '原卷版' not in fn:
                files.append(os.path.join(root, fn))

    if not files:
        print('未找到待处理文件'); return

    # 有子目录时不同课程可能各自编号，默认不用卷号兜底，优先按试卷名称匹配规划表。
    has_subdirs = any(os.path.isdir(os.path.join(base, d)) for d in os.listdir(base) if not d.startswith('.'))
    allow_seq_fallback = not has_subdirs

    print(f'模板: {os.path.basename(tpl)}')
    print(f'目录: {base}')
    print(f'共 {len(files)} 个文件  (卷号兜底: {"是" if allow_seq_fallback else "否"})\n')

    ok = err = skip = 0
    for fp in sorted(files):
        rel = os.path.relpath(fp, base)
        outp = os.path.join(out_dir, rel) if out_dir != base else fp
        if out_dir != base: os.makedirs(os.path.dirname(outp), exist_ok=True)
        print(f'处理: {rel[:80]}')
        try:
            if process_one(fp, outp, tpl, sep, volume_map, topic_map, allow_seq_fallback, cfg.get('exam_table_title', '')): ok += 1
            else: skip += 1
        except Exception as e:
            print(f'  失败: {e}'); import traceback; traceback.print_exc(); err += 1
    print(f'\n完成: {ok}成功 {skip}跳过 {err}失败')


if __name__ == '__main__':
    # 示例：按当前 wyy 项目结构套用模板。
    # 如需处理其他省份/考类，请改 planning_table、input_dir、output_dir。
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    run({
        'template':         os.path.join(project_root, '02_配置资源', '模板和资源', 'template.docx'),
        'separator_image':  '',
        'planning_table':   os.path.join(project_root, '04_生成输出', '生产规划', '重庆市 电子信息类', '重庆市_电子信息类_考点规划总表.xlsx'),
        'input_dir':        os.path.join(project_root, '04_生成输出', '生成结果', '解析版'),
        'output_dir':       os.path.join(project_root, '04_生成输出', '生成结果', '解析版'),
    })
