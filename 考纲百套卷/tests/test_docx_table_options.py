from pathlib import Path
import sys

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "01_工具脚本"))

from 生成器.paper_loader import load_manual_paper  # noqa: E402
from 生成器.paper_assembler import assemble_analysis_paper_text  # noqa: E402


def test_docx_table_options_stay_with_question_and_analysis_label_is_clean(tmp_path):
    source = tmp_path / "第1卷 表格选项.docx"
    doc = Document()
    doc.add_paragraph("一、单项选择题")
    doc.add_paragraph("1．下列说法正确的是（ ）")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A. 选项一"
    table.cell(0, 1).text = "B. 选项二"
    table.cell(1, 0).text = "C. 选项三"
    table.cell(1, 1).text = "D. 选项四"
    doc.add_paragraph("【答案】A")
    doc.add_paragraph("【详解】【详解】因为选项一正确。")
    doc.save(source)

    loaded = load_manual_paper(source)

    assert len(loaded.questions) == 1
    question = loaded.questions[0]
    assert question["options"] == ["A. 选项一", "B. 选项二", "C. 选项三", "D. 选项四"]
    assert question["option_table_indices"] == [0]
    assert question["analysis"] == "因为选项一正确。"

    text = assemble_analysis_paper_text(loaded.questions)
    assert "【详解】" not in text
    assert "【解析】因为选项一正确。" in text
