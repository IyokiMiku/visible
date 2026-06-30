"""DOCX 生成。"""
import copy
import os
import re
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from .paths import BASE_DIR, TEMPLATE_PATH
from .text_processing import _clean_paper_text

COMMON_SCRIPT_NAMES = ("class.py", "answer2none.py", "zip.py", "fix.py")
RESIDUAL_FORMULA_RE = re.compile(r"\{\{?math:|\\\(|\\\)|\$[^$]+\$|\\(?:frac|times|Omega|rho|Phi|Delta|mu|sqrt)\b")
A_BLIP = qn("a:blip")
R_EMBED = qn("r:embed")
R_ID = qn("r:id")
W_SECTPR = qn("w:sectPr")
NS_VML = "urn:schemas-microsoft-com:vml"
NS_OFFICE = "urn:schemas-microsoft-com:office:office"


def _iter_table_paragraphs(tables):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                yield from _iter_table_paragraphs(cell.tables)


def _find_residual_formula_text(doc):
    """扫描 DOCX 可见文本中的公式标记/LaTeX 残留，防止乱码进入成品。"""
    hits = []
    paragraphs = list(doc.paragraphs) + list(_iter_table_paragraphs(doc.tables))
    for paragraph in paragraphs:
        text = paragraph.text or ""
        if RESIDUAL_FORMULA_RE.search(text):
            hits.append(text.strip()[:120])
    return hits


def _make_unique_partname(dst_package, original_partname):
    name = str(original_partname)
    base, ext = os.path.splitext(name)
    existing = {str(part.partname) for part in dst_package.iter_parts()}
    if name not in existing:
        return PackURI(name)
    counter = 1
    while True:
        candidate = f"{base}_c{counter}{ext}"
        if candidate not in existing:
            return PackURI(candidate)
        counter += 1


def _remap_embedded_relationships(element, src_doc, dst_doc, rid_cache):
    src_part = src_doc.part
    dst_part = dst_doc.part
    dst_package = dst_part.package
    refs = []
    for blip in element.findall(".//" + A_BLIP):
        old_rid = blip.get(R_EMBED)
        if old_rid:
            refs.append((blip, R_EMBED, old_rid))
    for imgdata in element.findall(".//{" + NS_VML + "}imagedata"):
        old_rid = imgdata.get(R_ID)
        if old_rid:
            refs.append((imgdata, R_ID, old_rid))
    for ole in element.findall(".//{" + NS_OFFICE + "}OLEObject"):
        old_rid = ole.get(R_ID)
        if old_rid:
            refs.append((ole, R_ID, old_rid))

    for elem, attr, old_rid in refs:
        if old_rid in rid_cache:
            elem.set(attr, rid_cache[old_rid])
            continue
        if old_rid not in src_part.rels:
            continue
        rel = src_part.rels[old_rid]
        src_target = rel.target_part
        unique_name = _make_unique_partname(dst_package, src_target.partname)
        new_part = Part(unique_name, src_target.content_type, src_target.blob, dst_package)
        new_rid = dst_part.relate_to(new_part, rel.reltype)
        rid_cache[old_rid] = new_rid
        elem.set(attr, new_rid)


def _append_body_element(doc, element):
    body = doc.element.body
    sect_pr = body.find(W_SECTPR)
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)


# -- 受保护视觉/数学标记，含这些东西的段落应保持原样 --
# 注意：mc:(AlternateContent) 的命名空间前缀未在 python-docx nsmap 中注册，
# 需手动拼 Clark 名: {urn}AlternateContent
_NS_MC = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_PROTECTED_CONTENT_TAGS = (
    qn('w:drawing'), qn('w:pict'), qn('w:object'),
    qn('m:oMath'), qn('m:oMathPara'),
    f"{_NS_MC}AlternateContent",
)


def _paragraph_has_protected_content(element) -> bool:
    """检测段落 XML 元素中是否含图片/公式/OLE/备用内容等不可覆写的视觉对象。"""
    for tag in _PROTECTED_CONTENT_TAGS:
        if element.find(f'.//{tag}') is not None:
            return True
    return False


def _apply_text_run_format(new_element, color_hex: str, bold_tags: list[str] | None = None) -> None:
    """对段落中不含图片/OLE 对象的 w:r 元素覆写颜色，并对匹配 bold_tags 的文字加粗。
    bold_tags: 如 ["【答案】", "【解析】"] — 命中任一文字的 run 设 w:b。
    """
    from lxml import etree
    tags = bold_tags or []
    for r_elem in new_element.findall(qn('w:r')):
        if (r_elem.find(qn('w:drawing')) is not None
                or r_elem.find(qn('w:pict')) is not None
                or r_elem.find(qn('w:object')) is not None):
            continue
        # 获取文字内容判断是否需要加粗
        text_content = ""
        for t in r_elem.findall(qn('w:t')):
            text_content += (t.text or "")
        should_bold = any(tag and tag in text_content for tag in tags)

        rpr = r_elem.find(qn('w:rPr'))
        if rpr is None:
            rpr = etree.SubElement(r_elem, qn('w:rPr'))

        # 字体：拉丁 Times New Roman，中文 宋体
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), '宋体')

        # 颜色
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = etree.SubElement(rpr, qn('w:color'))
        color_el.set(qn('w:val'), color_hex)

        # 加粗 (w:b)
        if should_bold:
            b_el = rpr.find(qn('w:b'))
            if b_el is None:
                b_el = etree.SubElement(rpr, qn('w:b'))


def _format_copied_element(new_element, doc_rid_cache, dst_doc, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> None:
    """对刚复制到目标文档的 w:p 元素施加字体/字号/行距。
    含图片/公式/OLE 对象的段落只保留颜色+加粗覆写，不改变行距和字体防止渲染劣化。
    直接操作 lxml 元素，不依赖 python-docx 的 paragraph 缓存。

    text_color: None=不改变颜色; (R,G,B)=设为指定颜色（如 (255,0,0) 红色）。
    bold_tags: 【答案】/【解析】等需加粗的 tag 列表。
    """
    from lxml import etree

    # --- 清掉所有段落底纹（w:shd），防止源文档背景色污染输出 ---
    for scope in (new_element, *new_element.findall(qn('w:r'))):
        for parent_tag in (qn('w:pPr'), qn('w:rPr')):
            parent = scope.find(parent_tag)
            if parent is not None:
                shd = parent.find(qn('w:shd'))
                if shd is not None:
                    parent.remove(shd)

    color_hex = ''.join(f'{value:02X}' for value in text_color) if text_color is not None else None

    # 含图片/公式/OLE 对象的段落→跳过版式覆写，只处理颜色和加粗
    if _paragraph_has_protected_content(new_element):
        if color_hex is not None:
            _apply_text_run_format(new_element, color_hex, bold_tags)
        return

    # --- 段落属性：单倍行距，段前0段后2pt ---
    ppr = new_element.find(qn('w:pPr'))
    if ppr is None:
        ppr = etree.SubElement(new_element, qn('w:pPr'))
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = etree.SubElement(ppr, qn('w:spacing'))
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '40')  # 2pt = 40 twips

    for r_elem in new_element.findall(qn('w:r')):
        # 跳过含 drawing / pict 的图片 run
        if (r_elem.find(qn('w:drawing')) is not None
                or r_elem.find(qn('w:pict')) is not None
                or r_elem.find(qn('w:object')) is not None):
            continue

        # 字体/字号：宋体 10.5pt
        rpr = r_elem.find(qn('w:rPr'))
        if rpr is None:
            rpr = etree.SubElement(r_elem, qn('w:rPr'))

        # 西文字体
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = etree.SubElement(rpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), '宋体')

        # 字号 10.5pt = 21 half-points
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = etree.SubElement(rpr, qn('w:sz'))
        sz.set(qn('w:val'), '21')
        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = etree.SubElement(rpr, qn('w:szCs'))
        sz_cs.set(qn('w:val'), '21')

        # 颜色 + 加粗（标签）
        if color_hex is not None:
            color_el = rpr.find(qn('w:color'))
            if color_el is None:
                color_el = etree.SubElement(rpr, qn('w:color'))
            color_el.set(qn('w:val'), color_hex)

            text_content = ""
            for t in r_elem.findall(qn('w:t')):
                text_content += (t.text or "")
            if bold_tags and any(tag and tag in text_content for tag in bold_tags):
                b_el = rpr.find(qn('w:b'))
                if b_el is None:
                    etree.SubElement(rpr, qn('w:b'))


# -- 按段落索引从源 DOCX 复制段落的缓存 --
_SRC_DOC_CACHE: dict[str, Any] = {}


def _get_cached_src_doc(source_path: str):
    if source_path not in _SRC_DOC_CACHE:
        _SRC_DOC_CACHE[source_path] = Document(source_path)
    return _SRC_DOC_CACHE[source_path]


def _copy_source_paragraphs(doc, source_path: str, indices: list[int], rid_cache: dict | None = None, text_color: tuple | None = None, bold_tags: list[str] | None = None) -> int:
    """从源 DOCX 复制指定段落到目标文档（含图片/公式 remap）。
    共享 rid_cache 避免同一次生成中重复注入相同图片资源。
    text_color: None=不改变颜色; (R,G,B)=设为指定颜色。
    bold_tags: 需加粗的标签文本列表，如 ["【答案】"]。
    返回复制的段落数。
    """
    if not source_path or not indices:
        return 0
    src_path = Path(source_path)
    if not src_path.exists():
        print(f"  → 警告：源 DOCX 不存在，无法复制段落：{src_path}")
        return 0
    if rid_cache is None:
        rid_cache = {}
    try:
        src_doc = _get_cached_src_doc(str(src_path))
        paragraphs = list(src_doc.paragraphs)
        copied = 0
        sorted_indices = sorted(set(int(i) for i in indices if isinstance(i, int) or str(i).isdigit()))
        for idx in sorted_indices:
            if idx < 0 or idx >= len(paragraphs):
                continue
            new_element = copy.deepcopy(paragraphs[idx]._element)
            _remap_embedded_relationships(new_element, src_doc, doc, rid_cache)
            _format_copied_element(new_element, rid_cache, doc, text_color=text_color, bold_tags=bold_tags)
            _append_body_element(doc, new_element)
            copied += 1
        return copied
    except Exception as exc:
        print(f"  → 警告：复制源段落失败：{exc}")
        return 0


# 答案/解析颜色
_ANSWER_RED = (255, 0, 0)


def _append_original_question_block(doc, question: dict[str, Any]) -> int:
    """复制带图题的原始段落到目标文档，并按分区上色。
    题干 → 默认黑色；答案 → 红色；解析 → 红色。
    返回复制的段落数。
    """
    source_path = question.get("source_docx_path") or question.get("source_path")
    if not source_path:
        return 0
    rid_cache: dict = {}
    total = 0

    # 题干段落：黑色（text_color=None 即不覆写颜色）
    stem_indices = question.get("stem_paragraph_indices") or []
    total += _copy_source_paragraphs(doc, source_path, stem_indices, rid_cache, text_color=None)

    # 答案段落：红色 + 【答案】标签加粗
    answer_indices = question.get("answer_paragraph_indices") or []
    total += _copy_source_paragraphs(doc, source_path, answer_indices, rid_cache, text_color=_ANSWER_RED, bold_tags=["【答案】"])

    # 解析段落：红色 + 【解析】/【详解】/【分析】标签加粗
    analysis_indices = question.get("analysis_paragraph_indices") or []
    total += _copy_source_paragraphs(doc, source_path, analysis_indices, rid_cache, text_color=_ANSWER_RED, bold_tags=["【解析】", "【详解】", "【分析】"])

    # 兜底：如果有 source_paragraph_indices 但分区索引缺失，统一用黑色
    all_indices = question.get("source_paragraph_indices") or []
    covered = set(stem_indices + answer_indices + analysis_indices)
    fallback = [i for i in all_indices if i not in covered]
    if fallback:
        total += _copy_source_paragraphs(doc, source_path, fallback, rid_cache, text_color=None)

    return total



def _question_is_protected(question: dict[str, Any] | None) -> bool:
    return bool(question and question.get("protected_original_docx_block"))


def _question_has_structured_images(question: dict[str, Any] | None) -> bool:
    """判断题目结构中是否已有可排版的题干图/选项图字段。
    兼容 stem_images/option_images（API 丰富来源）和 image_refs（DOCX 解析来源）。
    """
    if not isinstance(question, dict):
        return False
    # API 丰富来源
    if question.get("stem_images") or question.get("option_images"):
        return True
    # DOCX 解析来源：image_refs.stem 存在即说明题干含图
    image_refs = question.get("image_refs") or {}
    if image_refs.get("stem"):
        return True
    options = question.get("options") or question.get("option") or []
    if isinstance(options, list):
        return any(isinstance(opt, dict) and opt.get("images") for opt in options)
    return False


def _strip_question_number(text: str) -> str:
    return re.sub(r"^\s*\d+[\.．、]\s*", "", str(text or "").strip())


def _parse_text_options_from_fallback(fallback_text: str) -> list[dict[str, Any]]:
    """从解析版文本兜底提取 A-D 选项，供图片字段缺少文字时补齐。"""
    lines = [line.strip() for line in str(fallback_text or "").splitlines() if line.strip()]
    text = "\n".join(lines[1:]) if len(lines) > 1 else ""
    pattern = re.compile(r"([A-D])[\.．]\s*(.*?)(?=(?:\s+[A-D][\.．])|$)", re.S)
    options = []
    for match in pattern.finditer(text.replace("\t", " ")):
        content = re.sub(r"\s+", " ", match.group(2)).strip()
        options.append({"label": match.group(1), "text": content, "images": []})
    return options


def _normalize_structured_choice_question(question: dict[str, Any], fallback_text: str) -> dict[str, Any] | None:
    """把不同来源的图片选择题字段归一化为 docx_utils1 可消费的结构。"""
    if not isinstance(question, dict):
        return None
    fallback_lines = [line.strip() for line in str(fallback_text or "").splitlines() if line.strip()]
    stem_text = str(
        question.get("stem")
        or question.get("stem_text")
        or question.get("question")
        or question.get("content")
        or (fallback_lines[0] if fallback_lines else "")
        or ""
    ).strip()
    if stem_text and not re.match(r"^\d+[\.．、]", stem_text) and fallback_lines and re.match(r"^\d+[\.．、]", fallback_lines[0]):
        # 保留解析版文本中的题号，避免结构化 stem 只有裸题干时丢失编号。
        stem_text = re.match(r"^\s*\d+[\.．、]", fallback_lines[0]).group(0) + stem_text
    elif stem_text:
        stem_text = _strip_question_number(stem_text)
        if fallback_lines and re.match(r"^\d+[\.．、]", fallback_lines[0]):
            stem_text = re.match(r"^\s*\d+[\.．、]", fallback_lines[0]).group(0) + stem_text

    stem_images = question.get("stem_images") or question.get("images") or []
    raw_options = question.get("options") or question.get("option") or []
    text_options = _parse_text_options_from_fallback(fallback_text)
    text_by_label = {opt["label"]: opt.get("text", "") for opt in text_options}

    options = []
    if isinstance(raw_options, list):
        for idx, opt in enumerate(raw_options):
            if isinstance(opt, dict):
                label = str(opt.get("label") or chr(ord("A") + idx)).strip().rstrip(".．")
                text = str(opt.get("text") or opt.get("content") or text_by_label.get(label, "") or "").strip()
                images = opt.get("images") or opt.get("image") or []
                if isinstance(images, (str, dict)):
                    images = [images]
                options.append({"label": label, "text": text, "images": images})
            elif idx < len(text_options):
                options.append(text_options[idx])
    elif isinstance(raw_options, dict):
        for idx, label in enumerate("ABCD"):
            opt = raw_options.get(label) or raw_options.get(label.lower())
            if isinstance(opt, dict):
                images = opt.get("images") or opt.get("image") or []
                if isinstance(images, (str, dict)):
                    images = [images]
                options.append({
                    "label": label,
                    "text": str(opt.get("text") or opt.get("content") or text_by_label.get(label, "") or "").strip(),
                    "images": images,
                })
    if not options:
        options = text_options

    option_images = question.get("option_images") or {}
    if isinstance(option_images, dict):
        by_label = {opt["label"]: opt for opt in options}
        for label, images in option_images.items():
            key = str(label).strip().rstrip(".．")
            if isinstance(images, (str, dict)):
                images = [images]
            by_label.setdefault(key, {"label": key, "text": text_by_label.get(key, ""), "images": []})["images"].extend(images or [])
        options = [by_label[label] for label in sorted(by_label.keys())]

    if not stem_images and not any(opt.get("images") for opt in options):
        return None
    return {"stem_text": stem_text, "stem_images": stem_images, "options": options}


def _copy_common_scripts_to_output(sub_dir):
    """将常用后处理脚本复制到当前考类/教材输出目录，便于直接双击使用。"""
    source_dir = BASE_DIR / "01_工具脚本" / "通用脚本"
    for script_name in COMMON_SCRIPT_NAMES:
        src = source_dir / script_name
        if not src.exists():
            print(f"  → 警告：未找到通用脚本 {src}")
            continue
        dst = sub_dir / script_name
        try:
            shutil.copy2(src, dst)
        except Exception as exc:
            print(f"  → 警告：复制通用脚本失败 {script_name}: {exc}")

def _get_value(source, *names, default=""):
    """兼容 dict 和 dataclass/object 的安全取值。"""
    if source is None:
        return default
    for name in names:
        if isinstance(source, dict):
            value = source.get(name)
        else:
            value = getattr(source, name, None)
        if value not in (None, ""):
            return value
    return default


def _normalize_province_name(name):
    """规范省市名称，自治区全称保持不变。"""
    text = str(name or "").strip()
    if not text:
        return ""
    direct = {
        "北京": "北京市", "天津": "天津市", "上海": "上海市", "重庆": "重庆市",
        "河北": "河北省", "山西": "山西省", "辽宁": "辽宁省", "吉林": "吉林省",
        "黑龙江": "黑龙江省", "江苏": "江苏省", "浙江": "浙江省", "安徽": "安徽省",
        "福建": "福建省", "江西": "江西省", "山东": "山东省", "河南": "河南省",
        "湖北": "湖北省", "湖南": "湖南省", "广东": "广东省", "海南": "海南省",
        "四川": "四川省", "贵州": "贵州省", "云南": "云南省", "陕西": "陕西省",
        "甘肃": "甘肃省", "青海": "青海省", "台湾": "台湾省",
        "内蒙古": "内蒙古自治区", "广西": "广西壮族自治区", "西藏": "西藏自治区",
        "宁夏": "宁夏回族自治区", "新疆": "新疆维吾尔自治区",
    }
    if text in direct:
        return direct[text]
    if text.endswith(("省", "市", "自治区")):
        return text
    return direct.get(text, text)


def _safe_path_part(text):
    """清理 Windows 文件名/目录名非法字符。"""
    cleaned = re.sub(r'[<>:"/\\|?*]+', "_", str(text or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or "未命名"


def _parse_edition(raw):
    """从简短版次说明中兜底提取出版社和版次。"""
    text = str(raw or "").strip()
    if not text:
        return "高教版", "第一版"
    if "·" in text:
        publisher, edition = text.split("·", 1)
        return publisher.strip() or "高教版", edition.strip() or "第一版"
    if "版" in text and len(text) <= 6:
        return text, "第一版"
    return "高教版", text or "第一版"


def _extract_textbook_detail(meta, topic):
    """提取教材/课程显示所需信息。"""
    detail = None
    course = _get_value(topic, "course", "module")
    course_map = _get_value(meta, "course_textbook_map", default={}) or {}
    if course_map and course in course_map:
        detail = course_map[course]
    elif _get_value(meta, "textbook_details", default=None):
        detail = _get_value(meta, "textbook_details")[0]

    if detail:
        if isinstance(detail, dict):
            name = detail.get("name") or detail.get("textbook_name") or ""
            publisher = detail.get("publisher") or "高教版"
            edition = detail.get("edition") or "第一版"
            display = detail.get("display") or f"{publisher}·{edition}"
        else:
            name = getattr(detail, "name", "")
            publisher = getattr(detail, "publisher", "高教版")
            edition = getattr(detail, "edition", "第一版")
            display = getattr(detail, "display", f"{publisher}·{edition}")
        return {"name": name or "教材", "publisher": publisher, "edition": edition, "display": display}

    textbook_list = _get_value(meta, "textbook_list", default=None) or []
    if textbook_list:
        first = textbook_list[0]
        if isinstance(first, (list, tuple)) and len(first) >= 2:
            publisher, edition = _parse_edition(first[1])
            return {"name": first[0] or "教材", "publisher": publisher, "edition": edition, "display": f"{publisher}·{edition}"}
        return {"name": str(first), "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}

    textbooks = _get_value(meta, "textbooks")
    match = re.search(r"《(.+?)》", str(textbooks or ""))
    if match:
        return {"name": match.group(1), "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}

    fallback = _get_value(meta, "subject") or _get_value(topic, "course", "module") or "教材"
    return {"name": fallback, "publisher": "高教版", "edition": "第一版", "display": "高教版·第一版"}


def _format_paper_type(topic):
    text = str(_get_value(topic, "paper_type", "paper_kind", "type", "卷型") or "").strip()
    aliases = {
        "考点": "考点训练卷", "考点卷": "考点训练卷", "point": "考点训练卷",
        "专题": "专题训练卷", "专题卷": "专题训练卷", "topic": "专题训练卷",
        "综合": "课程综合卷", "课程综合": "课程综合卷", "综合卷": "课程综合卷", "course": "课程综合卷",
    }
    if text in {"考点训练卷", "专题训练卷", "课程综合卷"}:
        return text
    lowered = text.lower()
    if lowered in aliases:
        return aliases[lowered]
    return aliases.get(text, text or "考点训练卷")


def _format_exam_category(meta):
    return str(_get_value(meta, "category", "exam_category", "major_category") or "").strip()


def _format_paper_name(topic):
    paper_type = _format_paper_type(topic)
    if paper_type == "考点训练卷":
        name = _get_value(topic, "point_name", "theme", "topic", "subject")
    elif paper_type == "专题训练卷":
        name = _get_value(topic, "topic", "theme", "point_name", "subject")
    else:
        name = _get_value(topic, "theme", "topic", "module", "subject") or "课程综合"
    return str(name or "未命名试卷").strip()


def _format_course_name(meta, topic, textbook_name=None):
    return str(
        _get_value(topic, "course", "module")
        or _get_value(meta, "subject", "course")
        or textbook_name
        or "课程"
    ).strip()


def _unique_nonempty(values):
    result = []
    seen = set()
    for value in values:
        text = str(value or "").strip()
        if text and text not in seen:
            result.append(text)
            seen.add(text)
    return result


def _iter_topic_rows(topic):
    rows = _get_value(topic, "rows", default=[]) or []
    return rows if isinstance(rows, (list, tuple)) else []


def _row_value(row, *names):
    for name in names:
        value = _get_value(row, name, default="")
        if value:
            return value
        raw = _get_value(row, "raw", default={}) or {}
        if isinstance(raw, dict):
            for key in names:
                value = raw.get(key)
                if value:
                    return value
    return ""


def _format_all_courses(meta, topic, course_name):
    """按规划总表收集的全考类课程列表生成专辑课程信息和计数。"""
    # 优先使用 meta.all_courses（从全部行收集的课程列表）
    courses = getattr(meta, "all_courses", None) or []
    if not courses:
        # 兜底：从传入的 rows 或 topic 中提取
        rows = _get_value(meta, "rows", default=None) or _get_value(topic, "all_rows", default=None) or _iter_topic_rows(topic)
        courses = _unique_nonempty(_row_value(row, "module", "知识模块") for row in rows)
    if not courses:
        explicit = _get_value(meta, "course_names", "courses", default=None)
        if explicit:
            courses = _unique_nonempty(explicit if isinstance(explicit, (list, tuple)) else re.split(r"[、,，]\s*", str(explicit)))
    if not courses:
        courses = [course_name]
    return "、".join(courses), len(courses)



def _strip_and_reindex_content(raw_text: str) -> str:
    """去掉原始序号（如 "2."、"（2）"、"一、"），从 1 重新编号，多条用分号分隔。"""
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return raw_text
    stripped = []
    for line in lines:
        # 先按行切分；行内若含多个序号则再切
        parts = re.split(r'(?:^|[（(])\s*\d+\s*[）).、．]\s*|^[一二三四五六七八九十]+\s*[、．.]\s*', line)
        for part in parts:
            cleaned = re.sub(r'^[、．.,，；;]+\s*', '', part).strip()
            if cleaned:
                stripped.append(cleaned)
    if not stripped:
        return raw_text
    if len(stripped) == 1:
        return f'1. {stripped[0]}'
    return '；'.join(f'{i}. {item}' for i, item in enumerate(stripped, 1))


def _format_specific_content(topic):
    """按卷型映射“具体内容为”：考点卷用考点内容，专题/综合卷用覆盖名称。"""
    paper_type = _format_paper_type(topic)
    rows = _iter_topic_rows(topic)
    if paper_type == "考点训练卷":
        content = _get_value(topic, "point_content", "knowledge", "content")
        if not content and rows:
            content = _row_value(rows[0], "point_content", "掌握/理解/了解考点内容", "knowledge", "content")
        raw = str(content or _format_paper_name(topic)).strip()
        return _strip_and_reindex_content(raw)

    if paper_type == "专题训练卷":
        names = _unique_nonempty(
            _row_value(row, "point_name", "考点名称（考点训练卷名称）", "knowledge_point")
            for row in rows
        )
        if not names:
            names = [_format_paper_name(topic)]
        return "、".join(names)

    names = _unique_nonempty(_row_value(row, "topic", "考纲一级标题（专题名称）") for row in rows)
    if not names:
        names = _unique_nonempty(
            _row_value(row, "point_name", "考点名称（考点训练卷名称）")
            for row in rows
        )
    if not names:
        names = [_format_paper_name(topic)]
    return "、".join(names)


def _format_exam_table_title(meta, province):
    year = _get_value(meta, "exam_year", "year", default="2025")
    return _get_value(
        meta,
        "exam_table_title",
        default=f"{province}{year}年中等职业学校毕业生进入普通高校学习专业基础课和专业课考试科目表",
    )


def _get_topic_output_base(meta, topic, output_dir):
    """按"省份 考类"组织输出目录（单级目录，省/市与考类合并）。"""
    province = _safe_path_part(_normalize_province_name(_get_value(meta, "province")))
    category = _safe_path_part(_format_exam_category(meta))
    combined = " ".join(part for part in (province, category) if part and part != "未命名").strip()
    sub_dir = Path(output_dir)
    if combined:
        sub_dir = sub_dir / combined
    return sub_dir


def _format_exam_type(meta):
    """生成文件名和编写说明使用的考试类型。"""
    return str(_get_value(meta, "exam_type_name", "exam_type", default="高职分类考试") or "高职分类考试").strip()


def _build_docx_filename(meta, topic, set_idx=1, variant="解析版"):
    """按考纲百套卷命名规范生成 DOCX 文件名。"""
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    category = _format_exam_category(meta)
    textbook = _extract_textbook_detail(meta, topic)
    course_name = _format_course_name(meta, topic, textbook["name"])
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    sets = int(_get_value(topic, "sets", default=1) or 1)
    set_suffix = f"-{set_idx}" if sets > 1 else ""
    filename = (
        f"第{seq}卷{set_suffix} {paper_name} {paper_type}《{course_name}》"
        f"{province}（{exam_type}）{category} 考纲百套卷（{variant}）.docx"
    )
    return _safe_path_part(filename)


def _build_docx_titles(meta, topic):
    """生成考纲百套卷三行标题，与模板脚本的标题区域保持一致。"""
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    category = _format_exam_category(meta)
    textbook = _extract_textbook_detail(meta, topic)
    course_name = _format_course_name(meta, topic, textbook["name"])
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    category_part = f"《{category}考纲百套卷》" if category else "《考纲百套卷》"
    return (
        f"{province}（{exam_type}）{category_part}第{seq}卷",
        f"《{course_name}》",
        f"{paper_name} {paper_type}",
    )


def _paper_has_blank_before_first_heading(paper_text: str) -> bool:
    """判断第一类大题前是否已经保留空行。"""
    lines = paper_text.split("\n")
    for idx, line in enumerate(lines):
        if re.match(r"^\s*[一二三四五六七八九十][、.．]", line):
            return idx > 0 and not lines[idx - 1].strip()
    return False


def _add_exam_info_header(doc, add_para, add_separator_blank=True) -> None:
    """在标题下方添加考试时间、满分和考生信息。"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    line1 = "考试时间：60分钟        满分：100分"
    line2 = "班级：________ 姓名：________ 学号：________ 成绩：________"

    add_para(doc, line1, font_name="宋体", font_size=10.5,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_para(doc, line2, font_name="宋体", font_size=10.5,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    if add_separator_blank:
        add_para(doc, " ", font_name="宋体", font_size=10.5,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)


def generate_docx(meta, topic, set_idx, paper_text, output_dir, needs_manual_review=False, questions=None):
    """将试卷文本生成为格式化的 DOCX 文件"""
    sys.path.insert(0, str(BASE_DIR / "01_工具脚本" / "核心脚本"))
    from docx_utils1 import (
        copy_template, add_editorial_note,
        add_paragraph_with_style, add_question_options_table,
        add_structured_choice_question,
        add_labeled_text, save_docx
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # 清理 AI 返回文本
    paper_text = _clean_paper_text(paper_text)

    # 确定文件名和标题元数据
    seq = _get_value(topic, "seq", "paper_no", default=1)
    province = _normalize_province_name(_get_value(meta, "province"))
    exam_type = _format_exam_type(meta)
    textbook = _extract_textbook_detail(meta, topic)
    textbook_name = textbook["name"]
    publisher = textbook["publisher"]
    edition = textbook["edition"]
    paper_name = _format_paper_name(topic)
    paper_type = _format_paper_type(topic)
    course_name = _format_course_name(meta, topic, textbook_name)

    filename = _build_docx_filename(meta, topic, set_idx)
    if needs_manual_review:
        filename = f"（待人工审核）{filename}"

    # 按"省份/考类/课程"组织子目录，与规划表位置保持一致
    sub_dir = _get_topic_output_base(meta, topic, output_dir)
    os.makedirs(sub_dir, exist_ok=True)
    _copy_common_scripts_to_output(sub_dir)
    output_path = sub_dir / filename

    # 使用 02_配置资源/模板和资源/template.docx 作为唯一模板来源，保留其页眉、页脚和样式。
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"未找到 DOCX 模板：{TEMPLATE_PATH}")
    doc = copy_template(str(TEMPLATE_PATH), str(output_path))

    # 不强制覆盖模板页边距，避免破坏 template.docx 中已配置的页眉页脚版式。

    # 编写说明；新版考纲百套卷模板不再插入 separator 图片。
    sep_img = None
    all_courses, course_count = _format_all_courses(meta, topic, course_name)
    specific_content = _format_specific_content(topic)
    add_editorial_note(
        doc,
        textbook_name=textbook_name,
        edition=edition,
        chapter_seq=seq,
        knowledge_scope=_get_value(topic, "knowledge", "point_content", "topic", default=paper_name),
        province=province,
        publisher=publisher,
        separator_image=sep_img,
        exam_type=exam_type,
        series_name="考纲百套卷",
        volume_label=f"第{seq}卷",
        paper_type=paper_type,
        paper_name=paper_name,
        course_name=course_name,
        category=_format_exam_category(meta),
        all_courses=all_courses,
        course_count=course_count,
        exam_table_title=_format_exam_table_title(meta, province),
        specific_content=specific_content,
    )

    # 标题
    title1, title2, title3 = _build_docx_titles(meta, topic)

    add_paragraph_with_style(doc, title1, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title2, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph_with_style(doc, title3, font_name="宋体", font_size=14,
                             bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # 考试信息行
    _add_exam_info_header(
        doc,
        add_paragraph_with_style,
        add_separator_blank=not _paper_has_blank_before_first_heading(paper_text),
    )

    # 答案/解析使用红色字体
    RED_COLOR = (255, 0, 0)

    # 逐行解析并输出
    lines = paper_text.split("\n")
    protected_questions = questions or []
    question_cursor = 0
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 大题标题（加粗，黑体）
        if re.match(r"^[一二三四五六七八九十][、.．]", line):
            add_paragraph_with_style(doc, line, font_name="黑体", font_size=12,
                                     bold=True, space_after=6)
            i += 1
            continue

        # 题目开头（数字编号）
        if re.match(r"^\d+[\.．、]", line):
            current_question = protected_questions[question_cursor] if question_cursor < len(protected_questions) else None
            question_cursor += 1
            if _question_is_protected(current_question):
                copied = _append_original_question_block(doc, current_question)
                if copied:
                    j = i + 1
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if re.match(r"^\d+[\.．、]", next_line):
                            break
                        if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                            break
                        j += 1
                    i = j
                    continue
                print(f"  → 警告：第{current_question.get('question_no')}题原始题块复制失败，回退为文本生成。")

            # 收集题目完整文本（直到下一个【答案】或下一题）
            q_lines = [line]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if next_line.startswith(("【答案】", "【解析】", "【详解】", "【分析】")):
                    break
                if not next_line:
                    # 防止题目文本和【答案】之间的偶发空行进入 DOCX。
                    lookahead = j + 1
                    while lookahead < len(lines) and not lines[lookahead].strip():
                        lookahead += 1
                    if lookahead < len(lines) and lines[lookahead].strip().startswith(("【答案】", "【解析】", "【详解】", "【分析】")):
                        j = lookahead
                        continue
                    j += 1
                    continue
                if re.match(r"^\d+[\.．、]", next_line):
                    break
                if re.match(r"^[一二三四五六七八九十][、.．]", next_line):
                    break
                q_lines.append(lines[j])
                j += 1

            q_text = "\n".join(q_lines)

            # 结构化图片题：题干文字 → 题干图 → 图片选项（protected 原始题块优先级更高）
            if _question_has_structured_images(current_question):
                normalized = _normalize_structured_choice_question(current_question, q_text)
                if normalized:
                    rendered = add_structured_choice_question(
                        doc,
                        normalized["stem_text"],
                        normalized["options"],
                        stem_images=normalized["stem_images"],
                    )
                    if rendered:
                        i = j
                        continue
                    q_no = current_question.get("question_no") if current_question else ""
                    print(f"  → 警告：第{q_no}题结构化图片排版失败，回退为文本生成。")

            # 判断是否有选项（含 A. B. C. D.）
            if re.search(r"[A-D][\.．]\s*\S", q_text):
                add_question_options_table(doc, q_text)
            else:
                add_paragraph_with_style(doc, q_text, font_size=10.5, space_after=2)

            i = j
            continue

        # 【答案】— 红色字体；多行答案继续归入同一红色段落
        if line.startswith("【答案】"):
            answer_lines = [line[4:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith(("【解析】", "【详解】", "【分析】", "【答案】"))
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    answer_lines.append(next_line)
                j += 1
            answer_content = "\n".join(answer_lines)

            # 答案图片：从 source DOCX 按段落索引复制原段（与题干图片处理方式一致）
            answer_copied = False
            if current_question and not _question_is_protected(current_question):
                answer_refs = (current_question.get("image_refs") or {}).get("answer") or []
                answer_indices = current_question.get("answer_paragraph_indices") or []
                if answer_refs and answer_indices:
                    source_path = current_question.get("source_docx_path") or current_question.get("source_path")
                    if source_path:
                        copied = _copy_source_paragraphs(doc, source_path, answer_indices, text_color=_ANSWER_RED, bold_tags=["【答案】"])
                        answer_copied = copied > 0

            if answer_copied:
                # 源段落已包含完整答案（标签+图片+文字），只需追加可能的纯文本后续行
                if answer_content:
                    add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            else:
                add_labeled_text(doc, "【答案】", answer_content, color=RED_COLOR)
            i = j
            continue

        # 【解析】/【详解】/【分析】— 红色字体，统一显示为【解析】
        if line.startswith(("【解析】", "【详解】", "【分析】")):
            label = "【解析】"
            explanation_lines = [line[line.index("】") + 1:]]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if (next_line.startswith(("【答案】", "【解析】", "【详解】"))
                        or re.match(r"^\d+[\.．、]", next_line)
                        or re.match(r"^[一二三四五六七八九十][、.．]", next_line)):
                    break
                if next_line:
                    explanation_lines.append(next_line)
                j += 1
            explanation_content = "\n".join(explanation_lines)

            # 解析图片：从 source DOCX 按段落索引复制原段（与题干图片处理方式一致）
            analysis_copied = False
            if current_question and not _question_is_protected(current_question):
                analysis_refs = (current_question.get("image_refs") or {}).get("analysis") or []
                analysis_indices = current_question.get("analysis_paragraph_indices") or []
                if analysis_refs and analysis_indices:
                    source_path = current_question.get("source_docx_path") or current_question.get("source_path")
                    if source_path:
                        copied = _copy_source_paragraphs(doc, source_path, analysis_indices, text_color=_ANSWER_RED, bold_tags=["【解析】", "【详解】", "【分析】"])
                        analysis_copied = copied > 0

            if analysis_copied:
                # 源段落已包含完整解析（标签+图片+文字），只需追加可能的纯文本后续行
                if explanation_content:
                    add_labeled_text(doc, label, explanation_content, color=RED_COLOR)
            else:
                add_labeled_text(doc, label, explanation_content, color=RED_COLOR)
            i = j
            continue

        # 普通文本行
        if line:
            add_paragraph_with_style(doc, line, font_size=10.5, space_after=2)

        i += 1

    save_docx(doc, str(output_path))
    residuals = _find_residual_formula_text(doc)
    if residuals:
        print(f"  [!] 公式转换后仍发现 {len(residuals)} 处疑似 LaTeX/math 残留，请检查 DOCX：{output_path}")
        for sample in residuals[:3]:
            print(f"      - {sample}")
    return str(output_path)
