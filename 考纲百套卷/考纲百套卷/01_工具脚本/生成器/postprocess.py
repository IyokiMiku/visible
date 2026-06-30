"""原卷版生成、打包和分类后处理。"""
import os
import re
import shutil
import zipfile
from pathlib import Path

SUBJECTIVE_TYPE_KEYWORDS = ("简答", "综合", "计算", "作图", "绘图", "画图", "分析", "应用")
OBJECTIVE_TYPE_KEYWORDS = ("选择", "判断", "填空")
QUESTION_RE = re.compile(r"^\s*\d+\s*[．.、]")
TYPE_HEADING_RE = re.compile(r"^[一二三四五六七八九十百千万]+[、.．]\s*([^（(\s]+)")

def _is_type_heading(text):
    """判断段落是否为大题标题。"""
    return bool(TYPE_HEADING_RE.match(text.strip()))

def _is_subjective_heading(text):
    """判断大题标题是否为需要留答题空间的主观题。"""
    match = TYPE_HEADING_RE.match(text.strip())
    if not match:
        return False
    name = match.group(1)
    if any(keyword in name for keyword in OBJECTIVE_TYPE_KEYWORDS):
        return False
    return any(keyword in name for keyword in SUBJECTIVE_TYPE_KEYWORDS)

def _insert_blank_paragraph_after(paragraph):
    """在指定段落后插入一个空白段落，并尽量沿用原段落格式。"""
    from copy import deepcopy
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement('w:p')
    pPr = paragraph._element.pPr
    if pPr is not None:
        new_p.append(deepcopy(pPr))
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def _add_subjective_answer_spaces(doc, blank_count=3):
    """在原卷版主观题每道题后插入空白答题区域。"""
    insert_after = []
    in_subjective = False
    current_question_last_para = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if _is_type_heading(text):
            if in_subjective and current_question_last_para is not None:
                insert_after.append(current_question_last_para)
            in_subjective = _is_subjective_heading(text)
            current_question_last_para = None
            continue

        if in_subjective and QUESTION_RE.match(text):
            if current_question_last_para is not None:
                insert_after.append(current_question_last_para)
            current_question_last_para = p
            continue

        if in_subjective and current_question_last_para is not None:
            current_question_last_para = p

    if in_subjective and current_question_last_para is not None:
        insert_after.append(current_question_last_para)

    inserted = 0
    for p in reversed(insert_after):
        anchor = p
        for _ in range(blank_count):
            anchor = _insert_blank_paragraph_after(anchor)
            inserted += 1
    return inserted

def _convert_to_blank(src_path, dst_path):
    """将解析版 docx 转换为原卷版（删除答案和解析段落）"""
    from docx import Document
    from docx.oxml.ns import qn as _qn

    shutil.copy2(src_path, dst_path)
    doc = Document(dst_path)
    paras = list(doc.paragraphs)

    # 标记需要删除的段落（含【答案】【解析】【详解】）
    rm_indices = []
    in_answer = False

    for i, p in enumerate(paras):
        text = p.text.strip()

        # 大类标题或新题号 → 重置答案区域标记
        if re.match(r'^[一二三四五六七八九十][、.．]', text):
            in_answer = False
            continue
        if re.match(r'^\d+[\.．、]', text) and '【答案】' not in text:
            in_answer = False
            continue

        # 混合行（题干+【答案】同行）→ 截掉答案部分
        if re.match(r'^\d+[\.．、]', text) and '【答案】' in text:
            idx = p.text.find('【答案】')
            if idx >= 0:
                pos = 0
                for run in p.runs:
                    rt = run.text or ''
                    rs = pos
                    re_ = pos + len(rt)
                    if rs >= idx:
                        run.text = ''
                    elif re_ > idx:
                        run.text = rt[:idx - rs]
                    pos = re_
            in_answer = True
            continue

        # 答案/解析行
        if '【答案】' in text or '【解析】' in text or '【详解】' in text:
            in_answer = True
            rm_indices.append(i)
            continue

        if in_answer:
            rm_indices.append(i)
            continue

    # 删除标记的段落（从后往前）
    for i in reversed(rm_indices):
        paras[i]._element.getparent().remove(paras[i]._element)

    # 去除底纹和高亮
    W_PPR = _qn('w:pPr')
    W_RPR = _qn('w:rPr')
    W_SHD = _qn('w:shd')
    W_HL = _qn('w:highlight')
    W_R = _qn('w:r')

    for p in doc.paragraphs:
        pPr = p._element.find(W_PPR)
        if pPr is not None:
            for e in pPr.findall(W_SHD):
                pPr.remove(e)
        for r in p._element.findall(W_R):
            rPr = r.find(W_RPR)
            if rPr is None:
                continue
            for tag in (W_SHD, W_HL):
                for e in rPr.findall(tag):
                    rPr.remove(e)

    _add_subjective_answer_spaces(doc, blank_count=4)
    doc.save(dst_path)

def _find_docx_pairs(directory):
    """递归查找（解析版）/（原卷版）配对文件，兼容已分类和未分类目录。"""
    pairs = {}
    pattern = re.compile(r'^(.+?)(?:[（(](解析版|原卷版)[）)])')

    for root, dirs, files in os.walk(directory):
        if '_原始文本' in root:
            continue
        root_path = Path(root)
        folder_variant = root_path.name if root_path.name in ('解析版', '原卷版') else None
        logical_root = root_path.parent if folder_variant else root_path

        for f in files:
            path = root_path / f
            if not path.is_file():
                continue
            if not f.endswith('.docx') or f.startswith('~'):
                continue
            match = pattern.match(f)
            if not match:
                continue
            base_name = match.group(1).strip()
            variant = match.group(2)
            key = (str(logical_root), base_name)
            if key not in pairs:
                pairs[key] = {}
            pairs[key][variant] = str(path)

    return {k: v for k, v in pairs.items() if '解析版' in v and '原卷版' in v}

def _iter_existing_dirs(paths):
    """按传入顺序去重并返回存在的处理目录。"""
    seen = set()
    for path in paths:
        dir_path = Path(path)
        key = str(dir_path.resolve())
        if key in seen or not dir_path.exists():
            continue
        seen.add(key)
        yield dir_path


def _post_process(output_dir, target_dirs=None):
    """后处理三步：1.生成原卷版 2.打包zip 3.分类到子文件夹。

    target_dirs 用于限制本次只处理当前规划表/教材对应的输出目录；
    未传入时保留旧行为，处理整个输出目录。
    """
    process_dirs = list(_iter_existing_dirs(target_dirs or [output_dir]))
    if not process_dirs:
        print("  未找到需要后处理的输出目录")
        return

    # === 第1步：解析版 → 原卷版 ===
    print("\n[1/3] 生成原卷版...")
    converted = 0
    for process_dir in process_dirs:
        for root, dirs, files in os.walk(process_dir):
            if '_原始文本' in root or root.endswith('原卷版'):
                continue
            root_path = Path(root)
            dst_dir = root_path.parent / '原卷版' if root_path.name == '解析版' else root_path
            for f in files:
                src = root_path / f
                if not src.is_file():
                    continue
                if f.endswith('.docx') and '解析版' in f and '待人工审核' not in f and not f.startswith('~'):
                    out_name = f.replace('解析版', '原卷版')
                    dst = dst_dir / out_name
                    if not dst.exists():
                        try:
                            dst_dir.mkdir(parents=True, exist_ok=True)
                            _convert_to_blank(str(src), str(dst))
                            converted += 1
                            print(f"  {out_name}")
                        except Exception as e:
                            print(f"  失败: {f} → {e}")
    print(f"  共生成 {converted} 个原卷版文件")

    # === 第2步：配对打包为 zip ===
    print("\n[2/3] 打包 zip（解析版+原卷版）...")
    pairs = {}
    for process_dir in process_dirs:
        pairs.update(_find_docx_pairs(process_dir))
    zipped = 0
    for (root, base_name), variants in pairs.items():
        zip_name = f"{base_name}.zip"
        zip_path = os.path.join(root, zip_name)
        if not os.path.exists(zip_path):
            try:
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for variant_type, filepath in variants.items():
                        zf.write(filepath, os.path.basename(filepath))
                zipped += 1
                print(f"  {zip_name}")
            except Exception as e:
                print(f"  失败: {zip_name} → {e}")
    print(f"  共打包 {zipped} 个 zip 文件")

    # === 第3步：分类到解析版/原卷版/压缩包子文件夹 ===
    print("\n[3/3] 分类文件...")
    moved_jiexi = 0
    moved_yuanjuan = 0
    moved_archive = 0

    archive_extensions = {".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz", ".tgz", ".tbz2", ".txz"}

    for process_dir in process_dirs:
        for root, dirs, files in os.walk(process_dir):
            if '_原始文本' in root or root.endswith('解析版') or root.endswith('原卷版') or root.endswith('压缩包'):
                continue

            has_target_file = any(
                (f.endswith('.docx') and ('解析版' in f or '原卷版' in f))
                or Path(f).suffix.lower() in archive_extensions
                for f in files
                if not f.startswith('~') and '待人工审核' not in f
            )
            if not has_target_file:
                continue

            dir_jiexi = os.path.join(root, '解析版')
            dir_yuanjuan = os.path.join(root, '原卷版')
            dir_archive = os.path.join(root, '压缩包')
            os.makedirs(dir_jiexi, exist_ok=True)
            os.makedirs(dir_yuanjuan, exist_ok=True)
            os.makedirs(dir_archive, exist_ok=True)

            for f in files:
                if f.startswith('~') or '待人工审核' in f:
                    continue
                src = os.path.join(root, f)
                if not os.path.isfile(src):
                    continue

                if Path(f).suffix.lower() in archive_extensions:
                    shutil.move(src, os.path.join(dir_archive, f))
                    moved_archive += 1
                    continue

                if not f.endswith('.docx'):
                    continue
                if '解析版' in f:
                    shutil.move(src, os.path.join(dir_jiexi, f))
                    moved_jiexi += 1
                elif '原卷版' in f:
                    shutil.move(src, os.path.join(dir_yuanjuan, f))
                    moved_yuanjuan += 1

    print(f"  解析版: {moved_jiexi} 个文件")
    print(f"  原卷版: {moved_yuanjuan} 个文件")
    print(f"  压缩包: {moved_archive} 个文件")
    print("\n后处理完成！")
