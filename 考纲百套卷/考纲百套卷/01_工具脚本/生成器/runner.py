"""考纲百套卷主流程入口。"""
from __future__ import annotations

import argparse
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import asdict, dataclass, field
import json
import os
from pathlib import Path
import re
import time
from typing import Any, Callable, Iterable

from .config_io import (
    _load_daily_usage,
    _new_usage_summary,
    _print_token_summary,
    _record_token_usage,
    call_api,
    load_config,
    load_spec,
)
from .docx_generation import generate_docx, _get_topic_output_base
from .paper_assembler import write_analysis_text
from .paper_loader import (
    find_manual_paper_for_plan,
    load_generated_paper_json,
    load_manual_paper_for_plan,
    loaded_paper_to_dict,
)
from .paths import (
    API_RAW_OUTPUT_DIR,
    CLEAN_OUTPUT_DIR,
    CONFIG_PATH,
    FINAL_OUTPUT_DIR,
    MANUAL_PAPER_DIR,
    PLAN_DIR,
    QA_REPORT_DIR,
    RUN_RECORD_DIR,
    ensure_output_dirs,
    manual_paper_dir_for_meta,
)
from .planning import (
    COURSE_PAPER_TYPE,
    POINT_PAPER_TYPE,
    TOPIC_PAPER_TYPE,
    PaperPlan,
    attach_blueprints_for_papers,
    blueprint_filename_for_paper,
    format_paper_label,
    group_by_paper,
    load_plan,
    load_planning_workbook,
    parse_paper_numbers,
    validate_paper_index,
)
from .postprocess import _post_process
from .prompts import build_generate_full_paper_prompt
from .regenerator import fix_answer_distribution, needs_regeneration, regenerate_question, normalize_all_judge_answers
from 质检.report import build_quality_report, write_markdown_report
from 质检.rules import run_quality_checks, make_issue, _is_duplicate_stem_pair, WARNING

@dataclass
class PreflightResult:
    selected_numbers: list[int]
    selected_papers: list[PaperPlan]
    missing_selected_numbers: list[int]
    missing_sequence_numbers: list[int]
    manual_docx_missing: list[int]
    blueprint_warnings: list[str]
    mapping_missing: list[int]
    config_warnings: list[str]
    mapping_loaded: bool = False
    mapping_table_missing: bool = False
    mapping_table_path: str = ""


@dataclass
class PaperRunResult:
    paper_label: str
    paper_type: str
    status: str = "pending"
    question_count: int = 0
    repaired_count: int = 0
    repair_rounds: int = 0
    manual_review_count: int = 0
    remaining_issue_count: int = 0
    paper_issues: list = field(default=None)  # 全卷级别质检问题列表
    cross_paper_issues: int = 0
    cross_paper_repeat_rate: float = 0.0
    questions_path: str = ""
    repaired_questions_path: str = ""
    qc_report_path: str = ""
    final_qc_report_path: str = ""
    text_path: str = ""
    docx_path: str = ""
    output_dir: str = ""
    error: str = ""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="考纲百套卷生成管线")
    parser.add_argument(
        "--plan",
        "-p",
        default=None,
        help="规划表 xlsx 路径；默认使用 04_生成输出/生产规划 下的规划表",
    )
    parser.add_argument(
        "--paper",
        "--sequence",
        "-n",
        dest="paper",
        default=None,
        help="要生成/查看的卷号，如 34、1-10、1,3,5；不传则进入交互选择",
    )
    parser.add_argument(
        "--range",
        "-r",
        dest="paper_range",
        default=None,
        help="兼容旧参数：卷号范围，如 all、1-10、1,3,5",
    )
    parser.add_argument("--list", action="store_true", help="列出规划表中的全部卷号和卷型后退出")
    parser.add_argument("--max-fix-rounds", type=int, default=2, help="不合格题最大修复轮数；0 表示不调用模型修复")
    parser.add_argument("--max-fix-workers", type=int, default=6, help="并发修复线程数；1 表示串行，建议 3-8（取决于 API 限流）")
    parser.add_argument(
        "--preview",
        "--dry-run",
        dest="preview",
        action="store_true",
        help="只解析规划表并展示任务，不执行拆题/质检处理",
    )
    parser.add_argument(
        "--generate",
        action="store_true",
        help="兼容旧参数；现在默认即执行拆题/质检处理",
    )
    parser.add_argument("--status", action="store_true", help="查看当前运行状态后退出")
    parser.add_argument("--rerun-failed", action="store_true", help="自动重跑失败卷后退出")
    parser.add_argument("--rerun-missing", action="store_true", help="自动重跑缺失卷后退出")
    return parser.parse_args()


def find_default_plan(plan_dir: Path = PLAN_DIR) -> Path:
    candidates = [
        path
        for path in plan_dir.rglob("*.xlsx")
        if path.is_file()
        and not path.name.startswith("~$")
        and "考点规划总表" in path.name
    ]
    if not candidates:
        raise FileNotFoundError(f"未在 {plan_dir} 及其子目录找到 xlsx 规划表")

    candidates.sort(key=lambda path: str(path.relative_to(plan_dir)))
    print("\n可用的生产规划表：")
    for idx, path in enumerate(candidates, 1):
        print(f"  {idx}. {path.relative_to(plan_dir)}")

    default_choice = "1" if len(candidates) == 1 else ""
    prompt = "\n请输入编号选择规划表"
    if default_choice:
        prompt += "（默认 1）"
    prompt += ": "
    choice = input(prompt).strip() or default_choice

    try:
        return candidates[int(choice) - 1]
    except (ValueError, IndexError):
        raise SystemExit("无效选择")


def parse_sequence_selector(selector: str | None, available: Iterable[int]) -> list[int]:
    if selector is None:
        return []
    text = str(selector).strip()
    if not text:
        return []
    if text.lower() == "all":
        return sorted(set(available))
    return parse_paper_numbers(text)


def prompt_for_sequence(paper_index: dict[int, PaperPlan]) -> list[int]:
    raw = input("请输入要生成/查看的卷号（如 34、1-3、all，直接回车退出）：").strip()
    if not raw:
        return []
    return parse_sequence_selector(raw, paper_index.keys())


def _print_generation_status(paper_index: dict[int, PaperPlan], meta) -> None:
    """打印已有生成文件的状态，帮助用户判断哪些卷还需要处理。"""
    from ._topic_pool import scan_generated_papers

    generated = scan_generated_papers(meta)
    all_labels = {paper.paper_label for paper in paper_index.values()}
    pending_labels = all_labels - generated

    if generated:
        print(f"\n已生成：{len(generated)} 卷")
    if pending_labels:
        pending_numbers = sorted(
            paper.paper_no for paper in paper_index.values()
            if paper.paper_label in pending_labels
        )
        # 找连续区间，简洁输出
        ranges = _compact_number_ranges(pending_numbers)
        print(f"待生成：{len(pending_labels)} 卷 → {ranges}")
    if not pending_labels:
        print("全部已生成，无需再次拉题。")


def _compact_number_ranges(numbers: list[int]) -> str:
    """将 [1,2,3,5,7,8,9] 格式化为 "1-3, 5, 7-9"。"""
    if not numbers:
        return "无"
    parts = []
    start = numbers[0]
    end = numbers[0]
    for n in numbers[1:]:
        if n == end + 1:
            end = n
        else:
            parts.append(f"{start}-{end}" if end > start else str(start))
            start = end = n
    parts.append(f"{start}-{end}" if end > start else str(start))
    return ", ".join(parts)


def _format_type_counts(paper_index: dict[int, PaperPlan]) -> str:
    counts = Counter(paper.paper_type for paper in paper_index.values())
    return "，".join(
        f"{paper_type} {counts.get(paper_type, 0)} 卷"
        for paper_type in [POINT_PAPER_TYPE, TOPIC_PAPER_TYPE, COURSE_PAPER_TYPE]
    )


def print_plan_summary(meta, rows, paper_index: dict[int, PaperPlan]) -> None:
    print(f"已读取规划表：{meta.path}")
    if meta.title:
        print(f"标题：{meta.title}")
    if meta.province or meta.exam_category:
        print(f"地区/考类：{meta.province} {meta.exam_category}".strip())
    print(f"规划考点行数：{len(rows)}")
    print(f"可选试卷数：{len(paper_index)}（{_format_type_counts(paper_index)}）")
    if paper_index:
        numbers = sorted(paper_index)
        print(f"卷号范围：{format_paper_label(numbers[0])} - {format_paper_label(numbers[-1])}")


def print_dry_run_summary(plan_path: Path) -> None:
    items = load_plan(plan_path)
    grouped = group_by_paper(items)
    preview = "、".join(list(grouped)[:10])
    print(f"已读取规划表：{plan_path}")
    print(f"规划任务数：{len(items)}")
    print(f"卷号数：{len(grouped)}")
    if preview:
        print(f"卷号预览：{preview}")
    print("规划表 dry-run 通过。")


def _paper_subject(paper: PaperPlan) -> str:
    if paper.paper_type == POINT_PAPER_TYPE:
        return paper.point_name
    if paper.paper_type == TOPIC_PAPER_TYPE:
        return paper.topic
    return "全课程综合"


def print_paper_list(paper_index: dict[int, PaperPlan]) -> None:
    print("\n卷号    类型        知识模块              专题/考点")
    print("-" * 72)
    for paper in paper_index.values():
        module = (paper.module[:18] + "…") if len(paper.module) > 18 else paper.module
        subject = _paper_subject(paper)
        print(f"{paper.paper_label:<7} {paper.paper_type:<10} {module:<20} {subject}")


def _row_range(rows) -> str:
    if not rows:
        return ""
    row_numbers = sorted(row.row_no for row in rows)
    if row_numbers[0] == row_numbers[-1]:
        return f"第{row_numbers[0]}行"
    return f"第{row_numbers[0]}-{row_numbers[-1]}行"


def _unique(values: Iterable[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        if value and value not in seen:
            result.append(value)
            seen.add(value)
    return result


def _print_blueprint_info(paper: PaperPlan) -> None:
    if paper.paper_type not in {TOPIC_PAPER_TYPE, COURSE_PAPER_TYPE}:
        return

    if paper.blueprint_path:
        print(f"细目表：{paper.blueprint_path}")
        print(f"细目表题数：{len(paper.blueprint_rows)}")
        if paper.blueprint_rows:
            print("细目表预览：")
            for row in paper.blueprint_rows[:5]:
                question_no = row.question_no_raw or (str(row.question_no) if row.question_no is not None else "")
                focus = row.point_name or row.knowledge_point
                content = row.content or row.knowledge_point
                print(f"  {question_no}｜{row.question_type}｜{row.difficulty}｜{focus}｜{content}")
            if len(paper.blueprint_rows) > 5:
                print(f"  ... 其余 {len(paper.blueprint_rows) - 5} 题略")
    else:
        print("细目表：未找到（生成专题训练卷/课程综合卷时必须提供）")
        print(f"期望文件：{blueprint_filename_for_paper(paper)}")
    for warning in paper.blueprint_warnings:
        print(f"细目表提示：{warning}")


def display_selected_papers(selected_papers: list[PaperPlan]) -> None:
    for paper in selected_papers:
        print("\n" + "=" * 72)
        print(f"{paper.paper_label}｜{paper.paper_type}")
        print(f"知识模块：{paper.module}")

        if paper.paper_type == POINT_PAPER_TYPE:
            print(f"专题：{paper.topic}")
            print(f"考点：{paper.point_name}")
            if paper.point_content:
                print("考点内容：")
                for line in paper.point_content.splitlines():
                    print(f"  {line}")
            print(f"规划行：{_row_range(paper.rows)}")
            print("后续动作：将围绕该考点生成单卷训练题。")

        elif paper.paper_type == TOPIC_PAPER_TYPE:
            point_names = _unique(row.point_name for row in paper.rows)
            print(f"专题：{paper.topic}")
            print(f"覆盖考点数：{len(point_names)}")
            print("覆盖考点：")
            for point_name in point_names:
                print(f"  - {point_name}")
            print(f"规划行：{_row_range(paper.rows)}")
            print("后续动作：将覆盖该专题下全部考点。")
            _print_blueprint_info(paper)

        elif paper.paper_type == COURSE_PAPER_TYPE:
            topics = _unique(row.topic for row in paper.rows)
            point_names = _unique(row.point_name for row in paper.rows)
            print(f"课程综合卷组：{paper.paper_ref}")
            print(f"覆盖专题数：{len(topics)}")
            print(f"覆盖考点数：{len(point_names)}")
            if topics:
                print("覆盖专题：")
                for topic in topics:
                    print(f"  - {topic}")
            print(f"规划行：{_row_range(paper.rows)}")
            print("后续动作：该课程综合卷围绕本课程全部考点生成综合训练题。")
            _print_blueprint_info(paper)


def resolve_selected_papers(
    selected_numbers: list[int], paper_index: dict[int, PaperPlan]
) -> tuple[list[PaperPlan], list[int]]:
    selected: list[PaperPlan] = []
    missing: list[int] = []
    for number in selected_numbers:
        paper = paper_index.get(number)
        if paper is None:
            missing.append(number)
        else:
            selected.append(paper)
    return selected, missing


def _format_paper_labels(numbers: Iterable[int], limit: int | None = None) -> str:
    unique_numbers = sorted(set(numbers))
    if limit is not None and len(unique_numbers) > limit:
        shown = unique_numbers[:limit]
        return "、".join(format_paper_label(number) for number in shown) + f" 等 {len(unique_numbers)} 卷"
    return "、".join(format_paper_label(number) for number in unique_numbers) or "无"


def _missing_sequence_numbers(paper_index: dict[int, PaperPlan]) -> list[int]:
    if not paper_index:
        return []
    numbers = sorted(paper_index)
    return [number for number in range(numbers[0], numbers[-1] + 1) if number not in paper_index]


def _extract_paper_no_from_filename(path: Path) -> int | None:
    match = re.search(r"第\s*(\d+)\s*卷", path.stem)
    return int(match.group(1)) if match else None


def _is_generated_output_file(path: Path) -> bool:
    """判断生成结果中的文件是否可视为某卷已经产出。"""
    if path.name.startswith("~$"):
        return False
    return path.suffix.lower() in {".docx", ".zip"}


def _final_output_search_dir(meta) -> Path:
    province = _safe_output_name(getattr(meta, "province", "") or "")
    category = _safe_output_name(getattr(meta, "exam_category", "") or "")
    if province and category:
        return FINAL_OUTPUT_DIR / _safe_output_name(f"{province} {category}")
    if province:
        return FINAL_OUTPUT_DIR / province
    return FINAL_OUTPUT_DIR


def _missing_output_numbers(meta, paper_index: dict[int, PaperPlan]) -> list[int]:
    """返回规划表中有、但生成结果里尚未产出的卷号。"""
    expected = set(paper_index)
    if not expected:
        return []

    output_dir = _final_output_search_dir(meta)
    generated: set[int] = set()
    if output_dir.exists():
        for path in output_dir.rglob("*"):
            if not path.is_file() or not _is_generated_output_file(path):
                continue
            paper_no = _extract_paper_no_from_filename(path)
            if paper_no in expected:
                generated.add(paper_no)

    return sorted(expected - generated)


def _check_manual_docx(selected_papers: list[PaperPlan]) -> list[int]:
    missing: list[int] = []
    for paper in selected_papers:
        search_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
        matched = find_manual_paper_for_plan(paper, search_dir, supported_suffixes={".docx"})
        if matched is None:
            missing.append(paper.paper_no)
    return missing


def _collect_blueprint_warnings(selected_papers: list[PaperPlan]) -> list[str]:
    warnings: list[str] = []
    attach_warnings = attach_blueprints_for_papers(selected_papers)
    warnings.extend(attach_warnings)
    for paper in selected_papers:
        if paper.paper_type not in {TOPIC_PAPER_TYPE, COURSE_PAPER_TYPE}:
            continue
        for warning in paper.blueprint_warnings:
            item = f"{paper.paper_label}：{warning}"
            if item not in warnings:
                warnings.append(item)
    return warnings


def _mapping_table_path(meta) -> Path:
    base = Path(__file__).resolve().parent.parent.parent
    return base / "04_生成输出" / "生产规划" / f"{meta.province} {meta.exam_category}" / f"{meta.province}_{meta.exam_category}_映射表.xlsx"


def _check_mapping_coverage(meta, paper_index: dict[int, PaperPlan]) -> tuple[list[int], bool]:
    """检查映射表覆盖率：只报告A列完全不存在的卷号，AI生成/聚合的不算缺失。"""
    try:
        import openpyxl

        xlsx_path = _mapping_table_path(meta)
        if not xlsx_path.exists():
            return [], False

        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        ws = wb["知识点映射"]
        # 收集 A 列全部已登记的试卷序号
        listed_labels: set[str] = set()
        for row in ws.iter_rows(min_row=2, values_only=True):
            key = str(row[0]).strip() if row[0] else ""
            # A 列可能是 "第34-36卷" 的区间格式，需展开
            if "-" in key and any(c.isdigit() for c in key):
                import re
                m = re.search(r"(\d+)\s*-\s*(\d+)", key)
                if m:
                    listed_labels.update(f"第{i}卷" for i in range(int(m.group(1)), int(m.group(2)) + 1))
            elif key and not key.startswith("#"):
                listed_labels.add(key)

        truly_missing = [
            paper.paper_no for paper in paper_index.values()
            if paper.paper_label not in listed_labels
        ]
        return truly_missing, True
    except Exception:
        return [], False


def _check_config_availability() -> list[str]:
    warnings: list[str] = []
    if not CONFIG_PATH.exists():
        return [f"配置文件不存在：{CONFIG_PATH}"]
    try:
        config = load_config()
    except Exception as exc:
        return [f"配置文件读取失败：{exc}"]

    if not config.get("api_key"):
        warnings.append("缺少模型 API key：config.json api_key")
    if not config.get("model"):
        warnings.append("缺少模型名称：config.json model")
    if not (os.environ.get("XKW_COOKIE") or config.get("xkw_cookie")):
        warnings.append("缺少学科网 Cookie：环境变量 XKW_COOKIE 或 config.json xkw_cookie")
    return warnings


def build_preflight_result(
    meta,
    paper_index: dict[int, PaperPlan],
    selected_numbers: list[int],
    selected_papers: list[PaperPlan],
    missing_selected_numbers: list[int],
) -> PreflightResult:
    blueprint_warnings = _collect_blueprint_warnings(selected_papers)
    mapping_missing, mapping_loaded = _check_mapping_coverage(meta, paper_index)
    mapping_table_path = str(_mapping_table_path(meta))

    return PreflightResult(
        selected_numbers=selected_numbers,
        selected_papers=selected_papers,
        missing_selected_numbers=missing_selected_numbers,
        missing_sequence_numbers=_missing_output_numbers(meta, paper_index),
        manual_docx_missing=_check_manual_docx(selected_papers),
        blueprint_warnings=blueprint_warnings,
        mapping_missing=mapping_missing,
        config_warnings=_check_config_availability(),
        mapping_loaded=mapping_loaded,
        mapping_table_missing=not mapping_loaded,
        mapping_table_path=mapping_table_path,
    )


def print_preflight_dialog(result: PreflightResult, meta, paper_index: dict[int, PaperPlan]) -> None:
    selected_count = len(result.selected_papers)
    blueprint_papers = [
        paper for paper in result.selected_papers
        if paper.paper_type in {TOPIC_PAPER_TYPE, COURSE_PAPER_TYPE}
    ]

    print("\n" + "=" * 60)
    print("生成前检查")
    print("=" * 60)
    print(f"地区/考类：{meta.province} {meta.exam_category}".strip())
    print(f"本次选择：{_format_paper_labels(result.selected_numbers, limit=30)}")
    missing_range = _compact_number_ranges(sorted(result.missing_sequence_numbers))
    print(f"当前考类缺失卷号：{missing_range}（共 {len(result.missing_sequence_numbers)} 卷）")
    print("\n检查项：")

    if result.manual_docx_missing:
        print(f"  [异常] 人工组卷 DOCX：缺失 {_format_paper_labels(result.manual_docx_missing, limit=30)}")
    else:
        print(f"  [通过] 人工组卷 DOCX：已匹配 {selected_count}/{selected_count}")

    if result.blueprint_warnings:
        print(f"  [异常] 细目表：{len(result.blueprint_warnings)} 项缺失或异常")
    elif blueprint_papers:
        print(f"  [通过] 细目表：已满足 {len(blueprint_papers)}/{len(blueprint_papers)}")
    else:
        print("  [通过] 细目表：不适用或已满足")

    total_papers = len(paper_index)
    if result.mapping_missing:
        prefix = "缺失或未覆盖" if not result.mapping_loaded else "未覆盖"
        print(f"  [异常] 映射表：{prefix} {_format_paper_labels(result.mapping_missing, limit=30)}")
    else:
        print(f"  [通过] 映射表：已覆盖全部 {total_papers} 卷")

    if result.config_warnings:
        print(f"  [异常] 配置/API：{len(result.config_warnings)} 项缺失或异常")
    else:
        print("  [通过] 配置/API：api_key、model、XKW cookie 可用")

    if result.mapping_table_missing:
        print(f"  [异常] 正式映射表：缺失 {result.mapping_table_path}")
    else:
        print(f"  [通过] 正式映射表：{result.mapping_table_path}")


    details: list[str] = []
    if result.missing_selected_numbers:
        details.append("本次选择中规划表不存在：" + _format_paper_labels(result.missing_selected_numbers, limit=30))
    if result.manual_docx_missing:
        details.append("人工组卷 DOCX 缺失：" + _format_paper_labels(result.manual_docx_missing, limit=30))
    details.extend(f"细目表：{warning}" for warning in result.blueprint_warnings)
    if result.mapping_missing:
        details.append("映射表未覆盖：" + _format_paper_labels(result.mapping_missing, limit=30))
    details.extend(f"配置/API：{warning}" for warning in result.config_warnings)
    if result.mapping_table_missing:
        details.append(f"正式映射表缺失：{result.mapping_table_path}")

    if details:
        print("\n问题明细：")
        for detail in details:
            print(f"  - {detail}")
    else:
        print("\n问题明细：无")
    print("=" * 60)


def _preflight_has_blocking_issues(result: PreflightResult, mode: int) -> bool:
    return bool(
        result.missing_selected_numbers
        or result.blueprint_warnings
        or result.mapping_missing
        or result.config_warnings
        or (mode == 2 and result.mapping_table_missing)
    )


def _confirm_continue_after_preflight() -> bool:
    answer = input("生成前检查存在异常，是否继续？[y/N]: ").strip().lower()
    return answer in {"y", "yes", "是"}


def _safe_output_name(text: str) -> str:
    return "".join("_" if char in '<>:"/\\|?*' else char for char in text).strip() or "未命名"


def _run_record_path(paper: PaperPlan) -> Path:
    meta = paper.meta
    province_category = " ".join(
        part for part in [getattr(meta, "province", "") if meta else "", getattr(meta, "exam_category", "") if meta else ""]
        if part
    ).strip()
    output_dir = RUN_RECORD_DIR / _safe_output_name(province_category or "未命名考类")
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{paper.paper_label}_{_safe_output_name(paper.paper_type)}_运行记录.json"


def _read_run_record(paper: PaperPlan) -> dict[str, Any]:
    path = _run_record_path(paper)
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _plain_status(status: str) -> str:
    return {
        "success": "已完成，可以交付",
        "manual_review": "已生成解析版，但还需要人工看一眼",
        "failed": "处理失败，需要重新处理或人工排查",
        "pending": "处理被中断或等待后续处理",
        "qc_passed": "质检已通过，等待生成最终文件",
        "missing": "还没有找到产物",
    }.get(status or "", status or "未知")


def _update_run_record(paper: PaperPlan, **sections: Any) -> Path:
    """写出给人工看的单卷运行记录。字段尽量用大白话，方便无代码基础的人回溯。"""
    meta = paper.meta
    record = _read_run_record(paper)
    record.setdefault("这是什么", "这是一份单卷运行记录，用来说明本卷拉题、质检、修复和输出文件的情况。")
    record["试卷信息"] = {
        "卷号": paper.paper_label,
        "卷型": paper.paper_type,
        "课程": paper.module,
        "专题": paper.topic,
        "考点": paper.point_name,
        "地区": getattr(meta, "province", "") if meta else "",
        "考类": getattr(meta, "exam_category", "") if meta else "",
        "规划表": str(getattr(meta, "path", "") or "") if meta else "",
    }
    for key, value in sections.items():
        if value is not None:
            record[key] = value
    record["最后更新时间"] = time.strftime("%Y-%m-%d %H:%M:%S")
    path = _run_record_path(paper)
    path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _paper_output_base(paper: PaperPlan) -> Path | None:
    if paper.meta is None:
        return None
    return _get_topic_output_base(paper.meta, _topic_context(paper), FINAL_OUTPUT_DIR)


def _paper_output_candidates(paper: PaperPlan) -> list[Path]:
    base = _paper_output_base(paper)
    if base is None or not base.exists():
        return []
    candidates: list[Path] = []
    for path in base.rglob("*"):
        if not path.is_file():
            continue
        if paper.paper_label not in path.name:
            continue
        if path.suffix.lower() in {".docx", ".zip", ".json", ".md"}:
            candidates.append(path)
    return candidates


def _final_output_summary(paper: PaperPlan) -> dict[str, Any]:
    candidates = _paper_output_candidates(paper)
    analysis_docx = [path for path in candidates if path.suffix.lower() == ".docx" and "解析版" in path.name]
    original_docx = [path for path in candidates if path.suffix.lower() == ".docx" and "原卷版" in path.name]
    zip_files = [path for path in candidates if path.suffix.lower() == ".zip"]
    return {
        "解析版Word": str(analysis_docx[0]) if analysis_docx else "",
        "原卷版Word": str(original_docx[0]) if original_docx else "",
        "压缩包": str(zip_files[0]) if zip_files else "",
        "本卷相关产物数量": len(candidates),
    }


def _infer_paper_status(paper: PaperPlan) -> str:
    candidates = _paper_output_candidates(paper)
    if not candidates:
        return "missing"

    if any("待人工审核" in path.name for path in candidates if path.suffix.lower() in {".docx", ".zip"}):
        return "manual_review"

    has_final_docx = any(path.suffix.lower() == ".docx" and "待人工审核" not in path.name for path in candidates)
    has_final_zip = any(path.suffix.lower() == ".zip" for path in candidates)
    if has_final_docx or has_final_zip:
        return "success"

    if any(path.suffix.lower() in {".json", ".md"} for path in candidates):
        return "failed"

    return "missing"


def _print_run_status(meta, paper_index: dict[int, PaperPlan]) -> None:
    counts: Counter[str] = Counter()
    rows: list[tuple[str, str, str]] = []
    for paper in paper_index.values():
        status = _infer_paper_status(paper)
        counts[status] += 1
        rows.append((paper.paper_label, paper.paper_type, status))

    print("\n" + "=" * 60)
    print("运行状态")
    print("=" * 60)
    print(f"地区/考类：{meta.province} {meta.exam_category}".strip())
    print(f"规划表：{getattr(meta, 'path', '')}")
    print(f"总卷数：{len(rows)}")
    print(f"成功：{counts.get('success', 0)}，待人工审核：{counts.get('manual_review', 0)}，失败：{counts.get('failed', 0)}，未运行：{counts.get('missing', 0)}")

    if any(status != "success" for _, _, status in rows):
        print("\n非成功卷：")
        for label, paper_type, status in rows:
            if status == "success":
                continue
            print(f"  - {label}｜{paper_type}｜{status}")
    print("=" * 60)


def _select_batch_papers(paper_index: dict[int, PaperPlan], *, rerun_failed: bool = False, rerun_missing: bool = False) -> list[PaperPlan]:
    selected: list[PaperPlan] = []
    for paper in paper_index.values():
        status = _infer_paper_status(paper)
        if rerun_failed and status == "failed":
            selected.append(paper)
        elif rerun_missing and status == "missing":
            selected.append(paper)
    return selected


def _province_category_parts(paper: PaperPlan) -> list[str]:
    meta = paper.meta
    if not meta:
        return []
    province_category = " ".join(part for part in [meta.province, meta.exam_category] if part).strip()
    return [_safe_output_name(province_category)] if province_category else []


def _manual_questions_output_path(paper: PaperPlan) -> Path:
    output_dir = CLEAN_OUTPUT_DIR.joinpath(*_province_category_parts(paper), "组卷拆题")
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{paper.paper_label}_{_safe_output_name(paper.paper_type)}_questions.json"


def _ai_generated_paper_output_path(paper: PaperPlan) -> Path:
    output_dir = CLEAN_OUTPUT_DIR.joinpath(*_province_category_parts(paper), "AI直接生成")
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{paper.paper_label}_{_safe_output_name(paper.paper_type)}_generated.json"


def _repaired_questions_output_path(paper: PaperPlan) -> Path:
    output_dir = CLEAN_OUTPUT_DIR.joinpath(*_province_category_parts(paper), "修复结果")
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / f"{paper.paper_label}_{_safe_output_name(paper.paper_type)}_questions_repaired.json"


def _api_trace_dir(paper: PaperPlan) -> Path:
    output_dir = API_RAW_OUTPUT_DIR.joinpath(*_province_category_parts(paper), _safe_output_name(paper.paper_label))
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def _write_api_trace_json(paper: PaperPlan, filename: str, data: dict[str, Any]) -> Path:
    path = _api_trace_dir(paper) / filename
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    return path


def _count_questions_by_type_and_source(questions: list[dict]) -> dict[str, dict[str, int]]:
    counts: dict[str, dict[str, int]] = {}
    for question in questions:
        qtype = _question_type_of(question) or "未识别题型"
        source = "AI生成" if question.get("_ai_generated") else "API"
        counts.setdefault(qtype, {"API": 0, "AI生成": 0})
        counts[qtype][source] = counts[qtype].get(source, 0) + 1
    return counts


def _cross_paper_check_structured(
    questions_a: list[dict[str, Any]],
    questions_b: list[dict[str, Any]],
    label_a: str = "卷A",
    label_b: str = "卷B",
) -> tuple[list[dict[str, str]], dict[str, Any]]:
    """对两组结构化题目执行跨卷查重，复用 rules.py 的题干相似度引擎。"""
    issues: list[dict[str, str]] = []
    if not questions_a or not questions_b:
        stats = {
            "paper_a_label": label_a,
            "paper_b_label": label_b,
            "paper_a_questions": len(questions_a),
            "paper_b_questions": len(questions_b),
            "high_similarity_pairs": 0,
            "repeat_rate": 0.0,
            "similar_pairs": [],
        }
        return issues, stats

    from 质检.rules import _is_duplicate_stem_pair

    similar_pairs: list[dict[str, Any]] = []
    for idx_a, qa in enumerate(questions_a):
        stem_a = qa.get("stem", "")
        for idx_b, qb in enumerate(questions_b):
            stem_b = qb.get("stem", "")
            duplicate, score, reason = _is_duplicate_stem_pair(stem_a, stem_b)
            if duplicate:
                similar_pairs.append({
                    "paper_a_question": qa.get("question_no") or idx_a + 1,
                    "paper_b_question": qb.get("question_no") or idx_b + 1,
                    "similarity": round(score, 3),
                    "match_types": ["stem_similarity"],
                    "reason": reason,
                })
                break  # 每题只算一次

    total = min(len(questions_a), len(questions_b)) or 1
    repeat_rate = len(similar_pairs) / total

    stats = {
        "paper_a_label": label_a,
        "paper_b_label": label_b,
        "paper_a_questions": len(questions_a),
        "paper_b_questions": len(questions_b),
        "high_similarity_pairs": len(similar_pairs),
        "repeat_rate": round(repeat_rate, 3),
        "similar_pairs": similar_pairs,
    }

    if repeat_rate > 0.2:
        issues.append(
            make_issue(
                "cross_paper_high_repeat",
                "跨卷重复率过高",
                "failed",
                f"{label_a} 与 {label_b} 之间相似题目 {len(similar_pairs)} 对，"
                f"重复率={repeat_rate:.0%}（超过 20% 限制）",
            )
        )
    elif repeat_rate > 0.1:
        issues.append(
            make_issue(
                "cross_paper_moderate_repeat",
                "跨卷重复率偏高",
                "warning",
                f"{label_a} 与 {label_b} 之间相似题目 {len(similar_pairs)} 对，"
                f"重复率={repeat_rate:.0%}",
            )
        )

    return issues, stats


def _partition_papers(papers: list[PaperPlan]) -> tuple[list[PaperPlan], dict[str, list[PaperPlan]], dict[str, list[PaperPlan]]]:
    """将试卷列表拆分为 独立卷(专题训练卷)、考点训练卷组(按module+topic) 和 课程综合卷组。"""
    singles: list[PaperPlan] = []
    topic_groups: dict[str, list[PaperPlan]] = defaultdict(list)
    course_groups: dict[str, list[PaperPlan]] = defaultdict(list)
    for paper in papers:
        if paper.paper_type == COURSE_PAPER_TYPE:
            course_groups[paper.paper_ref or paper.paper_label].append(paper)
        elif paper.paper_type == POINT_PAPER_TYPE and paper.topic:
            # 考点训练卷按 (module, topic) 分组
            group_key = f"{paper.module}|{paper.topic}"
            topic_groups[group_key].append(paper)
        else:
            singles.append(paper)
    return singles, dict(topic_groups), dict(course_groups)


def _save_loaded_paper_questions(paper: PaperPlan, loaded_paper, output_path: Path | None = None) -> Path:
    target = output_path or _manual_questions_output_path(paper)
    data = {
        "paper": {
            "paper_no": paper.paper_no,
            "paper_label": paper.paper_label,
            "paper_type": paper.paper_type,
            "module": paper.module,
            "topic": paper.topic,
            "point_name": paper.point_name,
        },
        "loaded_paper": loaded_paper_to_dict(loaded_paper),
    }
    target.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


def _print_loaded_paper_summary(loaded_paper) -> None:
    print(f"人工组卷文件：{loaded_paper.path}")
    print(f"解析题数：{len(loaded_paper.questions)}")
    type_counts = Counter(question.get("question_type") or "未识别题型" for question in loaded_paper.questions)
    if type_counts:
        print("题型统计：" + "，".join(f"{qtype} {count}题" for qtype, count in type_counts.items()))
    for warning in loaded_paper.warnings:
        print(f"拆题提示：{warning}")


def _renumber_questions(questions: list[dict]) -> None:
    for idx, question in enumerate(questions, 1):
        question["question_no"] = idx
        question["_question_no"] = idx
        qtype = _question_type_of(question)
        if qtype:
            question["question_type"] = qtype
            question["_question_type"] = qtype



def _repaired_to_manual_txt(paper, questions):
    """修复完成后，将题目文本回写到组卷待质检。"""
    try:
        from .paper_assembler import assemble_analysis_paper_text
        manual_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
        manual_dir.mkdir(parents=True, exist_ok=True)
        txt_path = manual_dir / f"{paper.paper_label}_repaired.txt"
        txt_path.write_text(assemble_analysis_paper_text(questions, paper), encoding="utf-8")
        print(f"  已回写修复结果到：{txt_path}")
    except Exception as e:
        print(f"  回写修复失败（不影响继续处理）：{e}")


def _save_complete_docx_to_manual_dir(paper, questions):
    """AI 补题后，将完整题目回写到 组卷待质检 的原始 DOCX 文件。
    覆盖旧的残缺文件，确保下次加载时不再显示题量不足。
    """
    try:
        from docx import Document
        from docx.shared import Pt

        manual_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
        manual_dir.mkdir(parents=True, exist_ok=True)
        # 构建与 load_manual_paper_for_plan 同名的文件
        safe_module = "".join(c if c.isalnum() or c in " _-" else "_" for c in (paper.module or "未知课程"))
        safe_point = "".join(c if c.isalnum() or c in " _-" else "_" for c in (paper.point_name or paper.topic or "未知"))
        filename = f"{paper.paper_label} {safe_module} {safe_point}.docx"[:200]
        filepath = manual_dir / filename

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)

        # 按题型分组，简单输出题干和答案
        type_order = ["单选题", "多选题", "判断题", "填空题", "简答题", "计算题", "综合题"]
        qno = 0
        for qtype in type_order:
            typed = [q for q in questions if (q.get("question_type") or q.get("_question_type") or "") == qtype]
            if not typed:
                continue
            doc.add_paragraph(qtype)
            for q in typed:
                qno += 1
                stem = q.get("stem", "") or ""
                doc.add_paragraph(f"{qno}. {stem}")
                options = q.get("options") or []
                if isinstance(options, list):
                    for opt in options:
                        doc.add_paragraph(f"    {opt}")
                ans = q.get("answer", "")
                doc.add_paragraph(f"【答案】{ans}")
                analysis = q.get("analysis", "") or q.get("explanation", "")
                if analysis:
                    doc.add_paragraph(f"【解析】{analysis}")
                doc.add_paragraph("")

        doc.save(str(filepath))
        print(f"  已同步更新组卷待质检文件：{filepath}")
    except Exception as e:
        print(f"  组卷待质检文件更新失败（不影响继续处理）：{e}")


def _trim_questions_to_plan(questions: list[dict], type_plans: dict[str, list[dict]]) -> list[dict]:
    """题量过多时按题型规划随机裁剪。"""
    import random
    if not type_plans:
        return questions
    selected: list[dict] = []
    used_ids: set[int] = set()
    for qtype, plans in type_plans.items():
        typed = [q for q in questions if _question_type_of(q) == qtype]
        picked = _select_candidates_with_constraints(typed, len(plans))
        if len(picked) < len(plans):
            rest = [q for q in typed if id(q) not in {id(x) for x in picked}]
            random.shuffle(rest)
            picked.extend(rest[: len(plans) - len(picked)])
        for q in picked[:len(plans)]:
            selected.append(q)
            used_ids.add(id(q))
    remaining_needed = _planned_total(type_plans) - len(selected)
    if remaining_needed > 0:
        rest = [q for q in questions if id(q) not in used_ids]
        random.shuffle(rest)
        selected.extend(rest[:remaining_needed])
    _renumber_questions(selected)
    return selected


def _supplement_loaded_paper_with_ai(paper: PaperPlan, loaded_paper, type_plans: dict[str, list[dict]]) -> int:
    """已有原文档题量不足时，只用 AI 按规划缺口补齐；低于 40% 时补到半量。"""
    planned = _planned_total(type_plans)
    current = len(loaded_paper.questions)
    if planned <= 0 or current >= planned:
        return 0
    target_total, score_multiplier = _decide_ai_target(planned, current)
    target_plans = _limit_type_plans(type_plans, target_total)
    _trim_overflow_simple(target_plans, loaded_paper.questions)
    shortfall = _compute_shortfall(target_plans, loaded_paper.questions)
    if not shortfall:
        return 0
    try:
        from 学科网API拉题移植版.kpoint_resolver import load_mapping_table, resolve_course
        mapping = load_mapping_table(paper.meta.province, paper.meta.exam_category) if paper.meta else {}
        kpoint_ids = mapping.get(paper.paper_label, [])
        course_id = resolve_course(paper.module) or 0
    except Exception:
        kpoint_ids = []
        course_id = 0
    if score_multiplier > 1:
        loaded_paper.warnings.append(f"题源不足 40%，AI 补至规划题量一半（{target_total}/{planned}），后续 DOCX 应按每题分值×2 处理。")
        print(f"  题源不足 40%，AI 只补至 {target_total} 题；请按每题分值×2 处理。")
    paper._score_multiplier = score_multiplier
    ai_questions = _ai_fill_shortfall(paper, shortfall, kpoint_ids, course_id)
    added = 0
    for item in ai_questions:
        qtype = _question_type_of(item) or item.get("question_type") or ""
        loaded_paper.questions.append({
            "question_no": len(loaded_paper.questions) + 1,
            "question_type": qtype,
            "heading": qtype,
            "stem": item.get("stem", ""),
            "options": item.get("options") or [],
            "answer": item.get("answer", ""),
            "analysis": item.get("analysis") or item.get("explanation", ""),
            "knowledge_points": [],
            "difficulty": item.get("difficulty") or item.get("_target_difficulty") or "",
            "source_path": "AI补题",
            "status": "ai_generated",
            "issues": [],
            "raw_text": item.get("stem", ""),
            "_ai_generated": True,
            "_question_type": qtype,
        })
        added += 1
    if len(loaded_paper.questions) > target_total:
        loaded_paper.questions = _trim_questions_to_plan(loaded_paper.questions, target_plans)
    else:
        _renumber_questions(loaded_paper.questions)
    _update_run_record(
        paper,
        **{
            "AI补题记录": {
                "一句话说明": "原材料题量不够，系统按规划表用 AI 补了一部分题。",
                "规划应有题数": planned,
                "补题前已有题数": current,
                "本次补了几题": added,
                "补题后题数": len(loaded_paper.questions),
                "是否触发半量规则": score_multiplier > 1,
                "半量规则说明": "题源命中率低于40%，只补到规划题量一半，后续每题分值加倍。" if score_multiplier > 1 else "未触发。",
            }
        },
    )
    return added


def _qa_report_output_path(paper: PaperPlan, stage: str = "修复前") -> Path:
    parts = list(_province_category_parts(paper))
    if paper.module:
        parts.append(_safe_output_name(paper.module))
    if stage:
        parts.append(_safe_output_name(stage))
    output_dir = QA_REPORT_DIR.joinpath(*parts) if parts else QA_REPORT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    base_name = f"{paper.paper_label}_质检报告"
    return output_dir / f"{base_name}.md"


def _apply_question_reports(questions: list[dict], question_reports: list[dict]) -> None:
    for report in question_reports:
        index = report.get("index")
        if not isinstance(index, int) or index < 0 or index >= len(questions):
            continue
        issues = report.get("issues") or []
        questions[index]["status"] = report.get("status") or "passed"
        questions[index]["issues"] = [issue.get("name", "") for issue in issues if issue.get("name")]
        questions[index]["issue_details"] = issues


def _mark_manual_review(questions: list[dict], question_reports: list[dict]) -> int:
    count = 0
    for report in question_reports:
        if report.get("status") == "passed" and not report.get("issues"):
            continue
        index = report.get("index")
        if isinstance(index, int) and 0 <= index < len(questions):
            questions[index]["needs_manual_review"] = True
            count += 1
    return count


def _print_quality_summary(report: dict, md_path: Path) -> None:
    summary = report.get("summary") or {}
    print(
        "质检完成："
        f"通过 {summary.get('passed', 0)}，"
        f"失败 {summary.get('failed', 0)}，"
        f"警告 {summary.get('warning', 0)}，"
        f"全卷问题 {summary.get('paper_issues', 0)}"
    )
    print(f"质检 Markdown：{md_path}")


def _issue_texts(report: dict) -> list[str]:
    result = []
    for issue in report.get("issues") or []:
        parts = [issue.get("name"), issue.get("message"), issue.get("code")]
        result.append("：".join(str(part) for part in parts if part))
    return result


def _blueprint_context(row: Any) -> dict[str, Any]:
    if row is None:
        return {}
    if hasattr(row, "__dataclass_fields__"):
        return asdict(row)
    if isinstance(row, dict):
        return dict(row)
    return {name: getattr(row, name) for name in dir(row) if not name.startswith("_") and not callable(getattr(row, name))}


def _plan_context_for_question(paper: PaperPlan, question: dict[str, Any]) -> dict[str, Any]:
    question_no = question.get("question_no")
    row = None
    for blueprint_row in paper.blueprint_rows or []:
        if getattr(blueprint_row, "question_no", None) == question_no:
            row = blueprint_row
            break
    return {
        "paper_no": paper.paper_no,
        "paper_label": paper.paper_label,
        "paper_type": paper.paper_type,
        "module": paper.module,
        "topic": paper.topic,
        "point_name": paper.point_name,
        "point_content": paper.point_content,
        "blueprint": _blueprint_context(row),
    }


def _is_protected_original_question(question: dict[str, Any]) -> bool:
    return bool(question.get("protected_original_docx_block"))


def _filter_repair_candidates(question_reports: list[dict], questions: list[dict]) -> list[dict]:
    candidates: list[dict] = []
    for report in question_reports:
        if not needs_regeneration(_issue_texts(report)):
            continue
        index = report.get("index")
        if isinstance(index, int) and 0 <= index < len(questions):
            question = questions[index]
            if _is_protected_original_question(question):
                print(
                    f"  第{question.get('question_no')}题：题目含图片/公式对象/OLE 对象，跳过自动重生成，最终保留原 DOCX 题块"
                )
                continue
        candidates.append(report)
    return candidates


def _make_llm_call(config: dict[str, Any]) -> Callable[[str], str]:
    from openai import OpenAI

    client = OpenAI(api_key=config["api_key"], base_url=config.get("base_url"))
    model = config.get("model")
    max_tokens = int(config.get("max_tokens") or 8000)
    temperature = float(config.get("temperature") or 0.1)
    system_prompt = "你是职业教育题目修复助手。只输出一个纯 JSON 对象（不带 ```json 包裹，不带任何解释或 Markdown 标记）。"

    def llm_call(prompt: str) -> str:
        text, _usage = call_api(client, model, system_prompt, prompt, max_tokens=max_tokens, temperature=temperature, json_mode=True)
        return text

    return llm_call


def _repair_one_question(
    question: dict[str, Any],
    issues: list[str],
    plan_context: dict[str, Any],
    spec_text: str,
    llm_call: Callable[[str], str],
) -> tuple[int, Any]:
    """单题修复包装函数，供线程池并发调用。"""
    result = regenerate_question(question, issues, plan_context, spec_text, llm_call)
    question_no = question.get("question_no", "?")
    if result.status == "success" and result.regenerated:
        print(f"  第{question_no}题：修复成功")
    else:
        print(f"  第{question_no}题：修复失败 {result.message}")
    return question.get("_repair_index", -1), result


def _repair_questions_with_qc(paper: PaperPlan, loaded_paper, args: argparse.Namespace, planned_total: int = 0, score_multiplier: int = 1) -> tuple[list[dict], list[dict], int, int]:
    question_reports, paper_issues = run_quality_checks(paper, loaded_paper.questions, planned_total, score_multiplier)
    repaired_count = 0
    rounds = 0
    before_summary = {
        "失败题数": sum(1 for item in question_reports if item.get("status") == "failed"),
        "警告题数": sum(1 for item in question_reports if item.get("status") == "warning"),
        "全卷问题数": len(paper_issues or []),
    }

    # 0. 本地修复答案分布失衡（选项互换，无需 AI）
    dist_fixed = fix_answer_distribution(loaded_paper.questions)
    if dist_fixed:
        print("  已执行答案分布本地修复：互换最高频/最低频选项")
        question_reports, paper_issues = run_quality_checks(paper, loaded_paper.questions, planned_total, score_multiplier)

    if args.max_fix_rounds <= 0:
        _update_run_record(
            paper,
            **{
                "修复记录": {
                    "一句话说明": "本次没有调用 AI 修复，因为最大修复轮数设置为 0。",
                    "修复前": before_summary,
                    "实际修复轮数": rounds,
                    "修复成功题数": repaired_count,
                    "修复后": before_summary,
                }
            },
        )
        return question_reports, paper_issues, repaired_count, rounds

    candidates = _filter_repair_candidates(question_reports, loaded_paper.questions)
    if not candidates:
        _update_run_record(
            paper,
            **{
                "修复记录": {
                    "一句话说明": "质检后没有发现需要 AI 修复的题目。",
                    "修复前": before_summary,
                    "实际修复轮数": rounds,
                    "修复成功题数": repaired_count,
                    "修复后": before_summary,
                }
            },
        )
        return question_reports, paper_issues, repaired_count, rounds

    config = load_config()
    spec_text = load_spec()
    llm_call = _make_llm_call(config)
    max_workers = max(1, int(args.max_fix_workers or 1))

    _t_repair_start = time.perf_counter()

    for round_no in range(1, args.max_fix_rounds + 1):
        rounds = round_no
        changed = 0
        concurrent = max_workers > 1
        worker_label = f"（并发 {max_workers} 线程）" if concurrent else ""
        print(f"开始第 {round_no} 轮问题题修复，共 {len(candidates)} 题候选{worker_label}。")
        _t_round = time.perf_counter()

        if concurrent:
            # 出队前打索引标记，方便 worker 回来后定位
            for report in candidates:
                index = report.get("index")
                if isinstance(index, int) and 0 <= index < len(loaded_paper.questions):
                    loaded_paper.questions[index]["_repair_index"] = index

            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures: dict[Any, dict] = {}
                for report in candidates:
                    index = report.get("index")
                    if not isinstance(index, int) or index < 0 or index >= len(loaded_paper.questions):
                        continue
                    question = loaded_paper.questions[index]
                    issues = _issue_texts(report)
                    future = executor.submit(
                        _repair_one_question,
                        question,
                        issues,
                        _plan_context_for_question(paper, question),
                        spec_text,
                        llm_call,
                    )
                    futures[future] = index

                _total = len(futures)
                _done = 0
                for future in as_completed(futures):
                    index, result = future.result()
                    _done += 1
                    if not isinstance(index, int) or index < 0 or index >= len(loaded_paper.questions):
                        continue
                    if result.status == "success" and result.regenerated:
                        loaded_paper.questions[index] = result.regenerated
                        repaired_count += 1
                        changed += 1
                    if _done % max(1, _total // 5) == 0 or _done == _total:
                        print(f"  修复进度：{_done}/{_total}（成功 {changed} 题）")

            # 清理临时标记
            for q in loaded_paper.questions:
                q.pop("_repair_index", None)
        else:
            _total = len(candidates)
            _done = 0
            for report in candidates:
                index = report.get("index")
                if not isinstance(index, int) or index < 0 or index >= len(loaded_paper.questions):
                    continue
                question = loaded_paper.questions[index]
                issues = _issue_texts(report)
                _done += 1
                _t_q = time.perf_counter()
                result = regenerate_question(question, issues, _plan_context_for_question(paper, question), spec_text, llm_call)
                _dt_q = time.perf_counter() - _t_q
                if result.status == "success" and result.regenerated:
                    loaded_paper.questions[index] = result.regenerated
                    repaired_count += 1
                    changed += 1
                    print(f"  [{_done}/{_total}] 第{question.get('question_no')}题：修复成功（{_dt_q:.1f}s）")
                else:
                    print(f"  [{_done}/{_total}] 第{question.get('question_no')}题：修复失败 {result.message}（{_dt_q:.1f}s）")

        _dt_round = time.perf_counter() - _t_round
        print(f"  第 {round_no} 轮修复完成：成功 {changed} 题，耗时 {_dt_round:.1f}s")
        question_reports, paper_issues = run_quality_checks(paper, loaded_paper.questions, planned_total, score_multiplier)
        candidates = _filter_repair_candidates(question_reports, loaded_paper.questions)
        if not candidates or changed == 0:
            break
    after_summary = {
        "失败题数": sum(1 for item in question_reports if item.get("status") == "failed"),
        "警告题数": sum(1 for item in question_reports if item.get("status") == "warning"),
        "全卷问题数": len(paper_issues or []),
    }
    _update_run_record(
        paper,
        **{
            "修复记录": {
                "一句话说明": "系统只对质检指出的可修复问题做有限轮次修复，不会无限重写整卷。",
                "修复前": before_summary,
                "实际修复轮数": rounds,
                "修复成功题数": repaired_count,
                "修复后": after_summary,
                "是否仍需人工看": (after_summary["失败题数"] + after_summary["警告题数"] + after_summary["全卷问题数"]) > 0,
            }
        },
    )

    # 回写修复后结果到组卷待质检，下次加载直接用
    _repaired_to_manual_txt(paper, loaded_paper.questions)

    _dt_repair = time.perf_counter() - _t_repair_start
    print(f"  [修复总耗时] {paper.paper_label}：{_dt_repair:.1f}s（{rounds} 轮，成功修复 {repaired_count} 题）")
    return question_reports, paper_issues, repaired_count, rounds


def _build_quality_report_and_write(paper: PaperPlan, loaded_paper, question_reports, paper_issues, questions_path: Path, stage: str = "修复前") -> tuple[dict, Path]:
    report = build_quality_report(
        paper,
        loaded_paper,
        question_reports,
        paper_issues,
        loaded_questions_path=questions_path,
    )
    md_path = _qa_report_output_path(paper, stage=stage)
    write_markdown_report(report, md_path)
    _print_quality_summary(report, md_path)
    return report, md_path


def _question_context_rows(paper: PaperPlan) -> list[dict[str, Any]]:
    return [_blueprint_context(row) for row in paper.rows]


def _full_paper_generation_context(paper: PaperPlan) -> dict[str, Any]:
    meta = paper.meta
    return {
        "paper_no": paper.paper_no,
        "paper_label": paper.paper_label,
        "paper_type": paper.paper_type,
        "province": meta.province if meta else "",
        "exam_category": meta.exam_category if meta else "",
        "title": meta.title if meta else "",
        "module": paper.module,
        "course": paper.module,
        "topic": paper.topic,
        "point_name": paper.point_name,
        "point_content": paper.point_content,
        "paper_ref": paper.paper_ref,
        "planning_rows": _question_context_rows(paper),
        "blueprint_path": str(paper.blueprint_path) if paper.blueprint_path else "",
        "blueprint_rows": [_blueprint_context(row) for row in paper.blueprint_rows],
        "blueprint_warnings": paper.blueprint_warnings,
    }


def _confirm_direct_ai_generation(paper: PaperPlan) -> bool:
    search_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
    print(f"未找到待质检试卷：{paper.paper_label}｜{paper.paper_type}")
    print(f"查找目录：{search_dir}")
    raw = input("是否调用 API 直接按编写规范生成整卷？[y/N]: ").strip().lower()
    return raw in {"y", "yes", "是", "直接生成", "生成"}


def _generate_full_paper_with_ai(paper: PaperPlan):
    if paper.paper_type in {TOPIC_PAPER_TYPE, COURSE_PAPER_TYPE} and not paper.blueprint_rows:
        raise RuntimeError("专题训练卷/课程综合卷缺少细目表，无法可靠直接生成整卷。")

    from openai import OpenAI

    config = load_config()
    spec_text = load_spec()
    client = OpenAI(api_key=config["api_key"], base_url=config.get("base_url"))
    model = config.get("model")
    max_tokens = int(config.get("max_tokens") or 8000)
    temperature = float(config.get("temperature") or 0.2)
    system_prompt = "你是职业教育试卷命题专家。只输出一个 JSON 对象，不要输出 Markdown、解释或代码块。"
    prompt = build_generate_full_paper_prompt(_full_paper_generation_context(paper), spec_text)

    print("开始调用 API 直接生成整卷...")
    text, usage = call_api(client, model, system_prompt, prompt, max_tokens=max_tokens, temperature=temperature, json_mode=True)
    session_usage = _new_usage_summary(config=config)
    daily_usage = _load_daily_usage()
    _record_token_usage(session_usage, daily_usage, usage, config)
    _print_token_summary(session_usage, daily_usage)

    output_path = _ai_generated_paper_output_path(paper)
    generated_data = {
        "paper": _full_paper_generation_context(paper),
        "raw_response": text,
        "usage": usage or {},
    }
    output_path.write_text(json.dumps(generated_data, ensure_ascii=False, indent=2), encoding="utf-8")
    loaded_paper = load_generated_paper_json(text, paper, output_path)
    normalized_data = json.loads(output_path.read_text(encoding="utf-8"))
    normalized_data["loaded_paper"] = loaded_paper_to_dict(loaded_paper)
    output_path.write_text(json.dumps(normalized_data, ensure_ascii=False, indent=2), encoding="utf-8")
    _update_run_record(
        paper,
        **{
            "AI整卷生成记录": {
                "一句话说明": f"【兜底路径】{paper.paper_label} 无待质检 DOCX 文件，跳过拆题/修复流程，由 AI 直接按编写规范生成整卷（非标准管线产物，未经质检）。",
                "生成了几题": len(getattr(loaded_paper, "questions", []) or []),
                "结果文件": str(output_path),
            }
        },
    )
    print(f"已保存 AI 直接生成结果：{output_path}")
    return loaded_paper


def _topic_context(paper: PaperPlan) -> dict[str, Any]:
    return {
        "seq": paper.paper_no,
        "paper_no": paper.paper_no,
        "paper_type": paper.paper_type,
        "module": paper.module,
        "course": paper.module,
        "topic": paper.topic,
        "theme": paper.topic or paper.point_name or paper.module,
        "point_name": paper.point_name,
        "point_content": paper.point_content,
        "rows": paper.rows,
    }


def _course_error_collection_path(paper: PaperPlan) -> Path:
    output_dir = _get_topic_output_base(paper.meta, _topic_context(paper), FINAL_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir / "错误收集.md"


def _append_course_error_collection(result: PaperRunResult, paper: PaperPlan) -> None:
    """将待人工审核试卷的问题题目追记到错误收集文档，只保留卷号/题号/原因。
    
    同一卷重新生成时，会先清除该卷的旧记录再写入，避免新旧混杂。
    同时提取逐题问题和全卷级别问题（如答案分布失衡），确保不遗漏。
    """
    if result.status != "manual_review":
        return

    collection_path = _course_error_collection_path(paper)
    is_new_file = not collection_path.exists() or collection_path.stat().st_size == 0

    error_lines: list[str] = []

    # 1. 提取全卷级别问题（如答案分布失衡等）
    paper_issues = result.paper_issues or []
    if paper_issues:
        issue_texts = []
        for pi in paper_issues:
            name = pi.get("name", "") if isinstance(pi, dict) else str(pi)
            msg = pi.get("message", "") if isinstance(pi, dict) else ""
            issue_texts.append(f"{name}：{msg}" if msg else name)
        if issue_texts:
            paper_issue_summary = "；".join(issue_texts)
            error_lines.append(f"| {result.paper_label} | 全卷 | {paper_issue_summary} |")

    # 2. 从修复后拆题 JSON 中读取每题状态
    q_path = Path(result.repaired_questions_path) if result.repaired_questions_path else (
        Path(result.questions_path) if result.questions_path else None
    )

    if q_path and q_path.exists():
        try:
            data = json.loads(q_path.read_text(encoding="utf-8"))
            loaded = data.get("loaded_paper") or data
            questions = loaded.get("questions", [])
            if not questions and isinstance(data, list):
                questions = data
            for q in questions:
                qno = q.get("question_no", "?")
                status = q.get("qc_status", q.get("status", ""))
                if status not in ("", "passed", None):
                    issues = q.get("qc_issues", q.get("issues", []))
                    if not issues:
                        error_lines.append(f"| {result.paper_label} | 第{qno}题 | 待人工审核 |")
                    for iss in issues:
                        name = iss.get("name", iss) if isinstance(iss, dict) else str(iss)
                        msg = iss.get("message", "") if isinstance(iss, dict) else ""
                        reason = f"{name}：{msg}" if msg else name
                        error_lines.append(f"| {result.paper_label} | 第{qno}题 | {reason} |")
        except Exception as e:
            error_lines.append(f"| {result.paper_label} | - | 解析失败：{e} |")

    # 3. 兜底：确有未解决问题但无法逐条提取时，给出有意义的汇总信息
    if not error_lines:
        qc_path = result.final_qc_report_path or result.qc_report_path or "未知"
        stage = "修复后" if result.repaired_questions_path else "修复前"
        error_lines.append(
            f"| {result.paper_label} | 全卷 | "
            f"({stage}质检) 仍有 {result.remaining_issue_count} 个未解决问题，"
            f"详见质检报告：{qc_path} |"
        )

    # 过滤掉本卷旧记录，避免重新生成时新旧混杂
    if not is_new_file:
        try:
            old_lines = collection_path.read_text(encoding="utf-8").split("\n")
            kept = []
            for line in old_lines:
                if line.startswith("| ") and result.paper_label in line:
                    continue  # 跳过本卷旧记录
                kept.append(line)
            # 如果只剩表头无数据行，视为新文件
            data_lines = [l for l in kept if l.startswith("| ")]
            if not data_lines:
                is_new_file = True
        except Exception:
            pass

    if not is_new_file:
        # 已有文件 → 先写回过滤后的内容，再追加新行
        filtered = "\n".join(
            l for l in kept
            if l or l.strip()  # keep all non-empty header lines too
        )
        collection_path.write_text(filtered.strip() + "\n", encoding="utf-8")
    else:
        collection_path.write_text("", encoding="utf-8")

    with collection_path.open("a", encoding="utf-8") as f:
        if is_new_file or collection_path.stat().st_size == 0:
            f.write("# 错误收集\n\n")
            f.write("| 试卷序号 | 题目序号 | 错误原因 |\n")
            f.write("|---|---|---|\n")
        for line in error_lines:
            f.write(line + "\n")

    print(f"错误收集已更新：{collection_path}")


def _try_load_questions(result: PaperRunResult, paper: PaperPlan) -> list[dict[str, Any]] | None:
    """从修复后或初始拆题 JSON 中重新读取题目列表。"""
    path_str = result.repaired_questions_path or result.questions_path
    if not path_str:
        return None
    path = Path(path_str)
    if not path.exists():
        return None
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        lp = data.get("loaded_paper") or {}
        return lp.get("questions") or []
    except Exception:
        return None


def _process_single_paper(
    paper: PaperPlan,
    args: argparse.Namespace,
    results: list[PaperRunResult],
) -> PaperRunResult:
    """处理单卷：加载、拆题、质检、修复。返回 PaperRunResult。"""
    result = PaperRunResult(paper_label=paper.paper_label, paper_type=paper.paper_type)
    print(f"\n准备处理：{paper.paper_label}｜{paper.paper_type}")
    try:
        search_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
        loaded_paper = load_manual_paper_for_plan(paper, search_dir)
        type_plans = _build_type_plans_for_paper(paper, paper.meta)
        planned_total = _planned_total(type_plans)

        if loaded_paper is None:
            print(f"未找到待质检试卷：{paper.paper_label}｜{paper.paper_type}")
            print(f"查找目录：{search_dir}")
            print("将先按规划表通过 API 拉题生成待质检 DOCX。")
            _update_run_record(
                paper,
                **{
                    "输入材料记录": {
                        "一句话说明": "没有在组卷待质检目录里找到这卷的 DOCX/TXT，系统会先尝试 API 拉题。",
                        "查找目录": str(search_dir),
                        "找到待质检文件": False,
                    }
                },
            )
            api_ok = _api_pull_for_paper(paper, paper.meta) if paper.meta else False
            if api_ok:
                loaded_paper = load_manual_paper_for_plan(paper, search_dir)
            if loaded_paper is None:
                print("API 拉题未生成可加载 DOCX，尝试 AI 兜底生成整卷。")
                if not _confirm_direct_ai_generation(paper):
                    result.status = "failed"
                    result.error = f"未找到人工组卷文件，且 API/AI 未生成：{search_dir}"
                    print(result.error)
                    print(f"文件名建议包含卷号，例如：{paper.paper_label} {paper.paper_type}.docx")
                    _update_run_record(
                        paper,
                        **{
                            "最终结果": {
                                "一句话说明": "没有找到原始文件，API 和 AI 也没有生成可用试卷，本卷失败。",
                                "状态": _plain_status(result.status),
                                "错误说明": result.error,
                            }
                        },
                    )
                    return result
                loaded_paper = _generate_full_paper_with_ai(paper)
        else:
            _update_run_record(
                paper,
                **{
                    "输入材料记录": {
                        "一句话说明": "已在组卷待质检目录找到这卷的材料，直接进入拆题和质检。",
                        "查找目录": str(search_dir),
                        "找到待质检文件": True,
                        "文件路径": str(getattr(loaded_paper, "path", "") or ""),
                    }
                },
            )

        _print_loaded_paper_summary(loaded_paper)
        result.question_count = len(loaded_paper.questions)

        if planned_total > 0:
            actual = len(loaded_paper.questions)
            if actual < planned_total:
                print(f"\n  题量不足：实际 {actual} 题，规划 {planned_total} 题，启动 AI 缺口补题...")
                added = _supplement_loaded_paper_with_ai(paper, loaded_paper, type_plans)
                if added:
                    print(f"  AI 已补充 {added} 题，当前 {len(loaded_paper.questions)} 题")
                    result.question_count = len(loaded_paper.questions)
                    # 同步更新组卷待质检 DOCX，避免下次仍加载旧的不完整文件
                    _save_complete_docx_to_manual_dir(paper, loaded_paper.questions)
                else:
                    print("  AI 未补充到新题，将继续使用现有题目进入质检。")
            elif actual > planned_total:
                print(f"\n  题量超出：实际 {actual} 题，规划 {planned_total} 题，按规划题型随机裁剪...")
                loaded_paper.questions = _trim_questions_to_plan(loaded_paper.questions, type_plans)
                result.question_count = len(loaded_paper.questions)
                print(f"  已裁剪为 {result.question_count} 题")
            else:
                _renumber_questions(loaded_paper.questions)

        # 判断题答案统一为 √ / ×
        normalize_all_judge_answers(loaded_paper.questions)

        initial_reports, initial_paper_issues = run_quality_checks(paper, loaded_paper.questions, planned_total)
        _apply_question_reports(loaded_paper.questions, initial_reports)
        initial_questions_path = _save_loaded_paper_questions(paper, loaded_paper)
        result.questions_path = str(initial_questions_path)
        print(f"已保存拆题结果：{initial_questions_path}")
        _initial_report, initial_md_path = _build_quality_report_and_write(
            paper,
            loaded_paper,
            initial_reports,
            initial_paper_issues,
            initial_questions_path,
            stage="修复前",
        )
        result.qc_report_path = str(initial_md_path)
        initial_summary = _initial_report.get("summary") or {}
        _update_run_record(
            paper,
            **{
                "拆题和初检记录": {
                    "一句话说明": (
                        f"【修复前】第1轮质检完成：拆出 {len(loaded_paper.questions)} 题，"
                        f"通过 {initial_summary.get('passed', 0)} 题，"
                        f"失败 {initial_summary.get('failed', 0)} 题，"
                        f"警告 {initial_summary.get('warning', 0)} 题，"
                        f"全卷问题 {initial_summary.get('paper_issues', 0)} 个"
                    ),
                    "拆出题数": len(loaded_paper.questions),
                    "拆题结果文件": str(initial_questions_path),
                    "修复前质检报告": str(initial_md_path),
                    "初检通过题数": initial_summary.get("passed", 0),
                    "初检失败题数": initial_summary.get("failed", 0),
                    "初检警告题数": initial_summary.get("warning", 0),
                    "初检全卷问题数": initial_summary.get("paper_issues", 0),
                }
            },
        )

        final_reports, final_paper_issues, repaired_count, repair_rounds = _repair_questions_with_qc(
            paper, loaded_paper, args, planned_total,
            getattr(paper, '_score_multiplier', 1),
        )
        result.repaired_count = repaired_count
        result.repair_rounds = repair_rounds
        result.paper_issues = final_paper_issues  # 保存全卷级别问题，供错误收集使用
        _apply_question_reports(loaded_paper.questions, final_reports)
        manual_review_count = _mark_manual_review(loaded_paper.questions, final_reports)
        result.manual_review_count = manual_review_count

        final_questions_path = initial_questions_path
        report_stage = "修复后" if repaired_count or repair_rounds or manual_review_count or final_paper_issues else "修复前"
        if repaired_count or repair_rounds or manual_review_count or final_paper_issues:
            final_questions_path = _save_loaded_paper_questions(paper, loaded_paper, _repaired_questions_output_path(paper))
            result.repaired_questions_path = str(final_questions_path)
            print(f"已保存修复后结果：{final_questions_path}")

        final_report, final_md_path = _build_quality_report_and_write(
            paper,
            loaded_paper,
            final_reports,
            final_paper_issues,
            final_questions_path,
            stage=report_stage,
        )
        result.final_qc_report_path = str(final_md_path)
        final_summary = final_report.get("summary") or {}
        result.remaining_issue_count = (
            int(final_summary.get("failed", 0) or 0)
            + int(final_summary.get("warning", 0) or 0)
            + int(final_summary.get("paper_issues", 0) or 0)
        )

        if result.remaining_issue_count > 0:
            result.status = "manual_review"
        else:
            result.status = "qc_passed"
        _update_run_record(
            paper,
            **{
                "质检修复后的结果": {
                    "一句话说明": (
                    f"【{report_stage}】质检完成："
                    f"逐题通过 {final_summary.get('passed', 0)} 题，"
                    f"失败 {final_summary.get('failed', 0)} 题，"
                    f"警告 {final_summary.get('warning', 0)} 题，"
                    f"全卷问题 {final_summary.get('paper_issues', 0)} 个"
                    f"{'，需人工复核' if result.remaining_issue_count > 0 else '，通过质检'}"
                ),
                    "最终质检报告": str(final_md_path),
                    "最终题目文件": str(final_questions_path),
                    "修复后仍有问题数": result.remaining_issue_count,
                    "需要人工复核题数": result.manual_review_count,
                    "当前状态": _plain_status(result.status),
                }
            },
        )

    except Exception as exc:
        result.status = "failed"
        result.error = str(exc)
        print(f"处理失败：{paper.paper_label}｜{exc}")
        _update_run_record(
            paper,
            **{
                "最终结果": {
                    "一句话说明": f"{paper.paper_label} 处理异常中断（{type(exc).__name__}），本卷未完成。",
                    "状态": _plain_status(result.status),
                    "错误说明": result.error,
                }
            },
        )

    return result

def _write_cross_paper_report(
    paper: PaperPlan,
    cross_paper_issues: list[dict[str, str]],
    cross_paper_stats: list[dict[str, Any]],
    group_questions: list[dict[str, Any]],
) -> Path:
    """写出课程综合卷组的跨卷质检报告。"""
    parts = list(_province_category_parts(paper))
    if paper.module:
        parts.append(_safe_output_name(paper.module))
    output_dir = QA_REPORT_DIR.joinpath(*parts) if parts else QA_REPORT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    group_ref = paper.paper_ref or paper.paper_label
    safe_ref = "".join("_" if c in '<>:"/\\|?*' else c for c in group_ref)
    report_path = output_dir / f"课程综合卷组_{safe_ref}_跨卷质检报告.md"

    lines = [
        f"# 课程综合卷组 {group_ref} 跨卷质检报告",
        "",
        "## 跨卷重复率",
        "",
    ]

    for stats in cross_paper_stats:
        label_a = stats.get("paper_a_label", "?")
        label_b = stats.get("paper_b_label", "?")
        rate = stats.get("repeat_rate", 0)
        pairs = stats.get("high_similarity_pairs", 0)
        status = "❌ 超标" if rate > 0.2 else ("⚠️ 偏高" if rate > 0.1 else "✅ 正常")
        lines.append(f"- {label_a} ↔ {label_b}：相似 {pairs} 对，重复率 {rate:.1%} {status}")

    lines.append("")

    # 汇总跨卷问题
    if cross_paper_issues:
        lines.extend(["## 跨卷问题", ""])
        for issue in cross_paper_issues:
            severity = issue.get("severity", "")
            name = issue.get("name", "")
            msg = issue.get("message", "")
            lines.append(f"- [{severity}] {name}：{msg}")
        lines.append("")

    # 各跨卷对的详细相似题
    for stats in cross_paper_stats:
        pairs = stats.get("similar_pairs") or []
        if not pairs:
            continue
        label_a = stats.get("paper_a_label", "?")
        label_b = stats.get("paper_b_label", "?")
        lines.append(f"## {label_a} ↔ {label_b} 相似题目明细")
        lines.append("")
        type_counts = stats.get("match_type_counts") or {}
        if type_counts:
            summary = "，".join(f"{key}={value}" for key, value in type_counts.items() if value)
            if summary:
                lines.append(f"- 分类型命中：{summary}")
                lines.append("")
        lines.append("| 卷A题号 | 卷B题号 | 相似度 | 类型 | 原因 |")
        lines.append("|---:|---:|---:|---|---|")
        for pair in pairs:
            match_types = pair.get("match_types") or []
            if isinstance(match_types, str):
                match_types_text = match_types
            else:
                match_types_text = ",".join(str(item) for item in match_types)
            lines.append(
                f"| {pair.get('paper_a_question', '')} "
                f"| {pair.get('paper_b_question', '')} "
                f"| {pair.get('similarity', 0):.1%} "
                f"| {match_types_text} "
                f"| {pair.get('reason', '')} |"
            )
        lines.append("")

    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return report_path


def _process_cross_paper_group(
    group_papers: list[PaperPlan],
    group_ref: str,
    args: argparse.Namespace,
    results: list[PaperRunResult],
) -> None:
    """处理一组需要跨卷查重的试卷（考点训练卷组或课程综合卷组）。
    包含两阶段查重：阶段1 MD5哈希精确匹配 + 阶段2 相似度引擎。
    """
    group_results: list[PaperRunResult] = []
    group_loaded: list[tuple[PaperPlan, Any]] = []

    # 第1阶段：逐卷加载 + 单卷质检 + 修复
    for idx, paper in enumerate(group_papers, 1):
        print(f"\n--- [{idx}/{len(group_papers)}] {paper.paper_label} ---")
        _t_paper = time.perf_counter()
        result = _process_single_paper(paper, args, results)
        _dt_paper = time.perf_counter() - _t_paper
        print(f"--- [{idx}/{len(group_papers)}] {paper.paper_label} 完成，耗时 {_dt_paper:.1f}s ---")
        group_results.append(result)
        results.append(result)
        if result.status in ("failed", "pending"):
            continue
        loaded = _try_load_questions(result, paper)
        if loaded:
            group_loaded.append((paper, loaded))

    # 第2阶段：跨卷查重（两阶段筛选）
    if len(group_loaded) >= 2:
        print(f"\n  开始跨卷查重：{group_ref}")
        _run_two_phase_cross_check(group_loaded, group_ref, group_papers, group_results)


def _run_two_phase_cross_check(
    group_loaded: list[tuple[PaperPlan, Any]],
    group_ref: str,
    group_papers: list[PaperPlan],
    group_results: list[PaperRunResult],
) -> None:
    """两阶段跨卷查重：Phase1 MD5哈希 + Phase2 相似度引擎 + 增强检测。"""
    _t_cross_start = time.perf_counter()
    import hashlib
    from 质检.rules import (
        _is_duplicate_stem_pair,
        _is_number_substitute_pair,
        _check_answer_stem_leak,
        _check_kpoint_density,
        _check_formula_signature_pair,
        _check_image_context_signature_pair,
        make_issue,
    )

    all_cross_issues: list[dict[str, str]] = []
    all_cross_stats: list[dict[str, Any]] = []

    _n_papers = len(group_loaded)
    _total_pairs = _n_papers * (_n_papers - 1) // 2
    _pair_idx = 0

    for i in range(len(group_loaded)):
        for j in range(i + 1, len(group_loaded)):
            paper_a, loaded_a = group_loaded[i]
            paper_b, loaded_b = group_loaded[j]
            _pair_idx += 1
            _t_pair = time.perf_counter()

            # Phase 1: MD5 哈希精确匹配（快速筛出完全相同的题）
            stems_a: dict[str, int] = {}
            for idx, q in enumerate(loaded_a):
                stem_raw = str(q.get("stem", "")).replace(" ", "").replace("\n", "")
                if stem_raw:
                    h = hashlib.md5(stem_raw.encode()).hexdigest()
                    stems_a[h] = idx + 1

            hash_matches = 0
            for q in loaded_b:
                stem_raw = str(q.get("stem", "")).replace(" ", "").replace("\n", "")
                if stem_raw:
                    h = hashlib.md5(stem_raw.encode()).hexdigest()
                    if h in stems_a:
                        hash_matches += 1

            # Phase 2: 相似度引擎（仅对哈希未匹配的做深度比对）
            # 对于同专题考点训练卷，phrase_score 阈值从 0.55 调至 0.60
            is_topic_group = "考点训练卷" in group_ref or any(
                p.paper_type == POINT_PAPER_TYPE for p in group_papers
            )

            # 执行标准跨卷查重（内部调用 _is_duplicate_stem_pair）
            cross_issues, cross_stats = _cross_paper_check_structured(
                loaded_a, loaded_b,
                paper_a.paper_label, paper_b.paper_label,
            )

            # 增强检测1：数字替换题
            num_sub_count = 0
            extra_pairs = cross_stats.setdefault("similar_pairs", [])
            matched_extra_pairs: set[tuple[str, str, str]] = set()

            def add_extra_pair(qa: dict[str, Any], qb: dict[str, Any], match_type: str, reason: str) -> bool:
                qa_no = qa.get('question_no') or '?'
                qb_no = qb.get('question_no') or '?'
                key = (str(qa_no), str(qb_no), match_type)
                if key in matched_extra_pairs:
                    return False
                matched_extra_pairs.add(key)
                extra_pairs.append({
                    "paper_a_question": qa_no,
                    "paper_b_question": qb_no,
                    "similarity": 1.0,
                    "match_types": [match_type],
                    "reason": reason,
                })
                return True

            for qa in loaded_a:
                for qb in loaded_b:
                    is_sub, reason = _is_number_substitute_pair(
                        qa.get("stem", ""), qb.get("stem", "")
                    )
                    if is_sub and add_extra_pair(qa, qb, "number_substitute", reason):
                        num_sub_count += 1
                        qa_no = qa.get('question_no') or '?'
                        qb_no = qb.get('question_no') or '?'
                        cross_issues.append(make_issue(
                            "number_substitute", "数字替换题",
                            WARNING,
                            f"{paper_a.paper_label}第{qa_no}题 ↔ "
                            f"{paper_b.paper_label}第{qb_no}题: {reason}",
                        ))
                        break  # 每题只记一次

            # 增强检测2：答案泄露
            leak_issues = _check_answer_stem_leak(
                loaded_a, loaded_b, paper_a.paper_label, paper_b.paper_label,
            )
            cross_issues.extend(leak_issues)
            # 双向检查
            leak_issues_rev = _check_answer_stem_leak(
                loaded_b, loaded_a, paper_b.paper_label, paper_a.paper_label,
            )
            cross_issues.extend(leak_issues_rev)

            # 增强检测3：知识点密度
            kp_issues = _check_kpoint_density(
                loaded_a, loaded_b, paper_a.paper_label, paper_b.paper_label,
            )
            cross_issues.extend(kp_issues)

            # 第二步检测1：公式指纹
            formula_sig_count = 0
            for qa in loaded_a:
                for qb in loaded_b:
                    matched, reason = _check_formula_signature_pair(qa, qb)
                    if matched and add_extra_pair(qa, qb, "formula_signature", reason):
                        formula_sig_count += 1
                        qa_no = qa.get('question_no') or '?'
                        qb_no = qb.get('question_no') or '?'
                        cross_issues.append(make_issue(
                            "formula_signature_match", "公式指纹重复",
                            WARNING,
                            f"{paper_a.paper_label}第{qa_no}题 ↔ "
                            f"{paper_b.paper_label}第{qb_no}题: {reason}",
                        ))
                        break

            # 第二步检测2：图表/图片上下文签名
            image_sig_count = 0
            for qa in loaded_a:
                for qb in loaded_b:
                    matched, reason = _check_image_context_signature_pair(qa, qb)
                    if matched and add_extra_pair(qa, qb, "image_signature", reason):
                        image_sig_count += 1
                        qa_no = qa.get('question_no') or '?'
                        qb_no = qb.get('question_no') or '?'
                        cross_issues.append(make_issue(
                            "image_context_signature_match", "图表签名重复",
                            WARNING,
                            f"{paper_a.paper_label}第{qa_no}题 ↔ "
                            f"{paper_b.paper_label}第{qb_no}题: {reason}",
                        ))
                        break

            combined_extra = (
                num_sub_count
                + len(leak_issues)
                + len(leak_issues_rev)
                + len(kp_issues)
                + formula_sig_count
                + image_sig_count
            )

            # 补充哈希阶段发现的精确重复
            if hash_matches > 0:
                cross_stats["hash_exact_matches"] = hash_matches
                cross_stats["high_similarity_pairs"] += hash_matches

            cross_stats["match_type_counts"] = {
                "题干相似": sum(1 for p in extra_pairs if "stem_similarity" in (p.get("match_types") or [])),
                "数字替换": num_sub_count,
                "公式指纹": formula_sig_count,
                "图表图片": image_sig_count,
                "精确哈希": hash_matches,
            }

            # 数字/公式/图表命中也进入明细表和统计，方便人工审核定位。
            extra_match_count = num_sub_count + formula_sig_count + image_sig_count
            if extra_match_count:
                cross_stats["high_similarity_pairs"] += extra_match_count
                total = min(len(loaded_a), len(loaded_b)) or 1
                cross_stats["repeat_rate"] = round(cross_stats["high_similarity_pairs"] / total, 3)

            all_cross_issues.extend(cross_issues)
            all_cross_stats.append(cross_stats)

            rate = cross_stats.get("repeat_rate", 0)
            pairs = cross_stats.get("high_similarity_pairs", 0)
            h_msg = f"（精确重复{hash_matches}对）" if hash_matches else ""
            extra_msg = (
                f"（数字替换{num_sub_count} + 答案泄露{len(leak_issues)+len(leak_issues_rev)} "
                f"+ 密度{len(kp_issues)} + 公式{formula_sig_count} + 图表{image_sig_count}）"
                if combined_extra else ""
            )
            if cross_issues or hash_matches or combined_extra:
                print(f"  ⚠️ {paper_a.paper_label} ↔ {paper_b.paper_label}：相似{pairs}对{h_msg}{extra_msg}，重复率{rate:.1%}")
            else:
                print(f"  ✅ {paper_a.paper_label} ↔ {paper_b.paper_label}：重复率{rate:.1%}，正常")

            _dt_pair = time.perf_counter() - _t_pair
            print(f"    [查重进度 {_pair_idx}/{_total_pairs}] 本对耗时：{_dt_pair:.1f}s")

    # 将跨卷问题分配到各卷的结果（均摊，避免每卷重复计数）
    if all_cross_issues:
        shared = max(1, len(group_results))
        share = len(all_cross_issues) // shared
        for i, result in enumerate(group_results):
            result.cross_paper_issues = len(all_cross_issues)
            result.remaining_issue_count += share + (1 if i == 0 and len(all_cross_issues) % shared else 0)

        cross_report_path = _write_cross_paper_report(
            group_papers[0], all_cross_issues, all_cross_stats,
            [q for _, loaded in group_loaded for q in loaded],
        )
        print(f"  跨卷质检报告已保存：{cross_report_path}")

        if any(issue.get("severity") == "failed" for issue in all_cross_issues):
            for result in group_results:
                result.status = "pending"
                result.remaining_issue_count += 1

    _dt_cross = time.perf_counter() - _t_cross_start
    print(f"  [跨卷查重总耗时] {_dt_cross:.1f}s")


def generate_selected_papers(selected_papers: list[PaperPlan], args: argparse.Namespace) -> list[PaperRunResult]:
    """读取人工组卷，执行拆题、质检、修复、组卷与交付输出。
    对于课程综合卷，整组处理并执行跨卷查重。"""
    singles, topic_groups, course_groups = _partition_papers(selected_papers)
    results: list[PaperRunResult] = []

    # 处理独立卷（专题训练卷）
    for paper in singles:
        result = _process_single_paper(paper, args, results)
        results.append(result)

    # 处理考点训练卷组（同专题跨卷查重）
    for group_ref, group_papers in topic_groups.items():
        group_papers.sort(key=lambda p: p.paper_no)
        print(f"\n{'=' * 50}")
        print(f"考点训练卷组：{group_ref}（共 {len(group_papers)} 卷）")
        print(f"{'=' * 50}")
        _process_cross_paper_group(group_papers, group_ref, args, results)

    # 处理课程综合卷组
    for group_ref, group_papers in course_groups.items():
        group_papers.sort(key=lambda p: p.paper_no)
        print(f"\n{'=' * 50}")
        print(f"课程综合卷组：{group_ref}（共 {len(group_papers)} 卷）")
        print(f"{'=' * 50}")
        _process_cross_paper_group(group_papers, group_ref, args, results)

    # 第3阶段：生成 DOCX 和后处理
    for result in results:
        if result.status in ("failed", "pending"):
            continue
        # find matching paper
        paper = None
        for p in selected_papers:
            if p.paper_label == result.paper_label:
                paper = p
                break
        if paper is None:
            continue

        loaded = _try_load_questions(result, paper)
        if loaded is None:
            continue

        text_path = write_analysis_text(paper, loaded)
        result.text_path = str(text_path)
        print(f"已生成解析版文本：{text_path}")

        needs_manual = result.status == "manual_review"
        docx_path = Path(
            generate_docx(
                paper.meta,
                _topic_context(paper),
                1,
                text_path.read_text(encoding="utf-8"),
                FINAL_OUTPUT_DIR,
                needs_manual_review=needs_manual,
                questions=loaded,
            )
        )
        result.docx_path = str(docx_path)
        result.output_dir = str(docx_path.parent)
        print(f"已生成解析版 DOCX：{docx_path}")

        if needs_manual:
            print("仍存在质检问题（含跨卷重复），解析版 DOCX 已标记待人工审核，跳过原卷版和 ZIP 自动交付。")
            _append_course_error_collection(result, paper)
        else:
            _post_process(FINAL_OUTPUT_DIR, target_dirs=[docx_path.parent])
            result.status = "success"

    success = sum(1 for item in results if item.status == "success")
    failed = sum(1 for item in results if item.status == "failed")
    manual = sum(1 for item in results if item.status == "manual_review")
    repaired = sum(item.repaired_count for item in results)
    remaining = sum(item.remaining_issue_count for item in results)


    print("\n最终汇总：")
    print(f"成功卷数：{success}，失败卷数：{failed}，待人工审核卷数：{manual}")
    print(f"修复题数：{repaired}，仍有问题题数：{remaining}")
    if any(item.status == "failed" for item in results):
        raise SystemExit(1)
    return results


def _ask_continue_generation() -> bool:
    """询问是否继续生成其他批次试卷。"""
    print()
    while True:
        cont = input("是否还需要生成别的批次的试卷？(y=继续 / n=退出): ").strip().lower()
        if cont in {"y", "yes", "是", "需要", "继续"}:
            print()
            return True
        if cont in {"n", "no", "否", "不用", "不需要", "退出", ""}:
            return False
        print("请输入 y 或 n")


def _write_meta_to_xlsx(plan_path: Path, meta) -> None:
    """将 exam_table_title 和 exam_type 回写到规划表 xlsx。"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(plan_path)
        ws = wb[wb.sheetnames[0]]
        ws.cell(2, 1).value = meta.exam_table_title
        ws.cell(2, 2).value = meta.exam_type
        wb.save(plan_path)
        print(f"已保存考试科目表名称和考试类型到规划表。")
    except Exception as exc:
        print(f"警告：无法回写规划表：{exc}")


# ============================================================
# API 拉题 → DOCX（模式 2 专用）
# ============================================================

# 题型中文名 → 拉题顺序和缓冲量
from .question_types import TYPE_ORDER as _QUESTION_TYPE_ORDER, normalize_question_type, type_display_name

# 难度映射：简单→API(0.80~1.00), 适中→API(0.50~0.79), 困难→API(0.00~0.49)
_DIFFICULTY_RANGES = {
    "简单": (0.80, 1.00),
    "适中": (0.50, 0.79),
    "困难": (0.00, 0.49),
}

# API 拉题超额系数：每卷每难度多拉 count × PULL_MULTIPLIER 道候选，供来源去重筛选
_PULL_MULTIPLIER = 3

# 来源去重：同一 source+year 组内最多取 N 道
_MAX_PER_SOURCE_GROUP = 3


def _question_type_of(question: dict[str, Any]) -> str:
    return normalize_question_type(
        question.get("_question_type")
        or question.get("question_type")
        or question.get("type")
        or ""
    )


def _build_type_plans_for_paper(paper: PaperPlan, meta=None) -> dict[str, list[dict]]:
    """按细目表或课程题型汇总生成本卷逐题计划。"""
    type_plans: dict[str, list[dict]] = {}
    if paper.blueprint_rows:
        for row in paper.blueprint_rows:
            qtype = normalize_question_type(getattr(row, "question_type", ""))
            if not qtype:
                continue
            diff = getattr(row, "difficulty", "简单")
            type_plans.setdefault(qtype, []).append({
                "question_no": getattr(row, "question_no", 0),
                "difficulty": diff if diff in ("简单", "适中", "困难") else "简单",
            })
        return type_plans

    resolved_meta = paper.meta or meta
    summaries = getattr(resolved_meta, "question_type_summaries", {}) if resolved_meta else {}
    module_summary = summaries.get(paper.module, [])
    qno = 0
    for ts in module_summary:
        qtype = normalize_question_type(ts.get("type", ""))
        if not qtype:
            continue
        count_text = str(ts.get("count", "0")).replace("题", "").strip()
        try:
            count = int(float(count_text or 0))
        except ValueError:
            continue
        for diff in _distribute_difficulty(count):
            qno += 1
            type_plans.setdefault(qtype, []).append({"question_no": qno, "difficulty": diff})
    return type_plans


def _planned_total(type_plans: dict[str, list[dict]]) -> int:
    return sum(len(plans) for plans in type_plans.values())


def _limit_type_plans(type_plans: dict[str, list[dict]], target_total: int) -> dict[str, list[dict]]:
    if target_total <= 0:
        return {}
    limited: dict[str, list[dict]] = {}
    remaining = target_total
    for qtype in _QUESTION_TYPE_ORDER:
        plans = type_plans.get(qtype, [])
        if not plans or remaining <= 0:
            continue
        take = min(len(plans), remaining)
        limited[qtype] = list(plans[:take])
        remaining -= take
    if remaining > 0:
        for qtype, plans in type_plans.items():
            if qtype in limited or remaining <= 0:
                continue
            take = min(len(plans), remaining)
            limited[qtype] = list(plans[:take])
            remaining -= take
    return limited


def _decide_ai_target(planned_total: int, selected_count: int) -> tuple[int, int]:
    """返回最终目标题量和分值倍率。"""
    if planned_total <= 0:
        return selected_count, 1
    hit_rate = selected_count / planned_total
    if hit_rate < 0.40:
        return max(1, planned_total // 2), 2
    return planned_total, 1


def _source_group_key(item: dict[str, Any]) -> str:
    year = str(item.get("year") or "").strip()
    source = str(
        item.get("source")
        or item.get("paperName")
        or item.get("paper_name")
        or item.get("paper")
        or item.get("paper_id")
        or item.get("source_id")
        or ""
    ).strip()
    if year or source:
        return f"{year}|{source}"
    qid = str(item.get("questionId") or item.get("question_id") or "").strip()
    if qid:
        return f"_unknown:{qid}"
    import hashlib
    stem = str(item.get("stem") or item.get("question") or item.get("content") or id(item))
    return "_unknown:" + hashlib.md5(stem.encode("utf-8")).hexdigest()[:12]


def _select_candidates_with_constraints(
    candidates: list[dict],
    needed: int,
    existing_questions: list[dict] | None = None,
) -> list[dict]:
    """按来源/年份/试卷组硬限制选题；不足时返回实际可选数量，不突破每组上限。"""
    import random
    if needed <= 0 or not candidates:
        return []
    group_counts: dict[str, int] = defaultdict(int)
    for question in existing_questions or []:
        group_counts[_source_group_key(question)] += 1
    shuffled = list(candidates)
    random.shuffle(shuffled)
    selected: list[dict] = []
    seen_ids: set[str] = set()
    for item in shuffled:
        qid = str(item.get("questionId") or item.get("question_id") or item.get("stem") or id(item))
        if qid in seen_ids:
            continue
        key = _source_group_key(item)
        if group_counts[key] >= _MAX_PER_SOURCE_GROUP:
            continue
        selected.append(item)
        seen_ids.add(qid)
        group_counts[key] += 1
        if len(selected) >= needed:
            break
    return selected


def _select_by_source_dedup(pool: list[dict], needed: int) -> list[dict]:
    """来源去重选取：同一 year/source/paper 组内最多取 3 道，且 fallback 不突破上限。"""
    return _select_candidates_with_constraints(pool, needed)


def _distribute_difficulty(count: int) -> list[str]:
    """按 简单80% 适中10% 困难10% 分配难度标签列表。"""
    if count <= 0:
        return []
    easy = round(count * 0.80)
    medium = round(count * 0.10)
    hard = count - easy - medium
    # 确保至少 0
    if hard < 0:
        easy += hard
        hard = 0
    if medium < 0:
        easy += medium
        medium = 0
    result = ["简单"] * easy + ["适中"] * medium + ["困难"] * hard
    return result


def _compute_shortfall(
    type_plans: dict[str, list[dict]],
    pulled: list[dict],
) -> dict[str, list[dict]]:
    """对比需求与已拉取，返回不足的 {(qtype, difficulty): [{question_no, difficulty}, ...]}。"""
    shortfall: dict[str, list[dict]] = {}

    for qtype, plans in type_plans.items():
        pulled_by_diff: dict[str, int] = {"简单": 0, "适中": 0, "困难": 0}
        for q in pulled:
            if _question_type_of(q) == qtype:
                d = q.get("_target_difficulty") or q.get("difficulty") or "简单"
                pulled_by_diff[d] = pulled_by_diff.get(d, 0) + 1

        needed_by_diff: dict[str, int] = {"简单": 0, "适中": 0, "困难": 0}
        for p in plans:
            d = p.get("difficulty", "简单")
            needed_by_diff[d] = needed_by_diff.get(d, 0) + 1

        for diff in ("简单", "适中", "困难"):
            need = needed_by_diff[diff]
            have = pulled_by_diff.get(diff, 0)
            missing = max(0, need - have)
            if missing > 0:
                key = f"{qtype}|{diff}"
                shortfall[key] = [{"question_type": qtype, "difficulty": diff}] * missing

    return shortfall


def _trim_overflow_simple(type_plans: dict[str, list[dict]], pulled: list[dict]) -> int:
    """裁掉超出需求的简单题，为居中/困难题腾位。返回移除数。"""
    import random
    removed = 0
    for qtype, plans in type_plans.items():
        need_simple = sum(1 for p in plans if p.get("difficulty", "简单") == "简单")
        type_qs = [(i, q) for i, q in enumerate(pulled) if _question_type_of(q) == qtype]
        simple_qs = [(i, q) for i, q in type_qs
                     if (q.get("_target_difficulty") or q.get("difficulty") or "简单") == "简单"]
        overflow = len(simple_qs) - need_simple
        if overflow <= 0:
            continue
        # 随机去掉溢出的简单题（从后往前删，避免索引错位）
        to_remove = random.sample(simple_qs, overflow)
        for idx, q in sorted(to_remove, key=lambda x: x[0], reverse=True):
            del pulled[idx]
            removed += 1
    if removed:
        print(f"    简单题溢出 {removed} 道，已裁减")
    return removed


def _ai_fill_shortfall(
    paper: PaperPlan,
    shortfall: dict[str, list[dict]],
    kpoint_ids: list[int],
    course_id: int,
) -> list[dict]:
    """AI 补全 API 拉题不足的题目。"""
    import re
    from openai import OpenAI
    from .config_io import load_config, load_spec, call_api

    config = load_config()
    spec_text = load_spec() or ""
    client = OpenAI(api_key=config["api_key"], base_url=config.get("base_url"))
    model = config.get("model")
    temperature = float(config.get("temperature") or 0.3)

    # 构建需求描述
    shortfall_desc_parts = []
    total_needed = 0
    for key, plans in shortfall.items():
        qtype, diff = key.split("|")
        shortfall_desc_parts.append(f"- {qtype}：{len(plans)} 道（{diff}）")
        total_needed += len(plans)

    # D 列知识点文本
    point_text = "\n".join(row.point_content for row in paper.rows if row.point_content)

    system = (
        "你是中职对口升学考试命题专家。严格根据给定的知识点范围出题，"
        "题型、难度不得偏离。只输出一个 JSON 数组，每个元素是一个题目对象，"
        "不要输出 Markdown、解释或代码块。"
    )

    prompt = (
        f"请为以下考试生成 {total_needed} 道题目。\n\n"
        f"【课程】{paper.module}\n"
        f"【考点】{paper.point_name}\n"
        f"【知识点范围】\n{point_text}\n\n"
        f"【需要补充的题目】\n" + "\n".join(shortfall_desc_parts) + "\n\n"
        f"【出题要求】\n"
        f"1. 题目必须严格在知识点范围内，不得超出\n"
        f"2. 选项应具有干扰性，正确选项不暴露\n"
        f"3. 每题必须包含完整题干、选项（单选题/多选题/判断题）、答案和解析\n"
        f"4. 判断题答案填\"A\"（正确）或\"B\"（错误）\n"
        f"5. 返回 JSON 数组格式：\n"
        f'[\n  {{\n    "stem": "题干文本",\n'
        f'    "options": ["A. 选项A", "B. 选项B", ...],  // 判断题不需要此项\n'
        f'    "answer": "A",\n    "explanation": "解析文本",\n'
        f'    "question_type": "{list(shortfall.keys())[0].split("|")[0] if shortfall else "单选题"}",\n'
        f'    "difficulty": "简单/适中/困难"\n  }}\n]\n\n'
        f'{spec_text}'
    )

    print(f"    正在调用 AI 生成 {total_needed} 道题...")
    text, usage = call_api(client, model, system, prompt, max_tokens=16000, temperature=temperature, json_mode=True)

    # 记录 token 用量
    session_usage = _new_usage_summary(config=config)
    daily_usage = _load_daily_usage()
    _record_token_usage(session_usage, daily_usage, usage, config)
    _print_token_summary(session_usage, daily_usage)

    # 解析响应
    try:
        data = json.loads(text)
        if isinstance(data, dict) and "questions" in data:
            data = data["questions"]
        if not isinstance(data, list):
            text_stripped = text.strip()
            if text_stripped.startswith("```"):
                text_stripped = re.sub(r"^```\w*\n|```$", "", text_stripped).strip()
                data = json.loads(text_stripped)
            else:
                data = []
    except (json.JSONDecodeError, ValueError):
        # 尝试提取 JSON 数组
        match = re.search(r"\[.*?\]", text, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group())
            except Exception:
                print(f"    AI 生成结果解析失败")
                return []
        else:
            print(f"    AI 生成结果非 JSON 格式")
            return []

    # 标注元数据
    result = []
    for item in (data or []):
        if isinstance(item, dict) and item.get("stem"):
            item["_ai_generated"] = True
            item["_question_type"] = item.get("question_type", "")
            item["_target_difficulty"] = item.get("difficulty", "简单")
            item["_course_id"] = course_id
            item["_kpoint_ids"] = kpoint_ids
            # 确保 options 是列表
            if isinstance(item.get("options"), str):
                item["options"] = re.split(r"(?=[A-H][.、．)])", item["options"])
                item["options"] = [o.strip() for o in item["options"] if o.strip()]
            # 清除 AI 产出的 LaTeX / Markdown 残留
            for field in ("stem", "answer", "explanation", "analysis", "options"):
                item[field] = _clean_latex_residue(item.get(field))
            result.append(item)

    return result






def _run_api_pull_phase(papers, meta):
    """逐卷独立拉题，只用各卷自己的 kpoints。"""
    from 学科网API拉题移植版.kpoint_resolver import load_mapping_table

    mapping = load_mapping_table(meta.province, meta.exam_category)

    for paper in papers:
        own_ids = mapping.get(paper.paper_label, [])
        print(f"\n正在为 {paper.paper_label}（{paper.paper_type}）拉取题目...")
        api_result = _api_pull_for_paper(paper, meta, own_ids)
        if api_result:
            print("  拉题完成，已保存到组卷待质检")
        else:
            print("  拉题失败，将尝试使用已有组卷文件")


def _clean_latex_residue(value):
    """清除 AI 产出的 LaTeX/Markdown 残留符号。"""
    import re
    if isinstance(value, list):
        return [_clean_latex_residue(v) for v in value]
    if not isinstance(value, str) or not value:
        return value
    value = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', value)
    value = re.sub(r'\$([^$]+)\$', r'\1', value)
    value = re.sub(r'\\\((.+?)\\\)', r'\1', value)
    value = re.sub(r'\\\[(.+?)\\\]', r'\1', value)
    value = re.sub(r'```\w*', '', value)
    value = re.sub(r'\*\*([^*]+)\*\*', r'\1', value)
    for cmd in ('\\times', '\\cdot', '\\Delta', '\\Omega', '\\rho', '\\alpha', '\\beta', '\\gamma'):
        value = value.replace(cmd, cmd.lstrip('\\'))
    value = re.sub(r'\\text\{([^{}]+)\}', r'\1', value)
    return value.strip()
def _api_pull_for_paper(paper: PaperPlan, meta, topic_kpoint_ids: list[int] | None = None) -> bool:
    """为单卷调用学科网 API 拉题，生成 DOCX 存入 组卷待质检。

    流程：
    1. 从映射表获取该卷的 kpointIds
    2. 从规划表题型汇总获取题型/题数/分值要求
    3. 从细目表获取逐题约束（题型、难度）
    4. 逐题型调用 API 拉题
    5. 组装成 DOCX 存入 组卷待质检
    """
    from 学科网API拉题移植版.kpoint_resolver import load_mapping_table, get_mapping_ai_generate_papers
    from 学科网API拉题移植版.query_questions import (
        build_payload,
        query as api_query,
        DEFAULT_APP_KEY,
        DEFAULT_SIGN,
        print_request_snapshot,
        print_response_diagnostics,
    )
    from 学科网API拉题移植版.api_pull_core import apply_content_filter
    # --- 1. 检查是否标记为 AI 生成 ---
    ai_gen_papers = get_mapping_ai_generate_papers(meta.province, meta.exam_category)
    if paper.paper_label in ai_gen_papers:
        print(f"  {paper.paper_label} 映射表标记为「AI生成」，跳过 API 拉题，将由 AI 直接命题")
        _update_run_record(
            paper,
            **{
                "API拉题记录": {
                    "一句话说明": "映射表把本卷标为 AI生成，所以没有调用题库 API。",
                    "是否调用API": False,
                    "跳过原因": "映射表标记为 AI生成",
                }
            },
        )
        return False

    # --- 2. 加载映射表，构建分层 kpointIds ---
    mapping = load_mapping_table(meta.province, meta.exam_category)
    own_kpoint_ids = mapping.get(paper.paper_label, [])
    if not own_kpoint_ids:
        print(f"  警告：{paper.paper_label} 在映射表中无知识点 ID，跳过 API 拉题")
        _update_run_record(
            paper,
            **{
                "API拉题记录": {
                    "一句话说明": "映射表里没有本卷可用的知识点 ID，所以没有调用题库 API。",
                    "是否调用API": False,
                    "跳过原因": "映射表没有知识点 ID",
                    "映射表路径": str(_mapping_table_path(meta)),
                }
            },
        )
        return False

    # --- 3. 获取课程 courseId + 题型映射 ---
    from 学科网API拉题移植版.kpoint_resolver import resolve_course, resolve_type
    course_id = resolve_course(paper.module)
    if not course_id:
        course_id = resolve_course(paper.module)
    if not course_id:
        print(f"  警告：无法解析课程「{paper.module}」的 courseId，跳过 API 拉题")
        _update_run_record(
            paper,
            **{
                "API拉题记录": {
                    "一句话说明": "没有解析到课程 ID，所以不能调用题库 API。",
                    "是否调用API": False,
                    "跳过原因": "无法解析课程 courseId",
                    "课程": paper.module,
                }
            },
        )
        return False

    # API 拉题用的 kpointIds：优先用专题级扩展，否则用自身的
    api_kpoint_ids = topic_kpoint_ids if topic_kpoint_ids else own_kpoint_ids
    print(f"  courseId={course_id}, 自身kpoints={len(own_kpoint_ids)}个, API搜索kpoints={len(api_kpoint_ids)}个")

    # --- 4. 获取逐题约束（统一使用规划表/细目表 helper） ---
    type_plans = _build_type_plans_for_paper(paper, meta)
    if not type_plans:
        print(f"  警告：{paper.paper_label} 无法确定题型约束，跳过 API 拉题")
        return False
    total_needed = _planned_total(type_plans)
    target_total = total_needed
    score_multiplier = 1

    # 计时：API 拉题阶段总耗时
    _t_pull_start = time.perf_counter()

    #     # --- 4. 逐题型拉题 ---
    pulled_questions: list[dict] = []
    question_no_offset = 0
    import math
    api_type_records: list[dict[str, Any]] = []

    for qtype in _QUESTION_TYPE_ORDER:
        plans = type_plans.get(qtype, [])
        if not plans:
            continue

        type_id_str = resolve_type(paper.module, qtype)
        if not type_id_str:
            print(f"  跳过题型「{qtype}」：课程「{paper.module}」无匹配的 typeId")
            continue
        type_ids = [int(type_id_str)]

        needed = len(plans)
        type_record: dict[str, Any] = {
            "题型": qtype,
            "typeId": type_ids,
            "计划题数": needed,
            "库存量": "?",
            "API返回题数": 0,
            "API入选题数": 0,
            "AI补题数": 0,
            "缺题数": 0,
            "按难度": {},
        }
        # 预检：轻量请求获取该题型库存量（pageSize=1，仅查 totalCount）
        try:
            probe_payload = build_payload(
                course_id=course_id,
                kpoint_ids=api_kpoint_ids,
                type_ids=type_ids,
                page_size=1,
            )
            print_request_snapshot(probe_payload, prefix="  [预检] ")
            _t_probe = time.perf_counter()
            probe_result = api_query(probe_payload, app_key=DEFAULT_APP_KEY, sign=DEFAULT_SIGN)
            _dt_probe = time.perf_counter() - _t_probe
            print(f"  [预检] API耗时：{_dt_probe:.2f}s")
            print_response_diagnostics(probe_result, prefix="  [预检] ")
            total_available = probe_result.get("result", {}).get("totalCount", "?") if probe_result and probe_result.get("valid") else "?"
        except Exception:
            total_available = "?"
        type_record["库存量"] = total_available
        print(f"  拉取 {qtype}：需 {needed} 题（库存 {total_available}）")

        # 4a. 统计各难度需求
        diff_needed: dict[str, int] = {"简单": 0, "适中": 0, "困难": 0}
        for p in plans:
            d = p.get("difficulty", "简单")
            diff_needed[d] = diff_needed.get(d, 0) + 1

        # 4b. 按难度拉题
        # 4c. API 补全不足（用扩展 kpointIds）
        for diff_label, count in list(diff_needed.items()):
            if count <= 0:
                continue
            pool_size = math.ceil(count * _PULL_MULTIPLIER)
            low, high = _DIFFICULTY_RANGES.get(diff_label, (0.80, 1.00))
            payload = build_payload(
                course_id=course_id,
                kpoint_ids=api_kpoint_ids,
                type_ids=type_ids,
                difficulty_low=low,
                difficulty_up=high,
                page_size=min(pool_size + 5, 50),
            )
            print_request_snapshot(payload, prefix="    [API拉题] ")
            _t_api = time.perf_counter()
            result = api_query(payload, app_key=DEFAULT_APP_KEY, sign=DEFAULT_SIGN)
            _dt_api = time.perf_counter() - _t_api
            print(f"    [API耗时] {qtype}（{diff_label}）：{_dt_api:.2f}s")
            print_response_diagnostics(result, prefix="    [API返回] ")
            api_items = []
            diff_record = type_record["按难度"].setdefault(
                diff_label,
                {"需题数": count, "API返回": 0, "API入选": 0, "仍缺": 0},
            )
            diff_record["需题数"] = max(diff_record.get("需题数", 0), count)
            raw_path = _write_api_trace_json(
                paper,
                f"{_safe_output_name(qtype)}_{diff_label}_api_response.json",
                {
                    "试卷": paper.paper_label,
                    "题型": qtype,
                    "难度": diff_label,
                    "courseId": course_id,
                    "kpointIds数量": len(api_kpoint_ids),
                    "typeIds": type_ids,
                    "计划补题数": count,
                    "请求payload": payload,
                    "API原始响应": result,
                },
            )
            diff_record["API原始响应文件"] = str(raw_path)
            if result and result.get("valid"):
                api_items = result.get("result", {}).get("list", [])[:pool_size]
                type_record["API返回题数"] += len(api_items)
                diff_record["API返回"] += len(api_items)
                print(f"    {diff_label}：API 拉取 {len(api_items)}/{pool_size} 题")
            else:
                print(f"    {diff_label}：API 无返回")

            if api_items:
                # 考点训练卷：kpoint 合并后防串知识点过滤
                if paper.paper_type == POINT_PAPER_TYPE and paper.point_name:
                    before = len(api_items)
                    api_items, _removed = apply_content_filter(api_items, paper.point_name)
                    if before != len(api_items):
                        print(f"    [内容过滤] 移除 {before - len(api_items)} 道串知识点题目，剩余 {len(api_items)} 题")

                selected = _select_candidates_with_constraints(api_items, count, pulled_questions)
                for item in selected:
                    item["_target_difficulty"] = diff_label
                    item["_question_type"] = qtype
                    question_no_offset += 1
                    item["_question_no"] = question_no_offset
                    pulled_questions.append(item)
                type_record["API入选题数"] += len(selected)
                diff_record["API入选"] += len(selected)
                print(f"    {diff_label}：来源去重后入选 {len(selected)}/{count} 题")
            remaining_after_api = max(0, diff_needed[diff_label] - (len(selected) if api_items else 0))
            diff_record["仍缺"] = remaining_after_api
            diff_needed[diff_label] = remaining_after_api
        type_record["缺题数"] = sum(item.get("仍缺", 0) for item in type_record["按难度"].values())
        api_type_records.append(type_record)

    # --- 4d. 不足题目 AI 补充生成 ---
    target_total, score_multiplier = _decide_ai_target(total_needed, len(pulled_questions))
    target_type_plans = _limit_type_plans(type_plans, target_total)
    if score_multiplier > 1:
        hit_rate = len(pulled_questions) / total_needed if total_needed else 1.0
        print(f"  API/题库命中率 {hit_rate:.0%}，低于 40%，AI 仅补至半量（{total_needed}→{target_total}），每题分值×2")
    paper._score_multiplier = score_multiplier
    shortfall_plans = _compute_shortfall(target_type_plans, pulled_questions)
    # 简单题溢出裁减 + 居中/困难补位
    _trim_overflow_simple(target_type_plans, pulled_questions)
    shortfall_plans = _compute_shortfall(target_type_plans, pulled_questions)
    if shortfall_plans:
        print(f"  API 拉题不足，AI 补充生成 {sum(len(v) for v in shortfall_plans.values())} 题...")
        ai_generated = _ai_fill_shortfall(paper, shortfall_plans, own_kpoint_ids, course_id)
        if ai_generated:
            qno_start = max((q.get("_question_no", 0) for q in pulled_questions), default=0)
            for item in ai_generated:
                qno_start += 1
                item["_question_no"] = qno_start
                item["_ai_generated"] = True
                item["_question_type"] = _question_type_of(item) or item.get("question_type", "")
                pulled_questions.append(item)
            print(f"  AI 已补充 {len(ai_generated)} 题")
    if len(pulled_questions) > target_total:
        pulled_questions = _trim_questions_to_plan(pulled_questions, target_type_plans)

    ai_counts = Counter(_question_type_of(q) or "未识别题型" for q in pulled_questions if q.get("_ai_generated"))
    for record in api_type_records:
        qtype = record.get("题型", "")
        record["AI补题数"] = int(ai_counts.get(qtype, 0))
        selected_total = int(record.get("API入选题数", 0) or 0) + int(record.get("AI补题数", 0) or 0)
        record["最终入选题数"] = selected_total
        record["最终缺题数"] = max(0, int(record.get("计划题数", 0) or 0) - selected_total)

    candidates_path = _write_api_trace_json(
        paper,
        "清洗后候选池和最终入选题.json",
        {
            "试卷": paper.paper_label,
            "说明": "这里记录本卷 API/AI 进入最终组卷前后的结构化信息，方便回溯缺题或转换问题。",
            "计划题数": total_needed,
            "目标题数": target_total,
            "分值倍率": score_multiplier,
            "按题型统计": api_type_records,
            "最终题数": len(pulled_questions),
            "最终题目": pulled_questions,
        },
    )

    _update_run_record(
        paper,
        **{
            "API拉题记录": {
                "一句话说明": "系统已经按题型和难度记录 API、AI 补题的命中情况；每次 API 原始响应也已落盘。",
                "是否调用API": True,
                "courseId": course_id,
                "自身kpointIds数量": len(own_kpoint_ids),
                "API搜索kpointIds数量": len(api_kpoint_ids),
                "计划题数": total_needed,
                "目标题数": target_total,
                "实际凑到题数": len(pulled_questions),
                "最终题目来源统计": _count_questions_by_type_and_source(pulled_questions),
                "按题型拉题明细": api_type_records,
                "清洗后候选池和最终入选题文件": str(candidates_path),
                "API原始结果目录": str(_api_trace_dir(paper)),
            }
        },
    )

    # --- 5. 保存 ---
    # 判断题答案统一为 √ / ×
    normalize_all_judge_answers(pulled_questions)

    # --- 6. 组装 DOCX ---
    docx_path = _save_pulled_questions_as_docx(paper, pulled_questions, score_multiplier)
    if docx_path:
        _update_run_record(
            paper,
            **{
                "API拉题产物": {
                    "一句话说明": "API/AI 凑题完成后，系统已经生成一份待质检 Word，后续流程会继续拆题、质检、修复和输出。",
                    "待质检DOCX路径": str(docx_path),
                    "题量减半分值倍率": score_multiplier,
                }
            },
        )
        if score_multiplier > 1:
            print(f"  已保存 DOCX：{docx_path}（题量减半，每题分值×{score_multiplier}）")
        else:
            print(f"  已保存 DOCX：{docx_path}")
        _dt_pull = time.perf_counter() - _t_pull_start
        print(f"  [拉题总耗时] {paper.paper_label}：{_dt_pull:.1f}s")
        return True
    _dt_pull = time.perf_counter() - _t_pull_start
    print(f"  [拉题总耗时] {paper.paper_label}：{_dt_pull:.1f}s（失败）")
    return False


def _download_image(url: str) -> Path | None:
    """从 URL 下载图片到系统临时目录，返回本地路径。"""
    if not url or not url.startswith("http"):
        return None
    try:
        import tempfile
        import urllib.request
        from pathlib import Path
        suffix = Path(url.split("?")[0]).suffix or ".png"
        tmp = Path(tempfile.gettempdir()) / f"xkw_img_{hash(url) & 0x7FFFFFFF}{suffix}"
        if tmp.exists():
            return tmp
        urllib.request.urlretrieve(url, str(tmp))
        return tmp
    except Exception:
        return None


def _safe_img_dim(value, default=200):
    """将图片尺寸值转为安全数值。"""
    try:
        v = int(value)
        return v if 10 <= v <= 2000 else default
    except (TypeError, ValueError):
        return default


def _save_pulled_questions_as_docx(paper: PaperPlan, questions: list[dict], score_multiplier: int = 1) -> Path | None:
    """将 API 拉取的题目组装成 DOCX，存入 组卷待质检/{省份} {考类}/ 目录。
    score_multiplier: 分值倍率（题量减半时=2，正常=1）。
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm, Inches
    from 学科网API拉题移植版.html_content_converter import (
        convert_stem_html, convert_answer_html, convert_explanation_html,
    )

    manual_dir = manual_paper_dir_for_meta(paper.meta) if paper.meta else MANUAL_PAPER_DIR
    manual_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(10.5)

    # 分值倍率提示（题量减半时）
    if score_multiplier > 1:
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run(f"【注意】本卷题量减半，每题分值×{score_multiplier}，总分仍为100分")
        run.font.size = Pt(9)
        run.font.color.rgb = None  # 保持默认色
        run.italic = True

    CN = "一二三四五六七八九十"

    # 统一规范化所有题目的 _question_type
    for q in questions:
        if q.get("_question_type"):
            q["_question_type"] = normalize_question_type(q["_question_type"])
        if q.get("question_type"):
            q["question_type"] = normalize_question_type(q["question_type"])

    # 按题型分组输出（动态序号，连续不跳号；仅输出有题目的题型）
    type_order = _QUESTION_TYPE_ORDER
    _TYPE_NAME_MAP = {t: type_display_name(t) for t in type_order}

    qno = 0
    heading_idx = 0
    for qtype in type_order:
        typed_questions = [q for q in questions if q.get("_question_type") == qtype]
        if not typed_questions:
            continue

        # 题型标题（动态序号）
        number = CN[heading_idx] if heading_idx < len(CN) else str(heading_idx + 1)
        type_name = _TYPE_NAME_MAP.get(qtype, qtype)
        heading_text = f"{number}、{type_name}"
        heading_idx += 1

        heading_para = doc.add_paragraph(heading_text)
        heading_para.style.font.bold = True if heading_para.runs else None
        for run in heading_para.runs:
            run.bold = True

        for q in typed_questions:
            qno += 1
            q["_question_no"] = qno

            # 题干（HTML→纯文本，同时提取选项和图片）
            stem_raw = q.get("stem", "") or ""
            stem_text, stem_images, parsed_options = convert_stem_html(stem_raw) if stem_raw else ("", [], None)
            stem = stem_text or stem_raw
            stem_para = doc.add_paragraph(f"{qno}．{stem}")

            # 题干图片
            if stem_images:
                for img in stem_images:
                    img_path = _download_image(img.get("url") or "")
                    if img_path:
                        try:
                            run = stem_para.add_run()
                            w = _safe_img_dim(img.get("width"), 240)
                            h = _safe_img_dim(img.get("height"), 180)
                            run.add_picture(str(img_path), width=Cm(min(w, 10)), height=Cm(min(h, 8)))
                        except Exception:
                            pass

            # 选项：优先使用 convert_stem_html 解析出的结构化选项
            if parsed_options:
                for opt in parsed_options:
                    label = opt.get("label", "")
                    text = opt.get("text", "")
                    opt_imgs = opt.get("images") or []
                    if text or opt_imgs:
                        opt_para = doc.add_paragraph(f"{label}. {text}" if text else f"{label}.")
                        for img in opt_imgs:
                            img_path = _download_image(img.get("url") or "")
                            if img_path:
                                try:
                                    run = opt_para.add_run()
                                    w = _safe_img_dim(img.get("width"), 200)
                                    h = _safe_img_dim(img.get("height"), 150)
                                    run.add_picture(str(img_path), width=Cm(min(w, 8)), height=Cm(min(h, 6)))
                                except Exception:
                                    pass
            else:
                options = q.get("options") or q.get("option") or []
                if isinstance(options, str):
                    import re
                    options = re.split(r"(?=[A-H][.、．)])", options)
                    options = [o.strip() for o in options if o.strip()]
                if isinstance(options, list):
                    for opt in options:
                        opt_text = str(opt).strip()
                        if opt_text:
                            doc.add_paragraph(opt_text)

            # 空行
            doc.add_paragraph("")

            # 答案和解析（HTML→纯文本）
            answer_raw = q.get("answer", "") or ""
            answer = convert_answer_html(answer_raw) if answer_raw else ""
            analysis_raw = q.get("explanation", "") or q.get("analysis", "") or ""
            analysis = convert_explanation_html(analysis_raw) if analysis_raw else ""
            # 回写转换后的值，供后续 paper_loader 拆题使用
            q["answer"] = answer
            q["analysis"] = analysis
            q["stem"] = stem
            if answer:
                ans_para = doc.add_paragraph(f"【答案】{answer}")
                for run in ans_para.runs:
                    run.font.color.rgb = None  # 黑色
            if analysis:
                ana_para = doc.add_paragraph(f"【详解】{analysis}")

            doc.add_paragraph("")

    # 保存到 组卷待质检
    safe_module = "".join(c for c in (paper.module or "") if c.isalnum() or c in " _-")
    safe_point = "".join(c for c in (paper.point_name or paper.topic or paper.paper_label) if c.isalnum() or c in " _-")
    filename = f"{paper.paper_label} {safe_module} {safe_point}.docx"[:200]
    output_path = manual_dir / filename
    doc.save(str(output_path))
    return output_path


def main(mode: int = 1) -> None:
    args = parse_args()
    ensure_output_dirs()

    plan_path = Path(args.plan) if args.plan else find_default_plan()
    if args.preview:
        print_dry_run_summary(plan_path)
        return

    meta, rows, paper_index = load_planning_workbook(plan_path)

    if args.status:
        _print_run_status(meta, paper_index)
        return

    # 如果规划表中未填写考试科目表名称，交互询问
    if not meta.exam_table_title:
        default_title = f"{meta.province}2025年中等职业学校毕业生进入普通高校学习专业基础课和专业课考试科目表"
        user_input = input(
            f"\n请输入考试科目表名称（直接回车使用默认值）：\n"
            f"  默认：{default_title}\n"
            f"  > "
        ).strip()
        meta.exam_table_title = user_input or default_title

    # 考试类型：已设置则确认，未设置则选择
    if meta.exam_type:
        print(f"\n当前考试类型：{meta.exam_type}")
        keep = input(f"  是否继续使用？（Y/n，n 则重新选择）：").strip().lower()
        if keep == "n":
            meta.exam_type = ""
    if not meta.exam_type:
        print(f"\n请选择考试类型（将用于 DOCX 文件名和编写说明）：")
        print(f"  1. 高职分类考试")
        print(f"  2. 对口招生")
        print(f"  3. 春季高考")
        print(f"  4. 普通高考")
        print(f"  5. 职教高考")
        print(f"  6. 其他（自定义）")
        while True:
            choice = input(f"  > ").strip()
            if choice == "1":
                meta.exam_type = "高职分类考试"
                break
            elif choice == "2":
                meta.exam_type = "对口招生"
                break
            elif choice == "3":
                meta.exam_type = "春季高考"
                break
            elif choice == "4":
                meta.exam_type = "普通高考"
                break
            elif choice == "5":
                meta.exam_type = "职教高考"
                break
            elif choice == "6":
                custom = input(f"  请输入自定义考试类型名称：").strip()
                if custom:
                    meta.exam_type = custom
                    break
                print(f"  名称不能为空，请重新输入")
            else:
                print(f"  无效选择，请输入 1-6")

    # 回写 xlsx，避免下次重复询问
    _write_meta_to_xlsx(plan_path, meta)

    if args.list:
        print_paper_list(paper_index)
        return

    if args.rerun_failed or args.rerun_missing:
        selected_papers = _select_batch_papers(
            paper_index,
            rerun_failed=args.rerun_failed,
            rerun_missing=args.rerun_missing,
        )
        kind = "失败卷" if args.rerun_failed else "缺失卷"
        if not selected_papers:
            print(f"未找到可重跑的{kind}。")
            return
        print(f"\n自动重跑{kind}：{_format_paper_labels([paper.paper_no for paper in selected_papers], limit=30)}")
        generate_selected_papers(selected_papers, args)
        return

    selector = args.paper or args.paper_range
    exit_code = 0
    while True:
        _print_generation_status(paper_index, meta)
        selected_numbers = parse_sequence_selector(selector, paper_index.keys()) if selector else prompt_for_sequence(paper_index)
        selector = None
        if not selected_numbers:
            print("未选择卷号，已退出。")
            break

        selected_papers, missing = resolve_selected_papers(selected_numbers, paper_index)
        preflight = build_preflight_result(meta, paper_index, selected_numbers, selected_papers, missing)
        print_preflight_dialog(preflight, meta, paper_index)

        if _preflight_has_blocking_issues(preflight, mode) and not _confirm_continue_after_preflight():
            if not _ask_continue_generation():
                break
            continue

        if selected_papers:
            # 模式 2：先 API 拉题，生成 DOCX 到 组卷待质检
            if mode == 2:
                print("\n" + "=" * 60)
                print("  【模式 2】第一步：API 拉题 → 生成 DOCX")
                print("=" * 60)
                _run_api_pull_phase(selected_papers, meta)
                print("=" * 60)

            try:
                generate_selected_papers(selected_papers, args)
            except SystemExit as exc:
                exit_code = exit_code or int(exc.code or 0)

        if not _ask_continue_generation():
            break

    if exit_code:
        raise SystemExit(exit_code)


if __name__ == "__main__":
    main()
